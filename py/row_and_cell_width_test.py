"""
Regression test for 2 real follow-up bugs found after applying the
"check every override level" lesson (registered in
CLAUDE_INSTRUCTIONS.md) to the WIDTH property specifically, after it
had already been applied once to INDENT (the tblPrEx/tblInd fix):

  1. _fix_ambiguous_table_width originally only resolved the
     TABLE-level <w:tblPr><w:tblW w:type="auto" w:w="0"/> ambiguity. A
     real user reported tables with a correctly-fixed LEFT position
     still showing right-edge content cutoff in real Word. Direct XML
     inspection of the real reported document found EVERY single
     <w:tblPrEx> (95 of them) ALSO carried the identical ambiguous
     tblW - the same "auto-size" vs "fixed columns" contradiction, just
     at the per-ROW level, which the original fix never touched.

  2. A comprehensive, real scan of every cell in every table in that
     same document found 2 cells (one row) whose own w:tcW didn't
     match the table's own w:gridCol at that position - a real, if
     more cosmetic, inconsistency (the row's own total still summed
     correctly, so this specific case wasn't the primary overflow
     cause, but it's still a genuine mismatch worth correcting).

Run: python3 row_and_cell_width_test.py
"""
import os
import sys

sys.path.insert(0, os.path.dirname(__file__))
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import aspose_test_pipeline as asp


def assert_(cond, label):
    if not cond:
        print("FAIL:", label)
        sys.exit(1)
    print("PASS:", label)


def _add_row_tblPrEx_width(table, row_idx, width_w="0", width_type="auto"):
    tr = table.rows[row_idx]._tr
    tblPrEx = OxmlElement("w:tblPrEx")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), width_w)
    tblW.set(qn("w:type"), width_type)
    tblPrEx.append(tblW)
    tr.insert(0, tblPrEx)
    return tblW


def run():
    print("=== Test 1 (THE REAL REPORTED CASE): real document - row-level tblPrEx width ambiguity resolved ===")
    real_docx = "/home/claude/work/width_check/final_output5.docx"
    if os.path.exists(real_docx):
        doc = Document(real_docx)
        fixed = asp._fix_ambiguous_table_width(doc)
        assert_(fixed > 0, "A real number of ambiguous widths (table + row level) were resolved (got " + str(fixed) + ")")

        remaining = 0
        for tbl_el in doc.element.body.iter(qn("w:tbl")):
            tblPr = tbl_el.find(qn("w:tblPr"))
            tblW = tblPr.find(qn("w:tblW")) if tblPr is not None else None
            if tblW is not None and tblW.get(qn("w:type")) == "auto" and tblW.get(qn("w:w")) in ("0", None):
                remaining += 1
            for tr_el in tbl_el.findall(qn("w:tr")):
                tblPrEx = tr_el.find(qn("w:tblPrEx"))
                if tblPrEx is None:
                    continue
                row_tblW = tblPrEx.find(qn("w:tblW"))
                if row_tblW is not None and row_tblW.get(qn("w:type")) == "auto" and row_tblW.get(qn("w:w")) in ("0", None):
                    remaining += 1
        assert_(remaining == 0, "Zero ambiguous widths remain anywhere (table or row level)")
    else:
        print("  (real document not present in this environment - skipped)")

    print("\n=== Test 2 (synthetic, exact mechanics): row-level tblPrEx width ambiguity resolved to the table's real width ===")
    doc2 = Document()
    t2 = doc2.add_table(rows=2, cols=2)
    grid2 = t2._tbl.find(qn("w:tblGrid"))
    cols2 = grid2.findall(qn("w:gridCol"))
    cols2[0].set(qn("w:w"), "3000")
    cols2[1].set(qn("w:w"), "4000")
    bad_row_tblW = _add_row_tblPrEx_width(t2, 1)

    fixed2 = asp._fix_ambiguous_table_width(doc2)
    assert_(fixed2 == 2, "Both the table-level AND the row-level ambiguous width were fixed (got " + str(fixed2) + ")")
    assert_(bad_row_tblW.get(qn("w:type")) == "dxa", "The row-level tblPrEx width type is now 'dxa' (explicit)")
    assert_(bad_row_tblW.get(qn("w:w")) == "7000", "The row-level tblPrEx width is now the table's own real column-width sum (7000)")

    print("\n=== Test 3: an already-explicit row-level width is left untouched (table-level default also explicitly fixed first, to isolate the row-level check) ===")
    doc3 = Document()
    t3 = doc3.add_table(rows=2, cols=1)
    grid3 = t3._tbl.find(qn("w:tblGrid"))
    grid3.findall(qn("w:gridCol"))[0].set(qn("w:w"), "5000")
    # python-docx's own default table already has an ambiguous
    # table-level tblW (type="auto", w="0") - make it explicit first,
    # so this test isolates the ROW-level behavior specifically.
    tblPr3 = t3._tbl.find(qn("w:tblPr"))
    tblPr3.find(qn("w:tblW")).set(qn("w:type"), "dxa")
    tblPr3.find(qn("w:tblW")).set(qn("w:w"), "5000")
    good_row_tblW = _add_row_tblPrEx_width(t3, 1, width_w="5000", width_type="dxa")
    fixed3 = asp._fix_ambiguous_table_width(doc3)
    assert_(fixed3 == 0, "Zero fixes - both table-level and row-level widths were already explicit and correct")
    assert_(good_row_tblW.get(qn("w:w")) == "5000", "Unchanged")

    print("\n=== Test 4 (THE REAL REPORTED CASE): cell-width-vs-column mismatch fix, real document ===")
    if os.path.exists(real_docx):
        doc4 = Document(real_docx)
        fixed4 = asp._fix_cell_width_vs_column_mismatch(doc4)
        assert_(fixed4 == 2, "Exactly 2 cell-width mismatches fixed (matches the real confirmed count)")

        remaining_mismatches = 0
        for tbl_el in doc4.element.body.iter(qn("w:tbl")):
            grid = tbl_el.find(qn("w:tblGrid"))
            cols = grid.findall(qn("w:gridCol")) if grid is not None else []
            col_widths = [c.get(qn("w:w")) for c in cols]
            for tr_el in tbl_el.findall(qn("w:tr")):
                cells = tr_el.findall(qn("w:tc"))
                if len(cells) != len(col_widths):
                    continue
                for tc, expected in zip(cells, col_widths):
                    tcPr = tc.find(qn("w:tcPr"))
                    tcW = tcPr.find(qn("w:tcW")) if tcPr is not None else None
                    if tcW is not None and tcW.get(qn("w:w")) != expected:
                        remaining_mismatches += 1
        assert_(remaining_mismatches == 0, "Zero cell-vs-column mismatches remain")
    else:
        print("  (real document not present in this environment - skipped)")

    print("\n=== Test 5 (synthetic, exact mechanics): a cell whose width doesn't match its column is corrected ===")
    doc5 = Document()
    t5 = doc5.add_table(rows=1, cols=2)
    grid5 = t5._tbl.find(qn("w:tblGrid"))
    cols5 = grid5.findall(qn("w:gridCol"))
    cols5[0].set(qn("w:w"), "3000")
    cols5[1].set(qn("w:w"), "4000")
    # python-docx's own add_table() already gives every cell a default
    # w:tcW - modify the EXISTING one directly (not append a second,
    # duplicate tcW) to set up the real mismatch scenario.
    cell0_tcPr = t5.rows[0].cells[0]._tc.get_or_add_tcPr()
    tcW0 = cell0_tcPr.find(qn("w:tcW"))
    tcW0.set(qn("w:w"), "6000")  # wrong - should be 3000
    tcW0.set(qn("w:type"), "dxa")
    # cell1 already correctly has its own default tcW - give it the
    # real, correct value directly so it does NOT count as a fix.
    cell1_tcPr = t5.rows[0].cells[1]._tc.get_or_add_tcPr()
    cell1_tcPr.find(qn("w:tcW")).set(qn("w:w"), "4000")

    fixed5 = asp._fix_cell_width_vs_column_mismatch(doc5)
    assert_(fixed5 == 1, "Exactly 1 cell fixed (got " + str(fixed5) + ")")
    assert_(tcW0.get(qn("w:w")) == "3000", "The cell's width now matches its real column (3000)")

    print("\n=== Test 6: a row with a genuinely merged/spanned cell (gridSpan) is never touched (1:1 comparison would be meaningless) ===")
    doc6 = Document()
    t6 = doc6.add_table(rows=1, cols=1)
    t6.add_column(width=1)  # make it 2 real grid columns total
    grid6 = t6._tbl.find(qn("w:tblGrid"))
    cols6 = grid6.findall(qn("w:gridCol"))
    for c, w in zip(cols6, ["2000", "3000"]):
        c.set(qn("w:w"), w)
    merged_cell = t6.rows[0].cells[0]
    tcPr6 = merged_cell._tc.get_or_add_tcPr()
    span = OxmlElement("w:gridSpan")
    span.set(qn("w:val"), "2")
    tcPr6.append(span)
    tcW6 = OxmlElement("w:tcW")
    tcW6.set(qn("w:w"), "5000")  # correctly spans both columns combined
    tcW6.set(qn("w:type"), "dxa")
    tcPr6.append(tcW6)

    fixed6 = asp._fix_cell_width_vs_column_mismatch(doc6)
    assert_(fixed6 == 0, "Zero fixes - a merged/spanned row is correctly skipped, not force-matched to a single column")

    print("\nALL TESTS PASSED")


if __name__ == "__main__":
    run()
