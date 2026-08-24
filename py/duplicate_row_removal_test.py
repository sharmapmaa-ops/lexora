"""
Regression test for _remove_duplicate_table_rows: confirmed real bug
(a real Appendix table had one row - "1 | Contract Execution Date |
This field indicates the date of documenting the lease contract by
both parties." - appearing twice, character-for-character identical)
and a confirmed real false-positive risk this fix must NOT trigger on
(the SAME real document also had "2 | Lessor | <clarification A>"
immediately followed by "2 | Lessor | <clarification B>" - same Field
name, genuinely DIFFERENT Clarification text - real, intentional
content that must never be removed).

Run: python3 duplicate_row_removal_test.py
"""
import os
import sys

sys.path.insert(0, os.path.dirname(__file__))
from docx import Document
import aspose_test_pipeline as asp


def assert_(cond, label):
    if not cond:
        print("FAIL:", label)
        sys.exit(1)
    print("PASS:", label)


def _add_row(table, cell_texts):
    row = table.add_row()
    for cell, text in zip(row.cells, cell_texts):
        cell.text = text


def run():
    doc = Document()

    print("=== Test 1: exact adjacent duplicate row IS removed (the real reported bug) ===")
    t1 = doc.add_table(rows=1, cols=3)
    t1.rows[0].cells[0].text = "Clause Number"
    t1.rows[0].cells[1].text = "Field"
    t1.rows[0].cells[2].text = "Clarification"
    _add_row(t1, ["1", "Contract Type", "The type of lease contract shall be one of the following: 1. New. 2. Renewed"])
    _add_row(t1, ["1", "Contract Execution Date", "This field indicates the date of documenting the lease contract by both parties."])
    _add_row(t1, ["1", "Contract Execution Date", "This field indicates the date of documenting the lease contract by both parties."])
    _add_row(t1, ["1", "Lease Commencement Date", "This is the actual commencement date of the lease and utilization of the rental unit."])
    before_rows = len(t1.rows)
    removed = asp._remove_duplicate_table_rows(doc)
    after_rows = len(doc.tables[0].rows)
    assert_(removed == 1, "Exactly 1 duplicate row removed (got " + str(removed) + ")")
    assert_(after_rows == before_rows - 1, "Table row count decreased by exactly 1 (" + str(before_rows) + " -> " + str(after_rows) + ")")
    remaining_texts = [c.text for c in doc.tables[0].rows[2].cells]
    assert_(remaining_texts == ["1", "Contract Execution Date", "This field indicates the date of documenting the lease contract by both parties."],
            "The surviving row still has the correct, real content")

    print("\n=== Test 2: the REAL false-positive case - same Field, different Clarification - is NEVER removed ===")
    doc2 = Document()
    t2 = doc2.add_table(rows=1, cols=3)
    t2.rows[0].cells[0].text = "Clause Number"
    t2.rows[0].cells[1].text = "Field"
    t2.rows[0].cells[2].text = "Clarification"
    _add_row(t2, ["2", "Lessor", "The lessor may be an individual, an establishment, or a company."])
    _add_row(t2, ["2", "Lessor", "The lessor is an individual having one of the following statuses: 1. Either representing himself. 2. Or being an agent under a legal power of attorney. 3. Or being the legal representative of the commercial establishment."])
    before_rows2 = len(t2.rows)
    removed2 = asp._remove_duplicate_table_rows(doc2)
    after_rows2 = len(doc2.tables[0].rows)
    assert_(removed2 == 0, "Zero rows removed - same Field name but different Clarification text is real, distinct content")
    assert_(after_rows2 == before_rows2, "Table row count is completely unchanged (" + str(before_rows2) + ")")

    print("\n=== Test 3: non-adjacent identical rows are NOT removed (adjacency is required) ===")
    doc3 = Document()
    t3 = doc3.add_table(rows=1, cols=2)
    t3.rows[0].cells[0].text = "H1"
    t3.rows[0].cells[1].text = "H2"
    _add_row(t3, ["3", "Same value, legitimately repeated far apart"])
    _add_row(t3, ["4", "Something else entirely in between"])
    _add_row(t3, ["3", "Same value, legitimately repeated far apart"])
    removed3 = asp._remove_duplicate_table_rows(doc3)
    assert_(removed3 == 0, "Non-adjacent identical rows are left alone - only immediate repeats are treated as the accidental-duplication pattern")

    print("\n=== Test 4: three consecutive identical rows -> only the extras are removed, one survives ===")
    doc4 = Document()
    t4 = doc4.add_table(rows=1, cols=2)
    t4.rows[0].cells[0].text = "H1"
    t4.rows[0].cells[1].text = "H2"
    for _ in range(3):
        _add_row(t4, ["5", "Triple-written by accident"])
    removed4 = asp._remove_duplicate_table_rows(doc4)
    after4 = len(doc4.tables[0].rows)
    assert_(removed4 == 2, "Both extra copies removed from a triple-duplicate (got " + str(removed4) + ")")
    assert_(after4 == 2, "Exactly one real copy survives (header + 1 data row = 2 total, got " + str(after4) + ")")

    print("\nALL TESTS PASSED")


if __name__ == "__main__":
    run()
