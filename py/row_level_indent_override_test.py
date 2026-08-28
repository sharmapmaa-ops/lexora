"""
Regression test for the ACTUAL root cause of a real reported bug: a
table showing 50%+ left margin in real MS Word screenshots, even
though every check of the table's own <w:tblPr><w:tblInd> showed a
correct value (771 twips). The real cause, found only after the user
pushed back multiple times and asked what method was actually used to
inspect the file: 37 of that table's ROWS carried their own
<w:tblPrEx><w:tblInd w:w="6766"/> (or 6767) - a separate OOXML
mechanism (Table Property Exceptions, one per <w:tr>) that overrides
the parent table's properties for that specific row. 6766 twips is
~56.8% of an 11900-twip page width - exactly matching the reported
visual bug. LibreOffice does not appear to honor this row-level
override (renders "fine"), which is why earlier LibreOffice-based
verification never caught it.

This incident and its lesson (checking only the top-level/obvious
property location is not enough - a structured format can override
the same property at a more specific level) are registered in
CLAUDE_INSTRUCTIONS.md.

Run: python3 row_level_indent_override_test.py
"""
import os
import sys

sys.path.insert(0, os.path.dirname(__file__))
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Twips
import aspose_test_pipeline as asp


def assert_(cond, label):
    if not cond:
        print("FAIL:", label)
        sys.exit(1)
    print("PASS:", label)


def _add_row_tblPrEx(table, row_idx, indent_twips, also_tblW=True):
    tr = table.rows[row_idx]._tr
    tblPrEx = OxmlElement("w:tblPrEx")
    if also_tblW:
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), "0")
        tblW.set(qn("w:type"), "auto")
        tblPrEx.append(tblW)
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), str(indent_twips))
    tblInd.set(qn("w:type"), "dxa")
    tblPrEx.append(tblInd)
    tr.insert(0, tblPrEx)
    return tblPrEx


def run():
    print("=== Test 1 (THE REAL REPORTED CASE): real document - row-level tblPrEx overrides are found and removed ===")
    real_docx = "/home/claude/work/tblprex_check/final_output.docx"
    if os.path.exists(real_docx):
        doc = Document(real_docx)
        table_fixed = asp._fix_table_overflow_indent(doc)
        row_fixed = asp._fix_row_level_table_indent_override(doc)
        assert_(table_fixed > 0, "Table-level indents were normalized first (got " + str(table_fixed) + ")")
        assert_(row_fixed > 0, "A real number of row-level tblInd overrides were removed (got " + str(row_fixed) + ")")

        # Confirm zero MISMATCHED row-level overrides remain (matching
        # ones are fine to leave, per the function's own documented
        # behavior).
        mismatched_remaining = 0
        for tbl_el in doc.element.body.iter(qn("w:tbl")):
            tblPr = tbl_el.find(qn("w:tblPr"))
            table_tblInd = tblPr.find(qn("w:tblInd")) if tblPr is not None else None
            table_val = table_tblInd.get(qn("w:w")) if table_tblInd is not None else "0"
            for tr_el in tbl_el.findall(qn("w:tr")):
                tblPrEx = tr_el.find(qn("w:tblPrEx"))
                if tblPrEx is None:
                    continue
                row_ind = tblPrEx.find(qn("w:tblInd"))
                if row_ind is not None and row_ind.get(qn("w:w")) != table_val:
                    mismatched_remaining += 1
        assert_(mismatched_remaining == 0, "Zero mismatched row-level tblInd overrides remain anywhere in the document")
    else:
        print("  (real document not present in this environment - skipped)")

    print("\n=== Test 2 (synthetic, exact mechanics): the exact reported numbers - 6766 twips row override on a 771-twip table ===")
    doc2 = Document()
    sec2 = doc2.sections[0]
    sec2.page_width = Twips(11900)
    sec2.left_margin = Twips(540)
    sec2.right_margin = Twips(500)
    t2 = doc2.add_table(rows=2, cols=1)
    tblPr2 = t2._tbl.find(qn("w:tblPr"))
    tblInd2 = OxmlElement("w:tblInd")
    tblInd2.set(qn("w:w"), "771")
    tblInd2.set(qn("w:type"), "dxa")
    tblPr2.append(tblInd2)
    bad_tblPrEx = _add_row_tblPrEx(t2, 1, 6766)

    fixed2 = asp._fix_row_level_table_indent_override(doc2)
    assert_(fixed2 == 1, "Exactly 1 row-level override removed")
    assert_(bad_tblPrEx.find(qn("w:tblInd")) is None, "The bad tblInd is gone from that row's tblPrEx")

    print("\n=== Test 3: a tblPrEx WITHOUT tblInd (only tblW/tblLayout) is left completely alone ===")
    doc3 = Document()
    t3 = doc3.add_table(rows=2, cols=1)
    tblPr3 = t3._tbl.find(qn("w:tblPr"))
    tblInd3 = OxmlElement("w:tblInd")
    tblInd3.set(qn("w:w"), "867")
    tblPr3.append(tblInd3)
    tr3 = t3.rows[1]._tr
    harmless_tblPrEx = OxmlElement("w:tblPrEx")
    harmless_tblW = OxmlElement("w:tblW")
    harmless_tblW.set(qn("w:w"), "5000")
    harmless_tblW.set(qn("w:type"), "dxa")
    harmless_tblPrEx.append(harmless_tblW)
    tr3.insert(0, harmless_tblPrEx)

    fixed3 = asp._fix_row_level_table_indent_override(doc3)
    assert_(fixed3 == 0, "Zero fixes - this tblPrEx never touched indent at all, so it's untouched")
    assert_(harmless_tblPrEx.find(qn("w:tblW")) is not None, "The unrelated tblW property inside tblPrEx survives completely intact")

    print("\n=== Test 4: a row-level override that ALREADY matches the table's own indent is left alone (no pointless edit) ===")
    doc4 = Document()
    t4 = doc4.add_table(rows=2, cols=1)
    tblPr4 = t4._tbl.find(qn("w:tblPr"))
    tblInd4 = OxmlElement("w:tblInd")
    tblInd4.set(qn("w:w"), "867")
    tblPr4.append(tblInd4)
    already_matching = _add_row_tblPrEx(t4, 1, 867, also_tblW=False)

    fixed4 = asp._fix_row_level_table_indent_override(doc4)
    assert_(fixed4 == 0, "Zero fixes - the row-level override already equals the table's own correct value")
    assert_(already_matching.find(qn("w:tblInd")) is not None, "The already-correct override is left in place, not stripped needlessly")

    print("\n=== Test 5: tblPrEx that becomes completely empty after removing tblInd is itself removed (no leftover clutter) ===")
    doc5 = Document()
    t5 = doc5.add_table(rows=2, cols=1)
    tblPr5 = t5._tbl.find(qn("w:tblPr"))
    tblInd5 = OxmlElement("w:tblInd")
    tblInd5.set(qn("w:w"), "771")
    tblPr5.append(tblInd5)
    only_indent_tblPrEx = _add_row_tblPrEx(t5, 1, 6766, also_tblW=False)
    tr5 = t5.rows[1]._tr

    fixed5 = asp._fix_row_level_table_indent_override(doc5)
    assert_(fixed5 == 1, "Exactly 1 fix applied")
    assert_(tr5.find(qn("w:tblPrEx")) is None, "The now-empty tblPrEx element itself was removed, not left as empty clutter")

    print("\nALL TESTS PASSED")


if __name__ == "__main__":
    run()
