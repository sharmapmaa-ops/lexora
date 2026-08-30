"""
Lexora - ASPOSE CLOUD PIPELINE (originally isolated/experimental,
now also the real Translation service's Aspose path)
==============================================================

This does NOT touch translate_pipeline.py or lease_engine.py's real
translation flow. Originally reachable only via its own admin-only test
route (/api/test/aspose-translate); run_full_test() is now ALSO called
directly by the real, user-facing Translation service for documents
with tables/background color (see server.py's
_handle_translation_process_aspose and py/ocr_router.py's
analyze_pdf_for_ocr_strategy, which decides when that applies) - one
Aspose+translate pipeline, not two, reused rather than duplicated for
the real service once it was confirmed working via the admin test
route. js/engine-*.js's browser-side pdf.js pipeline remains untouched
and is still the path for documents WITHOUT tables/background.

Needs NO self-hosting - Aspose runs the actual service, we just call it
with a Client Id/Secret. Genuinely just `pip install aspose-words-cloud
asposepdfcloud pycryptodome` (all three now in requirements.txt) and two
environment variables.

TWO TEST MODES, both useful for different questions:

  1. run_structure_only_test() - PDF -> DOCX via Aspose.Words Cloud's
     own native conversion, NO translation involved at all. This is the
     fastest way to answer "does Aspose's own table/format detection
     actually look better than ours" - zero LLM cost, zero extra
     moving parts, pure format-fidelity comparison.

  2. run_full_test() - extract text (Aspose.PDF Cloud, or our own
     pdfplumber extractor as a fallback), translate via Lexora's
     existing LLM call, then rebuild via Aspose.Words Cloud. This tests
     the REAL end-to-end scenario (translated content, not source-
     language content) but costs one real LLM call per run.

SETUP NEEDED (env vars - never hardcode these, never paste them in chat)
---------------------------------------------------------------------
  ASPOSE_CLIENT_ID
  ASPOSE_CLIENT_SECRET
Get both from https://dashboard.aspose.cloud -> Applications -> (create
an app) -> Client Id / Client Secret. Free tier: 150 API calls/month.
"""

# Item (DEPLOYMENT-VERSION-VERIFICATION) - confirmed real incident: a
# production run's clause-number reversal was 48/49 wrong, while the
# exact same function, called directly on byte-identical text extracted
# from that SAME production file, correctly returned the right answer
# every single time in this sandbox. Paragraph-direction (bidi) WAS
# being cleared correctly throughout that same file, proving
# _fix_paragraph_direction was genuinely running - so the only
# remaining explanation is that the code actually deployed in
# production differs from what was verified here. This constant exists
# so that can be checked DIRECTLY from a run's own log output, instead
# of guessing: bump it with every delivered change, and confirm it
# appears in the "Configured: ..." log line of a fresh test run before
# treating that run's results as reflecting current code.
PIPELINE_CODE_VERSION = "2026-08-18-v32-standalone-cell-padding-cap"

import os
import io
import re
import time

import lease_engine as le

try:
    from asposewordscloud import WordsApi
    from asposewordscloud.models.requests import ConvertDocumentRequest
    _WORDS_SDK_AVAILABLE = True
except ImportError:
    _WORDS_SDK_AVAILABLE = False

try:
    from asposepdfcloud import PdfApi
    from asposepdfcloud.api_client import ApiClient as _PdfApiClient
    _PDF_SDK_AVAILABLE = True
except ImportError:
    _PDF_SDK_AVAILABLE = False

try:
    import pdfplumber
    import pypdfium2 as pdfium
    _SIGNATURE_EXTRACT_AVAILABLE = True
except ImportError:
    _SIGNATURE_EXTRACT_AVAILABLE = False


ASPOSE_CLIENT_ID = os.environ.get("ASPOSE_CLIENT_ID", "")
ASPOSE_CLIENT_SECRET = os.environ.get("ASPOSE_CLIENT_SECRET", "")


class AsposeNotConfiguredError(Exception):
    """Distinct from a real API failure, so the test UI can show 'you
    still need to set these env vars' instead of a generic error."""
    pass


def is_configured():
    return bool(ASPOSE_CLIENT_ID and ASPOSE_CLIENT_SECRET and _WORDS_SDK_AVAILABLE)


def _require_configured():
    if not _WORDS_SDK_AVAILABLE:
        raise AsposeNotConfiguredError(
            "aspose-words-cloud (and pycryptodome, a dependency it needs) aren't "
            "installed yet - check requirements.txt was actually installed on this deploy."
        )
    if not (ASPOSE_CLIENT_ID and ASPOSE_CLIENT_SECRET):
        raise AsposeNotConfiguredError(
            "Aspose test pipeline isn't configured yet - set ASPOSE_CLIENT_ID and "
            "ASPOSE_CLIENT_SECRET as environment variables (get them from "
            "https://dashboard.aspose.cloud -> Applications), then restart the server."
        )


def _words_api():
    _require_configured()
    return WordsApi(client_id=ASPOSE_CLIENT_ID, client_secret=ASPOSE_CLIENT_SECRET)


def _fix_incomplete_header_bar_shading(doc):
    """Item (HEADER-BAR-BACKGROUND-GAP) - Aspose's PDF->DOCX conversion
    sometimes reconstructs a table's colored 'header bar' row as a
    STANDALONE paragraph sitting directly in the document body (NOT
    inside the table at all - confirmed by walking a real converted
    document's XML and finding the header text's ancestor chain was
    document->body->p->r->t, no enclosing w:tc whatsoever), with
    per-character green shading (w:rPr/w:shd on each run) and a huge
    artificial letter-spacing value (a single trailing space run with
    w:spacing val="5100" - about 3.5 inches) apparently trying to fake a
    full-width colored bar by stretching that one space out.

    DOCX renderers don't extend background shading through that
    artificial letter-spacing gap (confirmed by rendering a real
    affected document both before and after this fix, via LibreOffice)
    - only the actual glyphs get shaded, so the bar visibly stops short
    of the real table that follows it, leaving a white gap.
    Paragraph-level shading (w:pPr/w:shd) was tried first and also
    doesn't reliably close the gap for the same reason (still bounded by
    where the line's actual rendered content - including that huge
    spacing hack - ends, not the paragraph's nominal indent box).

    The fix that actually closes the gap 100%, with zero magic numbers
    or font-metric guessing: convert the standalone paragraph into a
    real single-cell, single-row table, sized and positioned to EXACTLY
    match the width and left-indent of the next real table in the
    document (read dynamically from that table's own tblGrid/tblInd),
    with cell-level shading - which this document's OTHER header rows
    already use and which reliably fills its whole cell regardless of
    font or renderer. Verified against a real 10-page Aspose-converted
    document: 12 such standalone header-bar paragraphs found and fixed,
    every one rendering as a clean full-width bar afterward with no
    regressions to surrounding content or reading order.

    Also handles the case where MORE than just the header row was left
    standalone (confirmed in the same real document: a section's first
    data row, e.g. "Name", was ALSO a standalone paragraph between the
    header and the real w:tbl) - looks a few siblings ahead for the
    first actual table, skipping over plain (non-shaded) paragraphs, and
    bails out without guessing if it hits another shaded paragraph
    first (a different header's own bar) or doesn't find a table nearby.

    Item (MERGED-TWO-SIDED-HEADER) - confirmed real, separate bug from
    the background-gap this function was originally written for: SOME
    of these standalone header paragraphs (6 out of the document's
    header-style paragraphs, confirmed by direct inspection) hold BOTH
    the English-side label AND the (originally Arabic) number+label side
    as ONE CONTINUOUS paragraph - e.g. "Lessor Representative Data" +
    "3 Lessor Representative Data" - instead of two separate cells like
    every OTHER working header in the document (e.g. "Contract Data" |
    "1 Contract Data", a genuine 2-cell table). Wrapping the whole thing
    into a single cell (the original fix here) correctly closed the
    background-color gap but left both halves crammed together on one
    side with nothing on the right, since there's no cell boundary
    between them.

    Root cause: Aspose represents the visual gap between the two halves
    using a single run with an artificially huge w:spacing (letter-
    spacing) value (~5100, confirmed on multiple real headers) instead
    of a real table-cell boundary - so this function now detects that
    specific spacer run (spacing far larger than any normal text
    kerning would ever use) and, when found, splits the paragraph into
    TWO paragraphs at that point (discarding the spacer run itself,
    since a real cell boundary replaces its job) and builds a proper
    2-column table instead of 1, matching the structure every other
    working header in the document already uses. Headers with no such
    spacer run (genuinely single-sided) keep the original 1-cell
    behavior unchanged.

    Returns the count of paragraphs fixed, for the caller's log."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    _SPACER_LETTER_SPACING_THRESHOLD = 500  # twips - normal kerning never gets remotely this large

    def _find_spacer_run_index(paragraph_el):
        runs = paragraph_el.findall(qn("w:r"))
        for idx, r in enumerate(runs):
            rpr = r.find(qn("w:rPr"))
            if rpr is None:
                continue
            spacing_el = rpr.find(qn("w:spacing"))
            if spacing_el is None:
                continue
            val = spacing_el.get(qn("w:val"))
            if val and val.lstrip("-").isdigit() and abs(int(val)) > _SPACER_LETTER_SPACING_THRESHOLD:
                return idx
        return None

    body = doc.element.body
    children = list(body.iterchildren())
    fixed = 0

    for i, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue
        p = Paragraph(child, doc)
        runs_with_text = [r for r in p.runs if (r.text or "").strip()]
        if not runs_with_text:
            continue

        fills = set()
        ok = True
        for r in runs_with_text:
            rpr = r._element.find(qn("w:rPr"))
            if rpr is None:
                ok = False
                break
            shd = rpr.find(qn("w:shd"))
            if shd is None:
                ok = False
                break
            fill = shd.get(qn("w:fill"))
            if not fill or fill.lower() in ("auto", "ffffff"):
                ok = False
                break
            fills.add(fill)
        if not ok or len(fills) != 1:
            continue
        fill_color = fills.pop()

        # Find the next TABLE among the following siblings to borrow its
        # exact width/indent - skip over plain (non-shaded) paragraphs
        # in between, bail out (skip, never guess) on hitting another
        # shaded paragraph or running past a small lookahead.
        next_table_el = None
        for sib in children[i + 1:i + 6]:
            if sib.tag == qn("w:tbl"):
                next_table_el = sib
                break
            if sib.tag == qn("w:p"):
                sib_runs_with_text = [r for r in Paragraph(sib, doc).runs if (r.text or "").strip()]
                sib_has_shading = False
                for r in sib_runs_with_text:
                    rpr = r._element.find(qn("w:rPr"))
                    if rpr is None:
                        continue
                    shd = rpr.find(qn("w:shd"))
                    if shd is None:
                        continue
                    if (shd.get(qn("w:fill")) or "").lower() not in ("", "auto", "ffffff"):
                        sib_has_shading = True
                        break
                if sib_has_shading:
                    break  # a different header's own bar - don't borrow its table
                continue  # plain paragraph (e.g. a data row Aspose also left standalone) - keep looking
            break  # anything else unexpected - don't guess past it
        if next_table_el is None:
            continue

        tblPr = next_table_el.find(qn("w:tblPr"))
        tblInd_el = tblPr.find(qn("w:tblInd")) if tblPr is not None else None
        indent = tblInd_el.get(qn("w:w")) if tblInd_el is not None else None
        grid = next_table_el.find(qn("w:tblGrid"))
        if grid is None:
            continue
        cols = grid.findall(qn("w:gridCol"))
        if not cols:
            continue
        total_width = sum(int(c.get(qn("w:w"))) for c in cols)
        if not total_width:
            continue

        spacer_idx = _find_spacer_run_index(child)

        new_tbl = OxmlElement("w:tbl")
        tblPr_new = OxmlElement("w:tblPr")
        tblW_new = OxmlElement("w:tblW")
        tblW_new.set(qn("w:w"), str(total_width))
        tblW_new.set(qn("w:type"), "dxa")
        tblPr_new.append(tblW_new)
        if indent:
            tblInd_new = OxmlElement("w:tblInd")
            tblInd_new.set(qn("w:w"), indent)
            tblInd_new.set(qn("w:type"), "dxa")
            tblPr_new.append(tblInd_new)
        tblBorders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"), "none")
            tblBorders.append(b)
        tblPr_new.append(tblBorders)
        new_tbl.append(tblPr_new)

        body.remove(child)
        pPr_orig = child.find(qn("w:pPr"))
        if pPr_orig is not None:
            ind_orig = pPr_orig.find(qn("w:ind"))
            if ind_orig is not None:
                pPr_orig.remove(ind_orig)

        if spacer_idx is None:
            # Original single-cell behavior - no two-sided split needed.
            tblGrid_new = OxmlElement("w:tblGrid")
            gridCol = OxmlElement("w:gridCol")
            gridCol.set(qn("w:w"), str(total_width))
            tblGrid_new.append(gridCol)
            new_tbl.append(tblGrid_new)

            tr = OxmlElement("w:tr")
            tc = OxmlElement("w:tc")
            tcPr = OxmlElement("w:tcPr")
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(total_width))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)
            shd_el = OxmlElement("w:shd")
            shd_el.set(qn("w:val"), "clear")
            shd_el.set(qn("w:color"), "auto")
            shd_el.set(qn("w:fill"), fill_color)
            tcPr.append(shd_el)
            tc.append(tcPr)
            tc.append(child)
            tr.append(tc)
            new_tbl.append(tr)
        else:
            # Split into two cells at the spacer run - it's discarded
            # (a real cell boundary now does its job).
            all_runs = child.findall(qn("w:r"))
            spacer_run = all_runs[spacer_idx]
            left_width = total_width // 2
            right_width = total_width - left_width

            tblGrid_new = OxmlElement("w:tblGrid")
            for w in (left_width, right_width):
                gridCol = OxmlElement("w:gridCol")
                gridCol.set(qn("w:w"), str(w))
                tblGrid_new.append(gridCol)
            new_tbl.append(tblGrid_new)

            pPr_source = child.find(qn("w:pPr"))
            left_p = OxmlElement("w:p")
            right_p = OxmlElement("w:p")
            if pPr_source is not None:
                from copy import deepcopy
                left_p.append(deepcopy(pPr_source))
                right_p.append(deepcopy(pPr_source))

            spacer_run.getparent().remove(spacer_run)
            for r in all_runs:
                if r is spacer_run:
                    continue
                (left_p if all_runs.index(r) < spacer_idx else right_p).append(r)

            tr = OxmlElement("w:tr")
            for w, para_el in ((left_width, left_p), (right_width, right_p)):
                tc = OxmlElement("w:tc")
                tcPr = OxmlElement("w:tcPr")
                tcW = OxmlElement("w:tcW")
                tcW.set(qn("w:w"), str(w))
                tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)
                shd_el = OxmlElement("w:shd")
                shd_el.set(qn("w:val"), "clear")
                shd_el.set(qn("w:color"), "auto")
                shd_el.set(qn("w:fill"), fill_color)
                tcPr.append(shd_el)
                tc.append(tcPr)
                tc.append(para_el)
                tr.append(tc)
            new_tbl.append(tr)

        # Insert new_tbl at the header paragraph's OWN original position
        # (right before whichever sibling immediately followed it,
        # tracked from the pre-removal children list) - NOT right before
        # next_table_el, which can be several siblings further down
        # (past other standalone paragraphs like a "Name" row found
        # above) and would silently reorder that in-between content to
        # come BEFORE this header instead of after it - confirmed by
        # testing: inserting before next_table_el moved a real data row
        # to appear above its own section's header bar, a regression
        # caught by re-rendering and comparing against the original.
        insertion_anchor = children[i + 1]
        insertion_anchor.addprevious(new_tbl)
        fixed += 1

    return fixed


def _fix_merged_two_sided_table_header(doc):
    """Item (MERGED-TWO-SIDED-HEADER, gridSpan variant) - same underlying
    defect as _fix_incomplete_header_bar_shading's two-sided-header fix
    above, but for a DIFFERENT structural shape: confirmed one real
    instance ("Tenant rights") where the English label and the (Arabic-
    origin) number+label are crammed into ONE paragraph inside a SINGLE
    table cell that already has gridSpan="2" (i.e. it's ALREADY a real,
    correctly-merged, full-width cell - unlike the standalone-paragraph
    cases above, so this does NOT need a new table built around it).

    Detected the same way: a run with an artificially large w:spacing
    (letter-spacing) value inside that gridSpan=2 cell's paragraph -
    Aspose's same positioning hack seen elsewhere, confirmed present
    here too (w:spacing val="4097" on a single-space run, immediately
    before the Arabic content begins).

    Fix: split that cell's paragraph at the spacer run into two
    paragraphs (discarding the spacer run, same as the other fix),
    un-merge the gridSpan=2 cell into two real adjacent cells sized
    50/50 of its own total width, and put the "before spacer" content
    in the first and "after spacer" content in the second - giving this
    header the same two-sided visual structure every other working
    header in the document already has.

    Returns the count of cells fixed, for the caller's log."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from copy import deepcopy

    _SPACER_LETTER_SPACING_THRESHOLD = 500

    fixed = 0
    for tbl_el in doc.element.body.iter(qn("w:tbl")):
        for tr in tbl_el.findall(qn("w:tr")):
            tcs = tr.findall(qn("w:tc"))
            if len(tcs) != 1:
                continue
            tc = tcs[0]
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                continue
            gridSpan = tcPr.find(qn("w:gridSpan"))
            if gridSpan is None or gridSpan.get(qn("w:val")) != "2":
                continue
            tcW_el = tcPr.find(qn("w:tcW"))
            if tcW_el is None:
                continue
            total_width = int(tcW_el.get(qn("w:w")) or 0)
            if total_width <= 0:
                continue

            # Only handle the simple, verified case: exactly one
            # paragraph in the cell, with a real spacer run in it.
            paras = tc.findall(qn("w:p"))
            if len(paras) != 1:
                continue
            p_el = paras[0]
            all_runs = p_el.findall(qn("w:r"))
            spacer_idx = None
            for idx, r in enumerate(all_runs):
                rpr = r.find(qn("w:rPr"))
                if rpr is None:
                    continue
                spacing_el = rpr.find(qn("w:spacing"))
                if spacing_el is None:
                    continue
                val = spacing_el.get(qn("w:val"))
                if val and val.lstrip("-").isdigit() and abs(int(val)) > _SPACER_LETTER_SPACING_THRESHOLD:
                    spacer_idx = idx
                    break
            if spacer_idx is None:
                continue

            left_width = total_width // 2
            right_width = total_width - left_width

            spacer_run = all_runs[spacer_idx]
            pPr_source = p_el.find(qn("w:pPr"))

            left_p = p_el  # reuse the existing paragraph element for the "before" half
            right_p = OxmlElement("w:p")
            if pPr_source is not None:
                right_p.append(deepcopy(pPr_source))

            spacer_run.getparent().remove(spacer_run)
            for r in list(all_runs):
                if r is spacer_run:
                    continue
                if all_runs.index(r) > spacer_idx:
                    left_p.remove(r)
                    right_p.append(r)

            # Resize the original cell (drop gridSpan, it's now just col0).
            tcW_el.set(qn("w:w"), str(left_width))
            tcPr.remove(gridSpan)

            # Build the new sibling cell, copying the shading so it
            # matches the header bar's own color.
            shd_source = tcPr.find(qn("w:shd"))
            new_tc = OxmlElement("w:tc")
            new_tcPr = OxmlElement("w:tcPr")
            new_tcW = OxmlElement("w:tcW")
            new_tcW.set(qn("w:w"), str(right_width))
            new_tcW.set(qn("w:type"), "dxa")
            new_tcPr.append(new_tcW)
            if shd_source is not None:
                new_tcPr.append(deepcopy(shd_source))
            new_tc.append(new_tcPr)
            new_tc.append(right_p)

            tc.addnext(new_tc)
            fixed += 1

    return fixed


def _fix_heading_merged_into_previous_clause(doc):
    """Item (ARTICLE-HEADING-MERGED-INTO-BODY) - confirmed real,
    pre-existing Aspose defect (present in Aspose's own untranslated
    conversion, not introduced by translation): one clause paragraph's
    ending text and the NEXT Article's heading text were concatenated
    into a single paragraph with no break between them (only a ". "
    separator) - e.g. "...بمراعاة تلك الإلتزامات أو الشروط. المادة
    الثانية عشرة: الإخلاء لعدم التجاوب" holds BOTH the end of clause
    5-1-11 AND the start of "Article Twelve: Eviction for Non-Response"
    as one paragraph, in two separate w:r runs. Every OTHER Article
    heading in the same document is its own correctly-separated
    paragraph with distinctive run-level gray shading (w:shd
    fill=DDDDDD) - this one pair just never got split apart by Aspose's
    conversion, so the heading rendered as an unstyled tail fragment of
    the previous clause's paragraph instead of its own highlighted
    heading bar.

    Detects a language-agnostic "Article <word>:" / "المادة <word>:"
    pattern appearing partway through a paragraph's text (not at the
    start - a paragraph starting with this pattern is already a normal,
    correctly-separated heading and is left alone). Splits the
    paragraph into two at that boundary and copies the run-level
    formatting (shading, color, outline level) from a REAL, already-
    correctly-separated heading paragraph elsewhere in the same
    document, so the newly split-out heading visually matches every
    other one instead of needing a hardcoded style guess.

    Must run BEFORE translation, same reasoning as the other structural
    split fixes in this file - the split itself is language-agnostic
    (based on where the sentence boundary falls), but doing it before
    translation means the LLM sees two clean, independent segments
    instead of one run-on paragraph mixing an end-of-clause and a
    heading.

    Returns the count of paragraphs split, for the caller's log."""
    import unicodedata
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    body = doc.element.body
    children = list(body.iterchildren())

    # Find a reference heading's formatting (a paragraph whose FULL text
    # is just an "Article X: ..." heading, own outlineLvl, own run
    # shading) to copy onto the newly split-out heading below - avoids
    # hardcoding a specific fill color or font here.
    #
    # NFKC-normalizes text before pattern matching throughout this
    # function - confirmed necessary, not optional: Aspose stores this
    # document's Arabic text using PRESENTATION FORM codepoints (e.g.
    # "اﻟﻤﺎدة" using contextual glyph-form characters) rather than
    # standard Arabic letters, so a plain literal match against
    # "المادة" silently finds nothing even though the text reads
    # identically to a human. NFKC normalization maps presentation
    # forms back to their standard base characters, so matching works
    # regardless of which form Aspose happened to use - a more general,
    # reusable fix than the document-specific presentation-form
    # searches used earlier in this project's development.
    reference_rpr = None
    reference_pPr_extras = {}
    for child in children:
        if child.tag != qn("w:p"):
            continue
        text = unicodedata.normalize(
            "NFKC", "".join((t.text or "") for t in child.findall(".//" + qn("w:t"))).strip()
        )
        if not re.match(r"^(Article|المادة)\s+\S+(?:\s+\S+)?\s*:", text):
            continue
        if len(text) > 80:
            continue  # a real heading is short - a long match here is probably body text that happens to start similarly
        runs = child.findall(qn("w:r"))
        real_runs = [r for r in runs if (r.find(qn("w:t")) is not None and (r.find(qn("w:t")).text or "").strip())]
        if not real_runs:
            continue
        rpr = real_runs[0].find(qn("w:rPr"))
        if rpr is None:
            continue
        reference_rpr = rpr
        pPr = child.find(qn("w:pPr"))
        if pPr is not None:
            outlineLvl = pPr.find(qn("w:outlineLvl"))
            if outlineLvl is not None:
                reference_pPr_extras["outlineLvl"] = outlineLvl
        break
    if reference_rpr is None:
        return 0  # no reference heading style found in this document - don't guess at formatting

    fixed = 0
    for child in children:
        if child.tag != qn("w:p"):
            continue
        all_runs = child.findall(qn("w:r"))
        if len(all_runs) < 2:
            continue
        # Find a RUN BOUNDARY where the heading actually starts.
        # Confirmed on the real document: this is NOT a mid-run split -
        # Aspose stored the clause-ending text and the Article-heading
        # text as two ENTIRELY SEPARATE w:r runs within one paragraph
        # (run 0 ends "...الشروط. ", run 1 begins fresh with "المادة
        # الثانية عشرة: ..." - no single run's own text contains both
        # the ". " and the heading start together). So this looks for
        # the first run (after the first) whose OWN normalized text
        # starts with the heading pattern, and splits the paragraph
        # right there - no character-position math needed, which also
        # sidesteps the presentation-form-normalization length-mismatch
        # edge case entirely.
        split_run_idx = None
        for idx, r in enumerate(all_runs):
            if idx == 0:
                continue  # a heading pattern in run 0 means this paragraph is ALREADY a normal, correctly-separated heading
            t = r.find(qn("w:t"))
            if t is None or not t.text or not t.text.strip():
                continue
            normalized = unicodedata.normalize("NFKC", t.text.strip())
            if re.match(r"^(Article|المادة)\s+\S+(?:\s+\S+)?\s*:", normalized):
                split_run_idx = idx
                break
        if split_run_idx is None:
            continue

        try:
            from copy import deepcopy

            pPr_source = child.find(qn("w:pPr"))
            new_p = OxmlElement("w:p")
            if pPr_source is not None:
                new_pPr = deepcopy(pPr_source)
                existing_outline = new_pPr.find(qn("w:outlineLvl"))
                if existing_outline is not None:
                    new_pPr.remove(existing_outline)
                if "outlineLvl" in reference_pPr_extras:
                    new_pPr.append(deepcopy(reference_pPr_extras["outlineLvl"]))
                new_p.append(new_pPr)

            # Every run from the split point onward moves to the new
            # heading paragraph, restyled to match the reference
            # heading's formatting (shading/color/etc.) instead of
            # inheriting whatever plain-body-text formatting it had as
            # part of the previous clause's run-on paragraph.
            for r in all_runs[split_run_idx:]:
                child.remove(r)
                r_rpr = r.find(qn("w:rPr"))
                if r_rpr is not None:
                    r.remove(r_rpr)
                r.insert(0, deepcopy(reference_rpr))
                new_p.append(r)

            child.addnext(new_p)
            fixed += 1
        except Exception as err:  # noqa: BLE001
            print(f"[heading-merge-fix] skipped one paragraph after an error, continuing with the rest: {err}")

    return fixed


def _fix_merged_numbered_subclause(doc):
    """Item (SUBCLAUSE-MERGED-INTO-PREVIOUS, real reported issue 2) -
    confirmed real, same class of pre-existing Aspose defect as
    _fix_heading_merged_into_previous_clause, just for a DIFFERENT
    boundary pattern: two adjacent numbered sub-clauses ended up in one
    single paragraph with no break between them. Confirmed directly in
    a real document: "13-1-5 The Tenant and its employees shall not
    smoke in the corridors and lobbies of the property and shall
    comply with smoking in designated areas. 14-1-5 The Tenant shall
    dispose of waste..." was ONE paragraph holding BOTH clause 13-1-5
    AND clause 14-1-5, while every other numbered clause in the same
    document (12-1-5, 15-1-5, 16-1-5, ...) was correctly its own
    separate paragraph.

    REAL DIFFERENCE FROM THE HEADING-MERGE CASE (confirmed directly,
    not assumed): the heading-merge bug had the two logical units
    stored as two SEPARATE w:r runs, so a run-boundary search worked.
    This bug's merge is a SINGLE run whose own text string contains
    BOTH clauses end-to-end - checked directly against the real
    document, this paragraph had exactly ONE run holding the entire
    "13-1-5 ... 14-1-5 ..." text. A run-boundary-only search (like the
    heading-merge fix uses) finds nothing here, because there is no
    second run to find. So this fix searches WITHIN each run's own
    text for a mid-string clause-number boundary and splits the RUN
    ITSELF (into two runs with identical formatting, since it's the
    same original run just cut in two) before moving the tail into a
    new paragraph - a necessary difference in mechanism from the
    heading-merge fix, even though the end RESULT (two clean
    paragraphs) is the same kind of fix.

    Clause numbers like "13-1-5" are language-agnostic digits/hyphens,
    unaffected by translation either way - runs at the same point in
    the pipeline as the other structural split fixes (before
    translation, on the raw Aspose conversion) for the same reason
    given there: the LLM sees two clean, independent segments to
    translate instead of one run-on paragraph mixing two clauses'
    content together. Also handles the case where the boundary DOES
    fall on an existing run boundary (checked first, cheaper) before
    falling back to the mid-run text search.

    Only matches a clause number following ". " (end of a completed
    sentence) - never mid-sentence text that merely contains digits and
    hyphens, which keeps this from false-triggering on real prose."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from copy import deepcopy

    body = doc.element.body
    children = list(body.iterchildren())
    # Real reported gap (found during a careful before/after image review
    # of a real document): the original pattern only matched 3-segment
    # clause codes (N-N-N, e.g. "13-1-5"), completely missing this same
    # document's OWN 2-segment codes (N-N, e.g. "10-1"/"10-2") used for
    # several Articles (6, 7, 10 confirmed real) - "10-1...contract.
    # 10-2 Non-compliance..." was genuinely merged into one paragraph
    # and this function reported 0 fixes on it. The (?:-\d+)? makes the
    # third segment optional, matching BOTH real numbering schemes this
    # document actually uses, not just the narrower one.
    clause_at_start_pattern = re.compile(r"^\d+-\d+(?:-\d+)?\s")
    clause_mid_pattern = re.compile(r"\.\s+(\d+-\d+(?:-\d+)?\s)")

    fixed = 0
    for child in children:
        if child.tag != qn("w:p"):
            continue
        all_runs = child.findall(qn("w:r"))
        if not all_runs:
            continue

        # Pass 1 (cheaper, matches the heading-merge fix's own
        # approach): a later run whose OWN text starts with the
        # pattern - a genuine separate-run boundary.
        split_run_idx = None
        for idx, r in enumerate(all_runs):
            if idx == 0:
                continue
            t = r.find(qn("w:t"))
            if t is None or not t.text or not t.text.strip():
                continue
            if clause_at_start_pattern.match(t.text.strip()):
                split_run_idx = idx
                break

        if split_run_idx is not None:
            try:
                pPr_source = child.find(qn("w:pPr"))
                new_p = OxmlElement("w:p")
                if pPr_source is not None:
                    new_p.append(deepcopy(pPr_source))
                for r in all_runs[split_run_idx:]:
                    child.remove(r)
                    new_p.append(r)
                child.addnext(new_p)
                fixed += 1
            except Exception as err:  # noqa: BLE001
                print(f"[subclause-merge-fix] skipped one paragraph (run-boundary case) after an error: {err}")
            continue

        # Pass 2: the merge is INSIDE a single run's own text (the real
        # case confirmed on the actual reported document) - find the
        # first run whose text contains a mid-string ". <clause> "
        # boundary, split that run's text in two, and move the tail
        # (plus any later runs) into a new paragraph.
        target_run = None
        match = None
        for r in all_runs:
            t = r.find(qn("w:t"))
            if t is None or not t.text:
                continue
            m = clause_mid_pattern.search(t.text)
            if m:
                target_run = r
                match = m
                break
        if target_run is None:
            continue

        try:
            t = target_run.find(qn("w:t"))
            full_text = t.text
            split_at = match.start(1)  # right before the clause number itself
            before_text = full_text[:split_at]
            after_text = full_text[split_at:]

            t.text = before_text

            tail_run = deepcopy(target_run)
            tail_t = tail_run.find(qn("w:t"))
            tail_t.text = after_text
            tail_t.set(qn("xml:space"), "preserve")

            target_run_idx = all_runs.index(target_run)
            pPr_source = child.find(qn("w:pPr"))
            new_p = OxmlElement("w:p")
            if pPr_source is not None:
                new_p.append(deepcopy(pPr_source))
            new_p.append(tail_run)
            # any runs AFTER the split run also belong to the new (second) paragraph
            for r in all_runs[target_run_idx + 1:]:
                child.remove(r)
                new_p.append(r)
            child.addnext(new_p)
            fixed += 1
        except Exception as err:  # noqa: BLE001
            print(f"[subclause-merge-fix] skipped one paragraph (mid-run case) after an error: {err}")

    return fixed


def run_structure_only_test(pdf_path, output_path):
    """PDF -> DOCX via Aspose.Words Cloud's own native conversion,
    source-language text (no translation). Answers "does Aspose's own
    table/format detection look better than ours" with zero LLM cost."""
    words_api = _words_api()
    with open(pdf_path, "rb") as f:
        request = ConvertDocumentRequest(document=f, format="docx")
        # Item - confirmed bug (verified by reading the SDK's own
        # deserialize_file() source, not guessed): convert_document()
        # returns the converted file's RAW BYTES directly, not a path
        # to a temp file on disk despite what its docstring's "return:
        # file" phrasing suggests. Treating that return value as if it
        # were a path and calling open(result, "rb") on it was the
        # actual root cause of the "'bytes' object has no attribute
        # 'seek'" error - Python's open() was being handed the docx's
        # raw content instead of a filename.
        result_bytes = words_api.convert_document(request)

    from docx import Document
    from io import BytesIO
    doc = Document(BytesIO(result_bytes))
    headers_fixed = 0
    try:
        headers_fixed = _fix_incomplete_header_bar_shading(doc)
        headers_fixed += _fix_merged_two_sided_table_header(doc)
    except Exception:
        pass  # non-fatal - test pipeline still produces its output even if this cosmetic pass fails

    headings_split = 0
    try:
        headings_split = _fix_heading_merged_into_previous_clause(doc)
    except Exception:
        pass  # non-fatal - test pipeline still produces its output even if this cosmetic pass fails

    subclauses_split = 0
    try:
        subclauses_split = _fix_merged_numbered_subclause(doc)
    except Exception:
        pass  # non-fatal - test pipeline still produces its output even if this cosmetic pass fails

    # Real reported issue 4 - runs BEFORE the other table-structural
    # fixes below, so those operate on the final, correctly-merged
    # table shape rather than on fragments that are about to be
    # combined anyway.
    appendix_tables_merged = 0
    try:
        appendix_tables_merged = _merge_adjacent_header_tables(doc)
    except Exception:
        pass  # non-fatal - test pipeline still produces its output even if this cosmetic pass fails

    tables_repositioned = 0
    try:
        tables_repositioned = _fix_table_overflow_indent(doc)
    except Exception:
        pass  # non-fatal - test pipeline still produces its output even if this cosmetic pass fails

    # Must run right after the table-level fix above, so this can
    # compare each row's own override against the ALREADY-corrected
    # table-level indent (see this function's own docstring for the
    # real reported bug this addresses - a table-level indent that
    # looked correct while row-level tblPrEx overrides silently pushed
    # the actual rendered position 50%+ into the page in real Word).
    row_level_indent_overrides_fixed = 0
    try:
        row_level_indent_overrides_fixed = _fix_row_level_table_indent_override(doc)
    except Exception:
        pass  # non-fatal - test pipeline still produces its output even if this cosmetic pass fails

    duplicate_rows_removed = 0
    try:
        duplicate_rows_removed = _remove_duplicate_table_rows(doc)
    except Exception:
        pass  # non-fatal - test pipeline still produces its output even if this cosmetic pass fails

    ambiguous_table_width_fixed = 0
    try:
        ambiguous_table_width_fixed = _fix_ambiguous_table_width(doc)
    except Exception:
        pass  # non-fatal - test pipeline still produces its output even if this cosmetic pass fails

    cell_width_mismatches_fixed = 0
    try:
        cell_width_mismatches_fixed = _fix_cell_width_vs_column_mismatch(doc)
    except Exception:
        pass  # non-fatal - test pipeline still produces its output even if this cosmetic pass fails

    leaked_names_fixed = 0
    try:
        leaked_names_fixed = _fix_leaked_internal_field_names(doc)
    except Exception:
        pass  # non-fatal - test pipeline still produces its output even if this cosmetic pass fails

    split_numeric_findings = []
    try:
        split_numeric_findings = _detect_split_numeric_values(doc)
    except Exception:
        pass  # non-fatal - a validation pass failing shouldn't block delivering the document

    doc.save(output_path)
    return {
        "output_path": output_path,
        "mode": "structure_only",
        "pipeline_code_version": PIPELINE_CODE_VERSION,
        "header_bars_fixed": headers_fixed,
        "headings_split": headings_split,
        "subclauses_split": subclauses_split,
        "appendix_tables_merged": appendix_tables_merged,
        "tables_repositioned": tables_repositioned,
        "row_level_indent_overrides_fixed": row_level_indent_overrides_fixed,
        "duplicate_rows_removed": duplicate_rows_removed,
        "ambiguous_table_width_fixed": ambiguous_table_width_fixed,
        "cell_width_mismatches_fixed": cell_width_mismatches_fixed,
        "leaked_names_fixed": leaked_names_fixed,
        "split_numeric_findings": split_numeric_findings,
        "aspose_words_calls": 1,
        "aspose_pdf_calls": 0,
        "llm_calls": 0,
        "llm_calls_by_provider": {},
    }


def _pdf_api():
    """PdfApi (unlike WordsApi) takes an ApiClient instance, not
    client_id/client_secret kwargs directly - confirmed by inspecting
    asposepdfcloud.apis.pdf_api.PdfApi.__init__'s actual signature
    (self, api_client=None), which is different from WordsApi's. Doesn't
    reuse _require_configured() since that also checks
    _WORDS_SDK_AVAILABLE (a different SDK/package) - callers here have
    already checked _PDF_SDK_AVAILABLE and the credentials themselves."""
    return PdfApi(_PdfApiClient(client_secret=ASPOSE_CLIENT_SECRET, client_id=ASPOSE_CLIENT_ID))


def extract_text_via_aspose(pdf_path):
    """Extracts text via Aspose.PDF Cloud if configured, otherwise falls
    back to Lexora's own pdfplumber-based extractor (le.extract_text) so
    the rest of this pipeline is still testable. The fallback is always
    clearly labeled in the result, never silently substituted.

    Item (ASPOSE.PDF-CLOUD-WIRING) - two things were actually broken here,
    confirmed by inspecting the installed asposepdfcloud SDK's real code
    (currently pinned >=25.0 in requirements.txt, which pip resolves to
    26.7.0 - a newer SDK generation than whatever version this file's
    original `from asposepdfcloud.models.requests import
    GetPdfInStorageToTextRequest` import was written against):

    1. That import itself FAILS on the installed SDK - there is no
       `asposepdfcloud.models.requests` module at all in 26.7.0 (the
       newer SDK generation takes plain kwargs instead of request
       objects). That ImportError was being silently caught by the
       try/except at the top of this file, setting _PDF_SDK_AVAILABLE =
       False - which is the actual reason "Aspose.PDF Cloud not
       configured" kept showing even with both env vars correctly set;
       the SDK "not being available" had nothing to do with the
       credentials.
    2. Real text extraction was never wired up at all - Aspose.PDF
       Cloud's text-extraction endpoint (get_pdf_in_storage_to_text)
       only operates on files already sitting in Aspose's own cloud
       storage, so it needs an upload_file() call first. That upload
       step is added below.

    Storage note: get_pdf_in_storage_to_text's response is a real file
    Aspose's SDK saves to a local temp path and returns the PATH to
    (confirmed by reading ApiClient.__deserialize_file's source) - this
    is DIFFERENT from asposewordscloud's convert_document(), which hands
    back raw bytes directly despite an identically-worded "return: file"
    docstring. Don't assume the two SDKs behave the same just because
    their docstrings read the same."""
    if not (_PDF_SDK_AVAILABLE and ASPOSE_CLIENT_ID and ASPOSE_CLIENT_SECRET):
        return {
            "text": le.extract_text(pdf_path),
            "source": "fallback:pdfplumber (Aspose.PDF Cloud not configured)",
            "aspose_pdf_calls": 0,
        }

    pdf_api = _pdf_api()
    folder = "lexora-aspose-test"
    filename = os.path.basename(pdf_path)
    storage_path = f"{folder}/{filename}"
    result_path = None
    api_calls = 0  # counts actual Aspose.PDF Cloud API calls MADE (attempted), even on failure - not assumed
    text, source = None, None
    try:
        # Item - confirmed bug (verified by reading the installed SDK's
        # own api_client.py __call_api source, not guessed): the `file`
        # argument to upload_file() is NOT meant to be an already-open
        # file object - the SDK does `open(n, 'rb')` on whatever gets
        # passed here ITSELF internally (see api_client.py's `if files:`
        # block), so it needs a PATH STRING it can open on its own.
        # Passing an open BufferedReader made the SDK call
        # open(<BufferedReader instance>, 'rb'), which raised exactly
        # "expected str, bytes or os.PathLike object, not
        # BufferedReader" - this was silently swallowed by the
        # try/except below and fell back to pdfplumber every single
        # time, meaning Aspose.PDF Cloud extraction never actually ran
        # even when correctly configured.
        api_calls += 1
        pdf_api.upload_file(storage_path, pdf_path)

        api_calls += 1
        result_path = pdf_api.get_pdf_in_storage_to_text(filename, folder=folder)
        with open(result_path, "rb") as rf:
            text = rf.read().decode("utf-8", errors="replace")
        source = "aspose_pdf_cloud"
    except Exception as err:  # noqa: BLE001
        text = le.extract_text(pdf_path)
        source = f"fallback:pdfplumber (Aspose.PDF Cloud call failed: {err})"
    finally:
        if result_path:
            try:
                os.remove(result_path)
            except OSError:
                pass  # local temp file cleanup - non-fatal if it's already gone
        try:
            api_calls += 1
            pdf_api.delete_file(storage_path)
        except Exception:
            pass  # cleanup best-effort - don't fail extraction over a leftover file in Aspose's cloud storage

    # Built AFTER the try/finally completes (not returned from inside it)
    # so api_calls reflects ALL three attempted calls, including
    # delete_file's - returning from inside try/except would have
    # captured api_calls before finally's own increment ran.
    return {"text": text, "source": source, "aspose_pdf_calls": api_calls}


_SIGNATURE_PLACEHOLDER_RE = re.compile(r"^_{3,}$")


def _extract_signature_images(pdf_path):
    """Item 2 (SIGNATURE-PRESERVATION) - Aspose's own PDF->DOCX conversion
    doesn't carry over digital-signature/annotation appearance content,
    it only leaves the blank underline that was drawn next to the
    signature widget in the original PDF - confirmed by testing against
    a real signed PDF (Agreement_-_Original.pdf), where a plain page
    render (no forms/annotations) produced just the underline, while
    rendering with pdfium's form environment initialized produced the
    actual "Firmato digitalmente da: ..." signature stamp.

    This mirrors the shared engine logic (extractOfflineImages ->
    signatureRects detection, identical in every js/engine-*.js copy -
    Widget/Sig annotations, skip ones with no appearance content i.e.
    hasAppearance === False / no 'AP' entry) but
    in Python: pdfplumber gives the annotation rects, pypdfium2 (with
    init_forms() + draw_annots=True + may_draw_forms=True - all THREE
    are required, confirmed by testing; render() without them silently
    renders only the blank underline, no error) does the actual pixel
    rendering the crop is taken from.

    Returns a list of (page_index, top_pt, png_bytes) tuples, sorted in
    reading order (page, then top-to-bottom), so callers can match them
    to placeholder text in that same order."""
    if not _SIGNATURE_EXTRACT_AVAILABLE:
        return []
    from PIL import Image  # noqa: F401 (import kept local; pdfium.to_pil() needs Pillow installed, this just fails loudly and early if it's missing)

    found = []
    pdf_doc = pdfium.PdfDocument(pdf_path)
    pdf_doc.init_forms()
    with pdfplumber.open(pdf_path) as pl_pdf:
        for page_idx, pl_page in enumerate(pl_pdf.pages):
            widget_rects = []
            for a in (pl_page.annots or []):
                data = a.get("data") or {}
                # pdfminer resolves these to PSLiteral objects whose
                # repr is "/'Widget'" (confirmed by testing against the
                # real annotation dict) - strip both '/' and "'" to get
                # the plain name.
                subtype = str(data.get("Subtype", "")).strip("/'")
                ft = str(data.get("FT", "")).strip("/'")
                if subtype != "Widget" and ft != "Sig":
                    continue
                if "AP" not in data:
                    continue  # no appearance stream - blank/unsigned field, nothing to crop
                x0, x1 = a["x0"], a["x1"]
                top, bottom = a["top"], a["bottom"]
                if (x1 - x0) < 10 or (bottom - top) < 10:
                    continue
                widget_rects.append((min(x0, x1), min(top, bottom), max(x0, x1), max(top, bottom)))
            if not widget_rects:
                continue

            fpage = pdf_doc[page_idx]
            scale = min(3.0, 2000 / max(pl_page.width, pl_page.height))
            bitmap = fpage.render(scale=scale, draw_annots=True, may_draw_forms=True)
            pil_img = bitmap.to_pil()

            for (x0, top, x1, bottom) in sorted(widget_rects, key=lambda r: (r[1], r[0])):
                px0, py0 = int(x0 * scale), int(top * scale)
                px1, py1 = int(x1 * scale), int(bottom * scale)
                crop = pil_img.crop((px0, py0, px1, py1))
                buf = io.BytesIO()
                crop.save(buf, format="PNG")
                found.append((page_idx, top, buf.getvalue()))
    return found


def _inject_signature_images(doc, signature_images):
    """Best-effort in-place placement: Aspose's converted docx leaves a
    paragraph containing just a run of underscores wherever the source
    PDF had a signature widget. Walk the document's paragraphs in order
    and swap each such placeholder for the matching extracted signature
    image, consuming signature_images in the same page/top-to-bottom
    order _extract_signature_images returned them in - this is a
    positional match (Nth placeholder <-> Nth signature found), not a
    coordinate-verified one, since Aspose's converted docx doesn't carry
    forward the original PDF coordinates to check against.

    Any signatures left over (more detected than placeholder paragraphs
    found - e.g. a placeholder Aspose rendered as something other than
    plain underscores) are appended as a clearly-labeled section at the
    end rather than silently dropped, same "don't lose real content"
    principle as the reference-page fallback below.

    Returns (placed_count, leftover_count) for the caller's log."""
    from docx.shared import Pt
    from io import BytesIO

    remaining = list(signature_images)
    placed = 0
    for para in doc.paragraphs:
        if not remaining:
            break
        if _SIGNATURE_PLACEHOLDER_RE.match((para.text or "").strip()):
            for run in list(para.runs):
                run.text = ""
            _page_idx, _top, img_bytes = remaining.pop(0)
            para.add_run().add_picture(BytesIO(img_bytes), height=Pt(50))
            placed += 1

    if remaining:
        doc.add_page_break()
        note = doc.add_paragraph()
        note_run = note.add_run(
            f"Signature image(s) detected in source PDF but not matched to a "
            f"placeholder line in Aspose's converted structure ({len(remaining)}):"
        )
        note_run.bold = True
        for _page_idx, _top, img_bytes in remaining:
            doc.add_paragraph().add_run().add_picture(BytesIO(img_bytes), height=Pt(50))

    return placed, len(remaining)


def _iter_paragraphs_in_order(container):
    """Yields every Paragraph inside `container` (a Document or a table
    _Cell) in true document order, descending into tables (and any
    tables nested inside table cells) recursively.

    python-docx's own doc.paragraphs and doc.tables are two SEPARATE flat
    lists - doc.paragraphs skips everything inside a table cell entirely,
    and neither list reflects where a table actually sits relative to
    surrounding paragraphs. That matters a lot here: Aspose's PDF->DOCX
    conversion of a table-heavy contract (e.g. the REGA-format Arabic
    lease contracts Lexora also handles, which are almost entirely
    bordered field/value tables) puts nearly all of its real content
    inside table cells, so a translation pass that only walked
    doc.paragraphs would silently skip almost the whole document.

    Standard recipe for this (walking the underlying XML body's direct
    children, matching each one back to a CT_P paragraph or CT_Tbl
    table) - verified against a real python-docx-built table (a plain
    doc.paragraphs/doc.tables walk would have returned the table's cell
    text in a completely different, disconnected list; this returns
    everything in one correctly-ordered stream)."""
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph
    from docx.document import Document as _DocxDocument

    if isinstance(container, _DocxDocument):
        parent_elm = container.element.body
    elif isinstance(container, _Cell):
        parent_elm = container._tc
    else:
        raise TypeError(f"unsupported container type: {type(container)}")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, container)
        elif isinstance(child, CT_Tbl):
            table = Table(child, container)
            for row in table.rows:
                for cell in row.cells:
                    yield from _iter_paragraphs_in_order(cell)


def _replace_paragraph_text(paragraph, new_text):
    """Puts `new_text` into a paragraph's FIRST run (keeping that run's
    own formatting - bold/italic/font/size/color - so a paragraph that
    was entirely bold, or in a particular font, stays that way) and
    blanks out any other runs in the paragraph rather than removing them
    (simpler and safer than XML surgery to delete run elements, and an
    empty run is harmless/invisible in the saved docx).

    Known, explicitly-accepted trade-off: if the ORIGINAL paragraph had
    MIXED formatting across multiple runs (e.g. "The rent is **500
    EUR**" where only "500 EUR" was bold within one paragraph), that
    finer-grained split is lost - the whole translated paragraph ends up
    under the first run's formatting only. Preserving per-run formatting
    across a translation (where word order and phrase boundaries change
    completely between languages) is a much harder alignment problem;
    paragraph-level formatting fidelity was judged the right scope for
    this pass."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return
    runs[0].text = new_text
    for extra in runs[1:]:
        extra.text = ""


_MAX_SEGMENT_CHARS_PER_BATCH = 9000  # Item (REDUCE-LLM-CALL-COUNT) - was 3500, chosen without
# testing against the model's real output-token ceiling. Confirmed the default OpenRouter model
# (OPENROUTER_MODEL, defaults to "openai/gpt-4o" - see lease_engine.load_llm_config()) supports up
# to 16384 output tokens, and the old 8000 max_tokens cap on each batch call was only using half of
# that. Tested against a real 575-segment/24063-char document: this threshold cuts 8 batches down to
# 3 (each batch's translated-JSON output comes to roughly 3500 output tokens by rough char/4 estimate
# - comfortably under the 12000 max_tokens now used below, leaving real headroom for target languages
# that expand more than English does). Batch failures already degrade gracefully (that batch's
# segments are left untranslated and logged, not a hard failure - see _translate_docx_segments_in_place
# below), so a wrong per-model assumption here costs some retranslatable segments, not a broken run.


_RTL_LANGUAGE_KEYWORDS = ("arabic", "hebrew", "persian", "farsi", "urdu", "pashto", "dhivehi", "divehi", "yiddish", "sindhi")


def _is_rtl_language(target_language):
    tl = (target_language or "").strip().lower()
    return any(kw in tl for kw in _RTL_LANGUAGE_KEYWORDS)


_NUMERIC_ID_PREFIX_RE = re.compile(
    r"^([\d\u0660-\u0669]+(?:[\-.][\d\u0660-\u0669]+){1,4})(\s+)(\S.{10,})", re.DOTALL
)


def _reverse_numeric_id_groups(token):
    """Reverses the ORDER of hyphen/dot-separated numeric groups in an
    identifier like '11-1-5' -> '5-1-11' - WITHOUT touching the digits
    within each group, and without touching which separator character
    sits where (handles mixed '-'/'.' use, though clause numbers
    typically use one consistently). Works with both Western (0-9) and
    Arabic-Indic (\u0660-\u0669) digits, since the source language isn't
    known/assumed here - see _fix_paragraph_direction's docstring for
    why this needs to be script-agnostic."""
    tokens = re.findall(r"[\d\u0660-\u0669]+|[\-.]", token)
    groups = [t for t in tokens if not re.match(r"[\-.]", t)]
    seps = [t for t in tokens if re.match(r"[\-.]", t)]
    if len(groups) < 2:
        return token
    rev_groups = groups[::-1]
    rev_seps = seps[::-1]
    out = rev_groups[0]
    for sep, g in zip(rev_seps, rev_groups[1:]):
        out += sep + g
    return out


def _fix_reversed_clause_number_prefix(paragraph):
    """Item (CLAUSE-NUMBER-REVERSAL) - confirmed root cause by comparing
    Aspose's own untranslated RTL conversion against the same paragraph
    post-translation: Aspose's PDF text extraction stores certain
    RTL-context numeric IDENTIFIER PREFIXES (clause/article numbers like
    "5-1-11") in REVERSED group order ("11-1-5") in the underlying XML
    text, relying on RTL bidi rendering to display them correctly -
    verified directly: Aspose's raw conversion held '\u0661\u0661-\u0661-\u0665'
    (logical order 11-1-5) for a clause whose SOURCE PDF plainly showed
    '\u0665-\u0661-\u0661\u0661' (5-1-11, confirmed against the original
    PDF's own text). RTL bidi rendering was silently "correcting" the
    display while bidi=1 was set - our _fix_paragraph_direction's
    RTL->LTR flip (needed for the now-translated LTR content) removes
    that compensating rendering, which is what exposes the reversal as
    literal, visible "11-1-5" text instead of "5-1-11".

    This is NOT a general numeric-integrity bug: Gregorian/Hijri dates
    and amounts elsewhere in the SAME originally-RTL paragraphs were
    verified to already be in CORRECT order (e.g. "2024-04-01" stayed
    "2024-04-01", never "01-04-2024") - only short numeric-group PREFIXES
    at the very start of a paragraph, followed by substantial prose,
    show this reversal. That "followed by substantial prose" condition
    is the safety gate used here: a standalone numeric VALUE that fills
    an entire cell (a date, an amount, an ID - confirmed several dozen
    of these also carry bidi=1 in the same document) must NEVER be
    reversed, since reversing "2025-06-17" would corrupt a real date
    into "17-06-2025". Requiring real text after the numeric prefix
    reliably tells clause-number labels ("11-1-5 The Tenant shall...")
    apart from bare data values ("2025-06-17" and nothing else).

    Only ever called for a paragraph that's about to flip from RTL to
    LTR (see call site in _fix_paragraph_direction) - a paragraph that
    stays RTL, or was already LTR, never had this compensating-reversal
    behavior baked in, so nothing to undo there. Operates on whichever
    run holds the paragraph's opening text (translation puts the whole
    translated paragraph into one run - see _replace_paragraph_text -
    so this is normally the first run with real text).

    Returns True if a reversal was applied."""
    for run in paragraph.runs:
        text = run.text or ""
        if not text.strip():
            continue
        m = _NUMERIC_ID_PREFIX_RE.match(text)
        if not m:
            # Item (DIAGNOSTIC-LOGGING-FOR-SPORADIC-FAILURES) - confirmed
            # real, unexplained pattern in a production run: clauses
            # 1-1-5 through 4-1-5 failed to reverse, 5-1-5 was a
            # neutral palindrome case, 6-1-5/8-1-5/9-1-5/10-1-5
            # succeeded, then 7-1-5 and 11-1-5 failed again - a
            # SCATTERED pattern, not a clean cascade (which the
            # fault-isolation fix already guards against), and every
            # isolated sandbox test of this exact function against
            # byte-identical text succeeds. This can't be reproduced
            # without the real LLM's actual output for those specific
            # segments, which isn't available here - so instead of
            # guessing further, log any near-miss so the NEXT real
            # production log shows exactly what the real text looked
            # like for these specific failures, turning this from a
            # guess into evidence.
            #
            # Scoped narrowly to avoid flooding the log: a real
            # near-miss needs BOTH multiple hyphen/dot-separated numeric
            # groups (like a genuine clause number would have) AND
            # substantial length (so a bare date or ID, which correctly
            # and intentionally never matches - see this function's
            # main docstring on why standalone values must not be
            # touched - doesn't get logged as if it were a failure).
            # Confirmed necessary: an unscoped version of this logged
            # over 100 lines for a single real document, mostly dates,
            # IDs, and short list items ("2. Renewed") that were never
            # supposed to match in the first place.
            looks_multi_group = bool(re.match(r"^[\d\u0660-\u0669]+[-.][\d\u0660-\u0669]+", text))
            if looks_multi_group and len(text) > 15:
                print(f"[clause-number-fix] near-miss - multi-group numeric prefix but didn't match the expected pattern: {text[:80]!r}")
            return False  # first real-text run doesn't start with a numeric-id-like prefix at all
        prefix, ws, rest = m.groups()
        fixed_prefix = _reverse_numeric_id_groups(prefix)
        if fixed_prefix == prefix:
            return False
        run.text = fixed_prefix + ws + rest
        return True
    return False


def _fix_body_paragraph_center_alignment(doc):
    """Item (BODY-PARAGRAPH-WRONGLY-CENTERED, real reported issues 1,
    2, and 3, follow-up after the right-indent fix) - confirmed real:
    9 genuinely long (multi-sentence, >=80 character) body clause
    paragraphs carried w:jc="center" in a real reported document -
    confirmed directly against that document's own real distribution:
    every OTHER long body paragraph (44 of them) correctly used
    w:jc="both" (justify), and EVERY short paragraph in the whole
    document (61 of them) used "both" or "left" - "center" appeared on
    zero short paragraphs and only on long ones, with no structural
    reason for those 9 specifically to be centered - clearly a
    conversion artifact, not an intentional design choice anywhere in
    this document.

    Real, confirmed visual effect (checked via an actual LibreOffice
    render): a long, multi-line paragraph rendered with w:jc="center"
    produces a "staircase" look - each wrapped line centers on ITS OWN
    width independently, so consecutive lines start at different
    horizontal positions instead of a consistent one. This directly
    explains all three symptoms the user reported: text not starting
    right after the left margin, text appearing to start "from the
    right" on some lines (short trailing lines center further right),
    and a wrapped line's start position not matching the paragraph's
    own first line - all three are the same root cause, not three
    separate bugs.

    Fix: reset w:jc from "center" to "both" (matching this document's
    own overwhelmingly dominant, clearly-intentional norm for body
    text) for non-table paragraphs. Left deliberately narrow - only
    touches "center", never touches "both" or "left" (both are valid,
    real choices already in use elsewhere in the same document) - so
    this cannot flip a paragraph that was already correctly aligned."""
    from docx.oxml.ns import qn

    fixed = 0
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        if p._p.getparent().tag == qn("w:tc"):
            continue  # table-cell paragraph - untouched
        pPr = p._p.find(qn("w:pPr"))
        if pPr is None:
            continue
        jc = pPr.find(qn("w:jc"))
        if jc is not None and jc.get(qn("w:val")) == "center":
            jc.set(qn("w:val"), "both")
            fixed += 1
    return fixed


def _find_reference_margin_table(doc):
    """Locates a real, early table in the document to use as the
    margin REFERENCE for non-table text, per explicit direction: check
    the page's own width and the table's own width to derive the real
    left/right margin area non-table text should also respect, rather
    than guessing or hardcoding a value.

    REAL BUG FOUND (via a direct, real-document XML/table comparison
    against a live-deployed pipeline's own output): the original
    version just grabbed the FIRST table anywhere in the document with
    a short first-cell label, assuming that was reliably the "Contract
    Data" table. A real deployed document turned out to structure its
    "Contract Data" section as plain paragraphs with tab stops, not a
    table at all - so this function silently picked a completely
    unrelated table ("Issuer:", the first genuinely short-labeled table
    it found) as the margin reference, which would have produced a
    wrong target margin for that whole document.

    Fix: specifically look for a table whose first-cell text
    SEMANTICALLY matches "Contract Data" (case-insensitive, tolerant of
    a leading section number like "1 " and of double-spaced OCR
    artifacts like "Contract  Data") - only a table that genuinely
    looks like the intended reference is used. If no such table exists
    anywhere in the document (confirmed real case above), returns None
    outright - the caller's own documented fallback behavior (skip the
    table-derived-margin fix entirely, or use the simpler right-only
    reset) then correctly applies, rather than silently substituting an
    unrelated table's numbers."""
    pattern = re.compile(r"^\d*\s*contract\s+data\b", re.IGNORECASE)
    for t in doc.tables:
        if not t.rows or not t.rows[0].cells:
            continue
        first_cell_text = t.rows[0].cells[0].text.strip()
        normalized = re.sub(r"\s+", " ", first_cell_text)
        if pattern.match(normalized):
            return t
    return None


def _compute_table_derived_margins(doc):
    """Computes the real left/right paragraph-indent values that make
    non-table text occupy EXACTLY the same horizontal span as the
    document's own reference table (see _find_reference_margin_table) -
    the user's own explicit methodology: (1) check the page's real
    width and the table's real width to get the total margin area,
    (2) from that, work out how much is on the left vs the right side
    specifically. Returns (target_left, target_right) in twips, or
    (None, None) if no reference table was found (caller should then
    leave indents alone rather than apply a guessed value).

    Confirmed against a real reported document: page width 11900,
    table indent 867, table width 9438 twips -> target_left=867,
    target_right=555 - independently verified these are the exact
    values that make a paragraph's usable width (10860-867-555=9438)
    equal the table's own real width."""
    from docx.oxml.ns import qn

    table = _find_reference_margin_table(doc)
    if table is None:
        return None, None

    sec = doc.sections[0]
    page_width = round(sec.page_width / 635)
    page_right_margin = round(sec.right_margin / 635)

    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        return None, None
    tblInd_el = tblPr.find(qn("w:tblInd"))
    table_indent = int(tblInd_el.get(qn("w:w"))) if tblInd_el is not None else 0
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is None:
        return None, None
    cols = grid.findall(qn("w:gridCol"))
    table_width = sum(int(c.get(qn("w:w"))) for c in cols)

    target_left = table_indent
    table_right_from_page_left_margin_start = round(sec.left_margin / 635) + table_indent + table_width
    target_right = (page_width - page_right_margin) - table_right_from_page_left_margin_start
    if target_right < 0:
        target_right = 0  # a genuinely overflowing reference table shouldn't push this negative
    return target_left, target_right


def _fix_body_paragraph_right_indent(doc):
    """Item (BODY-PARAGRAPH-WRONG-MARGIN, real reported issues 1, 2,
    and 4, refined after further real-document testing) - confirmed
    real: a body (non-table) paragraph's w:left AND w:right indent
    values are both essentially meaningless Aspose-conversion
    artifacts. The FIRST version of this fix only reset w:right to 0,
    leaving w:left completely untouched - confirmed against a real
    document this was NOT enough: paragraphs still carried scattered,
    random left-indent values (1069, 1742, ...), producing exactly the
    "text not starting after the left margin" and "wrapped line
    doesn't match the paragraph's own start" symptoms the user
    continued to report even after the right-indent-only fix shipped.

    REAL, USER-SPECIFIED METHODOLOGY (not guessed): derive the correct
    margin from the document's own real "Contract Data" reference
    table - page width minus table width gives the total margin area,
    and the table's own actual left/right position within the page
    gives the specific left vs right split (see
    _compute_table_derived_margins). Confirmed via an actual
    LibreOffice render, compared directly against a user-provided
    reference image, that applying these table-derived values (rather
    than a flat reset to 0) makes non-table paragraphs align exactly
    with the table's own visual margins.

    Also resets w:firstLine to 0 and removes any w:hanging - confirmed
    real: a leftover non-zero firstLine value was why one paragraph's
    OWN first line started at a different indent than its own wrapped
    continuation lines, even after left/right were corrected.

    Falls back to the simpler "reset right to 0 only" behavior if no
    reference table can be found in the document (keeps working on a
    document that doesn't have this document's specific "Contract
    Data"-style table, rather than skipping the fix or crashing)."""
    from docx.oxml.ns import qn

    target_left, target_right = _compute_table_derived_margins(doc)

    fixed = 0
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        if p._p.getparent().tag == qn("w:tc"):
            continue  # table-cell paragraph - untouched, uses cell width
        pPr = p._p.find(qn("w:pPr"))
        if pPr is None:
            continue
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            continue
        changed = False
        if target_left is not None and target_right is not None:
            if ind.get(qn("w:left")) != str(target_left):
                ind.set(qn("w:left"), str(target_left))
                changed = True
            if ind.get(qn("w:right")) != str(target_right):
                ind.set(qn("w:right"), str(target_right))
                changed = True
        else:
            # No reference table found - fall back to the original,
            # simpler behavior (right-only reset) rather than guessing
            # a left value with no real basis.
            current_right = ind.get(qn("w:right"))
            if current_right is not None and current_right != "0":
                ind.set(qn("w:right"), "0")
                changed = True
        if ind.get(qn("w:firstLine")) not in (None, "0"):
            ind.set(qn("w:firstLine"), "0")
            changed = True
        if ind.get(qn("w:hanging")) is not None:
            del ind.attrib[qn("w:hanging")]
            changed = True
        if changed:
            fixed += 1
    return fixed


def _split_first_numbered_list_item_onto_new_line(doc):
    """Item (NUMBERED-LIST-START-MERGED-WITH-PRECEDING-SENTENCE, real
    reported issue - "agar kisi jagah se numbering order start ho raha
    ho to wo new line se start hona chahiye") - confirmed real via a
    real document AND a real pdftotext -layout render: a cell's text
    "This field indicates the annual rent value...Tenant. 1. Cleaning.
    2. Operations. 3. Security and guarding...." already had items 2
    onward correctly on their own separate lines (confirmed 8 separate
    <w:p> elements in the real cell) - the ONLY genuinely merged part
    was item "1." itself, still attached to the END of the preceding
    intro sentence's own paragraph, instead of starting its own line
    like every item after it.

    Deliberately scoped to specifically "1." (not any digit) - a list
    beginning at 1 is a strong, low-false-positive signal that this is
    genuinely the START of a numbered list (unlike a bare mid-sentence
    digit, which could be anything and isn't safe to split on
    generally). Detects a run whose text contains ". 1. " (sentence end
    directly followed by the first list item) mid-string, splits the
    run's text there, and moves the "1. ..." tail (plus any later runs
    in the same paragraph) into a new paragraph - the exact same
    mechanism as _fix_merged_numbered_subclause's mid-run split, for a
    different real pattern."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from copy import deepcopy

    pattern = re.compile(r"\.\s+(1\.\s)")

    fixed = 0
    for tc in doc.element.body.iter(qn("w:tc")):
        for child in list(tc.iterchildren()):
            if child.tag != qn("w:p"):
                continue
            all_runs = child.findall(qn("w:r"))
            target_run, match = None, None
            for r in all_runs:
                t = r.find(qn("w:t"))
                if t is None or not t.text:
                    continue
                m = pattern.search(t.text)
                if m:
                    target_run, match = r, m
                    break
            if target_run is None:
                continue
            try:
                t = target_run.find(qn("w:t"))
                full_text = t.text
                split_at = match.start(1)
                before_text = full_text[:split_at]
                after_text = full_text[split_at:]
                t.text = before_text

                tail_run = deepcopy(target_run)
                tail_t = tail_run.find(qn("w:t"))
                tail_t.text = after_text
                tail_t.set(qn("xml:space"), "preserve")

                target_run_idx = all_runs.index(target_run)
                pPr_source = child.find(qn("w:pPr"))
                new_p = OxmlElement("w:p")
                if pPr_source is not None:
                    new_p.append(deepcopy(pPr_source))
                new_p.append(tail_run)
                for r in all_runs[target_run_idx + 1:]:
                    child.remove(r)
                    new_p.append(r)
                child.addnext(new_p)
                fixed += 1
            except Exception as err:  # noqa: BLE001
                print(f"[numbered-list-split-fix] skipped one paragraph after an error, continuing with the rest: {err}")
    return fixed


def _merge_continuation_paragraphs(doc):
    """Item (CONTINUATION-TEXT-WRONGLY-SPLIT-INTO-SEPARATE-PARAGRAPHS,
    real reported issue) - confirmed real: what reads as ONE flowing
    sentence/clause is often stored as MULTIPLE separate <w:p>
    paragraph elements in Aspose's conversion - e.g. a real document
    had "14-1 The lessor shall...and in these cases" as one paragraph,
    immediately followed by "the value of bills or rent...returned to
    the tenant." as a SEPARATE paragraph, with no real sentence break
    between them at all. Simply giving both paragraphs the same
    left/right margin (the first version of this whole fix) still left
    a visible seam between them - confirmed via a real render compared
    against a user-provided reference image - because they were still
    two separate paragraph elements with their own (small but nonzero)
    "before" spacing, not one truly continuous paragraph.

    Detects a paragraph that is very likely a continuation of the
    PRECEDING one, using several signals TOGETHER (confirmed via a
    real whole-document dry run that combining them, not any single
    one alone, is what keeps real "Label: Value" style fields - e.g.
    "Name: ...", "Unit Type: Office" - from being wrongly merged into
    unrelated neighboring text, since those don't reliably differ from
    genuine continuations by punctuation alone):
      - this paragraph does NOT start a new clause number ("14-2 ..."),
        a new "Article X:" heading, or is not itself a heading
      - this paragraph does NOT itself look like a "Label: Value" field
      - the PRECEDING real paragraph does NOT end with terminal
        punctuation (. ! ? :) - i.e. looks cut off mid-sentence
      - the PRECEDING paragraph does NOT itself look like a "Label:
        Value" field (a field is normally standalone, not something a
        later paragraph continues)
      - the PRECEDING paragraph is not itself a heading
      - the PRECEDING paragraph's text is not suspiciously short
        (<20 chars) - a short fragment right before this one is more
        likely another disconnected label than the start of a real
        sentence this paragraph continues

    Confirmed via a real whole-document dry run this combination drops
    false-positive candidates from an initial 42 down to a small
    residual few, concentrated in one specific type of content (dense,
    fragmented financial-data labels with no punctuation at all) - per
    explicit direction to implement document-wide anyway and verify
    carefully rather than pre-emptively scope this out, that residual
    risk is accepted, not unaddressed.

    Merges by moving the continuation paragraph's own runs onto the
    end of the preceding paragraph (with a single space between them),
    then removing the now-empty paragraph element - the preceding
    paragraph's own formatting (margins, justify, etc.) already
    applies to the newly-appended text, so no separate re-formatting
    step is needed for the merged-in runs themselves."""
    from docx.oxml.ns import qn
    from copy import deepcopy

    clause_start_pattern = re.compile(r"^\d+-\d+(-\d+)?\s")
    article_pattern = re.compile(r"^Article\s+\S+(?:\s+\S+)?\s*:")
    label_value_pattern = re.compile(r"^[A-Za-z][\w\s/]{0,40}:\s")

    def is_heading(p_el):
        pPr = p_el.find(qn("w:pPr"))
        if pPr is None:
            return False
        shd = pPr.find(qn("w:shd"))
        return shd is not None and shd.get(qn("w:fill")) not in (None, "auto", "FFFFFF")

    body = doc.element.body
    merged = 0
    i = 0
    while True:
        children = list(body.iterchildren())
        if i >= len(children):
            break
        child = children[i]
        if child.tag != qn("w:p") or child.getparent().tag == qn("w:tc"):
            i += 1
            continue
        text = "".join((t.text or "") for t in child.findall(".//" + qn("w:t"))).strip()
        if not text:
            i += 1
            continue

        prev_p, prev_text = None, ""
        for j in range(i - 1, -1, -1):
            if children[j].tag == qn("w:p") and children[j].getparent().tag != qn("w:tc"):
                candidate_text = "".join((t.text or "") for t in children[j].findall(".//" + qn("w:t"))).strip()
                if candidate_text:
                    prev_p, prev_text = children[j], candidate_text
                    break
            elif children[j].tag == qn("w:tbl"):
                break  # a table sits between - don't reach across it

        if prev_p is None:
            i += 1
            continue

        starts_new_unit = bool(clause_start_pattern.match(text)) or bool(article_pattern.match(text)) or is_heading(child)
        looks_like_field = bool(label_value_pattern.match(text))
        prev_looks_like_field = bool(label_value_pattern.match(prev_text))
        prev_ends_terminal = prev_text.endswith((".", "!", "?", ":"))
        prev_is_heading = is_heading(prev_p)
        prev_too_short = len(prev_text) < 20

        should_merge = (
            not starts_new_unit
            and not looks_like_field
            and not prev_ends_terminal
            and not prev_is_heading
            and not prev_looks_like_field
            and not prev_too_short
        )

        if not should_merge:
            i += 1
            continue

        try:
            runs = child.findall(qn("w:r"))
            if runs:
                space_r = deepcopy(runs[0])
                space_t = space_r.find(qn("w:t"))
                if space_t is not None:
                    space_t.text = " "
                    space_t.set(qn("xml:space"), "preserve")
                prev_p.append(space_r)
            for r in runs:
                child.remove(r)
                prev_p.append(r)
            child.getparent().remove(child)
            merged += 1
        except Exception as err:  # noqa: BLE001
            print(f"[continuation-merge-fix] skipped one paragraph after an error, continuing with the rest: {err}")
            i += 1
        # deliberately do NOT advance i - re-check from the same index,
        # since the next sibling has shifted into this position and
        # might ALSO be a continuation of what's now the (bigger) prev_p

    return merged


def _fix_table_row_height_autofit(doc):
    """Item (TABLE-ROW-HEIGHT-NOT-AUTOFIT, real reported issue) -
    confirmed real: every single trHeight in a real reported document
    (95 of them) used hRule="exact" - a FIXED height that clips/
    truncates content that doesn't fit, rather than growing the row to
    fit its actual content. Fix: reset hRule to "auto" everywhere,
    keeping the existing height VALUE as a minimum (auto-fit still
    respects w:val as a floor, per OOXML semantics, but never clips
    taller content the way "exact" does)."""
    from docx.oxml.ns import qn

    fixed = 0
    for trHeight in doc.element.body.iter(qn("w:trHeight")):
        if trHeight.get(qn("w:hRule")) != "auto":
            trHeight.set(qn("w:hRule"), "auto")
            fixed += 1
    return fixed


def _fix_table_vertical_alignment(doc):
    """Item (TABLE-CELL-VERTICAL-ALIGNMENT-INCONSISTENT, real reported
    issue) - confirmed real: a real document's table cells had wildly
    inconsistent vertical alignment (88 "bottom", 48 "center", 1 "top",
    and 246 cells with none set at all, defaulting to top). Fix: every
    non-empty table cell gets w:vAlign set to "center" (middle),
    unconditionally - matches the explicit rule given (table data
    should always be vertically centered)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    fixed = 0
    for tc in doc.element.body.iter(qn("w:tc")):
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            tc.insert(0, tcPr)
        vAlign = tcPr.find(qn("w:vAlign"))
        if vAlign is None:
            vAlign = OxmlElement("w:vAlign")
            tcPr.append(vAlign)
        if vAlign.get(qn("w:val")) != "center":
            vAlign.set(qn("w:val"), "center")
            fixed += 1
    return fixed


def _split_label_value_pairs_within_cell(doc):
    """Item (CELL-CONTAINS-MULTIPLE-DISTINCT-LABEL-VALUE-PAIRS, real
    reported general rule, per an explicit example - a cell whose
    original source text reads like "Nationality:     Indian" should
    render as two distinct pieces WITHIN that one cell: the label
    ("Nationality:") sitting at the cell's own left edge, and the value
    ("Indian") sitting at the cell's own right edge, on the same line -
    not as one continuous run of text with a large literal gap in the
    middle. A single cell can contain more than one such pair (each on
    its own line within the cell).

    Detects the pattern "<label>:<2+ spaces><value>" - the 2+-space gap
    is the real signal (an OCR/extraction artifact from what were
    originally two separately-positioned text elements in the source
    PDF merging into one run with excess whitespace between them) - a
    normal, single-space "Label: Value" is left completely alone, since
    that's ordinary continuous text, not this bug.

    Fix mechanism: for each paragraph inside a table cell whose text
    matches the pattern, sets a right-aligned tab stop at the cell's
    own real width (from its w:tcW, so the value lands exactly at that
    cell's own right edge, not some other cell's), and replaces the
    excess-whitespace gap with a single tab character - this keeps
    label and value on the SAME line, label naturally at the paragraph
    start and value pushed flush to the tab stop, matching the
    left-box/right-box description exactly. If a single paragraph
    contains MORE than one such pair, it is split into one paragraph
    per pair (each getting its own tab-stop treatment) so every pair
    still ends up on its own line within the cell, as described."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from copy import deepcopy

    pattern = re.compile(r"([A-Za-z][\w\s]{1,40}:)\s{2,}([^:]+?)(?=\s{2,}[A-Za-z][\w\s]{1,40}:|$)")

    def _cell_width_twips(tc):
        tcPr = tc.find(qn("w:tcPr"))
        tcW = tcPr.find(qn("w:tcW")) if tcPr is not None else None
        if tcW is not None and tcW.get(qn("w:w")):
            try:
                return int(tcW.get(qn("w:w")))
            except ValueError:
                pass
        return None

    def _set_right_tab(pPr, pos):
        tabs = pPr.find(qn("w:tabs"))
        if tabs is None:
            tabs = OxmlElement("w:tabs")
            pPr.insert(0, tabs)
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), str(pos))
        tabs.append(tab)

    fixed = 0
    for tc in doc.element.body.iter(qn("w:tc")):
        cell_width = _cell_width_twips(tc)
        if not cell_width:
            continue
        for p in list(tc.findall(qn("w:p"))):
            all_runs = p.findall(qn("w:r"))
            target_run, matches = None, []
            for r in all_runs:
                t = r.find(qn("w:t"))
                if t is None or not t.text:
                    continue
                found = list(pattern.finditer(t.text))
                if found:
                    target_run, matches = r, found
                    break
            if not matches:
                continue

            try:
                t = target_run.find(qn("w:t"))
                pPr = p.find(qn("w:pPr"))
                if pPr is None:
                    pPr = OxmlElement("w:pPr")
                    p.insert(0, pPr)

                pairs = [(m.group(1).strip(), m.group(2).strip()) for m in matches]
                # first pair replaces this paragraph's own text + gets the tab stop
                label0, value0 = pairs[0]
                t.text = label0
                tab_run = OxmlElement("w:r")
                tab_run.append(OxmlElement("w:tab"))
                # insert tab between label run and a new value run
                value_run = deepcopy(target_run)
                value_t = value_run.find(qn("w:t"))
                value_t.text = value0
                target_run.addnext(value_run)
                target_run.addnext(tab_run)
                _set_right_tab(pPr, cell_width)

                # any FURTHER pairs each become their own new paragraph,
                # inserted right after this one, each with the same
                # tab-stop treatment.
                insert_after = p
                for label, value in pairs[1:]:
                    new_p = OxmlElement("w:p")
                    new_pPr = OxmlElement("w:pPr")
                    new_p.append(new_pPr)
                    _set_right_tab(new_pPr, cell_width)
                    label_run = deepcopy(target_run)
                    label_t = label_run.find(qn("w:t"))
                    label_t.text = label
                    new_p.append(label_run)
                    new_tab_run = OxmlElement("w:r")
                    new_tab_run.append(OxmlElement("w:tab"))
                    new_p.append(new_tab_run)
                    val_run = deepcopy(target_run)
                    val_t = val_run.find(qn("w:t"))
                    val_t.text = value
                    new_p.append(val_run)
                    insert_after.addnext(new_p)
                    insert_after = new_p
                    fixed += 1

                fixed += 1
            except Exception as err:  # noqa: BLE001
                print(f"[label-value-split-fix] skipped one cell paragraph after an error, continuing with the rest: {err}")
    return fixed


def _fix_table_cell_alignment_by_length(doc, length_threshold=50):
    """Item (TABLE-CELL-ALIGNMENT-SHORT-VS-LONG, real reported rule -
    "table me agar short text he to wo cell me center me ayega lekin
    agar long text ho to wo left side se saru hona chahiye jaisa humne
    article ke liye set kiya tha") - mirrors the SAME short/long
    distinction already established for body Article/clause text
    elsewhere in this file, applied here to table DATA cells
    specifically (never header/title cells - those are always
    centered regardless of length, per _fix_table_header_alignment,
    which is a completely separate rule).

    Threshold picked from this project's own real data: a direct
    length-distribution check across a real reported document's 355
    non-empty cells showed a natural gap around 40-50 characters
    (66% of cells are pure short labels/values under 20 chars; content
    at or above ~50 chars is consistently genuine sentence-like
    explanatory text, not a label). Cells at or under the threshold get
    jc="center"; cells over it get jc="left" (not "both"/justify - per
    the explicit rule given, table cell long-text is LEFT-aligned,
    distinct from the "both" rule used for body Article/clause
    paragraphs elsewhere in this file). Header/title rows (detected the
    same way as _fix_table_header_alignment) are skipped entirely -
    they stay centered unconditionally regardless of their own length."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def is_header_row(tr_el):
        shd = tr_el.find(".//" + qn("w:shd"))
        return shd is not None and shd.get(qn("w:fill")) not in (None, "auto", "FFFFFF")

    fixed = 0
    for tbl_el in doc.element.body.iter(qn("w:tbl")):
        rows = tbl_el.findall(qn("w:tr"))
        for i, tr_el in enumerate(rows):
            if i == 0 and is_header_row(tr_el):
                continue  # header/title row - always centered, a separate rule
            for tc in tr_el.findall(qn("w:tc")):
                ts = tc.findall(".//" + qn("w:t"))
                text = "".join(t.text or "" for t in ts).strip()
                if not text:
                    continue
                target = "center" if len(text) <= length_threshold else "left"
                for p in tc.findall(qn("w:p")):
                    pPr = p.find(qn("w:pPr"))
                    if pPr is None:
                        pPr = OxmlElement("w:pPr")
                        p.insert(0, pPr)
                    jc = pPr.find(qn("w:jc"))
                    if jc is None:
                        jc = OxmlElement("w:jc")
                        pPr.append(jc)
                    if jc.get(qn("w:val")) != target:
                        jc.set(qn("w:val"), target)
                        fixed += 1
    return fixed


def _fix_table_header_alignment(doc):
    """Item (TABLE-HEADER-NOT-CENTERED, real reported issue) -
    confirmed real: a real document's shaded header/title rows used a
    scattered mix of "both" (dominant), "left", and "center" for their
    own jc value. Fix: every paragraph inside a header row (detected
    via the same dark-shading signature used elsewhere in this file -
    _merge_adjacent_header_tables's header_row_shading) gets jc set to
    "center", unconditionally - matches the explicit rule given (table
    headers/titles should always be centered)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def is_header_row(tr_el):
        shd = tr_el.find(".//" + qn("w:shd"))
        return shd is not None and shd.get(qn("w:fill")) not in (None, "auto", "FFFFFF")

    fixed = 0
    for tbl_el in doc.element.body.iter(qn("w:tbl")):
        rows = tbl_el.findall(qn("w:tr"))
        if not rows or not is_header_row(rows[0]):
            continue
        for p in rows[0].iter(qn("w:p")):
            pPr = p.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                p.insert(0, pPr)
            jc = pPr.find(qn("w:jc"))
            if jc is None:
                jc = OxmlElement("w:jc")
                pPr.append(jc)
            if jc.get(qn("w:val")) != "center":
                jc.set(qn("w:val"), "center")
                fixed += 1
    return fixed


_ARTICLE_NUMBER_WORDS = {
    "One": 1, "Two": 2, "Three": 3, "Four": 4, "Five": 5, "Six": 6,
    "Seven": 7, "Eight": 8, "Nine": 9, "Ten": 10, "Eleven": 11,
    "Twelve": 12, "Thirteen": 13, "Fourteen": 14, "Fifteen": 15,
    "Sixteen": 16, "Seventeen": 17, "Eighteen": 18, "Nineteen": 19,
    "Twenty": 20,
}


def _convert_article_word_numbers_to_digits(doc):
    """Item (ARTICLE-HEADING-WORD-NUMBER, real reported issue, per
    explicit direction to build a rule identifying word-form numbering
    in general, not hardcode one document's specific article count) -
    confirmed real: a real document had all 17 of its "Article <Word>:"
    headings spelled out ("Article One:" ... "Article Seventeen:")
    instead of digits ("Article 1:" ... "Article 17:"). Fix: replaces
    "Article <Word>" with "Article <digit>" wherever the word exactly
    matches a known English number-word (One through Twenty - covers
    real contract lengths with real headroom). Works per-run first (the
    common case, since a heading's "Article One" text usually lives in
    one run); if the pattern only resolves at the whole-paragraph text
    level (split across multiple runs - rare, but real for translated
    text), the fix is applied by rewriting the first run to the full
    corrected text and clearing the rest, rather than leaving it
    unfixed."""
    from docx.oxml.ns import qn

    pattern = re.compile(r"\bArticle (" + "|".join(_ARTICLE_NUMBER_WORDS.keys()) + r")\b")

    def _replace(m):
        return f"Article {_ARTICLE_NUMBER_WORDS[m.group(1)]}"

    fixed = 0
    for p in doc.paragraphs:
        if not pattern.search(p.text):
            continue
        resolved_per_run = True
        for r in p.runs:
            if r.text and pattern.search(r.text):
                r.text = pattern.sub(_replace, r.text)
        if pattern.search(p.text) and p.runs:
            # per-run replacement didn't fully resolve it (pattern
            # spanned multiple runs) - rebuild from full paragraph text.
            new_full = pattern.sub(_replace, p.text)
            p.runs[0].text = new_full
            for r in p.runs[1:]:
                r.text = ""
        fixed += 1
    return fixed


def _fix_clause_start_spacing(doc):
    """Item (NO-VISIBLE-GAP-BETWEEN-CLAUSES, real reported issue,
    follow-up to _merge_continuation_paragraphs) - confirmed real via
    an actual rendered comparison against a user-provided reference
    image: after continuation paragraphs are correctly merged back
    into one real paragraph per clause, adjacent CLAUSES (e.g. "14-1"
    ending and "14-2" beginning) still had no visible gap between them
    - the "before" spacing on a real document's own clause-start
    paragraphs varied inconsistently (134 twips on one, only 22 on the
    next), so some clauses got a visible gap and others didn't, for no
    structural reason.

    Fix: every paragraph that starts a new numbered clause (matches the
    same "<digits>-<digits>(-<digits>)? " pattern
    _merge_continuation_paragraphs and _fix_merged_numbered_subclause
    both use) gets the same "before" spacing value - the real value a
    working clause-start already used in the reported document (134
    twips), confirmed via an actual render to produce a clear, visible
    gap without looking exaggerated."""
    from docx.oxml.ns import qn

    clause_start_pattern = re.compile(r"^\d+-\d+(-\d+)?\s")
    GAP_BEFORE = "134"

    fixed = 0
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text or not clause_start_pattern.match(text):
            continue
        if p._p.getparent().tag == qn("w:tc"):
            continue
        pPr = p._p.find(qn("w:pPr"))
        spacing = pPr.find(qn("w:spacing")) if pPr is not None else None
        if spacing is not None and spacing.get(qn("w:before")) != GAP_BEFORE:
            spacing.set(qn("w:before"), GAP_BEFORE)
            fixed += 1
    return fixed


def _promote_uniform_run_shading_to_paragraph(doc):
    """Item (RUN-LEVEL-SHADING-NOT-FULL-WIDTH, real reported issue 3) -
    confirmed real: "Article Four: Rent" carried its DDDDDD background
    color on the RUN's own rPr (w:rPr/w:shd), not the paragraph's pPr -
    confirmed directly in the real document's XML. A run-level shading
    only highlights the exact characters of that run, not the full
    line/row between the page's margins - Word only extends a
    background across the full paragraph width when the shading lives
    on the paragraph (w:pPr/w:shd), not a run inside it.

    Per the user's own stated condition: only promote when the WHOLE
    paragraph shares ONE uniform shading color - if some runs have a
    different color (or no shading at all) while others do, promoting
    would incorrectly paint the whole row instead of just the
    genuinely-shaded portion, so that case is left alone.

    Fix: for each paragraph, if every run that has any text carries
    the exact same non-empty w:shd fill color, move that shading onto
    the paragraph's own pPr (so it spans the full margin-to-margin
    width) and remove it from the individual runs (avoiding a
    redundant, now-inconsistent duplicate). Table-cell paragraphs are
    untouched (a cell's own shading/width behaves differently - this is
    specifically about non-table body text)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    fixed = 0
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        if p._p.getparent().tag == qn("w:tc"):
            continue  # table-cell paragraph - untouched

        runs_with_text = [r for r in p.runs if r.text]
        if not runs_with_text:
            continue

        fills = []
        for r in runs_with_text:
            rPr = r._r.find(qn("w:rPr"))
            shd = rPr.find(qn("w:shd")) if rPr is not None else None
            fills.append(shd.get(qn("w:fill")) if shd is not None else None)

        uniform_fill = fills[0]
        if not uniform_fill or any(f != uniform_fill for f in fills):
            continue  # no shading, or not uniform across every run - leave alone

        pPr = p._p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p._p.insert(0, pPr)
        para_shd = pPr.find(qn("w:shd"))
        if para_shd is None:
            para_shd = OxmlElement("w:shd")
            pPr.append(para_shd)
        para_shd.set(qn("w:val"), "clear")
        para_shd.set(qn("w:color"), "auto")
        para_shd.set(qn("w:fill"), uniform_fill)

        for r in runs_with_text:
            rPr = r._r.find(qn("w:rPr"))
            run_shd = rPr.find(qn("w:shd")) if rPr is not None else None
            if run_shd is not None:
                rPr.remove(run_shd)

        fixed += 1
    return fixed


def _fix_paragraph_direction(doc, target_language):
    """Item (RTL/LTR-NOT-CORRECTED) - after in-place translation
    replaces a paragraph's text, the paragraph's own w:bidi (RTL
    paragraph flag) and each run's w:rtl flag were being left untouched
    - so a paragraph that used to hold right-to-left Arabic still told
    Word/LibreOffice "this paragraph is RTL" even though it now holds
    left-to-right English. Confirmed by scanning a real translated
    output: 328 paragraphs contained clearly-Latin-script (English) text
    while still carrying <w:bidi/> with no val attribute (which per
    OOXML defaults to true/on) - this was the direct cause of reported
    colon-placement, alignment, and unnecessary-wrapping bugs, since the
    bidi algorithm was still being applied to now-LTR content. Verified
    by rendering a real affected document before/after this fix: e.g.
    "Contract Sealing :Location" (broken) became "Contract Sealing
    Location:" (correct) with no other change.

    Sets every paragraph's w:bidi and every run's w:rtl to match the
    TARGET language's actual direction (RTL only for Arabic, Hebrew,
    Persian/Farsi, Urdu, Pashto, Dhivehi, Yiddish, Sindhi - LTR for
    everything else). Applied document-wide since a full-pipeline
    translation converts virtually the entire document to one target
    language, so the whole document should read in one consistent
    direction, not a per-paragraph mix left over from the source PDF's
    original bilingual (Arabic+English side-by-side) layout. Tested
    against an already-LTR source/target document too (Italian->English)
    to confirm no regression - it still found 27 stray bidi flags there
    and clearing them left the rendered output visually identical.

    Also fixes reversed clause/article-number prefixes exposed by the
    RTL->LTR flip - see _fix_reversed_clause_number_prefix()'s docstring
    for the full root-cause explanation.

    Also converts w:jc="both" (full justify) to "left" for paragraphs
    flipping from RTL to LTR - confirmed real complaint, with evidence:
    508 of 575 real translated paragraphs in an affected document kept
    jc="both" (inherited from the original Arabic layout, which
    routinely uses full justification), which - now holding English
    text of very different average word/line-length than the Arabic it
    replaced - produces visibly oversized, uneven gaps between words to
    stretch each line to full width. In table cells with short content
    (e.g. a single digit in a narrow "Clause Number" column), that same
    justify behavior can look like the text is oddly spaced or even
    reads as if the column wasn't resized at all, even when the
    column's own width is correct. Only touches paragraphs making the
    RTL->LTR transition (same evidence-gated condition as the clause-
    number fix) - "left" is a safe, unambiguous default for translated
    body text and table cells; jc="center" is left untouched since
    centering isn't direction-dependent.

    Returns (paragraphs_fixed, clause_numbers_fixed)."""
    import unicodedata
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    want_rtl = _is_rtl_language(target_language)
    fixed = 0
    clause_numbers_fixed = 0
    errors_encountered = 0

    # Item (WRONG-HEADING-SIGNAL, confirmed real bug) - the margin-
    # normalization fix below needs to tell a genuine Article heading
    # apart from ordinary content, to leave the heading's own indent
    # alone. The first attempt at this checked for "any non-white run
    # shading" - too broad: this same document ALSO uses a lighter gray
    # (EEEEEE) for plain alternating-row background striping on regular
    # data paragraphs (e.g. "Email aalmasoud@osoolre.com..."), which
    # got mistaken for a heading and left with its own leftover random
    # indent untouched. w:outlineLvl doesn't distinguish them either -
    # confirmed both a real heading AND a striped data row carry
    # outlineLvl=9 in this document's Aspose conversion. Rather than
    # hardcoding a specific color (which could easily be wrong for a
    # different document/color scheme - the same "don't hardcode this
    # document" principle used throughout this file), this instead
    # finds a REFERENCE heading by its TEXT pattern ("Article X:" /
    # "المادة X:" at the very start of a short paragraph - the one
    # reliable, language-agnostic signal available) and reads back
    # WHATEVER shading color that reference heading actually uses,
    # then matches against that exact value for the rest of this
    # function - correct for any document's own actual heading style,
    # not just this one's DDDDDD.
    _heading_shd_fill = None
    for p_ref in _iter_paragraphs_in_order(doc):
        ref_text = unicodedata.normalize("NFKC", (p_ref.text or "").strip())
        if not re.match(r"^(Article|المادة)\s+\S+(?:\s+\S+)?\s*:", ref_text) or len(ref_text) > 80:
            continue
        for r_ref in p_ref.runs:
            if not (r_ref.text and r_ref.text.strip()):
                continue
            rpr_ref = r_ref._element.find(qn("w:rPr"))
            shd_ref = rpr_ref.find(qn("w:shd")) if rpr_ref is not None else None
            if shd_ref is not None:
                fill = (shd_ref.get(qn("w:fill")) or "").lower()
                if fill not in ("", "auto", "ffffff"):
                    _heading_shd_fill = fill
            break
        if _heading_shd_fill:
            break

    # Item (STANDALONE-DATA-PARAGRAPH-BREAKS-OUT-OF-TABLE-BOX) - this
    # was previously handled by aligning a standalone data-row paragraph
    # to its neighboring table's own indent, as an exception to the old
    # flat-normalize default. Removed along with that default: per the
    # redesigned mirror-based rule below, a paragraph's OWN original
    # left/right margin is now preserved (mirrored, not replaced), so
    # there is no longer a flat generic value for a standalone data row
    # to diverge from in the first place - this table-adjacency lookup
    # is no longer needed.

    # Item (SILENT-CASCADING-FAILURE) - confirmed real bug pattern: in a
    # real production run, an ENTIRE block of 24 consecutive clause
    # numbers (all of Article Five's obligations) failed to get
    # reversed together, while clauses before and after that block were
    # unaffected - exactly the signature of ONE paragraph raising an
    # exception partway through this loop, which the caller's
    # try/except then silently swallows (rebuild_docx_with_translated_
    # text wraps this whole function in "except Exception: pass"),
    # leaving every paragraph from that point onward completely
    # unprocessed - not just its clause number, but its bidi/rtl flags
    # too. Extensive isolated testing (49/49 real clause numbers
    # correctly reversed) could not reproduce the specific paragraph
    # that raised, but the FIX for this class of bug doesn't require
    # reproducing it: wrapping each paragraph's own processing in its
    # own try/except means one paragraph's edge case can only skip that
    # one paragraph, never cascade into silently abandoning every
    # paragraph after it.
    for p in _iter_paragraphs_in_order(doc):
        try:
            pPr = p._p.find(qn("w:pPr"))
            changed = False

            was_rtl = pPr is not None and pPr.find(qn("w:bidi")) is not None
            if was_rtl and not want_rtl:
                if _fix_reversed_clause_number_prefix(p):
                    clause_numbers_fixed += 1
                if pPr is not None:
                    # Item (CENTER-ALIGNED-CLAUSE-BODY) - confirmed
                    # real, pixel-measured bug: a user-provided
                    # screenshot showed a clause paragraph's wrapped
                    # lines each starting at a WILDLY different x
                    # position (159 to 278px, a 119px spread) instead
                    # of a consistent left margin - the exact visual
                    # signature of CENTER alignment (each line
                    # centers based on its own width) rather than
                    # left. Confirmed root cause directly in Aspose's
                    # own untranslated conversion: this specific
                    # clause paragraph had jc="center" while every
                    # sibling clause paragraph around it (same
                    # Article, same structure) correctly had
                    # jc="both" - a one-off Aspose inconsistency, not
                    # something translation introduced. The earlier
                    # both->left fix deliberately left jc="center"
                    # alone project-wide (reasoned that centering
                    # isn't inherently RTL/LTR-dependent) - true in
                    # general, but wrong for this specific case: a
                    # numbered clause's BODY TEXT should never be
                    # centered regardless of source direction. Scoped
                    # narrowly here (only applies when the
                    # paragraph's own text starts with a digit, i.e.
                    # it looks like clause/list body content) so a
                    # genuinely-intentional centered short title or
                    # signature line elsewhere isn't touched.
                    first_run_text = next((r.text for r in p.runs if r.text and r.text.strip()), "")
                    looks_like_clause_body = bool(re.match(r"^\d", first_run_text.strip()))

                    # Item (CENTER-FIX-ALSO-TOO-NARROW, confirmed real,
                    # same gap as the margin-normalization fix below) -
                    # the original center->left fix only covered digit-
                    # prefixed clause paragraphs, so a real, non-
                    # numbered body paragraph ("All addresses,
                    # correspondence, notices...", Article Fifteen's
                    # body) kept jc="center" - invisible across most of
                    # a long paragraph's full-width lines, but visibly
                    # wrong on its short last line ("recognized legal
                    # means."), which centered instead of sitting at the
                    # left margin, looking like clipped/misplaced text.
                    # Computed once here (moved up from where it used to
                    # live, further below) so both the jc fix and the
                    # margin fix share the same real scope: any body-
                    # flow paragraph that is NOT inside a table cell and
                    # NOT a heading (by the document's own reference-
                    # heading shading color, found above).
                    is_table_cell_paragraph = p._p.getparent() is not None and p._p.getparent().tag == qn("w:tc")
                    _first_run_rpr_for_heading_check = None
                    for r in p.runs:
                        if r.text and r.text.strip():
                            _first_run_rpr_for_heading_check = r._element.find(qn("w:rPr"))
                            break
                    is_heading_styled = False
                    if _first_run_rpr_for_heading_check is not None and _heading_shd_fill is not None:
                        shd = _first_run_rpr_for_heading_check.find(qn("w:shd"))
                        if shd is not None and (shd.get(qn("w:fill")) or "").lower() == _heading_shd_fill:
                            is_heading_styled = True
                    is_normal_body_paragraph = not is_table_cell_paragraph and not is_heading_styled

                    jc = pPr.find(qn("w:jc"))
                    if jc is not None:
                        jc_val = jc.get(qn("w:val"))
                        if jc_val == "both" or (jc_val == "center" and is_normal_body_paragraph):
                            jc.set(qn("w:val"), "left")
                            changed = True

                    # Item (RAGGED-LEFT-MARGIN-FROM-SOURCE-POSITIONING) -
                    # confirmed real, user-reported bug with a clear root
                    # cause: measured a real document's clause paragraphs
                    # and found w:ind (left/right indent) values that
                    # vary essentially randomly between SIBLING clauses
                    # at the identical hierarchy depth - e.g. clauses
                    # 1-1-5 through 15-1-5 (all direct children of the
                    # same Article Five / Section One, so semantically
                    # equivalent) carry left-indent values ranging from
                    # 809 to 1661 twips and right-indent from -200 to
                    # 2439 twips, with no correlation to clause number,
                    # depth, or any other structural property. These
                    # aren't meaningful hierarchy markers - they're
                    # leftover artifacts from Aspose trying to reproduce
                    # the ORIGINAL ARABIC PDF's exact per-line visual
                    # positioning (which was tuned for Arabic word
                    # lengths and RTL layout), carried over unchanged
                    # onto English text with completely different word
                    # lengths - producing the jagged, inconsistent left
                    # margin.
                    #
                    # Item (NOT-ACTUALLY-GLOBAL, confirmed by direct
                    # user report with evidence) - the FIRST version of
                    # this fix only normalized paragraphs matching
                    # looks_like_clause_body (a leading digit) - but the
                    # SAME random-indent defect is not limited to
                    # numbered clauses. A real screenshot showed the
                    # plain prose paragraph directly under "Article
                    # Twelve" ("The parties agree to designate a third
                    # party...", no leading number at all) with the
                    # exact same symptom - ind left=1439/right=475/
                    # firstLine=534, a leftover Aspose value, while its
                    # numbered sibling clauses were already correctly
                    # normalized - producing a visibly inconsistent,
                    # oddly-indented paragraph sitting right next to
                    # properly-margined ones. Since the root cause
                    # (arbitrary per-paragraph source-positioning
                    # artifacts) applies to ANY body-flow paragraph, not
                    # just numbered ones, this now normalizes every
                    # paragraph in the normal document flow - EXCLUDING
                    # only: (a) paragraphs inside a table cell (handled
                    # separately, by the table-width/column fixes, which
                    # need cell-relative not body-relative margins), and
                    # (b) Article headings themselves, identified by
                    # matching the document's OWN reference heading
                    # shading color (_heading_shd_fill, found once
                    # above - see that note for why an exact-color match
                    # is used instead of "any non-white shading") - a
                    # heading's indent is part of its intentional
                    # heading style and must stay as-is. is_table_cell_
                    # paragraph/is_heading_styled were already computed
                    # above (shared with the jc fix) - reused as-is here.
                    #
                    # REDESIGNED per explicit direction, replacing the
                    # earlier "flatten every paragraph's left-indent to
                    # one hardcoded constant (446 twips, or 0 for
                    # headings), discard right/firstLine/hanging
                    # entirely" approach. That approach was based on a
                    # real observation (sibling clause left-indents
                    # varying 809-1661 twips with no obvious structural
                    # correlation) but the conclusion drawn from it -
                    # "this variance is meaningless Aspose noise, safe
                    # to discard" - was never actually verified against
                    # the ORIGINAL untranslated RTL document to confirm
                    # the variance wasn't genuine, meaningful original
                    # positioning. The corrected rule: an RTL->LTR
                    # paragraph's own original left/right margin is
                    # MIRRORED (new_left = old_right, new_right =
                    # old_left, width/content unaffected), not replaced
                    # by an external constant - preserving whatever
                    # real positioning data the original paragraph
                    # actually had, for both ordinary body paragraphs
                    # and headings alike (a heading's indent was
                    # ALSO previously flattened to a hardcoded 0,
                    # which is the same category of discard-instead-
                    # of-preserve mistake).
                    should_mirror_margin = is_normal_body_paragraph or (is_heading_styled and not is_table_cell_paragraph)

                    if should_mirror_margin:
                        ind = pPr.find(qn("w:ind"))
                        if ind is not None:
                            old_left = int(ind.get(qn("w:left")) or "0")
                            old_right = int(ind.get(qn("w:right")) or "0")
                            if old_left != old_right:
                                ind.set(qn("w:left"), str(old_right))
                                ind.set(qn("w:right"), str(old_left))
                                changed = True
                        # if there's no w:ind element at all, there is no
                        # original margin to mirror - correctly left
                        # untouched rather than inventing one.

            if want_rtl:
                if pPr is None:
                    pPr = OxmlElement("w:pPr")
                    p._p.insert(0, pPr)
                bidi = pPr.find(qn("w:bidi"))
                if bidi is None:
                    pPr.append(OxmlElement("w:bidi"))
                    changed = True
                elif bidi.get(qn("w:val")) == "0":
                    del bidi.attrib[qn("w:val")]
                    changed = True
            else:
                if pPr is not None:
                    bidi = pPr.find(qn("w:bidi"))
                    if bidi is not None:
                        pPr.remove(bidi)
                        changed = True

            for r in p.runs:
                rpr = r._element.find(qn("w:rPr"))
                if rpr is None:
                    if not want_rtl:
                        continue
                    rpr = OxmlElement("w:rPr")
                    r._element.insert(0, rpr)
                rtl_el = rpr.find(qn("w:rtl"))
                if want_rtl:
                    if rtl_el is None:
                        rpr.append(OxmlElement("w:rtl"))
                        changed = True
                    elif rtl_el.get(qn("w:val")) == "0":
                        del rtl_el.attrib[qn("w:val")]
                        changed = True
                else:
                    if rtl_el is not None:
                        rpr.remove(rtl_el)
                        changed = True

            if changed:
                fixed += 1
        except Exception as err:  # noqa: BLE001
            errors_encountered += 1
            print(f"[direction-fix] skipped one paragraph after an error, continuing with the rest: {err}")

    if errors_encountered:
        print(f"[direction-fix] {errors_encountered} paragraph(s) skipped due to errors - see messages above")

    return fixed, clause_numbers_fixed


def _fix_exact_row_heights(doc):
    """Item (VERTICAL-CLIPPING) - many of Aspose's table rows use
    <w:trHeight w:hRule="exact"/> - a FIXED height, sized for the
    original (often more compact) source-language text. Translated text
    needing more vertical space (longer English phrases, or text that
    now wraps to 2+ lines) gets visually clipped or overflows outside
    the fixed-height row instead of the row growing - confirmed on a
    real translated document showing headers like "Number of Parking
    Number of / Lots Elevators" overlapping garbled text at exact
    height. Switching hRule from "exact" to "atLeast" keeps the same
    MINIMUM height (short content still looks the same) but lets the
    row grow taller when its content genuinely needs more room. Found
    104 such rows in a real affected document; fixing this eliminated
    the clipping/overlap with no visible change to rows that didn't
    need to grow."""
    from docx.oxml.ns import qn

    fixed = 0
    for tr in doc.element.body.iter(qn("w:tr")):
        trPr = tr.find(qn("w:trPr"))
        if trPr is None:
            continue
        trHeight = trPr.find(qn("w:trHeight"))
        if trHeight is not None and trHeight.get(qn("w:hRule")) == "exact":
            trHeight.set(qn("w:hRule"), "atLeast")
            fixed += 1
    return fixed


def _remove_duplicate_table_rows(doc):
    """Item (DUPLICATE-TABLE-ROW, real reported bug): a real Appendix
    table ("Clause Number | Field | Clarification") had one row
    ("1 | Contract Execution Date | This field indicates the date of
    documenting the lease contract by both parties.") appearing TWICE
    in a row, character-for-character identical in every cell -
    confirmed directly in the real document's XML, not assumed.

    DELIBERATELY CONSERVATIVE, to avoid a real false-positive found in
    the SAME document: two rows can legitimately share the same Field
    name with GENUINELY DIFFERENT Clarification text (e.g. "Lessor"
    appeared twice with two different, real clarifications about who a
    lessor can be vs. what identity types apply) - that is real,
    intentional content, not a bug, and must never be removed. So this
    only removes a row when it is an EXACT match (every cell's text,
    in order) to the row IMMEDIATELY BEFORE it - never "same field
    name", never a fuzzy/partial match, and never a match to some
    earlier non-adjacent row (a document could legitimately repeat an
    identical short value like a single clause number far apart -
    adjacency is what makes this specific pattern look like an
    accidental double-write rather than real repeated content).

    Returns the number of rows removed."""
    from docx.oxml.ns import qn

    removed = 0
    for tbl_el in doc.element.body.iter(qn("w:tbl")):
        rows = tbl_el.findall(qn("w:tr"))
        prev_text = None
        for tr in rows:
            cell_texts = []
            for t_el in tr.iter(qn("w:t")):
                cell_texts.append(t_el.text or "")
            row_text = "\x1f".join(cell_texts)  # unlikely-to-collide separator
            if prev_text is not None and row_text == prev_text and row_text.strip("\x1f").strip():
                tr.getparent().remove(tr)
                removed += 1
                # prev_text stays the same - if a THIRD identical row
                # somehow followed, it would also be caught against the
                # same still-standing previous row.
                continue
            prev_text = row_text
    return removed


def _fix_mismatched_column_header_labels(doc):
    """Item (TABLE-HEADER-LABEL-DOESNT-MATCH-COLUMN-CONTENT, real
    reported issue - "table ke header check karo galat show ho rahe he
    aur translation sahi se nahi hua he") - confirmed real directly in
    a reported document: an appendix table's header row read
    ['Explanation', 'Field', 'Item Number'], but its own data rows
    showed column 0 holding short numeric IDs ("1", "2", ...) and
    column 2 holding the long explanatory sentences - i.e. the labels
    for columns 0 and 2 were swapped relative to what those columns
    actually contain (translation-injection apparently matched the
    header text to the wrong cell for this table).

    Detects this via each column's own real content: a header whose
    text matches a "count/number"-style label (contains "number") is
    expected to sit over a column of SHORT values; a header matching a
    "description"-style label (contains "explanation", "clarification",
    "description", or "details") is expected to sit over a column of
    LONG values. When a table's header row has exactly one column of
    each kind and their actual average content length is the OPPOSITE
    of what their own labels imply, the two header cells' text is
    swapped to match reality - the data itself is never touched, only
    the header labeling."""
    from docx.oxml.ns import qn

    NUMBER_WORDS = ("number", "no.", "no ")
    DESC_WORDS = ("explanation", "clarification", "description", "details")

    def _cell_text(tc):
        ts = tc.findall(".//" + qn("w:t"))
        return "".join(t.text or "" for t in ts).strip()

    fixed = 0
    for tbl_el in doc.element.body.iter(qn("w:tbl")):
        rows = tbl_el.findall(qn("w:tr"))
        if len(rows) < 3:
            continue
        header_cells = rows[0].findall(qn("w:tc"))
        if len(header_cells) < 2:
            continue

        header_texts = [_cell_text(tc) for tc in header_cells]
        number_col = next((i for i, t in enumerate(header_texts) if any(w in t.lower() for w in NUMBER_WORDS)), None)
        desc_col = next((i for i, t in enumerate(header_texts) if any(w in t.lower() for w in DESC_WORDS)), None)
        if number_col is None or desc_col is None or number_col == desc_col:
            continue

        data_rows = rows[1:]
        number_col_lengths, desc_col_lengths = [], []
        for row in data_rows:
            cells = row.findall(qn("w:tc"))
            if len(cells) <= max(number_col, desc_col):
                continue
            number_col_lengths.append(len(_cell_text(cells[number_col])))
            desc_col_lengths.append(len(_cell_text(cells[desc_col])))
        if not number_col_lengths:
            continue
        avg_number_col = sum(number_col_lengths) / len(number_col_lengths)
        avg_desc_col = sum(desc_col_lengths) / len(desc_col_lengths)

        # Mismatch: the column labeled "number" actually holds the
        # LONGER content, and the column labeled "explanation" holds
        # the SHORTER content - the opposite of what the labels imply.
        if avg_number_col > avg_desc_col and avg_number_col > 20 and avg_desc_col < 10:
            number_tc = header_cells[number_col]
            desc_tc = header_cells[desc_col]
            number_ts = number_tc.findall(".//" + qn("w:t"))
            desc_ts = desc_tc.findall(".//" + qn("w:t"))
            if number_ts and desc_ts:
                number_first_text = number_ts[0].text or ""
                desc_first_text = desc_ts[0].text or ""
                number_ts[0].text = desc_first_text
                for t in number_ts[1:]:
                    t.text = ""
                desc_ts[0].text = number_first_text
                for t in desc_ts[1:]:
                    t.text = ""
                fixed += 1
    return fixed


def _merge_adjacent_header_tables(doc):
    """Item (APPENDIX-TABLE-HEADER-WRONG-PAGE-POSITION, real reported
    issue 4) - confirmed real: a real document's Appendix content was
    split into 3 SEPARATE <w:tbl> elements (16, 14, and 15 rows), each
    carrying its OWN "Clause Number | Field | Explanation" header row,
    with NOTHING but empty filler between them (confirmed directly:
    the XML between consecutive appendix tables was empty paragraph
    markup, no real text). Word/LibreOffice paginates each of these as
    its own independent table, so a NEW table's own header can start
    partway down a page (right after the previous table's last row
    finishes) - which is what LOOKS like "the header is on the wrong
    row" but is actually a separate table's own correctly-positioned
    header, just visually indistinguishable from a genuine continuation
    of one bigger table.

    Detects consecutive tables whose header row matches the same
    generic "looks like a real column-header row" signature (checked
    via shading, since Aspose consistently uses a distinct dark fill
    for header rows - not hardcoded to specific header text, so this
    works for any language) with only empty/whitespace content between
    them, and merges them into ONE real table: keeps the FIRST table's
    header row, appends every OTHER row from the later tables (skipping
    each one's own duplicate header row), removes the now-empty later
    <w:tbl> elements and the empty filler between them.

    Also marks the surviving header row with <w:trPr><w:tblHeader/></w:trPr>
    - the real, standard OOXML "repeat this row at the top of every
    page this table spans" mechanism - so once these fragments are
    correctly merged into one continuous table, its header genuinely
    repeats at the top of EVERY page it spans, not just wherever it
    happened to start."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def header_row_shading(tbl_el):
        rows = tbl_el.findall(qn("w:tr"))
        if not rows:
            return None
        first_row = rows[0]
        shd = first_row.find(".//" + qn("w:shd"))
        return shd.get(qn("w:fill")) if shd is not None else None

    body = doc.element.body
    tables_merged = 0

    while True:
        children = list(body.iterchildren())
        merged_this_pass = False
        for i, child in enumerate(children):
            if child.tag != qn("w:tbl"):
                continue
            this_shading = header_row_shading(child)
            if not this_shading:
                continue

            # walk forward past ONLY empty (no real text) paragraphs to
            # find the next real element
            j = i + 1
            while j < len(children) and children[j].tag == qn("w:p"):
                text = "".join((t.text or "") for t in children[j].findall(".//" + qn("w:t"))).strip()
                if text:
                    break
                j += 1
            if j >= len(children) or children[j].tag != qn("w:tbl"):
                continue
            next_tbl = children[j]
            if header_row_shading(next_tbl) != this_shading:
                continue  # different header style - not the same appendix table family, don't merge

            # Merge: keep child's own header row, append next_tbl's
            # rows EXCEPT its own (duplicate) header row.
            next_rows = next_tbl.findall(qn("w:tr"))
            for r in next_rows[1:]:
                next_tbl.remove(r)
                child.append(r)

            # remove the now-empty filler paragraphs and the drained table
            for k in range(i + 1, j + 1):
                el = children[k]
                el.getparent().remove(el)

            tables_merged += 1
            merged_this_pass = True
            break  # restart the scan - the child list has changed

        if not merged_this_pass:
            break

    # Mark each surviving appendix table's header row to repeat on
    # every page it spans.
    for tbl_el in body.iter(qn("w:tbl")):
        if not header_row_shading(tbl_el):
            continue
        rows = tbl_el.findall(qn("w:tr"))
        if not rows:
            continue
        header_row = rows[0]
        trPr = header_row.find(qn("w:trPr"))
        if trPr is None:
            trPr = OxmlElement("w:trPr")
            header_row.insert(0, trPr)
        if trPr.find(qn("w:tblHeader")) is None:
            trPr.append(OxmlElement("w:tblHeader"))

    return tables_merged


def _remove_duplicate_field_labels_in_row(doc):
    """Item (DUPLICATE-FIELD-LABEL-IN-SAME-ROW, real reported issue,
    per a user-provided page-1 screenshot) - confirmed real directly in
    the reported document's own XML: rows in the Contract Data-style
    tables repeat the SAME field label twice within one row - e.g.
    ['Contract Sealing Date:', '2025-06-17', 'Contract Sealing Date',
    'Contract Sealing Location:', 'Riyadh', 'Contract Sealing
    Location'] - cells 0 and 2 (and 3 and 5) carry the identical label
    text, once with a trailing colon (the real field label) and once
    without (a leftover duplicate, likely from the original
    bilingual/dual-column source layout surviving translation).

    REAL FOLLOW-UP GAP (found during a careful, full page-by-page
    review of a real current output, comparing against what had
    already been fixed before): the original version only matched
    EXACT (post-colon-stripping) text, case-SENSITIVE - a real,
    confirmed, repeating pattern in that same document used
    case-VARYING duplicates instead ('Number of AC Units' / 'Number
    of AC units', 'Electricity Meter Number' / 'Electricity meter
    number', 'Current Reading' / 'Current reading' - appearing twice,
    in two different rows), which this fix silently missed every time,
    since "Number of AC Units" != "Number of AC units" as exact
    strings. Now compares case-INSENSITIVELY (the dict key used to
    detect a repeat is lowercased) while still PRESERVING the first
    occurrence's own original casing exactly as it was - only the
    later, redundant duplicate gets emptied out.

    Fix: within each row, group cells by their normalized text (strip
    a trailing colon and surrounding whitespace, then lowercase for
    comparison purposes only) - when two or more cells share the same
    normalized text, keep the FIRST occurrence exactly as it was
    (original casing preserved) and empty out every later duplicate.
    Scoped to matches of at least 5 characters containing a letter
    (not just digits/punctuation) - keeps this from ever accidentally
    treating a short, coincidentally-matching VALUE (a date, a number)
    as a duplicate label; genuine field labels are real English
    phrases, coincidental short/numeric matches are not this bug."""
    from docx.oxml.ns import qn

    def _normalize(text):
        t = text.strip()
        if t.endswith(":"):
            t = t[:-1].strip()
        return t

    fixed = 0
    for tbl_el in doc.element.body.iter(qn("w:tbl")):
        for tr_el in tbl_el.findall(qn("w:tr")):
            cells = tr_el.findall(qn("w:tc"))
            seen = {}
            for tc in cells:
                ts = tc.findall(".//" + qn("w:t"))
                raw_text = "".join(t.text or "" for t in ts)
                norm = _normalize(raw_text)
                if len(norm) < 5 or not any(ch.isalpha() for ch in norm):
                    continue
                key = norm.lower()
                if key in seen:
                    for t in ts:
                        t.text = ""
                    if ts:
                        fixed += 1
                else:
                    seen[key] = tc
    return fixed


def _fix_ambiguous_table_width(doc):
    """Item (TABLE-AMBIGUOUS-AUTO-WIDTH-REAL-WORD-DISCREPANCY, real
    reported issue) - confirmed real, but genuinely unverifiable in
    this sandbox (no real MS Word available to test against directly -
    flagged explicitly, not glossed over): a real user reported tables
    showing an unexpected left gap, or width overflowing the margin, in
    real MS Word screenshots, while the SAME document's XML (indent and
    gridCol widths) measured completely consistent and correct, and
    LibreOffice rendered those same tables correctly with no visible
    problem at all. This is the same class of LibreOffice/real-Word
    OOXML rendering discrepancy already documented elsewhere in this
    file (jc="distribute" was another confirmed instance) - LibreOffice
    confirmation is not equivalent to real Word confirmation for
    properties like this.

    Every affected table carried <w:tblW w:w="0" w:type="auto"/>
    TOGETHER WITH <w:tblLayout w:type="fixed"/> and explicit per-column
    <w:gridCol> widths - a genuinely ambiguous combination ("auto-size
    the table" and "use exactly these fixed column widths" stated at
    the same time). Different renderers can reasonably resolve that
    ambiguity differently; explicitly setting w:tblW to the table's own
    REAL width (the sum of its gridCol widths, w:type="dxa") removes
    the ambiguity entirely, telling every renderer the same, single,
    unambiguous width - the fixed columns already summed to.

    REAL FOLLOW-UP BUG (registered in CLAUDE_INSTRUCTIONS.md, same
    class as the tblPrEx/tblInd incident): this fix originally only
    corrected the TABLE-level <w:tblPr><w:tblW> - a real user then
    reported tables with a CORRECT left position (already fixed) still
    showing content cut off past the right margin in real Word.
    Checking again with the "check every override level, not just the
    obvious one" lesson applied: EVERY <w:tblPrEx> in the affected
    document ALSO carried its own <w:tblW w:w="0" w:type="auto"/> -
    the exact same ambiguity, just at the per-ROW level, which the
    original version of this fix never touched. If a renderer resolves
    a row's OWN ambiguous tblPrEx/tblW differently from the table's
    (now-corrected) tblPr/tblW, that specific row can render wider than
    the table's defined columns - explaining right-edge cutoff on some
    rows while the table's own overall position/width looks correct.
    Now also walks every row's <w:tblPrEx> and applies the identical
    correction (resolve w:type="auto"/w="0" to an explicit w:type="dxa"
    matching the table's own real column-width sum) wherever found.

    This is a defensive, low-risk correction (the computed value is
    never a guess - it is exactly what the table's own explicit column
    widths already summed to) rather than a confirmed root-cause fix,
    since the actual real-Word rendering behavior could not be directly
    tested here. Real MS Word verification of this specific fix is
    needed, not assumed."""
    from docx.oxml.ns import qn

    fixed = 0
    for tbl_el in doc.element.body.iter(qn("w:tbl")):
        tblPr = tbl_el.find(qn("w:tblPr"))
        if tblPr is None:
            continue
        grid = tbl_el.find(qn("w:tblGrid"))
        cols = grid.findall(qn("w:gridCol")) if grid is not None else []
        total_w = sum(int(c.get(qn("w:w"))) for c in cols) if cols else 0
        if total_w <= 0:
            continue

        def _resolve_ambiguous_width(width_el):
            nonlocal fixed
            if width_el is not None and width_el.get(qn("w:type")) == "auto" and width_el.get(qn("w:w")) in ("0", None):
                width_el.set(qn("w:type"), "dxa")
                width_el.set(qn("w:w"), str(total_w))
                fixed += 1

        _resolve_ambiguous_width(tblPr.find(qn("w:tblW")))

        # Row-level: every row's own tblPrEx can carry the SAME
        # ambiguous width, independently of the table-level one above -
        # must be resolved separately, per row, using this table's own
        # real width (not assumed to inherit the table-level fix).
        for tr_el in tbl_el.findall(qn("w:tr")):
            tblPrEx = tr_el.find(qn("w:tblPrEx"))
            if tblPrEx is not None:
                _resolve_ambiguous_width(tblPrEx.find(qn("w:tblW")))
    return fixed


def _fix_paragraph_shading_mismatch_within_cell(doc):
    """Item (WHITE-LINE-IN-SHADED-CELL, real reported issue, per a
    user-provided real MS Word screenshot with zoomed-in evidence) -
    confirmed real, directly in a reported document's XML: header
    cells like "Rent Value" / "Total Value" carry a dark background
    (w:tcPr/w:shd fill=666666) and consist of TWO separate paragraphs
    (not one paragraph with concatenated text - confirmed by directly
    reading the raw XML) - the FIRST paragraph correctly carries its
    own matching w:pPr/w:shd fill=666666, but the SECOND paragraph has
    NO w:pPr/w:shd at all, defaulting to no fill (visually white). That
    second paragraph's own line-height then renders as a visible white
    gap/seam cutting through what should be a solid dark header cell -
    exactly the "white lining" the user pointed out and confirmed via a
    zoomed screenshot. A systematic, whole-document scan (not scoped to
    the 3 example cells shown) found 196 real instances of this same
    mismatch - genuinely widespread, not a one-off.

    Fix: for every table cell whose OWN w:tcPr/w:shd has a real fill
    color (not "auto"/"FFFFFF"), every paragraph inside that cell gets
    its OWN w:pPr/w:shd set to that SAME fill color - adding the
    element if missing, correcting it if it's set to something else.
    This does not touch cells with no real background fill at all
    (ordinary white table cells are left completely alone - there is
    no "gap" to fix when there's no dark background for a paragraph to
    contrast against)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    fixed = 0
    for tc in doc.element.body.iter(qn("w:tc")):
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            continue
        cell_shd = tcPr.find(qn("w:shd"))
        if cell_shd is None:
            continue
        cell_fill = cell_shd.get(qn("w:fill"))
        if not cell_fill or cell_fill in ("auto", "FFFFFF"):
            continue
        for p in tc.findall(qn("w:p")):
            pPr = p.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                p.insert(0, pPr)
            para_shd = pPr.find(qn("w:shd"))
            if para_shd is None:
                para_shd = OxmlElement("w:shd")
                para_shd.set(qn("w:val"), "clear")
                para_shd.set(qn("w:color"), "auto")
                pPr.append(para_shd)
            if para_shd.get(qn("w:fill")) != cell_fill:
                para_shd.set(qn("w:fill"), cell_fill)
                fixed += 1
    return fixed


def _fix_within_cell_self_duplicated_label(doc):
    """Item (CELL-TEXT-SELF-DUPLICATED-NO-SEPARATOR, real reported
    issue, per a real user-provided screenshot) - confirmed real,
    directly in a reported document's XML: several table header cells
    carried their own label TEXT DUPLICATED, concatenated with no
    separator at all - 'Rent ValueRent value', 'VATVAT', 'Total
    ValueTotal value', 'Issue Date (AD)Issued Date(AD)', 'Due Date
    (AD)Due Date(AD)', 'Issue Date (AH)Issued Date(AH)', 'Due Date
    (AH)Due Date(AH)', 'Expiry DateExpiry Date' - 8 real, confirmed
    instances found via a systematic whole-document scan (not scoped
    to just the one example shown), each pair a near-identical repeat
    of the same label (sometimes exact, sometimes differing only in
    capitalization or a word like "Issue" vs "Issued").

    Detects this by trying every reasonable split point in a cell's
    text and measuring the similarity (difflib SequenceMatcher, on the
    normalized - lowercased, punctuation/whitespace-stripped - form of
    each half) between the two halves; the split with the highest
    similarity above 0.7 is treated as the real self-duplication
    boundary. A minimum half-length of 2 (after normalization) keeps
    this from firing on trivially short/coincidental overlaps.

    Once the split is found, keeps only ONE half (per real-example
    precedent, the SECOND half is consistently the more complete,
    properly-closed form when brackets are involved - e.g. h1="Issue
    Date (AD" would leave an unbalanced "(", while h2="Issued
    Date(AD)" is self-contained) - a light cleanup pass then strips
    any leading stray closing-bracket/colon left over from an
    imperfect split boundary (confirmed real: the highest-similarity
    split for "Issue Date (AD)Issued Date(AD)" lands one character off
    a "clean" word boundary, leaving h2=")Issued Date(AD)" with a
    stray leading ")" that needs stripping)."""
    from docx.oxml.ns import qn
    import difflib

    def _find_split(text, min_half=2):
        best = None
        n = len(text)
        for split in range(min_half, n - min_half + 1):
            h1, h2 = text[:split], text[split:]
            norm1 = re.sub(r"[^a-z0-9]", "", h1.lower())
            norm2 = re.sub(r"[^a-z0-9]", "", h2.lower())
            if len(norm1) < min_half or len(norm2) < min_half:
                continue
            ratio = difflib.SequenceMatcher(None, norm1, norm2).ratio()
            if ratio > 0.7 and (best is None or ratio > best[2]):
                best = (h1, h2, ratio, split)
        return best

    def _clean(text):
        text = text.strip()
        # strip a leading stray closing-bracket/colon left by an
        # imperfect split boundary (confirmed real case)
        while text and text[0] in ")]}:":
            text = text[1:].strip()
        return text

    fixed = 0
    for tc in doc.element.body.iter(qn("w:tc")):
        ts = tc.findall(".//" + qn("w:t"))
        full_text = "".join(t.text or "" for t in ts)
        if len(full_text.strip()) < 4:
            continue
        found = _find_split(full_text)
        if not found:
            continue
        h1, h2, ratio, split = found
        kept = _clean(h2) or _clean(h1)
        if not kept or kept == full_text.strip():
            continue
        # write the cleaned, de-duplicated text into the FIRST real
        # text-run, clear the rest - matches the same run-collapse
        # pattern used elsewhere in this file for similar text fixes.
        if not ts:
            continue
        ts[0].text = kept
        for t in ts[1:]:
            t.text = ""
        fixed += 1
    return fixed


def _fix_cell_padding_asymmetry_within_row(doc):
    """Item (CELL-CRAMPED-AGAINST-ROW-EDGE, real reported issue, per
    an explicit user-provided real MS Word screenshot) - confirmed real:
    a real document's "CR issue place" cell sat visually flush against
    its row's right boundary, with no visible breathing room, while
    every OTHER cell in that same row (CR date, 2016-10-05, CR issued
    date) clearly showed padding around its text. A systematic,
    whole-document scan (not just the one visually-flagged cell) found
    47 real instances of this same pattern across many different
    tables - a cell whose own w:tcMar/@right is drastically smaller
    (under 30% of its row's own average) than its row-siblings',
    producing exactly this "text touching the edge" visual symptom.

    Per explicit direction ("problem completely remove karo, sirf ek
    example nahi") - this is not scoped to the one reported cell. Fix:
    for every table row, compute that row's OWN average right-padding
    (w:tcMar/@right) across its real cells (skipping rows where the
    row itself deliberately uses near-zero padding throughout, i.e.
    average < 30 twips, since a uniformly-tight row is a real design
    choice, not this bug) - any cell whose own padding falls under 30%
    of that average gets reset to match the row's average, removing
    the asymmetry rather than guessing an arbitrary fixed value."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    fixed = 0
    for tbl_el in doc.element.body.iter(qn("w:tbl")):
        for tr_el in tbl_el.findall(qn("w:tr")):
            cells = tr_el.findall(qn("w:tc"))
            if len(cells) < 2:
                continue
            paddings = []
            for tc in cells:
                tcPr = tc.find(qn("w:tcPr"))
                tcMar = tcPr.find(qn("w:tcMar")) if tcPr is not None else None
                right_el = tcMar.find(qn("w:right")) if tcMar is not None else None
                val = int(right_el.get(qn("w:w"))) if right_el is not None and right_el.get(qn("w:w")) else 0
                paddings.append(val)
            avg_pad = sum(paddings) / len(paddings)
            if avg_pad < 30:
                continue  # this whole row is uniformly tight by design - not this bug
            target = round(avg_pad)
            for tc, pad in zip(cells, paddings):
                if pad >= avg_pad * 0.3:
                    continue
                tcPr = tc.find(qn("w:tcPr"))
                if tcPr is None:
                    continue
                tcMar = tcPr.find(qn("w:tcMar"))
                if tcMar is None:
                    tcMar = OxmlElement("w:tcMar")
                    tcPr.append(tcMar)
                right_el = tcMar.find(qn("w:right"))
                if right_el is None:
                    right_el = OxmlElement("w:right")
                    right_el.set(qn("w:type"), "dxa")
                    tcMar.append(right_el)
                right_el.set(qn("w:w"), str(target))
                fixed += 1
    return fixed


def _fix_cell_width_vs_column_mismatch(doc):
    """Item (CELL-WIDTH-NOT-MATCHING-TABLE-COLUMN, real reported
    follow-up, minor/secondary finding from the same investigation as
    the tblPrEx/tblW fix above) - confirmed real via a direct,
    comprehensive scan of every cell in every table in a real reported
    document: exactly one row (2 cells) had a w:tcW (cell width) that
    did not match the table's own <w:tblGrid> column width at that
    position - one cell was double its column's real width, the
    adjacent cell correspondingly narrower, with the row's own total
    still summing correctly (so this specific case does not itself
    cause right-margin overflow) but the row's internal column
    boundaries don't line up with the rest of the table's rows, a real
    (if more cosmetic) inconsistency.

    Fix: for every row where the cell count exactly matches the
    table's own column count (skipping rows with merged/spanned cells,
    identified via w:gridSpan, where a 1:1 comparison would be
    meaningless), reset each cell's own w:tcW to match its
    corresponding w:gridCol value exactly."""
    from docx.oxml.ns import qn

    fixed = 0
    for tbl_el in doc.element.body.iter(qn("w:tbl")):
        grid = tbl_el.find(qn("w:tblGrid"))
        cols = grid.findall(qn("w:gridCol")) if grid is not None else []
        if not cols:
            continue
        col_widths = [c.get(qn("w:w")) for c in cols]

        for tr_el in tbl_el.findall(qn("w:tr")):
            cells = tr_el.findall(qn("w:tc"))
            if len(cells) != len(col_widths):
                continue  # merged/spanned row - a 1:1 comparison doesn't apply
            has_span = any(tc.find(".//" + qn("w:gridSpan")) is not None for tc in cells)
            if has_span:
                continue

            for tc, expected_w in zip(cells, col_widths):
                tcPr = tc.find(qn("w:tcPr"))
                if tcPr is None:
                    continue
                tcW = tcPr.find(qn("w:tcW"))
                if tcW is None:
                    continue
                if tcW.get(qn("w:w")) != expected_w:
                    tcW.set(qn("w:w"), expected_w)
                    tcW.set(qn("w:type"), "dxa")
                    fixed += 1
    return fixed


def _reverse_table_column_order(doc, target_language):
    """Item (TABLE-COLUMN-ORDER-NOT-MIRRORED-FOR-DIRECTION, per explicit
    direction, with a user-provided reference image showing the exact
    expected before/after): a table converted from an RTL source
    document (Arabic) keeps its RTL-ordered column sequence even after
    translation into an LTR target language - visually wrong, since an
    LTR reader expects the FIRST column (left-to-right) to be the
    "first" logical column, not the last. Per the reference image's own
    worked example: Arabic table columns
    "Translation(T2,T1) | Data(D2,D1) | SR.No" become, in the English
    version, "SR.No | Data(D1,D2) | Translation(T1,T2)" - the ENTIRE
    column sequence is reversed end-to-end (not just the top-level
    groups mirrored while sub-columns stay put - "Translation 2" was
    first in Arabic and is LAST in English, "Translation 1" was second
    and is FIRST) - and every row (headers AND data rows alike) follows
    the exact same reversal, since a cell's own column position is
    positional, not content-dependent.

    Applies to EVERY table in the document, unconditionally in the
    sense that no table is special-cased or skipped - but only actually
    runs the reversal when the target language's direction differs from
    the source's (RTL source, non-RTL/LTR target) - reusing
    _is_rtl_language, the same direction-detection already used
    elsewhere in this file, rather than a new one. If the target
    language is ALSO RTL (translating between two RTL languages, or
    keeping the original), no reversal happens - the columns are
    already in the right order for an RTL reader.

    Mechanism: for each table, reverses the order of <w:gridCol>
    elements in its <w:tblGrid> (so column WIDTHS follow the same new
    positions), and for each row, reverses the order of its <w:tc>
    (cell) children - a cell that spans multiple columns (w:gridSpan)
    moves as one whole unit to its new position, its own span value
    unchanged, exactly matching how the reference image shows a merged
    header ("Translation" spanning 2 sub-columns) staying merged, just
    relocated as a block."""
    from docx.oxml.ns import qn

    want_rtl = _is_rtl_language(target_language)
    if want_rtl:
        return 0  # target is ALSO RTL - the existing column order is already correct

    fixed = 0
    for tbl_el in doc.element.body.iter(qn("w:tbl")):
        grid = tbl_el.find(qn("w:tblGrid"))
        if grid is not None:
            cols = grid.findall(qn("w:gridCol"))
            if len(cols) > 1:
                for c in cols:
                    grid.remove(c)
                for c in reversed(cols):
                    grid.append(c)

        for tr_el in tbl_el.findall(qn("w:tr")):
            cells = tr_el.findall(qn("w:tc"))
            if len(cells) <= 1:
                continue
            for c in cells:
                tr_el.remove(c)
            for c in reversed(cells):
                tr_el.append(c)
            fixed += 1
    return fixed


def _fix_article_heading_spacing(doc):
    """Item (ARTICLE-HEADING-EXCESS-BEFORE-SPACING, real reported
    issue - "white lining" in the title bar and excess heading height)
    - confirmed real: a real document's 17 "Article <N>:" headings had
    scattered w:spacing/@w:before values (401, 197, 244 x13, 1 x2,
    243) - the dominant, clearly-intentional value (244) appeared 13
    times, while "Article One" at 401 (nearly double) rendered with a
    visibly taller shaded bar and a gap of empty (white) space above
    its own text - exactly the reported symptom. Fix: every "Article
    <N>:" heading paragraph's own w:before is reset to match this
    document's own dominant real value among ALL such headings (falls
    back to leaving values alone if there's no real, non-trivial
    majority to learn from)."""
    from docx.oxml.ns import qn
    from collections import Counter

    article_pattern = re.compile(r"^Article\s+\d+\s*:")

    before_values = []
    target_paragraphs = []
    for p in doc.paragraphs:
        if article_pattern.match(p.text.strip()):
            pPr = p._p.find(qn("w:pPr"))
            spacing = pPr.find(qn("w:spacing")) if pPr is not None else None
            before = spacing.get(qn("w:before")) if spacing is not None else None
            target_paragraphs.append((p, spacing))
            if before is not None:
                before_values.append(before)

    if not before_values:
        return 0
    dominant, count = Counter(before_values).most_common(1)[0]
    if count < 2:
        return 0  # no real majority to learn from - don't guess

    fixed = 0
    for p, spacing in target_paragraphs:
        if spacing is not None and spacing.get(qn("w:before")) != dominant:
            spacing.set(qn("w:before"), dominant)
            fixed += 1
    return fixed


def _fix_anomalous_header_shading_on_data_cell(doc):
    """Item (DATA-CELL-CARRYING-HEADER-SHADING, real reported issue -
    "table data me formatting ka issue") - confirmed real: a real
    document had a data cell ("1928550.00", row 1 of a payment-schedule
    table) shaded with the exact same dark fill (666666) used by that
    table's OWN header row, while every other data cell in the same
    column/table used the normal light alternating-row shading
    (EEEEEE) - a real, visibly wrong "highlighted like a header" data
    value. Fix: for every table, identify its own header-shading color
    (from row 0, matching the detection used elsewhere in this file),
    then for every LATER row (not row 0 itself), any cell carrying that
    SAME header-shading color has it reset to "auto" (no shading) -
    removing the leftover/mis-assigned header-color rather than
    guessing what the "correct" row-shading should have been.

    REAL FOLLOW-UP GAP (found during a systematic, XML-driven re-check
    across an entire document rather than only chasing visually-obvious
    anomalies): the original version only checked w:tcPr/w:shd (the
    CELL's own shading). Direct XML inspection of that SAME real
    document found the actual "1928550.00" cell's own w:tcPr/w:shd was
    correctly "auto" - the header-colored fill was instead on that
    cell's PARAGRAPH (w:pPr/w:shd), a value this function never
    checked, so it silently reported 0 fixes on a document that still
    had 6 real instances of this exact bug. Now checks BOTH levels for
    every data cell - w:tcPr/w:shd (cell) and w:pPr/w:shd (each
    paragraph inside the cell) - resetting either to "auto" wherever it
    matches the table's own header color."""
    from docx.oxml.ns import qn

    fixed = 0
    for tbl_el in doc.element.body.iter(qn("w:tbl")):
        rows = tbl_el.findall(qn("w:tr"))
        if len(rows) < 2:
            continue
        header_shd = rows[0].find(".//" + qn("w:shd"))
        header_fill = header_shd.get(qn("w:fill")) if header_shd is not None else None
        if not header_fill or header_fill in ("auto", "FFFFFF"):
            continue
        for row in rows[1:]:
            for tc in row.findall(qn("w:tc")):
                tcPr = tc.find(qn("w:tcPr"))
                cell_shd = tcPr.find(qn("w:shd")) if tcPr is not None else None
                if cell_shd is not None and cell_shd.get(qn("w:fill")) == header_fill:
                    cell_shd.set(qn("w:fill"), "auto")
                    cell_shd.set(qn("w:val"), "clear")
                    fixed += 1
                for p in tc.findall(qn("w:p")):
                    pPr = p.find(qn("w:pPr"))
                    para_shd = pPr.find(qn("w:shd")) if pPr is not None else None
                    if para_shd is not None and para_shd.get(qn("w:fill")) == header_fill:
                        para_shd.set(qn("w:fill"), "auto")
                        para_shd.set(qn("w:val"), "clear")
                        fixed += 1
    return fixed


def _fix_row_level_table_indent_override(doc):
    """Item (ROW-LEVEL-TBLPREX-INDENT-OVERRIDE, real reported issue,
    the ACTUAL root cause of a table showing 50%+ left margin in real
    MS Word screenshots - registered in CLAUDE_INSTRUCTIONS.md as a
    real mistake this pipeline made repeatedly before finding it).

    CONFIRMED REAL, directly in the actual reported document's XML:
    every earlier version of the table-position fix
    (_fix_table_overflow_indent) only ever checked <w:tblPr><w:tblInd>
    - the TABLE-level indent - which was genuinely correct (771 twips,
    matching the document's own reference). But 37 of this table's
    rows ALSO carried their own <w:tblPrEx><w:tblInd w:w="6766".../>
    (or 6767) - a completely SEPARATE OOXML mechanism (Table Property
    Exceptions, one per <w:tr>) that can override the parent table's
    own properties for that specific row. 6766 twips is ~56.8% of an
    11900-twip page width - exactly the "50%+ left margin" the user
    reported seeing in real Word, while every XML/render check of only
    tblPr kept (wrongly) reporting the table as correctly positioned.

    LibreOffice apparently does not honor tblPrEx's row-level indent
    override (or handles it differently from real Word) - this is why
    LibreOffice renders of the same document never reproduced the bug,
    even though the real document's raw XML clearly carried it. This
    is a genuine, confirmed instance of a broader lesson: a structured
    format can override the SAME property at a MORE SPECIFIC level
    than the obvious/top-level one, and checking only the top-level
    property is not sufficient.

    Fix: for every row (<w:tr>) in every table, if its own <w:tblPrEx>
    contains a <w:tblInd> that differs from the PARENT table's own
    (already-corrected) <w:tblPr><w:tblInd>, remove the tblInd from
    that row's tblPrEx entirely - letting the row correctly inherit
    the table-level indent instead of overriding it. If tblPrEx becomes
    completely empty after removing tblInd, remove the tblPrEx element
    itself too (an empty property-exception carries no real information
    and is just clutter). tblPrEx elements that DON'T contain a tblInd
    at all (e.g., only tblW or tblLayout) are left completely alone -
    this fix is narrowly scoped to the confirmed real problem
    (indent), not a blanket "remove all tblPrEx" pass.

    Must run AFTER _fix_table_overflow_indent (so the parent table's
    own tblInd is already correct/known before comparing row-level
    overrides against it)."""
    from docx.oxml.ns import qn

    fixed = 0
    for tbl_el in doc.element.body.iter(qn("w:tbl")):
        tblPr = tbl_el.find(qn("w:tblPr"))
        if tblPr is None:
            continue
        table_tblInd = tblPr.find(qn("w:tblInd"))
        table_ind_val = table_tblInd.get(qn("w:w")) if table_tblInd is not None else "0"

        for tr_el in tbl_el.findall(qn("w:tr")):
            tblPrEx = tr_el.find(qn("w:tblPrEx"))
            if tblPrEx is None:
                continue
            row_tblInd = tblPrEx.find(qn("w:tblInd"))
            if row_tblInd is None:
                continue  # this row's tblPrEx doesn't touch indent at all - leave it alone
            if row_tblInd.get(qn("w:w")) == table_ind_val:
                continue  # already matches the (correct) table-level value - nothing to fix
            tblPrEx.remove(row_tblInd)
            fixed += 1
            if len(tblPrEx) == 0:
                tr_el.remove(tblPrEx)
    return fixed


def _fix_table_overflow_indent(doc):
    """Item (TABLE-LEFT-POSITION-NOT-MATCHING-REFERENCE-MARGIN, real
    reported issue, current design after a real user-identified gap in
    every earlier version of this fix) - the story of WHY this function
    looks the way it does now matters, because the earlier versions'
    mistake is exactly the trap to not fall back into:

    V1 only clamped an indent when it made the table overflow the raw
    page width. V2 (median-based) still ONLY corrected a table when
    "indent + width >= content_width" - i.e. still gated on overflow.
    A REAL document then showed a table sitting at indent=50%+ of the
    page width, entirely visible in a screenshot, that NEITHER version
    touched - because that table's own width was narrow enough that
    indent+width never reached the overflow threshold at all. The bug
    was not the formula; it was the CONDITION - "only fix it if it
    overflows" quietly assumes a table's horizontal position is only
    ever wrong when it visibly runs off the page, which is false: a
    table can sit anomalously far right while still being narrow enough
    to fit "inside the boundary" on paper.

    THE REAL RULE (given directly, not inferred): every table's left
    position must match the SAME reference margin non-table paragraphs
    already use (see _compute_table_derived_margins - derived from the
    document's own real "Contract Data"-style table) - unconditionally,
    not only when a table happens to overflow. So this function no
    longer asks "is this table's indent+width past the boundary?" at
    all - every table's w:tblInd is simply set to the reference value,
    full stop. A table that already happened to match needs no change
    (and costs nothing to re-set to the same value); a table sitting at
    any wrong position - whether that means overflow, or 50%+ into the
    page while technically still fitting - gets corrected the same way.

    Table-cell paragraphs' own NESTED tables are left untouched (a
    table inside a cell uses cell-relative width, not the page's own
    margins, so the page-level reference value doesn't apply to it).

    Falls back to the OLD median-of-siblings approach, restricted to
    genuinely overflowing tables, only when no reference table can be
    found at all (see _compute_table_derived_margins's own fallback
    behavior) - so a document without an identifiable reference table
    still gets SOME correction for the clearest, most damaging case
    (genuine off-page overflow) rather than none."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    target_left, _ = _compute_table_derived_margins(doc)

    fixed = 0
    if target_left is not None:
        # THE REAL RULE: unconditional - every table's left position
        # matches the reference, regardless of whether it currently
        # overflows, undershoots, or "technically fits".
        for tbl_el in doc.element.body.iter(qn("w:tbl")):
            if tbl_el.getparent().tag == qn("w:tc"):
                continue  # nested table inside a cell - cell-relative width, not page margin
            tblPr = tbl_el.find(qn("w:tblPr"))
            if tblPr is None:
                continue
            tblInd = tblPr.find(qn("w:tblInd"))
            if tblInd is None:
                tblInd = OxmlElement("w:tblInd")
                tblInd.set(qn("w:type"), "dxa")
                tblPr.append(tblInd)
            if tblInd.get(qn("w:w")) != str(target_left):
                tblInd.set(qn("w:w"), str(target_left))
                fixed += 1
        return fixed

    # Fallback: no reference table found anywhere in the document -
    # the old overflow-only, median-of-siblings correction, since some
    # real correction for the clearest case (genuine off-page overflow)
    # is still better than none.
    sec = doc.sections[0]
    content_width_twips = round((sec.page_width - sec.left_margin - sec.right_margin) / 635)

    entries = []
    for tbl_el in doc.element.body.iter(qn("w:tbl")):
        if tbl_el.getparent().tag == qn("w:tc"):
            continue
        tblPr = tbl_el.find(qn("w:tblPr"))
        if tblPr is None:
            continue
        grid = tbl_el.find(qn("w:tblGrid"))
        cols = grid.findall(qn("w:gridCol")) if grid is not None else []
        total_w = sum(int(c.get(qn("w:w"))) for c in cols) if cols else 0
        if total_w <= 0:
            continue
        tblInd = tblPr.find(qn("w:tblInd"))
        ind_val = int(tblInd.get(qn("w:w"))) if tblInd is not None else 0
        is_anomalous = (ind_val + total_w) >= content_width_twips
        entries.append({"tblPr": tblPr, "tblInd": tblInd, "ind_val": ind_val, "total_w": total_w, "anomalous": is_anomalous})

    normal_indents = sorted(e["ind_val"] for e in entries if not e["anomalous"])
    median_indent = normal_indents[len(normal_indents) // 2] if normal_indents else None

    for e in entries:
        if not e["anomalous"]:
            continue
        new_ind = median_indent if median_indent is not None else max(0, content_width_twips - e["total_w"])
        tblInd = e["tblInd"]
        if tblInd is None:
            tblInd = OxmlElement("w:tblInd")
            tblInd.set(qn("w:type"), "dxa")
            e["tblPr"].append(tblInd)
        tblInd.set(qn("w:w"), str(new_ind))
        fixed += 1
    return fixed


_RTL_SCRIPT_CHAR_RE = re.compile(
    r"[\u0590-\u05FF\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB1D-\uFDFF\uFE70-\uFEFF]"
)
_LATIN_CHAR_RE = re.compile(r"[A-Za-z]")


def _fix_table_column_order_for_target_direction(doc, target_language):
    """Item (TABLE-COLUMN-ORDER-NOT-MIRRORED, confirmed real via direct
    user report + visual evidence) - fixing paragraph-level w:bidi/w:rtl
    (see _fix_paragraph_direction) makes TEXT read correctly, but a
    table's own COLUMN ORDER is a separate, physical layout fact that
    bidi does nothing about: a table built for RTL reading naturally
    orders its columns right-to-left (e.g. a genuinely monolingual
    Arabic "Clause Number | Field | Clarification" table has "Clause
    Number" as the PHYSICALLY RIGHTMOST, narrowest column, since that's
    read first in RTL) - translating the text to English doesn't move
    the columns, so the result reads Clarification-first,
    Number-last, which is backwards for an LTR reader. Confirmed via a
    real screenshot: the Appendix table's "Clarification" column (long
    text) sat physically first/leftmost and "Clause Number" sat
    physically last/rightmost, exactly backwards from correct LTR
    reading order.

    This must NOT be applied to every table indiscriminately: most of
    this document's tables are a BILINGUAL side-by-side layout (an
    English label already sitting on the physical left, paired with the
    same field's Arabic-translated label on the physical right, e.g.
    "Brokerage Entity Address | - | عنوان منشأة الوساطة العقارية:") -
    these already have their English anchor correctly positioned on the
    left, and reversing them would break that. The distinguishing
    signal, confirmed by measuring actual character composition across
    this document's real tables: a table that still needs reordering is
    overwhelmingly (>85%) RTL-script content with no real Latin-script
    anchor column at all, while every bilingual-layout table sits
    comfortably below that (measured 0.09-0.74 RTL-character fraction
    across 28 real tables in this document, vs 0.96-1.0 for the 5 tables
    that do need reordering) - a wide, safe margin, not a coin-flip
    threshold.

    Runs BEFORE translation (like _fix_table_overflow_indent above) -
    this detection needs the ORIGINAL script to tell "genuinely RTL
    monolingual table" apart from "bilingual table with an English
    column already in place", which is impossible to tell apart once
    translation has turned everything into English. The actual column
    swap itself doesn't care about language and could run at any point;
    doing it here just keeps both structural table fixes together.

    Skips (does not touch) any table containing a row whose cell count
    doesn't match the column count (a sign of a gridSpan/merged cell in
    that row) - reordering cells positionally is only safe when every
    row has exactly one cell per column, and this document has at most
    one gridSpan cell total, so being conservative here costs
    essentially nothing.

    Only runs when the TARGET language is LTR (no evidence yet for the
    reverse LTR-source -> RTL-target case, so this deliberately doesn't
    guess at that scenario)."""
    from docx.oxml.ns import qn

    if _is_rtl_language(target_language):
        return 0

    fixed = 0
    for tbl_el in doc.element.body.iter(qn("w:tbl")):
        try:
            grid = tbl_el.find(qn("w:tblGrid"))
            cols = grid.findall(qn("w:gridCol")) if grid is not None else []
            num_cols = len(cols)
            if num_cols < 2:
                continue

            rows = tbl_el.findall(qn("w:tr"))
            if not rows:
                continue

            latin_count = 0
            rtl_count = 0
            skip_table = False
            for tr in rows:
                tcs = tr.findall(qn("w:tc"))
                if len(tcs) != num_cols:
                    skip_table = True
                    break
                for tc in tcs:
                    text = "".join((t.text or "") for t in tc.findall(".//" + qn("w:t")))
                    latin_count += len(_LATIN_CHAR_RE.findall(text))
                    rtl_count += len(_RTL_SCRIPT_CHAR_RE.findall(text))
            if skip_table:
                continue

            total_chars = latin_count + rtl_count
            if total_chars < 5:
                continue
            if (rtl_count / total_chars) <= 0.85:
                continue  # has a real Latin-script anchor column (bilingual layout) - don't reorder

            # Reverse the column widths - this applies uniformly across the
            # whole table (a single tblGrid can't have different widths per
            # row), matching the assumption that the "short identifier"
            # column becomes narrow-on-the-left and the "long text" column
            # becomes wide-on-the-right once semantically reordered.
            widths = [c.get(qn("w:w")) for c in cols]
            for c, w in zip(cols, reversed(widths)):
                c.set(qn("w:w"), w)

            # Reverse each row's CELL order - but only the rows that
            # actually need it. Confirmed by direct inspection of a real
            # affected table: its HEADER row and its DATA rows are NOT
            # consistently ordered relative to each other in Aspose's own
            # conversion - the header row already read "Clause Number |
            # Field | Clarification" in plain left-to-right order (already
            # correct for LTR), while every data row underneath had long
            # explanatory text on the LEFT and the short clause number on
            # the RIGHT (backwards). Blindly reversing every row uniformly
            # would have "fixed" the already-correct header into the WRONG
            # order while fixing the data rows - leaving header and data
            # mismatched either way.
            #
            # Per-row rule (evidence-based, not a guess): only reverse a row
            # whose first cell is markedly LONGER than its last cell (>10
            # chars and >1.5x) - this is exactly the signature of "long
            # content sitting on the left, short identifier on the right"
            # that every genuine data row in the real affected table showed,
            # while the header row (short labels in every cell, no strong
            # length asymmetry) correctly does NOT match this pattern and is
            # left alone. Verified against all 16 rows of a real appendix
            # table: correctly reversed all 15 data rows and correctly
            # skipped the 1 header row.
            # Item (ASYMMETRIC-EMPTY-CELL-NOT-REVERSED) - confirmed
            # second real bug in this same per-row rule: a row whose
            # trailing cell is genuinely EMPTY (not just short) - e.g. a
            # 2-column "Authority | " header row where the second column
            # holds nothing at all - never triggered the length-ratio
            # condition above if the first cell's own text was under 10
            # chars ("Authority" / "الصلاحية" is only 8), so its content
            # stayed in the PHYSICAL position that the table-wide width
            # reversal (which applies uniformly, unconditionally) had
            # just turned into the NARROW column - stranding real header
            # text in a column sized for nothing. Confirmed on the real
            # table: col0 width dropped from 9169 to 258 twips while
            # "Authority" stayed in col0, forcing character-by-character
            # wrapping. An empty trailing cell can never be "correctly"
            # holding content that belongs where it currently sits once
            # the table's width has been reversed - so this case is
            # reversed unconditionally, without needing the length-ratio
            # check (there's nothing in the last cell to compare against).
            for tr in rows:
                tcs = tr.findall(qn("w:tc"))
                first_text = "".join((t.text or "") for t in tcs[0].findall(".//" + qn("w:t"))).strip()
                last_text = "".join((t.text or "") for t in tcs[-1].findall(".//" + qn("w:t"))).strip()
                needs_reversal = (len(first_text) > 10 and len(first_text) > len(last_text) * 1.5) or (
                    first_text and not last_text
                )
                if not needs_reversal:
                    continue
                for tc in tcs:
                    tr.remove(tc)
                for tc in reversed(tcs):
                    tr.append(tc)

            # Item (INVISIBLE-HEADER-TEXT) - confirmed real bug this exact
            # fix introduced: reversing the table-wide gridCol widths
            # applies to EVERY row uniformly, but a row that got SKIPPED
            # by the per-row cell-reversal above (like the header row,
            # correctly left in place per the note above) keeps its
            # OWN per-cell w:tcW and w:tcMar exactly as Aspose originally
            # set them - which were calibrated for the OLD, PRE-REVERSAL
            # column width at that position. Confirmed on a real
            # document: the header row's "Clause Number" cell had its
            # own tcW="6756" (the WIDE column's original width) and a
            # tcMar left-margin of 6075 (a positioning hack that made
            # sense inside a 6756-wide cell) - but after this function
            # reversed the GRID to make that same position only 654
            # wide, the cell's own unchanged 6075 margin alone exceeded
            # its new 654-wide space, leaving zero or negative room for
            # the text, which the renderer then simply couldn't display
            # anywhere inside the cell's real boundaries - the "Clause
            # Number" header vanished, leaving a blank space in that
            # narrow column with no visible error.
            #
            # Fix: after all reordering is settled, walk every cell in
            # this table and correct its own tcW to match whatever the
            # (possibly-just-reversed) grid now says for that column
            # position - this is always safe since tcW is supposed to
            # track the grid for non-spanning cells. If a cell's tcMar
            # would still leave less than 100 twips of real content
            # width against that corrected tcW, shrink the margin(s)
            # proportionally down to a safe amount rather than leaving a
            # cell with no room for its own text.
            current_widths = [int(c.get(qn("w:w"))) for c in cols]
            for tr in rows:
                tcs = tr.findall(qn("w:tc"))
                for i, tc in enumerate(tcs):
                    if i >= len(current_widths):
                        break
                    tcPr = tc.find(qn("w:tcPr"))
                    if tcPr is None:
                        continue
                    tcW = tcPr.find(qn("w:tcW"))
                    target_w = current_widths[i]
                    if tcW is not None:
                        tcW.set(qn("w:w"), str(target_w))
                    tcMar = tcPr.find(qn("w:tcMar"))
                    if tcMar is not None:
                        left = tcMar.find(qn("w:left"))
                        right = tcMar.find(qn("w:right"))
                        lv = int(left.get(qn("w:w"))) if left is not None else 0
                        rv = int(right.get(qn("w:w"))) if right is not None else 0
                        if lv + rv > target_w - 100:
                            # shrink both margins proportionally so at least
                            # 100 twips of real content width remains
                            scale = max(0, target_w - 100) / max(lv + rv, 1)
                            if left is not None:
                                left.set(qn("w:w"), str(int(lv * scale)))
                            if right is not None:
                                right.set(qn("w:w"), str(int(rv * scale)))

            fixed += 1
        except Exception as err:  # noqa: BLE001
            # Item (SILENT-CASCADING-FAILURE) - one table's own quirk
            # must never abort processing for every table after it in
            # document order - see the matching note in
            # _fix_paragraph_direction for the confirmed real-world
            # failure pattern this guards against.
            print(f"[column-order-fix] skipped one table after an error, continuing with the rest: {err}")

    return fixed


_FONT_PATH_CACHE = {}


def _resolve_font_path(font_name, bold=False):
    """Maps a document's declared font name to an installed TrueType file
    for measurement purposes, using the same substitution LibreOffice
    itself uses (Times New Roman -> Liberation Serif, Arial -> Liberation
    Sans, etc - confirmed via `fc-match`, so measurements line up with
    how the document actually renders in this project's own PDF-preview
    pipeline). Falls back to Liberation Serif (the most common body font
    in these documents) for anything unrecognized. Cached since this is
    called per-run across potentially thousands of runs."""
    key = ((font_name or "").lower(), bold)
    if key in _FONT_PATH_CACHE:
        return _FONT_PATH_CACHE[key]
    name = (font_name or "").lower()
    base_dir = "/usr/share/fonts/truetype/liberation/"
    if "arial" in name or "helvetica" in name or "sans" in name:
        path = base_dir + ("LiberationSans-Bold.ttf" if bold else "LiberationSans-Regular.ttf")
    else:
        path = base_dir + ("LiberationSerif-Bold.ttf" if bold else "LiberationSerif-Regular.ttf")
    _FONT_PATH_CACHE[key] = path
    return path


def _measure_longest_word_twips(cell_el):
    """Item (CONTENT-AWARE-COLUMN-WIDTH) - measures the actual pixel
    width (converted to twips) of the WIDEST SINGLE WORD across every
    run in this cell, using each run's own declared font and size via
    real font metrics (PIL), not a guess or a fixed constant. A word is
    the real unit that matters here: normal word-wrap can always break
    a LINE at a space, but it can never break a single WORD shorter than
    the column - that's exactly the mechanism that forces character-by-
    character wrapping when a column is narrower than its longest word.
    Returns 0 for an empty/whitespace-only cell (nothing to measure)."""
    from docx.oxml.ns import qn
    from PIL import ImageFont

    max_width_twips = 0
    for r in cell_el.findall(".//" + qn("w:r")):
        t = r.find(qn("w:t"))
        if t is None or not t.text or not t.text.strip():
            continue
        rpr = r.find(qn("w:rPr"))
        size_half_points = 20  # default 10pt if not specified
        bold = False
        font_name = None
        if rpr is not None:
            sz = rpr.find(qn("w:sz"))
            if sz is not None and sz.get(qn("w:val")):
                try:
                    size_half_points = int(sz.get(qn("w:val")))
                except ValueError:
                    pass
            b = rpr.find(qn("w:b"))
            bold = b is not None and b.get(qn("w:val")) != "0"
            rFonts = rpr.find(qn("w:rFonts"))
            if rFonts is not None:
                font_name = rFonts.get(qn("w:ascii")) or rFonts.get(qn("w:hAnsi"))
        size_pt = max(size_half_points / 2, 6)  # never measure below 6pt - avoids near-zero degenerate widths
        font_path = _resolve_font_path(font_name, bold)
        try:
            font = ImageFont.truetype(font_path, int(round(size_pt)))
        except Exception:
            continue
        for word in t.text.split():
            try:
                bbox = font.getbbox(word)
                width_pt = bbox[2] - bbox[0]
            except Exception:
                continue
            width_twips = width_pt * 20
            if width_twips > max_width_twips:
                max_width_twips = width_twips
    return max_width_twips


_MAX_SANE_PADDING_TWIPS = 400  # combined left+right cell padding beyond this is always Aspose's own
# positioning-hack margin (confirmed multiple real instances up to 8561 twips), never a genuine
# reading-comfort requirement - shared between the width-measurement and the width-application steps
# below so both always agree on what "real" padding looks like.


def _cell_actual_padding_twips(cell_el, table_default_padding_twips):
    """Reads the REAL combined left+right padding for this specific
    cell - its own w:tcMar if set, otherwise the table's own
    w:tblCellMar default, otherwise a conservative fallback. Confirmed
    real bug this replaces: a fixed 200-twip padding assumption
    under-measured a real cell whose actual tcMar was 553 twips
    (left=540, right=13 - Aspose's own positioning-hack cells often
    carry unusually large asymmetric margins, as found repeatedly
    elsewhere in this file), causing a column to be sized ~350 twips
    too narrow and still wrap "Clause"/"Number" mid-word. Always read
    the real value instead of assuming one.

    Capped at _MAX_SANE_PADDING_TWIPS: confirmed a SEPARATE real bug
    this cap prevents - some cells carry tcMar as large as 8561 twips,
    which is Aspose's OWN positioning hack (the same "push content to a
    specific spot inside an oversized cell" trick documented elsewhere
    in this file), not a genuine minimum-padding requirement. Treating
    that number as real caused a table's needed-width sum to balloon
    past its total available width, tripping this function's own "not
    enough room even in principle" bailout and leaving the ENTIRE table
    (including a genuinely-narrow "Tenant Rights" column right next to
    it) completely unfixed. A real reading-comfort padding is never
    remotely this large, so anything past the cap is clamped down to
    it - the cell's own content still gets full credit via
    _measure_longest_word_twips, only the inherited-but-obsolete hack
    padding is bounded."""
    from docx.oxml.ns import qn

    tcPr = cell_el.find(qn("w:tcPr"))
    if tcPr is not None:
        tcMar = tcPr.find(qn("w:tcMar"))
        if tcMar is not None:
            left = tcMar.find(qn("w:left"))
            right = tcMar.find(qn("w:right"))
            lv = int(left.get(qn("w:w"))) if left is not None and left.get(qn("w:w")) else 0
            rv = int(right.get(qn("w:w"))) if right is not None and right.get(qn("w:w")) else 0
            if lv or rv:
                return min(lv + rv, _MAX_SANE_PADDING_TWIPS)
    return table_default_padding_twips


def _table_default_cell_padding_twips(tbl_el):
    """Reads a table's own w:tblCellMar (default cell padding applied
    to any cell that doesn't override it) - falls back to a
    conservative 200-twip estimate only if the table has no default set
    at all."""
    from docx.oxml.ns import qn

    tblPr = tbl_el.find(qn("w:tblPr"))
    if tblPr is not None:
        tblCellMar = tblPr.find(qn("w:tblCellMar"))
        if tblCellMar is not None:
            left = tblCellMar.find(qn("w:left"))
            right = tblCellMar.find(qn("w:right"))
            lv = int(left.get(qn("w:w"))) if left is not None and left.get(qn("w:w")) else 0
            rv = int(right.get(qn("w:w"))) if right is not None and right.get(qn("w:w")) else 0
            if lv or rv:
                return lv + rv
    return 200


def _fix_duplicate_translated_runs(doc):
    """Item (DUPLICATE-HEADER-AFTER-TRANSLATION) - confirmed real,
    root-caused bug: some table headers in Aspose's OWN original
    conversion ALREADY carry two labels with the same meaning in two
    languages within one cell - e.g. a real header cell had paragraph 1
    = "تاريخ الاستحقاق (هـ)" (Arabic, "Due Date (AH)") and paragraph 2
    already in English, "Due Date(AH)" - a pre-existing bilingual
    dual-label design, same idea as the "Contract Data" / "١ بيانات
    العقد" two-sided headers seen elsewhere in this document, just
    packed into ONE cell instead of two.

    That's fine for a genuinely bilingual document. But when translating
    the WHOLE document INTO English, the Arabic paragraph ALSO gets
    translated to English - and since both mean the same thing, the LLM
    naturally produces near-identical text for both, so the cell ends
    up visibly showing the same label twice ("Due Date (AH)"
    immediately followed by "Due Date(AH)").

    Must run AFTER translation - this can't be detected beforehand,
    since the duplication only exists once the previously-different-
    language content on both sides becomes the same target language.

    Confirmed the duplication can show up at TWO different structural
    levels depending on how Aspose originally built the cell, so both
    are checked: (a) two separate RUNS within the same paragraph, and
    (b) two separate PARAGRAPHS within the same table cell (the actual
    pattern found in the real "Due Date (AH)" case - Aspose had put the
    Arabic and the already-English label in their own paragraphs, not
    runs, inside one cell). Compares text with whitespace collapsed and
    case ignored, and only considers matches over 2 characters to avoid
    false positives on short shared tokens like "-" or ":".

    Returns the count of duplicates removed, for the caller's log."""
    from docx.oxml.ns import qn

    def _normalize(s):
        return re.sub(r"\s+", "", s).lower()

    removed = 0

    # (a) duplicate runs within the same paragraph.
    for p in _iter_paragraphs_in_order(doc):
        seen = []
        runs_to_remove = []
        for r in p.runs:
            text = r.text or ""
            if len(text.strip()) <= 2:
                continue
            key = _normalize(text)
            if key in seen:
                runs_to_remove.append(r)
            else:
                seen.append(key)
        for r in runs_to_remove:
            try:
                r._element.getparent().remove(r._element)
                removed += 1
            except Exception:
                pass

    # (b) duplicate paragraphs within the same table cell.
    for tc in doc.element.body.iter(qn("w:tc")):
        seen = []
        paras_to_remove = []
        for p_el in tc.findall(qn("w:p")):
            text = "".join((t.text or "") for t in p_el.findall(".//" + qn("w:t")))
            if len(text.strip()) <= 2:
                continue
            key = _normalize(text)
            if key in seen:
                paras_to_remove.append(p_el)
            else:
                seen.append(key)
        for p_el in paras_to_remove:
            try:
                p_el.getparent().remove(p_el)
                removed += 1
            except Exception:
                pass

    return removed


def _fix_oversized_cell_padding(doc):
    """Item (PADDING-STARVES-CELL-REGARDLESS-OF-COLUMN-WIDTH) -
    confirmed real, DIFFERENT gap from the tcMar-capping already done
    inside _fix_narrow_column_word_wrap: that fix only rechecks a
    cell's tcMar when its COLUMN's raw width gets changed as part of a
    deficit/surplus redistribution. A real "First year / Second year /
    Third year" table had a column that was ALREADY comfortably wide on
    paper (2730 twips - far more than "Third year" could ever need),
    so it was never flagged as deficient and never went through that
    redistribution path at all - yet the cell's own inherited tcMar
    was left=1216/right=1413 (2629 twips combined), leaving only ~101
    twips of real content room inside its own 2730-twip cell, and
    "Third year" still wrapped character-by-character despite the
    column width looking perfectly adequate.

    This is therefore a standalone, general check: for EVERY table
    cell, regardless of whether its column's width changes at all,
    verify the cell's own combined tcMar doesn't exceed a sane fraction
    of its own tcW - capping it down to _MAX_SANE_PADDING_TWIPS (the
    same constant used elsewhere in this file for the identical
    Aspose-positioning-hack pattern) whenever it does. Independent of
    and complementary to the column-width-redistribution fix - this
    one exists specifically for cells whose width was never touched at
    all.

    Returns the count of cells corrected, for the caller's log."""
    from docx.oxml.ns import qn

    fixed = 0
    for tc in doc.element.body.iter(qn("w:tc")):
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            continue
        tcW = tcPr.find(qn("w:tcW"))
        if tcW is None or not tcW.get(qn("w:w")):
            continue
        try:
            cell_width = int(tcW.get(qn("w:w")))
        except ValueError:
            continue
        if cell_width <= 0:
            continue
        tcMar = tcPr.find(qn("w:tcMar"))
        if tcMar is None:
            continue
        left = tcMar.find(qn("w:left"))
        right = tcMar.find(qn("w:right"))
        lv = int(left.get(qn("w:w"))) if left is not None and left.get(qn("w:w")) else 0
        rv = int(right.get(qn("w:w"))) if right is not None and right.get(qn("w:w")) else 0
        if lv + rv <= 0:
            continue
        capped = min(lv + rv, _MAX_SANE_PADDING_TWIPS)
        target_padding = min(capped, max(0, cell_width - 100))
        if target_padding == lv + rv:
            continue
        # Only ever shrinks - never enlarges a cell's own margin beyond
        # what Aspose already had, so this can't introduce new spacing
        # that wasn't there before.
        if target_padding >= lv + rv:
            continue
        scale = target_padding / (lv + rv)
        if left is not None:
            left.set(qn("w:w"), str(int(lv * scale)))
        if right is not None:
            right.set(qn("w:w"), str(int(rv * scale)))
        fixed += 1

    return fixed


def _fix_negative_usable_width_from_paragraph_indent(doc):
    """Item (CELL-PARAGRAPH-RIGHT-INDENT-EXCEEDS-CELL-WIDTH, real
    reported issue - catastrophic character-by-character vertical text
    wrap) - confirmed real, directly in a reported document's XML: one
    cell had w:tcW=1136 twips, but its OWN paragraph carried
    w:ind/@w:right=1677 twips - MORE than the cell's entire width.
    Combined with the cell's own w:tcMar/@right=400, the real usable
    width was 1136 - 400 - 1677 = -941 twips - genuinely NEGATIVE. Word
    has no valid space to lay out even a single character at its normal
    size within that width, so it falls back to breaking every single
    character onto its own line - confirmed via a real render showing
    "Brokerage Fee (Not included..." rendered as one letter per line
    for the entire remaining page height.

    A table-cell paragraph does not generally need its own right-indent
    at all - unlike body/Article paragraphs (which use w:ind
    deliberately, per the document's own real margin methodology), cell
    content relies on the cell's own w:tcMar for padding. A large,
    cell-width-exceeding right-indent here is a leftover/erroneous
    artifact, not an intentional design choice - no OTHER cell in the
    same real document showed anything close to this pattern.

    Fix: for every table-cell paragraph, if that paragraph's own
    w:ind/@right, combined with the cell's own w:tcW and w:tcMar,
    would leave less than a minimal readable width (min_readable_twips,
    default 200 - enough for a handful of characters, well below
    what any real word needs), reset that paragraph's w:ind/@right to
    0 - letting the cell's own tcMar handle real spacing instead."""
    from docx.oxml.ns import qn

    fixed = 0
    for tc in doc.element.body.iter(qn("w:tc")):
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            continue
        tcW = tcPr.find(qn("w:tcW"))
        if tcW is None or not tcW.get(qn("w:w")):
            continue
        try:
            cell_width = int(tcW.get(qn("w:w")))
        except ValueError:
            continue
        tcMar = tcPr.find(qn("w:tcMar"))
        tcMar_right_el = tcMar.find(qn("w:right")) if tcMar is not None else None
        tcMar_right = int(tcMar_right_el.get(qn("w:w"))) if tcMar_right_el is not None and tcMar_right_el.get(qn("w:w")) else 0
        tcMar_left_el = tcMar.find(qn("w:left")) if tcMar is not None else None
        tcMar_left = int(tcMar_left_el.get(qn("w:w"))) if tcMar_left_el is not None and tcMar_left_el.get(qn("w:w")) else 0

        for p in tc.findall(qn("w:p")):
            pPr = p.find(qn("w:pPr"))
            if pPr is None:
                continue
            ind = pPr.find(qn("w:ind"))
            if ind is None or not ind.get(qn("w:right")):
                continue
            try:
                right_indent = int(ind.get(qn("w:right")))
            except ValueError:
                continue
            if right_indent <= 0:
                continue
            usable = cell_width - tcMar_left - tcMar_right - right_indent
            if usable < 200:
                ind.set(qn("w:right"), "0")
                fixed += 1
    return fixed


def _fix_narrow_column_word_wrap(doc, min_readable_twips=350):
    """Item (CONTENT-AWARE-COLUMN-WIDTH, the actual root cause behind
    "Tenant Rights"/"Clause Number"/"Authority" rendering character-by-
    character) - confirmed real, reported multiple times: earlier fixes
    (_fix_table_column_order_for_target_direction) only REPOSITIONED a
    table's existing column widths (reversing which physical column got
    which value) to match LTR reading order - they never recalculated
    those widths for what the TRANSLATED content actually needs. A
    source Arabic label like "الحقل" is short enough to fit in a narrow
    column; its English translation "Clause Number" or "Authority" is
    physically wider and doesn't fit that same narrow width, so Word/
    LibreOffice's line-breaking algorithm - which can only break at
    space characters - has no valid break point and falls back to
    breaking the word itself, character by character.

    This function measures the REAL width every column's content needs
    (via _measure_longest_word_twips, real font metrics, not a guess),
    and redistributes width FROM columns that have slack (current width
    well beyond what their own longest word needs) TO columns that don't
    (current width below their longest word's actual size) - entirely
    generically, driven by measured content, with no hardcoded column
    names, table names, or document names anywhere in this function.
    Works identically for any language pair since it measures whatever
    text is actually in the cells at the time it runs (call it AFTER
    translation).

    Safety (learned the hard way from an earlier regression): before
    shrinking any column, this checks every cell in that column for a
    w:tcMar (cell padding) that would eat more than the new width minus
    a minimum content allowance - and proportionally shrinks that
    padding first, exactly like the tcW/tcMar consistency fix elsewhere
    in this file. A column is only ever shrunk down to
    max(its own longest-word requirement, min_readable_twips) - never
    below what its own content needs, so this can't recreate the
    invisible-text regression from before.

    Only redistributes within a table whose columns all have a matching
    cell count per row (skips tables with any gridSpan/merged-cell
    irregularity, consistent with the other structural fixes in this
    file - safer to leave a rare edge case untouched than guess at it).

    Returns the count of tables adjusted, for the caller's log."""
    from docx.oxml.ns import qn

    fixed = 0
    for tbl_el in doc.element.body.iter(qn("w:tbl")):
        try:
            grid = tbl_el.find(qn("w:tblGrid"))
            cols = grid.findall(qn("w:gridCol")) if grid is not None else []
            num_cols = len(cols)
            if num_cols < 2:
                continue
            rows = tbl_el.findall(qn("w:tr"))
            if not rows:
                continue

            skip_table = False
            for tr in rows:
                if len(tr.findall(qn("w:tc"))) != num_cols:
                    skip_table = True
                    break
            if skip_table:
                continue

            current_widths = [int(c.get(qn("w:w")) or 0) for c in cols]
            if not all(current_widths):
                continue
            total_width = sum(current_widths)

            table_default_padding = _table_default_cell_padding_twips(tbl_el)

            # Measure the real per-column requirement: the widest single
            # word found in ANY cell of that column, across every row,
            # plus that SPECIFIC cell's own real padding (not a guess).
            needed = [0] * num_cols
            for tr in rows:
                tcs = tr.findall(qn("w:tc"))
                for i, tc in enumerate(tcs):
                    padding = _cell_actual_padding_twips(tc, table_default_padding)
                    w = _measure_longest_word_twips(tc) + padding
                    if w > needed[i]:
                        needed[i] = w
            needed = [max(n, min_readable_twips) for n in needed]

            if sum(needed) > total_width:
                # Not enough room even in principle to satisfy every
                # column's own longest word (e.g. a genuinely dense,
                # many-column data table) - don't guess at a partial
                # solution, leave this table's widths untouched rather
                # than risk making an unrelated column worse.
                continue

            deficits = {i: needed[i] - current_widths[i] for i in range(num_cols) if needed[i] > current_widths[i]}
            if not deficits:
                continue  # every column already has at least what its own content needs

            surplus_total = sum(
                max(0, current_widths[i] - needed[i]) for i in range(num_cols) if i not in deficits
            )
            total_deficit = sum(deficits.values())
            if surplus_total < total_deficit:
                continue  # not enough slack elsewhere in this table to safely cover it - skip rather than guess

            new_widths = list(current_widths)
            for i, deficit in deficits.items():
                new_widths[i] = needed[i]
            # Take the needed amount from surplus columns, proportional
            # to how much slack each one has.
            for i in range(num_cols):
                if i in deficits:
                    continue
                own_surplus = max(0, current_widths[i] - needed[i])
                if surplus_total <= 0:
                    continue
                share = own_surplus / surplus_total
                new_widths[i] = current_widths[i] - round(total_deficit * share)

            # Rounding can drift the total by a twip or two - correct it
            # on the single largest column rather than leave the table
            # wider or narrower than the page allows.
            drift = total_width - sum(new_widths)
            if drift:
                largest_idx = max(range(num_cols), key=lambda i: new_widths[i])
                new_widths[largest_idx] += drift

            if any(w <= 0 for w in new_widths):
                continue  # a safety net that should never trigger given the checks above - skip rather than corrupt

            for c, w in zip(cols, new_widths):
                c.set(qn("w:w"), str(w))

            # Item (GROWING-COLUMN-STILL-INVISIBLE-TEXT) - confirmed
            # real bug: this used to only recheck tcW/tcMar for columns
            # that SHRANK, reasoning that a column given MORE width
            # could only ever have MORE room than before. That's true
            # relative to the column's OLD width, but not necessarily
            # true in absolute terms - confirmed on a real cell (the
            # "Authority" cell in a "Tenant Rights" table): its column
            # grew to 8670 twips, comfortably wide by any normal
            # measure, yet the cell's own inherited tcMar was
            # left=4214/right=4347 (8561 twips combined - the same
            # Aspose positioning-hack padding pattern found repeatedly
            # elsewhere in this file), leaving only ~109 twips of real
            # content room - nowhere near enough for "Authority",
            # which kept wrapping character-by-character despite
            # sitting in a "wide" column. A pre-existing oversized
            # tcMar doesn't care whether its column grew or shrank, so
            # every column touched by this function now gets its tcMar
            # rechecked against its NEW width, not just the ones that
            # got narrower.
            for tr in rows:
                for i, tc in enumerate(tr.findall(qn("w:tc"))):
                    tcPr = tc.find(qn("w:tcPr"))
                    if tcPr is None:
                        continue
                    tcW = tcPr.find(qn("w:tcW"))
                    if tcW is not None:
                        tcW.set(qn("w:w"), str(new_widths[i]))
                    tcMar = tcPr.find(qn("w:tcMar"))
                    if tcMar is not None:
                        left = tcMar.find(qn("w:left"))
                        right = tcMar.find(qn("w:right"))
                        lv = int(left.get(qn("w:w"))) if left is not None else 0
                        rv = int(right.get(qn("w:w"))) if right is not None else 0
                        # Item (LOOSE-THRESHOLD-STILL-INVISIBLE) -
                        # confirmed real bug in this exact safety check:
                        # comparing against "new_widths[i] - 100" let a
                        # genuinely oversized padding (8561 twips)
                        # through untouched because it was JUST under
                        # that threshold (8561 < 8670-100=8570) even
                        # though 100 twips (0.07") is nowhere near
                        # enough room for real text - "Authority" kept
                        # wrapping character-by-character despite
                        # "passing" this check. A combined padding this
                        # large is never a legitimate design choice (see
                        # _cell_actual_padding_twips's docstring - this
                        # is always Aspose's own positioning hack) so it
                        # gets unconditionally capped at the same sane
                        # maximum used when MEASURING content needs
                        # (_MAX_SANE_PADDING_TWIPS), rather than only
                        # being reduced just enough to clear a small
                        # fixed buffer.
                        capped = min(lv + rv, _MAX_SANE_PADDING_TWIPS)
                        target_padding = min(capped, max(0, new_widths[i] - 100))
                        if lv + rv > 0 and target_padding != lv + rv:
                            scale = target_padding / (lv + rv)
                            if left is not None:
                                left.set(qn("w:w"), str(int(lv * scale)))
                            if right is not None:
                                right.set(qn("w:w"), str(int(rv * scale)))

            fixed += 1
        except Exception as err:  # noqa: BLE001
            print(f"[narrow-column-width-fix] skipped one table after an error, continuing with the rest: {err}")

    return fixed


def _fix_missing_font_size_in_table_row(doc):
    """Item (MISSING-SIZE-LOOKS-BIGGER, a different bug from
    FONT-SIZE-OUTLIERS above) - confirmed real on a payment-schedule
    table's data row: most cells had an explicit w:sz="16" (8pt), one
    had w:sz="13" (6.5pt), but THREE cells (Total Value, Fixed Amounts,
    VAT) had NO w:sz AT ALL - meaning they fall back to the document's
    default paragraph-style size, which is noticeably LARGER than the
    rest of the row, so those three numbers visually look like a
    different, bigger font sitting in the same row as everything else.
    The other fix in this file only catches sizes that ARE present but
    anomalously tiny - a run with no size element at all is a distinct
    case that needs its own detection.

    For each table row, finds the most common EXPLICIT w:sz among its
    own cells (ignoring cells with no size set), then applies that same
    size to any cell in the SAME row that has visible text but no
    explicit w:sz - matching the row's own established size rather
    than a fixed guess, so this adapts to whatever size each table
    actually uses.

    Returns the count of runs given an explicit size, for the caller's
    log."""
    from collections import Counter
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    fixed = 0
    for tbl_el in doc.element.body.iter(qn("w:tbl")):
        for tr in tbl_el.findall(qn("w:tr")):
            row_sizes = Counter()
            for r in tr.iter(qn("w:r")):
                t = r.find(qn("w:t"))
                if t is None or not (t.text or "").strip():
                    continue
                rpr = r.find(qn("w:rPr"))
                sz = rpr.find(qn("w:sz")) if rpr is not None else None
                if sz is not None and sz.get(qn("w:val")):
                    row_sizes[sz.get(qn("w:val"))] += 1
            if not row_sizes:
                continue
            common_size = row_sizes.most_common(1)[0][0]

            for r in tr.iter(qn("w:r")):
                t = r.find(qn("w:t"))
                if t is None or not (t.text or "").strip():
                    continue
                rpr = r.find(qn("w:rPr"))
                if rpr is None:
                    rpr = OxmlElement("w:rPr")
                    r.insert(0, rpr)
                sz = rpr.find(qn("w:sz"))
                if sz is None:
                    sz = OxmlElement("w:sz")
                    rpr.append(sz)
                    sz.set(qn("w:val"), common_size)
                    fixed += 1
                szCs = rpr.find(qn("w:szCs"))
                if szCs is None:
                    szCs = OxmlElement("w:szCs")
                    rpr.append(szCs)
                    szCs.set(qn("w:val"), common_size)

    return fixed


def _fix_tiny_font_outliers(doc, min_readable_half_points=10):
    """Item (FONT-SIZE-OUTLIERS) - a small number of runs in Aspose's OWN
    conversion output carry an anomalously tiny w:sz (font size in
    half-points) - confirmed two real instances at w:sz="2" (1pt -
    practically invisible) on header text ("5 Tenant Representative
    Data", the Appendix section intro), while every sibling run in the
    same paragraph and the rest of the document sits at 16-19
    half-points (8-9.5pt). Pre-existing in Aspose's conversion, not
    introduced by translation, but inherited unchanged into translated
    output.

    For any run with visible text and a w:sz below
    min_readable_half_points, replaces it with the paragraph's own most
    common OTHER run size (falling back to the most common size seen
    anywhere in the document if the paragraph has no other sized runs) -
    normalizing the outlier to match its actual surrounding context
    instead of guessing a fixed replacement number."""
    from collections import Counter
    from docx.oxml.ns import qn

    doc_wide_sizes = Counter()
    for r in doc.element.body.iter(qn("w:r")):
        t = r.find(qn("w:t"))
        if t is None or not (t.text or "").strip():
            continue
        rpr = r.find(qn("w:rPr"))
        if rpr is None:
            continue
        sz = rpr.find(qn("w:sz"))
        if sz is not None:
            val = sz.get(qn("w:val"))
            if val and val.isdigit() and int(val) >= min_readable_half_points:
                doc_wide_sizes[val] += 1
    fallback_size = doc_wide_sizes.most_common(1)[0][0] if doc_wide_sizes else "20"

    fixed = 0
    for p_el in doc.element.body.iter(qn("w:p")):
        para_sizes = Counter()
        outlier_runs = []
        for r in p_el.findall(qn("w:r")):
            t = r.find(qn("w:t"))
            if t is None or not (t.text or "").strip():
                continue
            rpr = r.find(qn("w:rPr"))
            if rpr is None:
                continue
            sz = rpr.find(qn("w:sz"))
            if sz is None:
                continue
            val = sz.get(qn("w:val"))
            if not val or not val.isdigit():
                continue
            if int(val) < min_readable_half_points:
                outlier_runs.append(r)
            else:
                para_sizes[val] += 1
        if not outlier_runs:
            continue
        replacement = para_sizes.most_common(1)[0][0] if para_sizes else fallback_size
        for r in outlier_runs:
            rpr = r.find(qn("w:rPr"))
            for tag in ("w:sz", "w:szCs"):
                el = rpr.find(qn(tag))
                if el is not None:
                    el.set(qn("w:val"), replacement)
            fixed += 1

    return fixed


# Item (LEAKED-INTERNAL-FIELD-NAMES) - confirmed these tokens are already
# present, verbatim, in Aspose's OWN untranslated structure-only conversion of
# the SOURCE PDF (i.e. baked into the source contract-generation template
# itself - the Saudi Ejar/REGA system's own output - not something this
# pipeline's extraction or translation introduced). Still worth cleaning up
# before delivering a final document to a reader, since a raw
# programmer-style token like "tin_number" sitting next to real data looks
# unprofessional regardless of which upstream system produced it. Kept as an
# explicit, reviewable mapping (not a guess-based heuristic) so it only ever
# touches known-confirmed tokens - extend this dict if future documents
# surface other leaked field names, rather than trying to auto-detect
# arbitrary snake_case-looking substrings (which risks false positives on
# real content, e.g. genuine hyphenated/underscored reference numbers).
_KNOWN_LEAKED_FIELD_NAME_MAP = {
    "tin_number": "TIN",
}


def _fix_leaked_internal_field_names(doc):
    """Replaces known internal/template field-name tokens (see
    _KNOWN_LEAKED_FIELD_NAME_MAP above) with their proper display label,
    wherever they appear verbatim in a run's text. Returns the count of
    runs changed."""
    from docx.oxml.ns import qn

    fixed = 0
    for run in doc.element.body.iter(qn("w:r")):
        t = run.find(qn("w:t"))
        if t is None or not t.text:
            continue
        new_text = t.text
        changed = False
        for bad, good in _KNOWN_LEAKED_FIELD_NAME_MAP.items():
            if bad in new_text:
                new_text = new_text.replace(bad, good)
                changed = True
        if changed:
            t.text = new_text
            fixed += 1
    return fixed


def _detect_split_numeric_values(doc):
    """Item (NUMERIC-INTEGRITY VALIDATION - split amounts) - confirmed
    real, pre-existing defect: Aspose's OWN PDF->DOCX table-structure
    detection sometimes splits a single financial figure across two
    ADJACENT table cells (verified: a source PDF value of 3474876.00
    came out of Aspose's OWN untranslated conversion as two separate
    cells holding "347" and "4876.00" - confirmed against Aspose's raw
    structure-only output, so this is not something translation
    introduced). The actual digits are correct and in the right order
    when read left-to-right (347 then 4876.00 = 3474876.00), so this
    isn't a corruption of the VALUE - but per this project's numeric-
    integrity requirements, a "split amount" like this is exactly the
    kind of defect that should be surfaced for review rather than
    silently shipped, since a reader could misread it as two separate
    numbers.

    Deliberately does NOT attempt to auto-merge the cells: an earlier
    attempt at a related fix (redistributing column widths) hit real,
    confirmed regressions from this same document's inconsistent
    per-cell tcMar (padding) overrides - Aspose's table reconstruction
    has enough undocumented per-cell quirks that blindly restructuring
    table cells carries real risk of a worse defect (e.g. merging two
    cells that are NOT actually a split number, corrupting a legitimate
    two-column layout). This is a targeted-repair-only situation per the
    "do not regenerate/restructure broadly" principle - so this function
    only DETECTS and reports candidates for human review, it does not
    modify the document.

    Detection heuristic (generic, not tied to any specific document or
    number): within each table row, flag any adjacent cell pair where
    the first cell is ALL DIGITS (no separators - i.e. looks like a
    truncated integer fragment, not a complete formatted number) and the
    very next non-empty cell STARTS with digits. Genuine standalone
    numbers in this kind of document are formatted with a decimal point
    or thousands separators (e.g. "40", "8", "2025-06-17", "445050.00"),
    so a bare, separator-free digit run sitting immediately next to
    another digit-led cell is the distinguishing signature of a split
    fragment rather than two unrelated real values.

    Returns a list of dicts: {"row_text": ..., "fragment_1": ...,
    "fragment_2": ..., "combined": ...} for logging/review."""
    from docx.oxml.ns import qn

    findings = []
    for tbl_el in doc.element.body.iter(qn("w:tbl")):
        for tr in tbl_el.findall(qn("w:tr")):
            cells = tr.findall(qn("w:tc"))
            cell_texts = [
                "".join((t.text or "") for t in tc.findall(".//" + qn("w:t"))).strip() for tc in cells
            ]
            for i in range(len(cell_texts) - 1):
                frag1 = cell_texts[i]
                if not re.fullmatch(r"\d+", frag1):
                    continue
                # find the next NON-EMPTY cell (skip blank spacer cells, common in this table style)
                j = i + 1
                while j < len(cell_texts) and not cell_texts[j]:
                    j += 1
                if j >= len(cell_texts):
                    continue
                frag2 = cell_texts[j]
                if not re.match(r"^\d", frag2):
                    continue
                findings.append(
                    {
                        "fragment_1": frag1,
                        "fragment_2": frag2,
                        "combined": frag1 + frag2,
                    }
                )
    return findings


def _iter_paragraphs_with_location(doc):
    """Walks every paragraph in the document IN ORDER, including those
    inside table cells (recursively, so a table cell that itself
    contains a nested table is still fully covered), yielding
    (location_label, paragraph) pairs. location_label is a short,
    human-readable position description (e.g. "body paragraph 4",
    "table 1 > row 2 > cell 1 > paragraph 1") for real issue messages -
    not just an opaque index - so a real reviewer output can tell a
    person WHERE a problem is without them having to count elements
    themselves. Shared by both the fingerprint-extraction step and any
    future consumer that needs the same real walk order."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body_para_n = 0
    table_n = 0

    def _walk_cell(cell, path_prefix):
        for block in cell._element:
            if block.tag == qn_local("p"):
                p = Paragraph(block, cell)
                nonlocal_para_n[0] += 1
                yield (path_prefix + " > paragraph " + str(nonlocal_para_n[0]), p)
            elif block.tag == qn_local("tbl"):
                nonlocal_table_n[0] += 1
                t = Table(block, cell)
                yield from _walk_table(t, "table " + str(nonlocal_table_n[0]))

    def _walk_table(table, label):
        for r_idx, row in enumerate(table.rows, start=1):
            for c_idx, cell in enumerate(row.cells, start=1):
                nonlocal_para_n[0] = 0
                yield from _walk_cell(cell, label + " > row " + str(r_idx) + " > cell " + str(c_idx))

    nonlocal_para_n = [0]
    nonlocal_table_n = [table_n]

    for block in doc.element.body:
        if block.tag == qn_local("p"):
            body_para_n += 1
            from docx.text.paragraph import Paragraph as _P
            yield ("body paragraph " + str(body_para_n), _P(block, doc))
        elif block.tag == qn_local("tbl"):
            nonlocal_table_n[0] += 1
            from docx.table import Table as _T
            yield from _walk_table(_T(block, doc), "table " + str(nonlocal_table_n[0]))


def qn_local(tag):
    from docx.oxml.ns import qn
    return qn("w:" + tag)


def _extract_paragraph_fingerprint(paragraph):
    """Captures the structural (non-text) properties that translation
    should NEVER change on its own - shading/background, alignment,
    indent, and bidi (RTL/LTR direction flag) - plus a coarse run-style
    summary (any bold, any italic, first run's font/size/color, per
    _replace_paragraph_text's own documented behavior of keeping
    formatting from the first run only). Deliberately does NOT capture
    the paragraph's actual text - text is EXPECTED to differ between
    the original and translated documents (that's the whole point of
    translation), so comparing text would be comparing the wrong
    thing."""
    from docx.oxml.ns import qn

    p = paragraph._p
    pPr = p.find(qn("w:pPr"))

    shd = pPr.find(qn("w:shd")) if pPr is not None else None
    shading_fill = shd.get(qn("w:fill")) if shd is not None else None

    jc = pPr.find(qn("w:jc")) if pPr is not None else None
    alignment = jc.get(qn("w:val")) if jc is not None else None

    ind = pPr.find(qn("w:ind")) if pPr is not None else None
    indent_left = ind.get(qn("w:left")) if ind is not None else None
    indent_right = ind.get(qn("w:right")) if ind is not None else None

    has_bidi = (pPr is not None) and (pPr.find(qn("w:bidi")) is not None)

    runs = paragraph.runs
    any_bold = any(bool(r.bold) for r in runs)
    any_italic = any(bool(r.italic) for r in runs)
    first_run = runs[0] if runs else None
    first_font = first_run.font.name if first_run is not None else None
    first_size = first_run.font.size if first_run is not None else None
    first_color = None
    if first_run is not None and first_run.font.color is not None and first_run.font.color.rgb is not None:
        first_color = str(first_run.font.color.rgb)

    return {
        "shading_fill": shading_fill,
        "alignment": alignment,
        "indent_left": indent_left,
        "indent_right": indent_right,
        "has_bidi": has_bidi,
        "any_bold": any_bold,
        "any_italic": any_italic,
        "first_font": first_font,
        "first_size": first_size,
        "first_color": first_color,
        "has_text": bool(paragraph.text.strip()),
    }


def review_and_fix_translation(original_docx_path, translated_docx_path, target_language, output_path):
    """STEP 3 (document reviewer), per explicit direction: identifies
    every line/object in the ORIGINAL file (step 1's Aspose-converted,
    untranslated output) and the TRANSLATED file (step 2's output),
    stores their structural fingerprints, compares them position-by-
    position (translation is a strict in-place text replacement - see
    _replace_paragraph_text - so the Nth paragraph in one document is
    always the Nth paragraph in the other; no reordering/insertion/
    deletion happens between them), finds every real issue (formatting,
    style, background, ordering, LTR/RTL direction), builds a full
    issue+solution list for the WHOLE document first, THEN applies every
    solution - matching the two worked examples given: a missing
    background color gets restored, and a paragraph whose direction
    doesn't match the target language gets corrected to RTL/LTR as
    appropriate.

    RTL/LTR scope note: per earlier explicit direction, this does NOT
    do table-column-order reversal or left/right margin-mirroring (that
    remains a separate, not-yet-built step) - it DOES fix the bidi
    (text-direction) flag itself when it doesn't match what the target
    language requires, since that's squarely a "does the output look
    like the target language" issue this reviewer is responsible for,
    and is exactly the reviewer's own worked example 2.

    Returns (issues, output_path) where issues is a list of
    {location, issue, solution} dicts, in document order - the full,
    itemized review list a human reviewer would have produced."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    original_doc = Document(original_docx_path)
    translated_doc = Document(translated_docx_path)

    original_items = list(_iter_paragraphs_with_location(original_doc))
    translated_items = list(_iter_paragraphs_with_location(translated_doc))

    issues = []

    if len(original_items) != len(translated_items):
        issues.append({
            "location": "whole document",
            "issue": f"Original has {len(original_items)} paragraph(s)/cell(s) but the translated output has "
                     f"{len(translated_items)} - some content was added, dropped, or the document structure "
                     f"otherwise changed during translation.",
            "solution": "Could not safely auto-fix a structural count mismatch - needs manual review; "
                        "the comparisons below only cover the paragraphs that DO line up positionally.",
        })

    want_rtl = _is_rtl_language(target_language)

    for (orig_loc, orig_p), (trans_loc, trans_p) in zip(original_items, translated_items):
        orig_fp = _extract_paragraph_fingerprint(orig_p)
        trans_fp = _extract_paragraph_fingerprint(trans_p)
        if not trans_fp["has_text"] and not orig_fp["has_text"]:
            continue  # both sides genuinely empty - nothing to review here

        trans_pPr = trans_p._p.find(qn("w:pPr"))

        # --- Background/shading (worked example 1) ---
        if orig_fp["shading_fill"] and orig_fp["shading_fill"] != trans_fp["shading_fill"]:
            issues.append({
                "location": trans_loc,
                "issue": f"Background color is missing/changed (original: {orig_fp['shading_fill']}, "
                         f"translated: {trans_fp['shading_fill'] or 'none'}).",
                "solution": f"Set background color to {orig_fp['shading_fill']} to match the original.",
            })
            if trans_pPr is None:
                trans_pPr = OxmlElement("w:pPr")
                trans_p._p.insert(0, trans_pPr)
            shd = trans_pPr.find(qn("w:shd"))
            if shd is None:
                shd = OxmlElement("w:shd")
                trans_pPr.append(shd)
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), orig_fp["shading_fill"])

        # --- LTR/RTL direction (worked example 2): must match the
        # TARGET language's direction, not necessarily the original's -
        # a document translated FROM an RTL language INTO an LTR one (or
        # vice versa) is SUPPOSED to end up with the new direction. ---
        if trans_fp["has_text"] and want_rtl != trans_fp["has_bidi"]:
            issues.append({
                "location": trans_loc,
                "issue": f"Text direction doesn't match the target language ({target_language}, which is "
                         f"{'RTL' if want_rtl else 'LTR'}) - this line is currently "
                         f"{'RTL' if trans_fp['has_bidi'] else 'LTR'}.",
                "solution": f"Set this line's direction to {'RTL' if want_rtl else 'LTR'}.",
            })
            if trans_pPr is None:
                trans_pPr = OxmlElement("w:pPr")
                trans_p._p.insert(0, trans_pPr)
            existing_bidi = trans_pPr.find(qn("w:bidi"))
            if want_rtl and existing_bidi is None:
                trans_pPr.append(OxmlElement("w:bidi"))
            elif not want_rtl and existing_bidi is not None:
                trans_pPr.remove(existing_bidi)
            for run in trans_p.runs:
                rPr = run._r.find(qn("w:rPr"))
                if rPr is None:
                    rPr = OxmlElement("w:rPr")
                    run._r.insert(0, rPr)
                existing_rtl = rPr.find(qn("w:rtl"))
                if want_rtl and existing_rtl is None:
                    rPr.append(OxmlElement("w:rtl"))
                elif not want_rtl and existing_rtl is not None:
                    rPr.remove(existing_rtl)

        # --- Alignment (formatting/ordering-adjacent - should not
        # change on its own since we're not doing margin-mirroring yet) ---
        # CONFIRMED REAL BUG: this check used to force the translated
        # paragraph's alignment to match the ORIGINAL's alignment on
        # ANY mismatch, assuming the original was always correct. That
        # assumption broke once _fix_body_paragraph_center_alignment
        # started legitimately fixing a real pre-existing Aspose defect
        # (body clauses wrongly center-aligned) in step 2 - the
        # ORIGINAL (step 1's untranslated output) still had the SAME
        # uncorrected "center" bug, so this reviewer was comparing
        # against a known-buggy baseline and "fixing" the translated
        # paragraph BACK to center, undoing step 2's real fix.
        # Confirmed directly against a real reported document: several
        # paragraphs correctly fixed to "both" in step 2 came back as
        # "center" again in the Final Output, exactly matching this
        # mechanism. Fix: an original alignment of "center" is never
        # treated as the correct value to revert to - if the
        # translated side is "both" (the known-correct target), that's
        # left alone rather than flagged as an issue.
        orig_alignment_is_known_bug = orig_fp["alignment"] == "center"
        if (
            orig_fp["alignment"]
            and orig_fp["alignment"] != trans_fp["alignment"]
            and not (orig_alignment_is_known_bug and trans_fp["alignment"] == "both")
        ):
            issues.append({
                "location": trans_loc,
                "issue": f"Alignment changed (original: {orig_fp['alignment']}, translated: {trans_fp['alignment'] or 'default'}).",
                "solution": f"Reset alignment to {orig_fp['alignment']} to match the original.",
            })
            if trans_pPr is None:
                trans_pPr = OxmlElement("w:pPr")
                trans_p._p.insert(0, trans_pPr)
            jc = trans_pPr.find(qn("w:jc"))
            if jc is None:
                jc = OxmlElement("w:jc")
                trans_pPr.append(jc)
            jc.set(qn("w:val"), orig_fp["alignment"])

        # --- Style (bold/italic lost entirely - beyond the already-
        # accepted "mixed-run formatting collapses to the first run"
        # trade-off; this only flags the paragraph LOSING style it had
        # everywhere in the original) ---
        if orig_fp["any_bold"] and not trans_fp["any_bold"] and trans_fp["has_text"]:
            issues.append({
                "location": trans_loc,
                "issue": "Original text was bold, translated text is not.",
                "solution": "Set the translated text to bold to match the original.",
            })
            for run in trans_p.runs:
                run.bold = True
        if orig_fp["any_italic"] and not trans_fp["any_italic"] and trans_fp["has_text"]:
            issues.append({
                "location": trans_loc,
                "issue": "Original text was italic, translated text is not.",
                "solution": "Set the translated text to italic to match the original.",
            })
            for run in trans_p.runs:
                run.italic = True

    translated_doc.save(output_path)
    return issues, output_path


def translate_existing_docx(docx_path, target_language, output_path, llm_config=None):
    """STEP 2 ONLY, per explicit direction: takes a docx that has
    ALREADY been through step 1 (Aspose PDF->DOCX conversion + the
    structural cleanup passes - header bars, headings, table indent,
    duplicate rows - all already applied there), and does ONLY the
    translation-injection step - _translate_docx_segments_in_place(),
    unchanged, the same real function rebuild_docx_with_translated_text
    uses. Deliberately does NOT re-run the structural fixes (this docx
    already has them from step 1 - repeating them here would be
    redundant, not harmful, but there is no reason to redo work already
    done) and deliberately does NOT run _fix_paragraph_direction (the
    RTL/LTR column-order and margin-mirroring step) - explicitly
    excluded per direction, to be wired in only later as its own step.

    Takes an existing DOCX path (not a PDF) as input - this is the key
    difference from rebuild_docx_with_translated_text, which converts
    from a PDF itself. Here, conversion has already happened elsewhere
    (step 1) and this function's only job is translating what's
    already there.

    CONFIRMED REAL BUG (first version of this function): left
    llm_config defaulting to None with no caller ever loading a real
    one, unlike every other real caller of _translate_docx_segments_in_
    place (run_full_test explicitly does "llm_config = le.load_llm_
    config()" before calling it) - the server handler for this function
    called it with no llm_config argument at all, so every real
    translation attempt got None, every LLM batch call failed silently
    (caught by _translate_docx_segments_in_place's own try/except,
    counted as a failed batch), and EVERY segment was left in its
    original source-language text - confirmed directly against a real
    reported document that came back from this endpoint completely
    untranslated. Fixed by loading a real config here whenever the
    caller doesn't already have one, so this function is correct by
    itself and not dependent on every future caller remembering to load
    and pass one."""
    from docx import Document

    if llm_config is None:
        llm_config = le.load_llm_config()

    doc = Document(docx_path)

    (translated_count, skipped_count, failed_batches, total_batches,
     llm_calls_by_provider) = _translate_docx_segments_in_place(doc, target_language, llm_config)

    # Real reported issues 1, 2, 3, 4 - confirmed real formatting bugs
    # that only matter/are checkable once the FINAL translated text is
    # in place. All non-fatal cosmetic passes, same pattern as every
    # other fix in this pipeline.
    #
    # ORDER MATTERS, confirmed via a real bug found during testing:
    # shading MUST be promoted to paragraph-level BEFORE the
    # continuation-merge step runs, because that step's own heading
    # detection (is_heading()) checks PARAGRAPH-level shading - running
    # it first left headings still only shaded at the run level,
    # making is_heading() wrongly return False for them and letting a
    # sub-heading's text get incorrectly merged into the wrong
    # preceding paragraph. Merging must then happen BEFORE the margin
    # and clause-start-spacing fixes, since those operate on whatever
    # paragraphs exist AFTER merging (not the pre-merge fragments).
    shading_promoted = 0
    try:
        shading_promoted = _promote_uniform_run_shading_to_paragraph(doc)
    except Exception:
        pass  # non-fatal - the translated document still ships even if this cosmetic pass fails

    continuation_paragraphs_merged = 0
    try:
        continuation_paragraphs_merged = _merge_continuation_paragraphs(doc)
    except Exception:
        pass  # non-fatal - the translated document still ships even if this cosmetic pass fails

    right_indent_fixed = 0
    try:
        right_indent_fixed = _fix_body_paragraph_right_indent(doc)
    except Exception:
        pass  # non-fatal - the translated document still ships even if this cosmetic pass fails

    center_alignment_fixed = 0
    try:
        center_alignment_fixed = _fix_body_paragraph_center_alignment(doc)
    except Exception:
        pass  # non-fatal - the translated document still ships even if this cosmetic pass fails

    clause_spacing_fixed = 0
    try:
        clause_spacing_fixed = _fix_clause_start_spacing(doc)
    except Exception:
        pass  # non-fatal - the translated document still ships even if this cosmetic pass fails

    # Batch of 11 real reported issues, wired in dependency order:
    # content fixes first (article-numbering must resolve to digits
    # BEFORE the spacing-fix can match "Article <digit>:" headings;
    # header-label and duplicate-label fixes are independent but kept
    # here for locality), then STRUCTURAL splits (each creates NEW
    # paragraphs that the LATER alignment-by-length fix needs to see),
    # then the remaining formatting-only fixes.
    article_numbers_converted = 0
    try:
        article_numbers_converted = _convert_article_word_numbers_to_digits(doc)
    except Exception:
        pass  # non-fatal - the translated document still ships even if this cosmetic pass fails

    header_labels_fixed = 0
    try:
        header_labels_fixed = _fix_mismatched_column_header_labels(doc)
    except Exception:
        pass  # non-fatal - the translated document still ships even if this cosmetic pass fails

    duplicate_field_labels_removed = 0
    try:
        duplicate_field_labels_removed = _remove_duplicate_field_labels_in_row(doc)
    except Exception:
        pass  # non-fatal - the translated document still ships even if this cosmetic pass fails

    numbered_list_starts_split = 0
    try:
        numbered_list_starts_split = _split_first_numbered_list_item_onto_new_line(doc)
    except Exception:
        pass  # non-fatal - the translated document still ships even if this cosmetic pass fails

    label_value_pairs_split = 0
    try:
        label_value_pairs_split = _split_label_value_pairs_within_cell(doc)
    except Exception:
        pass  # non-fatal - the translated document still ships even if this cosmetic pass fails

    # Must run AFTER _convert_article_word_numbers_to_digits - its own
    # detection pattern requires "Article <digit>:", not the word form
    # (confirmed real ordering dependency during testing).
    article_spacing_fixed = 0
    try:
        article_spacing_fixed = _fix_article_heading_spacing(doc)
    except Exception:
        pass  # non-fatal - the translated document still ships even if this cosmetic pass fails

    row_height_autofit_fixed = 0
    try:
        row_height_autofit_fixed = _fix_table_row_height_autofit(doc)
    except Exception:
        pass  # non-fatal - the translated document still ships even if this cosmetic pass fails

    anomalous_header_shading_fixed = 0
    try:
        anomalous_header_shading_fixed = _fix_anomalous_header_shading_on_data_cell(doc)
    except Exception:
        pass  # non-fatal - the translated document still ships even if this cosmetic pass fails

    cell_vertical_alignment_fixed = 0
    try:
        cell_vertical_alignment_fixed = _fix_table_vertical_alignment(doc)
    except Exception:
        pass  # non-fatal - the translated document still ships even if this cosmetic pass fails

    header_alignment_fixed = 0
    try:
        header_alignment_fixed = _fix_table_header_alignment(doc)
    except Exception:
        pass  # non-fatal - the translated document still ships even if this cosmetic pass fails

    # Real reported issue, per an explicit user-provided screenshot -
    # table header cells carrying their own label text duplicated with
    # no separator ("Rent ValueRent value", "VATVAT" etc). Runs BEFORE
    # the alignment/padding-cosmetic fixes below so those operate on
    # the FINAL, cleaned-up text (a shorter, de-duplicated label may
    # change which alignment/length bucket the cell falls into).
    within_cell_duplicate_label_fixed = 0
    try:
        within_cell_duplicate_label_fixed = _fix_within_cell_self_duplicated_label(doc)
    except Exception:
        pass  # non-fatal - the translated document still ships even if this cosmetic pass fails

    # Real reported issue, per a user-provided zoomed real-Word
    # screenshot - a visible white line cutting through an otherwise-
    # solid dark header cell background. Runs right after the
    # duplicate-label fix above, since that fix can leave a cell's
    # second paragraph empty of NEW text while this fix ensures
    # whatever paragraphs remain still carry consistent shading.
    paragraph_shading_mismatch_fixed = 0
    try:
        paragraph_shading_mismatch_fixed = _fix_paragraph_shading_mismatch_within_cell(doc)
    except Exception:
        pass  # non-fatal - the translated document still ships even if this cosmetic pass fails

    # Runs AFTER the two structural splits above, so newly-created
    # paragraphs (the list-item split, the label/value split) also get
    # a real alignment decision, not just the paragraphs that existed
    # before this pass ran.
    cell_alignment_by_length_fixed = 0
    try:
        cell_alignment_by_length_fixed = _fix_table_cell_alignment_by_length(doc)
    except Exception:
        pass  # non-fatal - the translated document still ships even if this cosmetic pass fails

    # Real reported issue, per an explicit user-provided real MS Word
    # screenshot - a cell visually cramped against its own row's edge,
    # with no breathing room, unlike its row-siblings. Confirmed via a
    # comprehensive, whole-document scan this is not a one-off - 47
    # real instances across many tables in the reported document.
    cell_padding_asymmetry_fixed = 0
    try:
        cell_padding_asymmetry_fixed = _fix_cell_padding_asymmetry_within_row(doc)
    except Exception:
        pass  # non-fatal - the translated document still ships even if this cosmetic pass fails

    # Real reported issue 3 - runs LAST, after every other fix, since
    # Real reported issue - catastrophic character-by-character
    # vertical text wrap. Must run BEFORE _fix_narrow_column_word_wrap
    # below, since a cell-paragraph's own poisonous right-indent (this
    # fix's target) would otherwise distort that function's real-width
    # measurements for the same cell.
    negative_usable_width_fixed = 0
    try:
        negative_usable_width_fixed = _fix_negative_usable_width_from_paragraph_indent(doc)
    except Exception:
        pass  # non-fatal - the translated document still ships even if this cosmetic pass fails

    # it needs to measure the FINAL translated text's actual width
    # requirements (an English translation is often physically wider
    # than the Arabic source it replaced, which is the real root cause
    # of a column rendering character-by-character - see this
    # function's own docstring for the confirmed mechanism).
    narrow_columns_fixed = 0
    try:
        narrow_columns_fixed = _fix_narrow_column_word_wrap(doc)
        narrow_columns_fixed += _fix_oversized_cell_padding(doc)
    except Exception:
        pass  # non-fatal - the translated document still ships even if this cosmetic pass fails

    # Real reported issue, per an explicit user-provided reference
    # image: only runs when the target language's direction differs
    # from the source's (see the function's own docstring) - reverses
    # every table's ENTIRE column sequence, header and data rows alike.
    table_columns_reversed = 0
    try:
        table_columns_reversed = _reverse_table_column_order(doc, target_language)
    except Exception:
        pass  # non-fatal - the translated document still ships even if this cosmetic pass fails

    doc.save(output_path)
    return {
        "output_path": output_path,
        "mode": "translate_only",
        "pipeline_code_version": PIPELINE_CODE_VERSION,
        "segments_translated": translated_count,
        "segments_skipped": skipped_count,
        "translation_batches_failed": failed_batches,
        "translation_batches_total": total_batches,
        "translation_providers": sorted(llm_calls_by_provider.keys()),
        "llm_calls_by_provider": llm_calls_by_provider,
        "right_indent_fixed": right_indent_fixed,
        "shading_promoted": shading_promoted,
        "center_alignment_fixed": center_alignment_fixed,
        "continuation_paragraphs_merged": continuation_paragraphs_merged,
        "clause_spacing_fixed": clause_spacing_fixed,
        "negative_usable_width_fixed": negative_usable_width_fixed,
        "narrow_columns_fixed": narrow_columns_fixed,
        "table_columns_reversed": table_columns_reversed,
        "article_numbers_converted": article_numbers_converted,
        "header_labels_fixed": header_labels_fixed,
        "duplicate_field_labels_removed": duplicate_field_labels_removed,
        "numbered_list_starts_split": numbered_list_starts_split,
        "label_value_pairs_split": label_value_pairs_split,
        "article_spacing_fixed": article_spacing_fixed,
        "row_height_autofit_fixed": row_height_autofit_fixed,
        "anomalous_header_shading_fixed": anomalous_header_shading_fixed,
        "cell_vertical_alignment_fixed": cell_vertical_alignment_fixed,
        "header_alignment_fixed": header_alignment_fixed,
        "cell_alignment_by_length_fixed": cell_alignment_by_length_fixed,
        "cell_padding_asymmetry_fixed": cell_padding_asymmetry_fixed,
        "within_cell_duplicate_label_fixed": within_cell_duplicate_label_fixed,
        "paragraph_shading_mismatch_fixed": paragraph_shading_mismatch_fixed,
    }


def _translate_docx_segments_in_place(doc, target_language, llm_config):
    """Item (IN-PLACE-MERGE) - replaces the old approach of translating
    the PDF's extracted text as one big blob and appending it as a
    separate reference page (see the removed docstring note that used to
    be on rebuild_docx_with_translated_text - Aspose's Words Cloud API
    has no "swap all body text for this translated version, keep
    structure" call, so a real per-paragraph mapping had to be built by
    hand).

    Works directly on the ALREADY-Aspose-converted docx's own paragraphs
    (via _iter_paragraphs_in_order above, so table cells are included) -
    this sidesteps the harder problem of mapping PDF-extracted text back
    onto Aspose's converted structure entirely, since we're translating
    and replacing the very same paragraph objects Aspose already built,
    not reconciling two independently-extracted versions of the text.

    Batches segments into multiple LLM calls (id-keyed JSON in, JSON
    out) rather than one call for the whole document, since a long
    contract's full paragraph list can exceed a single call's practical
    output-token budget; each batch still gets the same Translation
    Rules block (le._fetch_translation_rules_block()) that translate_text()
    uses, so a saved rule applies identically here as it does in the
    single-call path.

    Returns (translated_count, skipped_count, failed_batch_count,
    llm_calls_total, llm_calls_by_provider) for the caller's log/API-call
    accounting. A batch whose JSON response fails to parse, or is
    missing some ids, is logged and its paragraphs are simply left in
    the source language rather than aborting the whole run - a partial
    translation is far more useful to a reviewer than no output at all."""
    import json
    from collections import Counter

    segments = []  # list of (id, paragraph, original_text)
    next_id = 1
    for para in _iter_paragraphs_in_order(doc):
        text = (para.text or "").strip()
        if not text or _SIGNATURE_PLACEHOLDER_RE.match(text):
            continue  # empty paragraphs and signature-placeholder underline lines aren't translatable content
        segments.append((next_id, para, text))
        next_id += 1

    if not segments:
        return 0, 0, 0, 0, {}

    rules_block = le._fetch_translation_rules_block()
    system_prompt = (
        f"You are a professional document translator. Translate each text segment below into "
        f"{target_language}, preserving meaning, tone, and register. Each segment is a single "
        f"paragraph or table cell from a legal document, already correctly structured - do NOT "
        f"add markdown, bullets, or numbering, just translate the text of each segment as-is.\n\n"
        f"Pronoun rule: if a segment refers to a contracting party (e.g. 'the Tenant', 'the "
        f"Lessor') that is a company, institution, establishment, or other organizational entity "
        f"rather than a named individual person - look for cues like 'Company', 'Corporation', "
        f"'Establishment', 'Institution', 'Single Person Company', or a commercial registration "
        f"number attached to that party - use 'its' rather than 'his'/'her' for that party's "
        f"possessive pronouns in {target_language} (where the target language distinguishes "
        f"entity vs. personal pronouns). If genuinely uncertain whether a party is an individual "
        f"or an organization from the segment's own content, prefer the organizational form, "
        f"since legal contracting parties styled with a company/entity name are far more often "
        f"organizations than individuals.\n\n"
        f"Legal-English register: when the target language is English, prefer plain, standard "
        f"legal-English phrasing over literal or archaic-sounding constructions. For example "
        f"(illustrative only, not an exhaustive list, and only apply when it fits the actual "
        f"sentence): 'delay' rather than 'procrastination'; 'the defaulting party' rather than "
        f"'the procrastinating party'; 'shall be vacated' rather than 'are evicted' when "
        f"describing units being returned/emptied (reserve 'evict' for a landlord expelling a "
        f"tenant); 'from the date the dispute arises' rather than 'from the arising of the "
        f"dispute'; 'any implied right' rather than 'his/its unmentioned right'. Never invent "
        f"or guess a specific meaning that isn't in the source text - this rule is about "
        f"choosing the more natural of two accurate phrasings, not changing what's being said.\n\n"
        f"Input is a JSON array of {{\"id\": <int>, \"text\": <string>}} objects. Respond with "
        f"ONLY a JSON array of {{\"id\": <int>, \"text\": <translated string>}} objects, one per "
        f"input segment, same ids, no other text, no markdown code fences."
    ) + rules_block

    translated_by_id = {}
    failed_batches = 0
    llm_calls_by_provider = Counter()  # counts by provider that actually SUCCEEDED per batch (see note below)

    batch, batch_chars = [], 0
    batches = []
    for seg in segments:
        seg_len = len(seg[2])
        if batch and batch_chars + seg_len > _MAX_SEGMENT_CHARS_PER_BATCH:
            batches.append(batch)
            batch, batch_chars = [], 0
        batch.append(seg)
        batch_chars += seg_len
    if batch:
        batches.append(batch)

    # Note on call counting: each batch below is ONE logical translation
    # request from this pipeline's perspective, but
    # le._call_chat_completion_with_failover() can itself make up to TWO
    # real HTTP calls internally (primary provider, then a fallback
    # provider if the primary's call errors) before returning - it
    # doesn't currently expose that internal count. llm_calls_total
    # below is therefore the number of batches ATTEMPTED (a solid lower
    # bound on real network calls, and the number that matters for "how
    # many times did this pipeline ask an LLM to translate something"),
    # while llm_calls_by_provider counts only the provider that actually
    # SUCCEEDED for each batch (a failed batch's provider, if any was
    # tried internally, isn't visible to us here).
    for batch in batches:
        user_content = json.dumps([{"id": sid, "text": text} for sid, _para, text in batch], ensure_ascii=False)
        try:
            content, provider = le._call_chat_completion_with_failover(
                llm_config, system_prompt, user_content, max_tokens=12000
            )
            if content is None:
                failed_batches += 1
                continue
            llm_calls_by_provider[provider] += 1
            cleaned = content.strip()
            if cleaned.startswith("```"):
                cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", cleaned.strip())
            parsed = json.loads(cleaned)
            for item in parsed:
                translated_by_id[item["id"]] = item["text"]
        except Exception as err:  # noqa: BLE001
            print(f"[in-place-merge] batch translation failed, leaving {len(batch)} segment(s) untranslated: {err}")
            failed_batches += 1

    translated_count = 0
    skipped_count = 0
    for sid, para, _original_text in segments:
        if sid in translated_by_id:
            _replace_paragraph_text(para, translated_by_id[sid])
            translated_count += 1
        else:
            skipped_count += 1  # left in source language - batch failed or id missing from response

    return translated_count, skipped_count, failed_batches, len(batches), dict(llm_calls_by_provider)


def rebuild_docx_with_translated_text(pdf_path, target_language, output_path, llm_config=None):
    """Converts the source PDF to DOCX via Aspose (preserving its native
    structure/table detection), then translates and replaces every
    paragraph's (and table cell's) text IN PLACE via
    _translate_docx_segments_in_place() above - the translated document
    keeps Aspose's own structure/formatting, no separate reference page."""
    words_api = _words_api()
    with open(pdf_path, "rb") as f:
        request = ConvertDocumentRequest(document=f, format="docx")
        # Item - confirmed bug: add_heading(text, level=2) looks up a style
        # named "Heading 2" in the document's style gallery - this document
        # is NOT a fresh python-docx template, it's Aspose's OWN PDF->DOCX
        # conversion output, which apparently doesn't define that built-in
        # style at all, so the lookup failed ("no style with name 'Heading
        # 2'"). A plain paragraph with manual bold+size formatting achieves
        # the same visual result without depending on any assumption about
        # which named styles happen to exist in whatever document Aspose
        # hands back. (This note is kept for context even though the
        # in-place merge below no longer adds a heading of its own - the
        # same "don't assume named styles exist in Aspose's output" lesson
        # still applies anywhere this module touches doc styles.)
        result_bytes = words_api.convert_document(request)

    from docx import Document
    from io import BytesIO
    doc = Document(BytesIO(result_bytes))

    # Item (HEADER-BAR-BACKGROUND-GAP) - see
    # _fix_incomplete_header_bar_shading()'s docstring above. Runs BEFORE
    # translation: it wraps certain standalone paragraphs into a new
    # single-cell table, and _translate_docx_segments_in_place's
    # _iter_paragraphs_in_order() already knows how to walk into ANY
    # table (including this newly-created one) to find and translate
    # that paragraph's text, so running this first doesn't lose or skip
    # anything - it just gives translation a cleaner, table-consistent
    # structure to work from.
    headers_fixed = 0
    try:
        headers_fixed = _fix_incomplete_header_bar_shading(doc)
        headers_fixed += _fix_merged_two_sided_table_header(doc)
    except Exception:
        pass  # non-fatal - test pipeline still produces its output even if this cosmetic pass fails

    headings_split = 0
    try:
        headings_split = _fix_heading_merged_into_previous_clause(doc)
    except Exception:
        pass  # non-fatal - test pipeline still produces its output even if this cosmetic pass fails

    tables_repositioned = 0
    try:
        tables_repositioned = _fix_table_overflow_indent(doc)
    except Exception:
        pass  # non-fatal - test pipeline still produces its output even if this cosmetic pass fails

    # Runs BEFORE translation, same reasoning as the column-order fix
    # right below: removing an exact-duplicate row here means the LLM
    # never wastes a translation call on a row that's about to be
    # deleted anyway, and exact-duplicate detection is more reliable
    # against the ORIGINAL source text than against independently
    # translated text (two separate LLM calls on identical source text
    # are not guaranteed to produce byte-identical translations, which
    # would make the same real duplicate harder to detect afterward).
    duplicate_rows_removed = 0
    try:
        duplicate_rows_removed = _remove_duplicate_table_rows(doc)
    except Exception:
        pass  # non-fatal - test pipeline still produces its output even if this cosmetic pass fails

    # Item (TABLE-COLUMN-ORDER-NOT-MIRRORED) - see
    # _fix_table_column_order_for_target_direction()'s docstring. MUST
    # run BEFORE translation - it tells "genuinely RTL monolingual
    # table, needs column reversal" apart from "bilingual table with an
    # English column already in place" by looking at the ORIGINAL
    # script, which is impossible to check once everything has been
    # translated to English.
    columns_reordered = 0
    try:
        columns_reordered = _fix_table_column_order_for_target_direction(doc, target_language)
    except Exception:
        pass  # non-fatal - test pipeline still produces its output even if this cosmetic pass fails

    llm_config = llm_config if llm_config is not None else le.load_llm_config()
    translated_count, skipped_count, failed_batches, llm_calls_total, llm_calls_by_provider = (
        _translate_docx_segments_in_place(doc, target_language, llm_config)
    )

    # Item (RTL/LTR-NOT-CORRECTED, VERTICAL-CLIPPING, FONT-SIZE-OUTLIERS)
    # - see each function's own docstring above. Run after translation so
    # the direction fix corresponds to the document's actual final
    # language; the row-height and font-size fixes are independent
    # Aspose-conversion cleanups that are safe to run alongside it.
    direction_fixed, rows_fixed, fonts_fixed, clause_numbers_fixed = 0, 0, 0, 0
    try:
        direction_fixed, clause_numbers_fixed = _fix_paragraph_direction(doc, target_language)
        rows_fixed = _fix_exact_row_heights(doc)
        fonts_fixed = _fix_tiny_font_outliers(doc)
        fonts_fixed += _fix_missing_font_size_in_table_row(doc)
    except Exception:
        pass  # non-fatal - test pipeline still produces its output even if these cosmetic passes fail

    duplicate_runs_removed = 0
    try:
        duplicate_runs_removed = _fix_duplicate_translated_runs(doc)
    except Exception:
        pass  # non-fatal - test pipeline still produces its output even if this cosmetic pass fails

    # Item (CONTENT-AWARE-COLUMN-WIDTH) - MUST run AFTER translation (and
    # after direction-fix/row-height normalization) since it measures the
    # actual TRANSLATED text's real word widths, not the source
    # language's - see _fix_narrow_column_word_wrap()'s docstring for the
    # full root-cause explanation (this is what actually fixes
    # "Tenant Rights"/"Clause Number"/"Authority" rendering character-by-
    # character instead of as readable words).
    narrow_columns_fixed = 0
    try:
        narrow_columns_fixed = _fix_narrow_column_word_wrap(doc)
        narrow_columns_fixed += _fix_oversized_cell_padding(doc)
    except Exception:
        pass  # non-fatal - test pipeline still produces its output even if this cosmetic pass fails

    # Item (LEAKED-INTERNAL-FIELD-NAMES) - run AFTER translation so this
    # catches the token regardless of whether the LLM carried
    # "tin_number" through untouched or not - see
    # _fix_leaked_internal_field_names()'s docstring for the confirmed
    # root cause (this is baked into the SOURCE PDF's own template, not
    # something this pipeline's extraction introduced).
    leaked_names_fixed = 0
    try:
        leaked_names_fixed = _fix_leaked_internal_field_names(doc)
    except Exception:
        pass  # non-fatal - test pipeline still produces its output even if this cosmetic pass fails

    # Item (NUMERIC-INTEGRITY VALIDATION) - detection-only, see
    # _detect_split_numeric_values()'s docstring for why this doesn't
    # attempt an automatic merge.
    split_numeric_findings = []
    try:
        split_numeric_findings = _detect_split_numeric_values(doc)
    except Exception:
        pass  # non-fatal - a validation pass failing shouldn't block delivering the document

    # Item 2 (SIGNATURE-PRESERVATION) - runs after translation, on the
    # same doc object, so it's placing signature images into the
    # now-translated paragraphs (the underscore placeholder text itself
    # was never sent for translation - see the skip in
    # _translate_docx_segments_in_place above - so it's still there for
    # this to match against).
    sig_placed, sig_leftover = 0, 0
    try:
        signature_images = _extract_signature_images(pdf_path)
        if signature_images:
            sig_placed, sig_leftover = _inject_signature_images(doc, signature_images)
    except Exception:
        pass  # non-fatal - test pipeline still produces its output even if signature extraction fails

    doc.save(output_path)
    return {
        "output_path": output_path,
        "mode": "in_place_translation",
        "segments_translated": translated_count,
        "segments_skipped": skipped_count,
        "translation_batches_failed": failed_batches,
        "translation_providers": sorted(llm_calls_by_provider.keys()),
        "signatures_placed": sig_placed,
        "signatures_leftover": sig_leftover,
        "header_bars_fixed": headers_fixed,
        "headings_split": headings_split,
        "tables_repositioned": tables_repositioned,
        "duplicate_rows_removed": duplicate_rows_removed,
        "columns_reordered": columns_reordered,
        "narrow_columns_fixed": narrow_columns_fixed,
        "direction_fixed": direction_fixed,
        "clause_numbers_fixed": clause_numbers_fixed,
        "split_numeric_findings": split_numeric_findings,
        "row_heights_fixed": rows_fixed,
        "tiny_fonts_fixed": fonts_fixed,
        "duplicate_runs_removed": duplicate_runs_removed,
        "aspose_words_calls": 1,
        "aspose_pdf_calls": 0,
        "llm_calls": llm_calls_total,
        "llm_calls_by_provider": llm_calls_by_provider,
    }


def run_full_test(pdf_path, target_language, output_path):
    log = []
    t0 = time.time()
    log.append(f"Pipeline code version: {PIPELINE_CODE_VERSION}")
    log.append(f"Configured: {is_configured()}")

    # Item (ASPOSE.PDF-CLOUD-WIRING) - kept as a standalone check here,
    # NOT used to build the translated output below. The in-place merge
    # (rebuild_docx_with_translated_text) translates Aspose's own
    # converted-docx paragraphs directly, which is a more reliable
    # source of truth than reconciling two independently-extracted
    # versions of the same text would be (see that function's docstring)
    # - so this call's only purpose here is to confirm, in the log, that
    # the Aspose.PDF Cloud wiring itself actually works end to end
    # (upload + convert-to-text + cleanup), independent of the merge.
    extraction = extract_text_via_aspose(pdf_path)
    log.append(
        f"Text extraction check via {extraction['source']} ({len(extraction['text'])} chars) "
        f"- informational only, not used for the translation below"
    )

    t1 = time.time()
    llm_config = le.load_llm_config()
    rebuild_result = rebuild_docx_with_translated_text(pdf_path, target_language, output_path, llm_config=llm_config)
    providers = ", ".join(rebuild_result["translation_providers"]) or "n/a"
    log.append(
        f"Structure converted via Aspose + {rebuild_result['segments_translated']} segment(s) "
        f"translated in-place via {providers} in {time.time()-t1:.1f}s"
    )
    if rebuild_result["segments_skipped"]:
        log.append(
            f"{rebuild_result['segments_skipped']} segment(s) left untranslated "
            f"({rebuild_result['translation_batches_failed']} batch(es) failed)"
        )
    if rebuild_result.get("header_bars_fixed"):
        log.append(f"Header-bar background gaps fixed: {rebuild_result['header_bars_fixed']}")
    if rebuild_result.get("duplicate_runs_removed"):
        log.append(f"Duplicate translated labels removed: {rebuild_result['duplicate_runs_removed']}")
    if rebuild_result.get("headings_split"):
        log.append(f"Article headings split out from a merged previous clause: {rebuild_result['headings_split']}")
    if rebuild_result.get("tables_repositioned"):
        log.append(f"Off-page tables repositioned (were overflowing page boundary): {rebuild_result['tables_repositioned']}")
    if rebuild_result.get("duplicate_rows_removed"):
        log.append(f"Exact-duplicate table rows removed: {rebuild_result['duplicate_rows_removed']}")
    if rebuild_result.get("columns_reordered"):
        log.append(f"Table column order mirrored for LTR reading (was RTL physical order): {rebuild_result['columns_reordered']}")
    if rebuild_result.get("narrow_columns_fixed"):
        log.append(f"Narrow columns widened to fit translated content (prevents character-by-character wrap): {rebuild_result['narrow_columns_fixed']}")
    if rebuild_result.get("leaked_names_fixed"):
        log.append(f"Leaked internal field-name tokens cleaned up: {rebuild_result['leaked_names_fixed']}")
    if rebuild_result.get("direction_fixed"):
        log.append(f"RTL/LTR direction corrected: {rebuild_result['direction_fixed']} paragraph(s)")
    if rebuild_result.get("clause_numbers_fixed"):
        log.append(f"Reversed clause/article numbers corrected: {rebuild_result['clause_numbers_fixed']}")
    if rebuild_result.get("row_heights_fixed"):
        log.append(f"Fixed-height rows relaxed (prevents clipping): {rebuild_result['row_heights_fixed']}")
    if rebuild_result.get("tiny_fonts_fixed"):
        log.append(f"Tiny-font outliers normalized: {rebuild_result['tiny_fonts_fixed']}")
    split_findings = rebuild_result.get("split_numeric_findings") or []
    if split_findings:
        examples = ", ".join(f"{f['fragment_1']}|{f['fragment_2']}" for f in split_findings[:5])
        log.append(
            f"\u26a0\ufe0f NEEDS REVIEW: {len(split_findings)} possible split numeric value(s) found "
            f"(Aspose's own table conversion, not introduced by translation) - not auto-merged, "
            f"review recommended: {examples}"
            + (f" (+{len(split_findings)-5} more)" if len(split_findings) > 5 else "")
        )

    sig_placed = rebuild_result.get("signatures_placed", 0)
    sig_leftover = rebuild_result.get("signatures_leftover", 0)
    if sig_placed or sig_leftover:
        log.append(
            f"Signatures: {sig_placed} placed in-line at placeholder line(s)"
            + (f", {sig_leftover} appended as a labeled section (no placeholder match)" if sig_leftover else "")
        )

    # API-call accounting: combines the informational extraction check
    # above with the rebuild/translate step below - so this reflects
    # EVERY Aspose and LLM call this one test run actually made, useful
    # for tracking usage against Aspose's free-tier call limits and
    # LLM provider costs.
    aspose_words_calls = rebuild_result.get("aspose_words_calls", 0)
    aspose_pdf_calls = extraction.get("aspose_pdf_calls", 0) + rebuild_result.get("aspose_pdf_calls", 0)
    aspose_calls_total = aspose_words_calls + aspose_pdf_calls
    llm_calls_by_provider = rebuild_result.get("llm_calls_by_provider", {})
    openrouter_calls = llm_calls_by_provider.get("openrouter", 0)
    llm_calls_total = rebuild_result.get("llm_calls", 0)
    log.append(
        f"API calls this run \u2014 Aspose: {aspose_calls_total} "
        f"(Words Cloud: {aspose_words_calls}, PDF Cloud: {aspose_pdf_calls}), "
        f"LLM: {llm_calls_total} (OpenRouter: {openrouter_calls})"
    )

    return {
        "output_path": output_path,
        "log": log,
        "extraction_source": extraction["source"],
        "translation_providers": rebuild_result["translation_providers"],
        "segments_translated": rebuild_result["segments_translated"],
        "segments_skipped": rebuild_result["segments_skipped"],
        "signatures_placed": sig_placed,
        "signatures_leftover": sig_leftover,
        "total_seconds": round(time.time() - t0, 1),
        "aspose_words_calls": aspose_words_calls,
        "aspose_pdf_calls": aspose_pdf_calls,
        "aspose_calls_total": aspose_calls_total,
        "llm_calls_total": llm_calls_total,
        "llm_calls_by_provider": llm_calls_by_provider,
        "openrouter_calls": openrouter_calls,
    }
