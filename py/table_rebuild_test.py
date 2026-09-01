"""
Regression test for _rebuild_table_from_ocr_structure - Solution
Variant D for table-structure issues, per an exact user-provided
manual methodology: structure (rows/cols/background) from the OCR
(pre-translation) table, content (translated words + formatting) from
the already-translated output table, width/position computed directly
from the page's own margins, cleanup pass at the end (autofit, remove
empty rows/cols, wrap, alignment-by-length, vertical-center).

Run: python3 table_rebuild_test.py
"""
import os
import sys

sys.path.insert(0, os.path.dirname(__file__))
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import aspose_test_pipeline as asp


def assert_(cond, label):
    if not cond:
        print("FAIL:", label)
        sys.exit(1)
    print("PASS:", label)


def run():
    print("=== Test 1 (THE REAL REPORTED DOCUMENT PAIR): rebuild the real Payment Schedule table ===")
    ocr_path = "/home/claude/work/two_docs_review/ocr_trace/ocr_step1.docx"
    output_path = "/home/claude/work/two_docs_review/doc1.docx"
    if os.path.exists(ocr_path) and os.path.exists(output_path):
        ocr_doc = Document(ocr_path)
        output_doc = Document(output_path)
        ocr_table = ocr_doc.tables[29]
        output_table = output_doc.tables[29]

        result = asp._rebuild_table_from_ocr_structure(output_doc, output_table, ocr_table, "English")
        assert_(result["rebuilt"] is True, "Rebuild reports success")
        assert_(result["row_count"] == 3, "Real row count (3) matches the OCR table's own real rows")
        assert_(result["col_count"] == 9, "Real column count (9) matches the OCR table's own real columns")
        assert_(result["content_width"] == 10860, "Content width is a real, sane twips value (not raw EMU) - the exact real units bug this test also guards against")

        output_doc.save("/tmp/table_rebuild_test_output.docx")
        reopened = Document("/tmp/table_rebuild_test_output.docx")
        found = None
        for t in reopened.tables:
            if not t.rows:
                continue
            first_row_text = [c.text for c in t.rows[0].cells]
            if any("Serial" in x for x in first_row_text):
                found = t
                break
        assert_(found is not None, "The rebuilt table is present and findable after a real save+reopen")
        assert_(found.rows[0].cells[0].text == "Serial Number\n.No", "No stray leading blank paragraph/newline artifact (the real confirmed bug from new_cell.text = \"\")")
        assert_(found.rows[1].cells[1].text == "1928550.00", "Real translated data value landed in the correct rebuilt cell")

        vAlign = found.rows[0].cells[0]._tc.find(qn("w:tcPr")).find(qn("w:vAlign"))
        assert_(vAlign is not None and vAlign.get(qn("w:val")) == "center", "Vertical alignment is centered, per the real reported process's final step")
    else:
        print("  (real document pair not present in this environment - skipped)")

    print("\n=== Test 2 (synthetic, exact mechanics): background comes from OCR, text comes from output, matched by RTL-reversed position ===")
    ocr_doc2 = Document()
    ocr_t2 = ocr_doc2.add_table(rows=1, cols=2)
    ocr_t2.rows[0].cells[0].text = "Arabic label"
    shd_a = OxmlElement("w:shd")
    shd_a.set(qn("w:fill"), "AAAAAA")
    ocr_t2.rows[0].cells[0]._tc.get_or_add_tcPr().append(shd_a)
    ocr_t2.rows[0].cells[1].text = "Arabic value"
    shd_b = OxmlElement("w:shd")
    shd_b.set(qn("w:fill"), "BBBBBB")
    ocr_t2.rows[0].cells[1]._tc.get_or_add_tcPr().append(shd_b)

    output_doc2 = Document()
    output_t2 = output_doc2.add_table(rows=1, cols=2)
    output_t2.rows[0].cells[0].text = "English Value"
    output_t2.rows[0].cells[1].text = "English Label"

    result2 = asp._rebuild_table_from_ocr_structure(output_doc2, output_t2, ocr_t2, "English")
    assert_(result2["rebuilt"] is True, "Second rebuild also succeeds")
    reopened2_tbl = output_doc2.tables[-1] if len(output_doc2.tables) == 1 else None
    # after rebuild the old table is removed and a new one inserted - re-find it
    found2 = None
    for t in output_doc2.tables:
        if t.rows and t.rows[0].cells[0].text:
            found2 = t
            break
    assert_(found2 is not None, "Rebuilt 2-column table found")
    # want_rtl=False (English target) -> columns reversed: new col0 = OCR col1, new col1 = OCR col0
    cell0_shd = found2.rows[0].cells[0]._tc.find(qn("w:tcPr")).find(qn("w:shd"))
    cell1_shd = found2.rows[0].cells[1]._tc.find(qn("w:tcPr")).find(qn("w:shd"))
    assert_(cell0_shd.get(qn("w:fill")) == "BBBBBB", "New cell 0's background comes from OCR's LAST column (BBBBBB), matching the reversed order")
    assert_(cell1_shd.get(qn("w:fill")) == "AAAAAA", "New cell 1's background comes from OCR's FIRST column (AAAAAA)")
    assert_(found2.rows[0].cells[0].text == "English Value", "New cell 0's TEXT comes from output's own cell 0 (position-matched, not re-derived)")
    assert_(found2.rows[0].cells[1].text == "English Label", "New cell 1's TEXT comes from output's own cell 1")

    print("\n=== Test 3: an OCR table with no real rows is safely rejected, not crashed on ===")
    ocr_doc3 = Document()
    ocr_t3 = ocr_doc3.add_table(rows=0, cols=0)
    output_doc3 = Document()
    output_t3 = output_doc3.add_table(rows=1, cols=1)
    output_t3.rows[0].cells[0].text = "x"
    result3 = asp._rebuild_table_from_ocr_structure(output_doc3, output_t3, ocr_t3, "English")
    assert_(result3["rebuilt"] is False, "Correctly refuses to rebuild from an empty OCR table, rather than crashing")

    print("\nALL TESTS PASSED")


if __name__ == "__main__":
    run()
