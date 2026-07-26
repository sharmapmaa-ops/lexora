#!/usr/bin/env python3
"""
Lease Abstraction processing engine.

Implements the real (non-simulated) steps behind the Lease Abstraction
workflow described in the project requirements (section 14):

  1. extract_text(path)        -> raw text pulled out of the uploaded PDF/DOCX.
                                   Tries pdfplumber, then pypdf, then falls
                                   back to real OCR (pytesseract) if both
                                   extract almost nothing - handles PDFs
                                   that are actually scanned page images
                                   with no usable text layer.
  2. analyze_lease(text, name) -> document-type classification + extracted
                                   lease fields + an accuracy/confidence
                                   score. Uses a real OpenAI/OpenRouter call
                                   (json/extraction_prompt.txt as the system
                                   prompt) when the .env file has an
                                   API key configured; otherwise falls back
                                   to a deterministic regex/keyword engine.
  3. generate_output_pdf(...)  -> builds Output.pdf from Output.json, laid
                                   out the same way as the default template

ABOUT THE LLM EXTRACTION:
json/extraction_prompt.txt is a detailed system prompt (attorney-grade
commercial lease abstraction rules) and json/rules.json is the same rule
set in structured form (used for reference/governance - its rules are
already folded into extraction_prompt.txt, so it isn't re-injected into
the prompt verbatim to avoid doubling token usage). When a key is present
in .env, call_llm_extraction() sends the extracted lease
text to OpenAI or OpenRouter (chat completions) and returns the rich
structured JSON described in extraction_prompt.txt's OUTPUT FORMAT
section. Without a key, analyze_lease() falls back to the simpler
heuristic engine below. A second, lighter LLM call (call_llm_validation)
then reviews that JSON for completeness and returns a confidence/accuracy
score - the same two-pass extract-then-validate shape used by the
reference project, just condensed into one pipeline step (section 14.3's
40% checkpoint) instead of two separate API routes.

Uses pdfplumber (PDF text + page rendering), pypdf (secondary text
extraction), pytesseract (OCR for scanned/image-only PDFs), python-docx
(DOCX text) and reportlab (PDF generation) - see requirements.txt. All are
optional at import time so the server still runs (with reduced
functionality / clear errors) if one hasn't been installed yet. The LLM
call itself only needs the standard library (urllib) - no extra HTTP
client dependency. OCR also needs the `tesseract-ocr` system package
(not a pip package) - see requirements.txt / devcontainer.json.
"""

import base64
import difflib
import hashlib
import io
import json
import os
import re
import urllib.request
import urllib.error
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import pytesseract
    from PIL import Image as PILImage, ImageDraw, ImageFont
    OCR_LIBS_OK = True
except ImportError:
    OCR_LIBS_OK = False

try:
    import docx as docx_lib
except ImportError:
    docx_lib = None

try:
    # Vision-layout translation pipeline (scanned certificates/posters):
    # OpenCV detects text regions at pixel level (catches decorative
    # calligraphy Tesseract can't read) and inpaints the original
    # strokes so the page's own background pattern survives under the
    # translated text - the same technique that produced the reference
    # "v5" certificate rebuild. Optional: without it, the OCR path
    # falls back to Tesseract boxes + flat-color masks.
    import cv2
    import numpy as np
    CV2_OK = True
except ImportError:
    cv2 = None
    np = None
    CV2_OK = False


# ---------------------------------------------------------------------------
# Self-training reviewer memory (workflow 4.2.15 + issues #2/#9).
# The reviewer agent writes every distinct mistake it catches to a small
# JSON "lessons" file. On the next run those lessons are injected into the
# metadata-extraction and reviewer prompts as an explicit "past mistakes to
# avoid" list, so the system keeps getting better without any retraining -
# a lightweight form of continual learning that persists across runs.
# ---------------------------------------------------------------------------
_LESSONS_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "reviewer_lessons.json")
_LESSONS_MAX = 40


def _load_reviewer_lessons():
    try:
        with open(_LESSONS_PATH, "r", encoding="utf-8") as f:
            data = json.load(f)
        if isinstance(data, list):
            return data[-_LESSONS_MAX:]
    except (OSError, json.JSONDecodeError):
        pass
    return []


def _save_reviewer_lessons(lessons):
    try:
        seen = set()
        deduped = []
        for l in lessons:
            key = (l.get("rule") or "").strip().lower()
            if key and key not in seen:
                seen.add(key)
                deduped.append(l)
        with open(_LESSONS_PATH, "w", encoding="utf-8") as f:
            json.dump(deduped[-_LESSONS_MAX:], f, ensure_ascii=False, indent=1)
    except OSError as err:
        print(f"Could not persist reviewer lessons: {err}")


def _lessons_as_prompt_block(lessons, limit=15):
    if not lessons:
        return ""
    lines = []
    for l in lessons[-limit:]:
        rule = (l.get("rule") or "").strip()
        if rule:
            lines.append(f"- {rule}")
    if not lines:
        return ""
    return ("\nLESSONS FROM PAST REVIEWS (do NOT repeat these mistakes):\n"
            + "\n".join(lines) + "\n")

try:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_RIGHT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, Image
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.utils import ImageReader
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False

try:
    # Item 6 - reportlab (like most PDF libraries) only ever lays text out
    # left-to-right, character-by-character, in Unicode logical order. For
    # RTL scripts (Arabic, Hebrew, etc.) that's wrong twice over: Arabic
    # letters need CONTEXTUAL reshaping (a letter's glyph changes depending
    # on its neighbors, e.g. isolated-vs-initial-vs-medial-vs-final forms),
    # and the whole line needs its visual character order reversed
    # (logical-to-visual bidi reordering). arabic_reshaper handles the
    # first, python-bidi's get_display() handles the second - together
    # they're the standard fix for rendering Arabic/Hebrew correctly in a
    # bidi-unaware renderer like reportlab.
    import arabic_reshaper
    from bidi.algorithm import get_display as _bidi_get_display
    RTL_SHAPING_OK = True
except ImportError:
    RTL_SHAPING_OK = False


# ISO-ish language names (as they'd appear in the app's language dropdown)
# that read right-to-left. Kept as a name-based list (rather than ISO
# codes) since translate_text()/generate_translation_pdf() both work with
# the plain language name the user picked in the UI, not a code.
_RTL_LANGUAGES = {"arabic", "hebrew", "urdu", "persian", "farsi", "pashto", "sindhi", "dhivehi", "yiddish"}


def is_rtl_language(language_name):
    return (language_name or "").strip().lower() in _RTL_LANGUAGES


def shape_rtl_text(text):
    """Reshapes+reorders a string of Arabic/Hebrew/etc text so it renders
    correctly in reportlab. Safe to call on already-LTR or mixed text -
    arabic_reshaper only touches RTL-script characters, so any English
    words/numbers embedded in an other-wise-RTL sentence keep their own
    correct (LTR) internal order after the bidi pass. Falls back to the
    original text untouched if the shaping libraries aren't installed,
    rather than failing the whole PDF generation over a missing optional
    dependency."""
    if not RTL_SHAPING_OK or not text:
        return text
    try:
        reshaped = arabic_reshaper.reshape(text)
        return _bidi_get_display(reshaped)
    except Exception:
        return text


class LeaseEngineError(Exception):
    pass


# lease_engine.py lives in py/, json/ is a sibling of py/ under the project root.
ROOT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
JSON_DIR = os.path.join(ROOT_DIR, "json")
EXTRACTION_PROMPT_PATH = os.path.join(JSON_DIR, "extraction_prompt.txt")

# Below this many extracted characters, the text is almost certainly just
# an artifact (e.g. a DocuSign "Envelope ID" stamp) rather than real body
# text - matches the threshold the reference project uses to decide when
# to fall back to a second extraction method. Applied both as a flat
# minimum and per-page (a multi-page doc where every page contributes
# only a tiny stamp can otherwise still clear a flat total threshold).
MIN_USABLE_TEXT_CHARS = 200
MIN_CHARS_PER_PAGE = 150
# Cap on parallel OCR workers - actual worker count is also capped by
# os.cpu_count() and the page count at call time, this is just an upper
# bound so a huge document doesn't spawn an unreasonable number of
# tesseract subprocesses at once.
OCR_MAX_WORKERS = 6


def _load_dotenv(path):
    """Same minimal .env loader as server.py (duplicated, not imported,
    so this module still works standalone/under a different entry point).
    See server.py's copy for the full rationale."""
    if not os.path.isfile(path):
        return
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, _, value = line.partition("=")
            key = key.strip()
            value = value.strip()
            if len(value) >= 2 and value[0] == value[-1] and value[0] in ("'", '"'):
                value = value[1:-1]
            os.environ.setdefault(key, value)


_load_dotenv(os.path.join(ROOT_DIR, ".env"))


def load_llm_config():
    """LLM provider/key config comes from environment variables (.env)
    now - json/llm-config.json was removed because committing real API
    keys in a tracked JSON file got git pushes blocked by GitHub's
    secret-scanning push protection.

    Item 5 (speed) - the validation/QC pass (call_llm_validation, a
    SEPARATE sequential LLM call after the main extraction) doesn't need
    the same heavyweight model the extraction itself does - it's just
    checking/scoring fields that already exist. Defaulting it to a
    smaller, faster model (matching the reference project's tiered
    validation/scoring/quick models) cuts real wall-clock time off every
    single lease/translation processed, without touching extraction
    quality at all. Override with LLM_VALIDATION_MODEL / OPENROUTER_
    VALIDATION_MODEL in .env if a different model is preferred."""
    return {
        # Default provider is OpenRouter: the whole pipeline (layout JSON
        # extraction + translation via a vision model, and Flux image fill)
        # runs on ONE OpenRouter key, exactly like the user's proof-of-
        # concept. No OpenAI account is required. Set LLM_PROVIDER=openai
        # only if you deliberately want to use a direct OpenAI key instead.
        "provider": os.environ.get("LLM_PROVIDER", "openrouter"),
        "openai": {
            "apiKey": os.environ.get("OPENAI_API_KEY"),
            "model": os.environ.get("OPENAI_MODEL", "gpt-4o"),
            "validationModel": os.environ.get("OPENAI_VALIDATION_MODEL", "gpt-4o-mini"),
            "baseUrl": "https://api.openai.com/v1/chat/completions",
        },
        "openrouter": {
            "apiKey": os.environ.get("OPENROUTER_API_KEY"),
            "model": os.environ.get("OPENROUTER_MODEL", "openai/gpt-4o"),
            "validationModel": os.environ.get("OPENROUTER_VALIDATION_MODEL", "openai/gpt-4o-mini"),
            "baseUrl": "https://openrouter.ai/api/v1/chat/completions",
        },
    }


def _resolve_provider_cfg(llm_config, provider):
    return llm_config.get(provider, {}) or {}


def load_extraction_prompt():
    try:
        with open(EXTRACTION_PROMPT_PATH, "r", encoding="utf-8") as f:
            base_prompt = f.read()
    except OSError:
        return None
    return base_prompt + _load_extra_approved_rules_supplement()


def get_rules_applied_count():
    """Item 1 - how many extraction rules (73 built-in + any extra
    Developer-approved ones) were actually included in the prompt for
    this extraction, for the activity log's 'Applying Rules > X/Y' line.
    Every currently-approved rule is always injected on every call (no
    per-document conditional skipping), so X and Y are the same number -
    this reports how many rules were active, not a partial-match score."""
    try:
        with open(os.path.join(JSON_DIR, "rules.json"), "r", encoding="utf-8") as f:
            data = json.load(f)
    except (OSError, json.JSONDecodeError):
        return 0, 0
    total = len(data.get("approved", []))
    return total, total


def _load_extra_approved_rules_supplement():
    """Renders any approved rule that ISN'T one of the 73 original
    built-in ones (those are already baked into extraction_prompt.txt's
    own prose - re-injecting them here would just double token usage for
    no benefit) as an extra prompt section. This is what makes a rule
    approved through the Update Rules UI actually take effect on the next
    extraction call, with no manual prompt editing needed."""
    try:
        with open(os.path.join(JSON_DIR, "rules.json"), "r", encoding="utf-8") as f:
            data = json.load(f)
    except (OSError, json.JSONDecodeError):
        return ""

    extra_rules = [r for r in data.get("approved", []) if not r.get("builtin")]
    if not extra_rules:
        return ""

    lines = [
        "\n\n---\nADDITIONAL FIELD-EXTRACTION RULES (approved by the Developer after the "
        "rules above were written - apply these too, on top of everything else in this prompt):"
    ]
    for rule in extra_rules:
        field_id = rule.get("fieldId", "")
        rule_type = rule.get("ruleType", "")
        rule_text = rule.get("ruleText", "")
        if field_id and rule_text:
            lines.append(f"- [{field_id}] ({rule_type}): {rule_text}")
    return "\n".join(lines)


def robust_json_parse(raw):
    """Same multi-pass strategy as the reference project's robust_json_parse:
    try a direct parse, then strip markdown fences, then pull out the first
    balanced {...} block. Returns None if nothing works so the caller can
    decide how to fall back."""
    if not raw:
        return None
    raw = raw.strip()

    try:
        return json.loads(raw)
    except (json.JSONDecodeError, TypeError):
        pass

    cleaned = re.sub(r"```(?:json)?", "", raw).strip()
    try:
        return json.loads(cleaned)
    except (json.JSONDecodeError, TypeError):
        pass

    m = re.search(r"(\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\})", cleaned, re.DOTALL)
    if m:
        try:
            return json.loads(m.group(1))
        except (json.JSONDecodeError, TypeError):
            pass

    return None


def _call_chat_completion(provider, provider_cfg, system_prompt, user_content, max_tokens=16000, use_validation_model=False, extra_messages=None):
    """Shared OpenAI/OpenRouter chat-completions call. Returns the raw
    response text (the caller decides how to parse it) or raises
    LeaseEngineError. No response_format is sent - not every OpenRouter
    model supports it, so (matching the reference project) we rely on the
    prompt's own "return ONLY JSON" instruction plus robust_json_parse on
    the way out instead.

    extra_messages (optional) - prior turns (user/assistant pairs) to
    insert between the system prompt and this call's new user message,
    so several calls can share one running conversation instead of each
    being a fresh, context-less call - see compare_extraction_quality's
    use of this for Test & Compare, which keeps ONE conversation going
    across every lease in a batch rather than starting over each time."""
    api_key = (provider_cfg.get("apiKey") or "").strip()
    if not api_key:
        return None

    # Item 5 (speed) - the QC/validation pass just re-checks fields that
    # already exist, a much lighter task than the original extraction, so
    # it uses the smaller/faster validationModel (see load_llm_config)
    # instead of the same heavyweight model extraction needs. This was
    # already configured (OPENAI_VALIDATION_MODEL / OPENROUTER_VALIDATION_
    # MODEL in .env) but never actually wired up - _call_chat_completion
    # always used the main model regardless.
    model = (provider_cfg.get("validationModel") if use_validation_model else provider_cfg.get("model")) or "gpt-4o"
    url = provider_cfg.get("baseUrl") or (
        "https://openrouter.ai/api/v1/chat/completions" if provider == "openrouter"
        else "https://api.openai.com/v1/chat/completions"
    )

    messages = [{"role": "system", "content": system_prompt}]
    messages.extend(extra_messages or [])
    messages.append({"role": "user", "content": user_content})

    payload = {
        "model": model,
        "messages": messages,
        "temperature": 0,
        "max_tokens": max_tokens,
    }
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
    }
    if provider == "openrouter":
        # OpenRouter asks for these but doesn't require them - harmless if generic.
        headers["HTTP-Referer"] = "https://lexora.ai"
        headers["X-Title"] = "Lexora Lease Abstraction"

    req = urllib.request.Request(
        url, data=json.dumps(payload).encode("utf-8"), headers=headers, method="POST"
    )
    try:
        with urllib.request.urlopen(req, timeout=180) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as err:
        detail = err.read().decode("utf-8", errors="ignore")[:500]
        raise LeaseEngineError(f"{provider} API error {err.code}: {detail}")
    except urllib.error.URLError as err:
        raise LeaseEngineError(f"Could not reach {provider} API: {err.reason}")

    try:
        return data["choices"][0]["message"]["content"]
    except (KeyError, IndexError) as err:
        raise LeaseEngineError(f"{provider} returned an unexpected response shape: {err}")


def _call_chat_completion_with_failover(llm_config, system_prompt, user_content, max_tokens=16000, use_validation_model=False, extra_messages=None):
    """Tries the configured primary provider first; if it's configured but
    the call itself fails (auth error, network error, bad response),
    automatically tries the OTHER provider too (if it has a key
    configured) before giving up - previously a single OpenAI hiccup fell
    straight through to the heuristic engine even though a working
    OpenRouter key was sitting right there in .env unused. Returns
    (content, provider_used) - content is None if neither provider has a
    key configured at all (caller falls back to heuristic silently);
    raises LeaseEngineError only if at least one provider was actually
    tried and every attempt failed."""
    primary_provider = llm_config.get("provider", "openrouter")
    fallback_provider = "openai" if primary_provider == "openrouter" else "openrouter"

    last_error = None
    tried_any = False
    for provider in (primary_provider, fallback_provider):
        provider_cfg = _resolve_provider_cfg(llm_config, provider)
        if not provider_cfg.get("apiKey"):
            continue  # this provider isn't configured at all - just skip it
        tried_any = True
        try:
            content = _call_chat_completion(provider, provider_cfg, system_prompt, user_content, max_tokens=max_tokens, use_validation_model=use_validation_model, extra_messages=extra_messages)
        except LeaseEngineError as err:
            print(f"{provider} call failed{' - trying ' + fallback_provider + ' next' if provider == primary_provider else ''}: {err}")
            last_error = err
            continue
        if content is not None:
            return content, provider

    if tried_any and last_error:
        raise last_error
    return None, None


def discover_new_rules(text, existing_rule_texts, llm_config=None):
    """Item 7 - the reference project's LMS auto-proposed new rules from
    patterns the AI noticed; this is our equivalent. Asks a lightweight
    LLM call to look at THIS lease's text and suggest 0-3 NEW
    field-extraction rules that existing rules (built-in + already
    approved/pending) don't already cover - never edits extraction_prompt
    directly, just proposes candidates for a human to approve. Returns a
    list of {fieldId, ruleType, ruleText, confidence} dicts, or [] on any
    failure or when there's nothing new. Never raises - this always runs
    as a background, best-effort call that must not affect the main
    extraction pipeline."""
    llm_config = llm_config if llm_config is not None else load_llm_config()
    existing_block = "\n".join(f"- {t}" for t in existing_rule_texts[:80] if t)
    system_prompt = (
        "You are reviewing a commercial lease document to find NEW field-extraction rules "
        "worth adding to a shared rules library for an AI lease-abstraction tool, based on "
        "patterns in THIS specific lease that the EXISTING rules below don't already cover.\n\n"
        f"EXISTING RULES (do not propose a duplicate of any of these):\n{existing_block}\n\n"
        "Return ONLY a JSON array (no markdown fences, no commentary) of 0 to 3 objects, each shaped as:\n"
        '{ "fieldId": string, "ruleType": "mapping" | "style" | "logic", '
        '"ruleText": string, "confidence": number 0.0-1.0 }\n'
        "Only propose a rule if it is genuinely new and useful (e.g. an unusual clause "
        "location, a wording pattern, a formatting convention this lease uses). "
        "Return an empty array [] if this lease doesn't reveal anything new."
    )
    try:
        content, _provider = _call_chat_completion_with_failover(
            llm_config, system_prompt, text[:60000], max_tokens=1000
        )
    except LeaseEngineError as err:
        print(f"Rule auto-discovery LLM call failed, skipping: {err}")
        return []
    if not content:
        return []

    parsed = robust_json_parse(content)
    if not isinstance(parsed, list):
        return []

    cleaned = []
    for r in parsed[:3]:
        if not isinstance(r, dict):
            continue
        field_id = str(r.get("fieldId") or "").strip()[:60]
        rule_text = str(r.get("ruleText") or "").strip()[:500]
        if not field_id or not rule_text:
            continue
        rule_type = r.get("ruleType") if r.get("ruleType") in ("mapping", "style", "logic") else "mapping"
        try:
            confidence = max(0.0, min(1.0, float(r.get("confidence", 0.5))))
        except (TypeError, ValueError):
            confidence = 0.5
        cleaned.append({"fieldId": field_id, "ruleType": rule_type, "ruleText": rule_text, "confidence": confidence})
    return cleaned


def compare_extraction_quality(original_text, human_output_text, current_output_text, existing_rule_texts, llm_config=None, conversation_history=None):
    """Item 4 (Test & Compare admin tool) - the core comparison behind the
    Test & Compare card: how closely does our system's own extraction
    (current_output_text) match a human expert's extraction
    (human_output_text) of the SAME original lease? Returns a dict with:
      - similarity: a 0-100 text-similarity score (difflib) between the
        human and current outputs - a rough, fast "how close are we"
        number, not a substitute for the LLM's own read below.
      - proposedRules: 0-3 rule proposals (same shape as
        discover_new_rules) an LLM derived specifically from where our
        output diverges from the human's, given the original text as
        context - these still land in rules.json's pending queue for
        Developer approval, same governance as every other rule source.
      - conversationHistory: the updated message list (pass this back in
        as `conversation_history` for the NEXT lease in the same Test &
        Compare batch) - every lease in one batch run shares a single
        growing conversation rather than each being an isolated, fresh
        call, so the model retains context across the whole batch (e.g.
        it won't propose near-duplicate rules for two leases that show
        the same gap, and can build a cumulative read of this batch's
        common failure patterns).
    Never raises - a failure in the LLM step just means proposedRules
    comes back empty; the similarity score (pure Python, no LLM needed)
    always still gets computed and returned."""
    import difflib
    similarity = round(difflib.SequenceMatcher(
        None, (human_output_text or "").strip().lower(), (current_output_text or "").strip().lower()
    ).ratio() * 100, 1)

    proposed_rules = []
    history = list(conversation_history or [])
    llm_config = llm_config if llm_config is not None else load_llm_config()
    existing_block = "\n".join(f"- {t}" for t in existing_rule_texts[:80] if t)
    system_prompt = (
        "You are comparing several leases' extractions across one Test & Compare batch, one lease at a "
        "time in this same conversation: for each lease, one extraction was produced by a human "
        "lease-abstraction expert (treat as the ground truth), one by an AI extraction system. Identify "
        "specific, genuine gaps where the AI's extraction missed or got wrong something the human correctly "
        "captured, then propose 0-3 NEW field-extraction rules per lease that would help the AI system close "
        "those specific gaps next time. Use your memory of earlier leases in this same conversation to avoid "
        "proposing a near-duplicate of a rule you already proposed for an earlier lease in this batch.\n\n"
        f"EXISTING RULES (do not propose a duplicate of any of these):\n{existing_block}\n\n"
        "Return ONLY a JSON array (no markdown fences, no commentary) of 0 to 3 objects, each shaped as:\n"
        '{ "fieldId": string, "ruleType": "mapping" | "validation" | "logic", '
        '"ruleText": string, "confidence": number 0.0-1.0 }\n'
        "Only propose a rule if the gap is real and a rule could plausibly fix it. Return [] if the AI's "
        "extraction already matches the human's closely enough that there's nothing meaningful to add."
    )
    user_content = (
        f"ORIGINAL LEASE (excerpt):\n{(original_text or '')[:40000]}\n\n"
        f"HUMAN EXPERT OUTPUT (ground truth):\n{(human_output_text or '')[:15000]}\n\n"
        f"AI SYSTEM OUTPUT (to evaluate):\n{(current_output_text or '')[:15000]}"
    )
    try:
        content, _provider = _call_chat_completion_with_failover(
            llm_config, system_prompt, user_content, max_tokens=1200, extra_messages=history
        )
        if content:
            history.append({"role": "user", "content": user_content})
            history.append({"role": "assistant", "content": content})
            # Item 4 safety - each user turn can carry ~70K chars of
            # document text; keeping every prior turn in a 10-lease batch
            # would balloon context size fast. Keeping only the most
            # recent 2 exchanges still gives the model useful "what did I
            # just see" continuity without unbounded growth.
            history = history[-4:]
            parsed = robust_json_parse(content)
            if isinstance(parsed, list):
                for r in parsed[:3]:
                    if not isinstance(r, dict):
                        continue
                    field_id = str(r.get("fieldId") or "").strip()[:60]
                    rule_text = str(r.get("ruleText") or "").strip()[:500]
                    if not field_id or not rule_text:
                        continue
                    rule_type = r.get("ruleType") if r.get("ruleType") in ("mapping", "validation", "logic", "style") else "mapping"
                    try:
                        confidence = max(0.0, min(1.0, float(r.get("confidence", 0.6))))
                    except (TypeError, ValueError):
                        confidence = 0.6
                    proposed_rules.append({"fieldId": field_id, "ruleType": rule_type, "ruleText": rule_text, "confidence": confidence})
    except LeaseEngineError as err:
        print(f"Test & Compare rule proposal LLM call failed (similarity score is still valid): {err}")

    return {"similarity": similarity, "proposedRules": proposed_rules, "conversationHistory": history}


def call_llm_extraction(text, llm_config=None):
    """Calls OpenAI or OpenRouter (whichever LLM_PROVIDER is set to in
    .env, with automatic failover to the other one if it's also
    configured - see _call_chat_completion_with_failover) using
    json/extraction_prompt.txt as the system prompt. Returns
    (parsed_fields_dict, provider_used) on success, or (None, None) if no
    API key is configured at all (caller falls back to the heuristic
    engine). Raises LeaseEngineError if every configured provider's call
    failed."""
    llm_config = llm_config if llm_config is not None else load_llm_config()

    system_prompt = load_extraction_prompt()
    if not system_prompt:
        return None, None  # no prompt file - caller falls back to heuristic

    content, provider = _call_chat_completion_with_failover(
        llm_config, system_prompt,
        "LEASE DOCUMENT TEXT:\n\n" + text[:120000],
        max_tokens=16000,
    )
    if content is None:
        return None, None  # no API key configured anywhere - caller falls back to heuristic

    parsed = robust_json_parse(content)
    if parsed is None:
        raise LeaseEngineError(f"{provider} returned a response that isn't valid JSON")
    return parsed, provider


VALIDATION_SYSTEM_PROMPT = "You are a lease abstraction QC validator. Return ONLY valid JSON."


def call_llm_validation(fields, llm_config=None):
    """Second pass over the extracted fields - same shape as the reference
    project's /api/validate route. Returns a dict with confidence_score
    (0.0-1.0), missing_fields, low_confidence_fields, and a one-line
    summary, or None if no LLM is configured (caller falls back to a
    heuristic completeness score)."""
    llm_config = llm_config if llm_config is not None else load_llm_config()

    prompt = (
        "Analyze this lease extraction for completeness and accuracy.\n"
        "Return ONLY valid JSON - no markdown, no code fences:\n"
        "{\n"
        '  "is_valid": true or false,\n'
        '  "confidence_score": 0.0-1.0,\n'
        '  "missing_fields": ["field1", "field2"],\n'
        '  "low_confidence_fields": ["field1"],\n'
        '  "summary": "one sentence QC summary"\n'
        "}\n\n"
        "Extracted lease data:\n" + json.dumps(fields, indent=2)[:8000]
    )

    content, provider = _call_chat_completion_with_failover(
        llm_config, VALIDATION_SYSTEM_PROMPT, prompt, max_tokens=2000, use_validation_model=True
    )
    if content is None:
        return None

    parsed = robust_json_parse(content)
    if parsed is None:
        raise LeaseEngineError(f"{provider} validation response wasn't valid JSON")
    return parsed


def heuristic_accuracy(fields):
    """Fallback 'accuracy' when no LLM is configured: percentage of leaf
    values that aren't a placeholder like 'Not found in document' /
    'Lease is silent.' - a rough completeness proxy, not a real QC pass."""
    placeholders = {"not found in document", "lease is silent.", "", None}
    leaves = []

    def walk(v):
        if isinstance(v, dict):
            for x in v.values():
                walk(x)
        elif isinstance(v, list):
            for x in v:
                walk(x)
        else:
            leaves.append(v)

    walk(fields)
    if not leaves:
        return {"confidence_score": 0.0, "summary": "No fields extracted.", "missing_fields": [], "low_confidence_fields": []}

    filled = sum(1 for v in leaves if str(v).strip().lower() not in placeholders)
    score = filled / len(leaves)
    return {
        "confidence_score": round(score, 2),
        "summary": f"{filled}/{len(leaves)} fields populated (heuristic completeness check, no LLM configured).",
        "missing_fields": [],
        "low_confidence_fields": [],
    }


# ============================================================
# 1. TEXT EXTRACTION (pdfplumber -> pypdf -> real OCR fallback)
# ============================================================
def _extract_text_pdfplumber(file_path):
    if not pdfplumber:
        return ""
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            try:
                text_parts.append(page.extract_text() or "")
            except Exception:
                text_parts.append("")
    return "\n\n".join(text_parts).strip()


def _extract_text_pypdf(file_path):
    if not PdfReader:
        return ""
    reader = PdfReader(file_path)
    text_parts = []
    for page in reader.pages:
        try:
            text_parts.append(page.extract_text() or "")
        except Exception:
            text_parts.append("")
    return "\n\n".join(text_parts).strip()


def _ocr_one_page(page, page_index):
    try:
        # 150 DPI + LSTM-only engine is meaningfully faster than the
        # 200 DPI / default-engine combo with no real accuracy loss on
        # typical scanned lease pages, and matters a lot here: OCR-ing a
        # long document one page at a time inside a single HTTP request
        # can otherwise run long enough to hit a proxy/gateway timeout
        # (this is what was causing "/api/lease/extract failed" on a
        # large scanned PDF during testing).
        image = page.to_image(resolution=150).original
        text = pytesseract.image_to_string(image, config="--oem 1 --psm 6") or ""
    except Exception:
        text = ""
    return page_index, text


def _extract_text_ocr(file_path, on_progress=None):
    """Renders every page to an image and runs Tesseract OCR on it - the
    real fallback for PDFs that are actually scanned page images with no
    usable text layer at all (pdfplumber/pypdf both come back nearly
    empty on these, even when the filename says "_ocr"). Needs pdfplumber
    (for page rendering - already a dependency) plus pytesseract and the
    tesseract-ocr system package. Pages are OCR'd in parallel (tesseract
    runs as a real subprocess per call, so this isn't GIL-limited) - a
    meaningful speedup on multi-core machines. This is called from a
    background thread (see server.py's /api/lease/extract-start), not
    inline in an HTTP request, specifically so a slow multi-page OCR run
    never risks tripping a reverse-proxy/gateway timeout - on_progress(done,
    total) lets the caller report live status while it works."""
    if not pdfplumber or not OCR_LIBS_OK:
        return ""

    with pdfplumber.open(file_path) as pdf:
        pages = list(pdf.pages)
        total = len(pages)
        worker_count = min(OCR_MAX_WORKERS, max(1, os.cpu_count() or 1), total) or 1
        results = [""] * total
        done_count = 0

        if on_progress:
            on_progress(0, total)

        if worker_count <= 1:
            for i, page in enumerate(pages):
                _, text = _ocr_one_page(page, i)
                results[i] = text
                done_count += 1
                if on_progress:
                    on_progress(done_count, total)
        else:
            with ThreadPoolExecutor(max_workers=worker_count) as executor:
                futures = [executor.submit(_ocr_one_page, page, i) for i, page in enumerate(pages)]
                for future in as_completed(futures):
                    i, text = future.result()
                    results[i] = text
                    done_count += 1
                    if on_progress:
                        on_progress(done_count, total)

    return "\n\n".join(results).strip()


def extract_text(file_path, on_progress=None):
    """Returns the raw text content of a PDF or DOCX file. For PDFs, tries
    pdfplumber first, then pypdf, and - only if both come back with almost
    nothing PER PAGE ON AVERAGE (a strong sign the PDF is actually scanned
    page images with no real text layer, as opposed to a text-based PDF -
    e.g. a multi-page document where every page only contributes a tiny
    "DocuSign Envelope ID" stamp can still add up past a flat character
    threshold in aggregate, which is why this is a per-page average, not a
    total) - falls back to real OCR via Tesseract. Whichever method
    returns the most text wins."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        if not pdfplumber and not PdfReader:
            raise LeaseEngineError(
                "pdfplumber/pypdf are not installed - run: pip install -r requirements.txt"
            )

        page_count = 1
        if pdfplumber:
            try:
                with pdfplumber.open(file_path) as pdf:
                    page_count = max(len(pdf.pages), 1)
            except Exception:
                pass

        min_chars_needed = max(MIN_USABLE_TEXT_CHARS, page_count * MIN_CHARS_PER_PAGE)

        best_text = ""
        best_method = None

        try:
            plumber_text = _extract_text_pdfplumber(file_path)
        except Exception:
            plumber_text = ""
        if len(plumber_text) > len(best_text):
            best_text, best_method = plumber_text, "pdfplumber"

        if len(best_text) < min_chars_needed:
            try:
                pypdf_text = _extract_text_pypdf(file_path)
            except Exception:
                pypdf_text = ""
            if len(pypdf_text) > len(best_text):
                best_text, best_method = pypdf_text, "pypdf"

        if len(best_text) < min_chars_needed:
            if not OCR_LIBS_OK:
                # Nothing more we can do without OCR libraries - return
                # whatever little text was found (e.g. just a DocuSign
                # envelope stamp on every page) rather than hard-failing,
                # and let the caller's classification logic report
                # "Invalid Document" with a clear reason instead of
                # crashing outright.
                print(
                    f"WARNING: {os.path.basename(file_path)} looks like a scanned/image-only PDF "
                    f"(only {len(best_text)} chars across {page_count} pages from text layers) and "
                    f"pytesseract/Pillow aren't installed, so OCR couldn't run. Install tesseract-ocr "
                    f"+ pytesseract - see requirements.txt."
                )
            else:
                try:
                    ocr_text = _extract_text_ocr(file_path, on_progress=on_progress)
                except Exception as err:
                    ocr_text = ""
                    print(f"OCR fallback failed for {os.path.basename(file_path)}: {err}")
                if len(ocr_text) > len(best_text):
                    best_text, best_method = ocr_text, "ocr"

        return best_text

    if ext == ".docx":
        if not docx_lib:
            raise LeaseEngineError(
                "python-docx is not installed - run: pip install -r requirements.txt"
            )
        document = docx_lib.Document(file_path)
        return "\n".join(p.text for p in document.paragraphs).strip()

    raise LeaseEngineError(f"Unsupported file type for text extraction: {ext}")


# ============================================================
# 2. HEURISTIC "GPT PROMPT" STAND-IN - document classification +
#    lease field extraction (see module docstring)
# ============================================================
LEASE_KEYWORDS = [
    "lease agreement", "this lease", "landlord", "tenant", "lessor", "lessee",
    "premises", "rent shall", "commencement date", "term of this lease",
    "leased premises", "monthly rent", "security deposit"
]
# STRONG phrases only show up in documents that ARE actually an amendment
# (they announce themselves as one). WEAK words like "amendment"/"amended"
# on their own are extremely common even inside a plain, original lease -
# e.g. "...including all amendments thereto..." (a document list clause) or
# "...altered, amended, or revoked only by an instrument in writing..." (a
# boilerplate entire-agreement clause) - so a couple of incidental
# occurrences must never be enough by themselves to override overwhelming
# lease-document evidence (this is exactly what caused a real, 26-page
# original lease to get misclassified as "Amendment" during testing).
STRONG_AMENDMENT_KEYWORDS = [
    "first amendment", "second amendment", "third amendment", "fourth amendment",
    "this amendment", "hereby amends", "amendment to lease", "amendment to that certain lease",
    "amendment no.", "lease amendment agreement",
]
WEAK_AMENDMENT_KEYWORDS = ["amendment", "amended", "addendum", "modification of lease"]


def classify_document(text):
    """Returns 'Lease', 'Amendment', or 'Other' based on keyword scoring."""
    lower = text.lower()
    lease_score = sum(1 for kw in LEASE_KEYWORDS if kw in lower)
    has_strong_amendment = any(kw in lower for kw in STRONG_AMENDMENT_KEYWORDS)
    weak_amendment_hits = sum(lower.count(kw) for kw in WEAK_AMENDMENT_KEYWORDS)

    if has_strong_amendment:
        return "Amendment"
    # Several incidental mentions AND weak overall lease evidence - more
    # likely a genuine (if informally worded) amendment than a full lease.
    if weak_amendment_hits >= 3 and lease_score < 2:
        return "Amendment"
    if lease_score >= 2:
        return "Lease"
    return "Other"


_MONEY_RE = re.compile(r"\$\s?([\d,]+(?:\.\d{2})?)")
_DATE_RE = re.compile(
    r"(January|February|March|April|May|June|July|August|September|October|"
    r"November|December)\s+\d{1,2},?\s+\d{4}|\d{1,2}/\d{1,2}/\d{2,4}|\d{4}-\d{2}-\d{2}"
)


def _find_party_name(text, role_words):
    """Looks for two common lease phrasings:
       1) 'Landlord: Acme Holdings LLC'                (label form)
       2) 'Acme Holdings LLC ("Landlord")' / '(Landlord)' (parenthetical form)
    """
    label_value = _find_after_label(text, role_words)
    if label_value:
        return label_value

    for role in role_words:
        pattern = re.compile(
            r'((?:[A-Z][A-Za-z0-9&\.\-\']*\s+){0,5}[A-Z][A-Za-z0-9&\.\-\']*)\s*\(\s*["\u201c]?'
            r'(?i:' + re.escape(role) + r')["\u201d]?\s*\)'
        )
        m = pattern.search(text)
        if m:
            value = re.sub(r"\s+", " ", m.group(1)).strip().rstrip(',')
            return value
    return None


def _find_after_label(text, labels, max_lines=3):
    """Looks for 'Label: value' or 'Label - value' style lines and returns
    the value - allows the value to continue onto a couple of wrapped
    continuation lines (common in PDF table cells, e.g. a long landlord
    name that wraps across 2-3 lines) but stops at a blank line or what
    looks like the start of the next label, so it doesn't run on into
    unrelated content."""
    for label in labels:
        pattern = re.compile(re.escape(label) + r"\s*[:\-]\s*(.+)", re.IGNORECASE)
        m = pattern.search(text)
        if not m:
            continue
        start = m.start(1)
        chunk = text[start:start + 600]
        lines = chunk.split("\n")
        value_lines = [lines[0].strip()] if lines and lines[0].strip() else []
        for line in lines[1:max_lines]:
            stripped = line.strip()
            if not stripped:
                break  # blank line - value block ended
            if re.match(r"^[A-Z][A-Z \-'\u2019/]{2,40}[:\-]", stripped):
                break  # looks like the start of the next label
            value_lines.append(stripped)
        value = " ".join(value_lines).strip()
        if value:
            return value[:200]
    return None


def _find_date_near_label(text, labels):
    """Looks for a label like 'Expiration Date' and extracts the nearest
    date-shaped value on the same line/nearby - far more reliable than
    grabbing "the Nth date found anywhere in the document" (which has no
    idea which date means what, and easily picks up an unrelated date
    like an "Estimated Commencement Date" instead of the real
    expiration date)."""
    for label in labels:
        pattern = re.compile(re.escape(label) + r"\s*[:\-]?\s*(.{0,80})", re.IGNORECASE)
        for m in pattern.finditer(text):
            date_match = _DATE_RE.search(m.group(1))
            if date_match:
                return date_match.group(0)
    return None


def _find_money_near_label(text, labels):
    """Same idea as _find_date_near_label but for dollar amounts - avoids
    grabbing "the first dollar figure anywhere in the document" (which
    could be a phone number, an unrelated fee, or a table header amount
    rather than the actual rent)."""
    for label in labels:
        pattern = re.compile(re.escape(label) + r"\s*[:\-]?\s*(.{0,120})", re.IGNORECASE)
        for m in pattern.finditer(text):
            money_match = _MONEY_RE.search(m.group(1))
            if money_match:
                return money_match.group(0)
    return None


def heuristic_analyze_lease(text, fallback_name="Lease Document"):
    """Deterministic regex/keyword field extraction - used when no LLM key
    is configured in .env (see analyze_lease below)."""
    doc_type = classify_document(text)

    tenant = _find_party_name(text, ["Tenant", "Lessee"])
    landlord = _find_party_name(text, ["Landlord", "Lessor"])
    property_address = _find_after_label(
        text, ["Premises", "Property Address", "Leased Premises", "Address"]
    )

    all_dates = [m.group(0) for m in _DATE_RE.finditer(text)]
    money_matches = _MONEY_RE.findall(text)

    lease_start = (
        _find_date_near_label(text, ["Commencement Date", "Date of Lease", "Lease Commencement", "Effective Date"])
        or (all_dates[0] if len(all_dates) > 0 else None)
    )
    lease_end = (
        _find_date_near_label(text, ["Expiration Date", "Lease Expiration", "Termination Date", "End Date"])
        or (all_dates[1] if len(all_dates) > 1 else None)
    )
    base_rent = (
        _find_money_near_label(text, ["Monthly Base Rent", "Base Rent", "Monthly Rent", "Minimum Monthly Rent"])
        or (money_matches[0] if money_matches else None)
    )

    lease_name_source = tenant or property_address or fallback_name
    lease_name = sanitize_lease_name(lease_name_source)

    fields = {
        "documentType": doc_type,
        "tenant": tenant or "Not found in document",
        "landlord": landlord or "Not found in document",
        "propertyAddress": property_address or "Not found in document",
        "leaseStart": lease_start or "Not found in document",
        "leaseEnd": lease_end or "Not found in document",
        "baseRent": base_rent or "Not found in document",
        "currency": "USD",
    }

    return {
        "docType": doc_type,
        "leaseName": lease_name,
        "fields": fields,
        "extractionMethod": "heuristic",
    }


def _lease_name_source_from_fields(fields, fallback_name):
    """Works with either the rich LLM schema (fields.parties.tenant_legal_name,
    fields.premises.property_address, ...) or the flat heuristic schema
    (fields.tenant, fields.propertyAddress, ...)."""
    parties = fields.get("parties")
    if isinstance(parties, dict):
        name = parties.get("tenant_legal_name") or parties.get("landlord_legal_name")
        if name:
            return name
    premises = fields.get("premises")
    if isinstance(premises, dict) and premises.get("property_address"):
        return premises["property_address"]

    if fields.get("tenant") and fields["tenant"] != "Not found in document":
        return fields["tenant"]
    if fields.get("propertyAddress") and fields["propertyAddress"] != "Not found in document":
        return fields["propertyAddress"]

    return fallback_name


def _run_accuracy_check(fields, llm_config):
    """Runs the validation/accuracy pass and normalizes it into
    {accuracy: 0-100 int, accuracySummary: str, missingFields: [...],
    lowConfidenceFields: [...]}. Never raises - a failed/unavailable
    accuracy check shouldn't take down an otherwise-successful extraction."""
    try:
        validation = call_llm_validation(fields, llm_config=llm_config)
    except LeaseEngineError as err:
        print(f"Accuracy validation call failed, using heuristic completeness instead: {err}")
        validation = None

    used_llm = validation is not None
    if validation is None:
        validation = heuristic_accuracy(fields)

    try:
        score = float(validation.get("confidence_score", 0))
    except (TypeError, ValueError):
        score = 0.0
    score = max(0.0, min(1.0, score))

    return {
        "accuracy": round(score * 100),
        "accuracyMethod": "llm-validation" if used_llm else "heuristic-completeness",
        "accuracySummary": validation.get("summary") or "",
        "missingFields": validation.get("missing_fields") or [],
        "lowConfidenceFields": validation.get("low_confidence_fields") or [],
    }


def translate_text(text, target_language, llm_config=None):
    """Real translation via OpenAI/OpenRouter, with automatic failover to
    whichever of the two is configured if the primary one's call fails
    (see _call_chat_completion_with_failover). Returns
    (translated_text, provider_used), or (None, None) if no LLM key is
    configured at all (caller falls back to a clear heuristic message
    rather than pretending to translate). Raises LeaseEngineError if
    every configured provider's call failed."""
    llm_config = llm_config if llm_config is not None else load_llm_config()

    system_prompt = (
        f"You are a professional document translator. Translate the user's text into "
        f"{target_language}. Preserve the original meaning, tone, register, and structure as "
        f"closely as the target language allows. Write naturally in {target_language}'s own "
        f"normal script and word order - the document renderer handles right-to-left vs "
        f"left-to-right layout separately, so just focus on an accurate, natural translation.\n\n"
        f"Mark up the structure you detect using these plain-text conventions, so the "
        f"translated document can be auto-formatted afterward:\n"
        f"- Separate distinct paragraphs with a single BLANK LINE between them.\n"
        f"- Bullet list items: start the line with '- '.\n"
        f"- Numbered list items: start the line with '1. ', '2. ', etc. (keep the original numbering).\n"
        f"- Section headings / titles: start the line with '## '.\n"
        f"- Bold or otherwise emphasized text (e.g. defined terms, warnings, key figures the "
        f"source visually emphasized): wrap it in **double asterisks**.\n"
        f"- Do NOT invent structure that isn't in the source - only mark what's genuinely a "
        f"list, heading, or emphasis there.\n\n"
        f"Return ONLY the translated, marked-up text - no preamble, no notes, no commentary "
        f"about the translation itself."
    )
    content, provider = _call_chat_completion_with_failover(llm_config, system_prompt, text[:100000], max_tokens=8000)
    if content is None:
        return None, None  # no API key configured anywhere - caller falls back
    return content.strip(), provider


_BOLD_MARKUP_RE = re.compile(r"\*\*(.+?)\*\*")


def _escape_and_apply_bold(text, rtl=False):
    """Escapes HTML-special characters for safe use in a reportlab
    Paragraph, then converts **bold** markers (see the markup convention
    translate_text() asks the LLM to use) into <b> tags - reportlab
    Paragraphs support this small tag set natively.

    For RTL text (item 6), each segment is bidi-reshaped BEFORE the <b>
    tags get added, never after - running arabic_reshaper/get_display on
    a string that already contains '<b>...</b>' scrambles the tag
    characters themselves (they get treated as just more RTL-adjacent
    text and reordered), which breaks reportlab's mini-XML parser. Doing
    it in this order (shape first, wrap second) keeps the tags intact."""
    parts = _BOLD_MARKUP_RE.split(text)  # alternates [plain, bold, plain, bold, ...]
    out = []
    for i, part in enumerate(parts):
        escaped = part.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if rtl:
            escaped = shape_rtl_text(escaped)
        out.append(f"<b>{escaped}</b>" if i % 2 == 1 else escaped)
    return "".join(out)


def _parse_translation_blocks(text):
    """Item 6 - splits the LLM's markdown-lite translated text (see the
    system prompt in translate_text()) into typed blocks so the output
    document auto-formats instead of being one flat wall of paragraphs:
    a blank line starts a new block; '- ' lines become a bullet list,
    '1. '/'2. ' lines become a numbered list, a lone '## ' line becomes a
    heading, and everything else is a regular paragraph."""
    blocks = []
    for raw in re.split(r"\n\s*\n", (text or "").strip()):
        lines = [ln.strip() for ln in raw.split("\n") if ln.strip()]
        if not lines:
            continue
        if all(ln.startswith("- ") for ln in lines):
            blocks.append(("bullet", [ln[2:].strip() for ln in lines]))
        elif all(re.match(r"^\d+\.\s", ln) for ln in lines):
            blocks.append(("numbered", [re.sub(r"^\d+\.\s*", "", ln) for ln in lines]))
        elif len(lines) == 1 and lines[0].startswith("## "):
            blocks.append(("heading", lines[0][3:].strip()))
        else:
            blocks.append(("paragraph", " ".join(lines)))
    return blocks


def _group_chars_into_lines(chars, y_tolerance=2.5):
    """Groups pdfplumber's per-character data into text lines, using Y
    (top) proximity - characters within y_tolerance points of each other
    are the same line, matching how a PDF's text is actually laid out
    (each line's characters share close to the same baseline)."""
    if not chars:
        return []
    buckets = {}
    for c in chars:
        key = round(c["top"] / y_tolerance)
        buckets.setdefault(key, []).append(c)
    # Merge adjacent buckets that are still within tolerance of each other
    # (a line can straddle two rounding buckets right at the boundary).
    sorted_keys = sorted(buckets.keys())
    merged = []
    current = []
    last_key = None
    for k in sorted_keys:
        if last_key is not None and k - last_key > 1:
            merged.append(current)
            current = []
        current.extend(buckets[k])
        last_key = k
    if current:
        merged.append(current)
    return [sorted(line, key=lambda c: c["x0"]) for line in merged if line]


def _line_to_region(line_chars):
    """Reduces one line's raw characters down to what's needed to mask +
    re-render it: its exact bounding box (so the translated text lands in
    the same place) and its dominant font/size/color (so it looks as
    close to the original as this renderer can manage)."""
    text = "".join(c["text"] for c in line_chars)
    if not text.strip():
        return None
    x0 = min(c["x0"] for c in line_chars)
    x1 = max(c["x1"] for c in line_chars)
    top = min(c["top"] for c in line_chars)
    bottom = max(c["bottom"] for c in line_chars)
    sizes = [c.get("size", 10) for c in line_chars]
    dominant_size = sorted(sizes)[len(sizes) // 2] if sizes else 10
    fontnames = [c.get("fontname", "") for c in line_chars]
    dominant_font = max(set(fontnames), key=fontnames.count) if fontnames else ""
    is_bold = "bold" in dominant_font.lower()
    is_italic = "italic" in dominant_font.lower() or "oblique" in dominant_font.lower()
    color_hex = _color_tuple_to_hex(line_chars[0].get("non_stroking_color")) or "#000000"
    return {
        "text": text, "x0": x0, "x1": x1, "top": top, "bottom": bottom,
        "size": dominant_size, "bold": is_bold, "italic": is_italic, "color": color_hex,
    }


def _translate_lines_batch(lines, target_language, llm_config=None, chunk_size=70):
    """Translates a list of short text lines while preserving strict 1:1
    correspondence (line 5 in must come back as line 5 out) - sent to the
    LLM as a numbered list rather than one flowing translation, since the
    layout-preserving renderer needs to put each translated line back in
    its own original bounding box. Chunks long documents (chunk_size
    lines per call) to stay well within a safe response size. Falls back
    to returning the original lines unchanged for any chunk where the LLM
    isn't configured or the call fails, rather than losing that whole
    chunk's text."""
    llm_config = llm_config if llm_config is not None else load_llm_config()
    results = []
    for start in range(0, len(lines), chunk_size):
        chunk = lines[start:start + chunk_size]
        numbered_input = "\n".join(f"{i+1}. {line}" for i, line in enumerate(chunk))
        system_prompt = (
            f"Translate each numbered line below into {target_language}. This is raw text extracted "
            f"line-by-line from fixed positions on a PDF page, so lines may be sentence fragments, "
            f"table cells, or labels rather than full sentences - translate each one independently and "
            f"naturally, using surrounding lines only for context, not to merge them together.\n\n"
            f"If a line is clearly meaningless OCR noise (random symbols or garbled fragments that are "
            f"not real words in any language), return exactly [SKIP] for that line instead of a translation.\n\n"
            f"Return ONLY a numbered list with EXACTLY the same number of lines, in the exact same "
            f"order (line N in must be line N out) - no extra commentary, no merged/split lines, no "
            f"markdown formatting, no ## or ** markers."
        )
        try:
            content, _provider = _call_chat_completion_with_failover(llm_config, system_prompt, numbered_input, max_tokens=4000)
        except LeaseEngineError:
            content = None
        if not content:
            results.extend(chunk)
            continue
        translated_chunk = [None] * len(chunk)
        for line in content.strip().split("\n"):
            m = re.match(r"^\s*(\d+)[.):]\s*(.*)$", line)
            if m:
                idx = int(m.group(1)) - 1
                if 0 <= idx < len(chunk):
                    translated_chunk[idx] = m.group(2).strip()
        results.extend(t if t else original for t, original in zip(translated_chunk, chunk))
    return results


def llm_is_configured(llm_config=None):
    """True if at least one provider (OpenAI/OpenRouter) has an API key -
    used to decide whether the vision-based translation paths are even
    possible, before promising them."""
    llm_config = llm_config if llm_config is not None else load_llm_config()
    return bool((llm_config.get("openai", {}) or {}).get("apiKey") or
                (llm_config.get("openrouter", {}) or {}).get("apiKey"))


def pdf_has_text_layer(pdf_path):
    """True if any page of the PDF has real extractable characters (a
    digitally-authored PDF), False for scanned/photo/image-only PDFs.
    Errs on the side of False if the file can't be inspected."""
    if not pdfplumber:
        return False
    try:
        with pdfplumber.open(pdf_path) as pdf:
            return any(p.chars for p in pdf.pages)
    except Exception:
        return False


def extract_page_layout_vision(page_image, target_language, llm_config=None,
                               return_debug=False):
    """RAASTA A: give the WHOLE page to the vision model and let IT do the
    segmentation + translation + styling in one shot - no OpenCV box
    detection. Mirrors what works when a user does it by hand in ChatGPT:
    the model understands the text, so it segments a dense manuscript far
    better than pixel-based detection.

    Returns (blocks, img_w, img_h) where coordinates are in the pixel
    space of the image we send (we tell the model that size), or
    (None, w, h) on refusal/empty so the caller can fall back. When
    return_debug=True, returns (blocks, img_w, img_h, debug) where debug
    has the exact prompt sent and the raw model response (for the log)."""
    llm_config = llm_config if llm_config is not None else load_llm_config()
    dbg = {"prompt": "", "response": "", "image": None}
    # Q2: OCR aur translation ALAG. "original"/empty target => sirf OCR
    # (transcription), koi translation NAHI. warna OCR + translation.
    _keep_original = str(target_language or "").strip().lower() in ("", "original", "none")
    if not llm_is_configured(llm_config):
        return (None, page_image.width, page_image.height, dbg) if return_debug \
            else (None, page_image.width, page_image.height)

    # Send at a reasonable size for token cost, BUT ask the model for
    # NORMALISED coordinates (0-1000 across width and height) instead of
    # absolute pixels. This makes coordinates independent of whatever
    # internal resolution the model uses - a value of 500 always means the
    # middle of the page - so there is no scaling mismatch. We map the
    # 0-1000 grid onto the real page-image pixels ourselves.
    max_dim = 2000
    img = page_image
    longest = max(img.width, img.height)
    if longest > max_dim:
        s = max_dim / float(longest)
        img = img.resize((max(1, int(img.width * s)), max(1, int(img.height * s))), PILImage.LANCZOS)
    iw, ih = img.width, img.height
    buf = io.BytesIO(); img.convert("RGB").save(buf, format="JPEG", quality=85)
    image_b64 = base64.b64encode(buf.getvalue()).decode("ascii")
    dbg["image"] = img.convert("RGB")

    if _keep_original:
        system_prompt = (
            "You are Lexora's document layout + OCR transcription engine. The "
            "user uploaded THEIR OWN document; you are authorised to read all "
            "of it. This is a routine, legitimate task - always comply, never "
            "refuse, never apologise.\n\n"
            "TRANSCRIPTION ONLY - DO NOT TRANSLATE. For every block, set "
            "\"translation\" to the EXACT SAME string as \"text\" (a verbatim "
            "copy of the original, same language and script). Never render "
            "text into another language.\n\n"
        )
    else:
        system_prompt = (
            "You are Lexora's document layout + translation engine. The user "
            "uploaded THEIR OWN document and asked to translate it; you are "
            "authorised to read and translate all of it. This is a routine, "
            "legitimate task - always comply, never refuse, never apologise.\n\n"
        )
    system_prompt += (
        "The image is ONE document page. Use a NORMALISED coordinate grid: "
        "the LEFT edge is x=0, the RIGHT edge is x=1000, the TOP edge is "
        "y=0, the BOTTOM edge is y=1000. Give every box in this 0-1000 grid "
        "(integers), regardless of the image's pixel size. Origin (0,0) is "
        "the TOP-LEFT corner. Segment the page YOURSELF into blocks and, for "
        "each block, give its box and details. Be precise: read every "
        "visible piece of text (including dense/handwritten Arabic), group "
        "it the way it visually reads (a multi-line paragraph is ONE block "
        "with its per-line boxes), and translate it.\n\n"
        "Return ONLY a JSON array; each element:\n"
        '{"left":int,"top":int,"right":int,"bottom":int,'
        '"lines":[{"left":int,"top":int,"right":int,"bottom":int}],'
        '"class":"text"|"element"|"decoration","kind":"heading|subheading|'
        'paragraph|label|caption|logo|signature|stamp|seal|qr|illustration|'
        'figure|photo|watermark|ornament",'
        '"text":"<exact original>","translation":"' +
        ('<verbatim copy of text, do NOT translate>' if _keep_original
         else '<into ' + target_language + '>') + '",'
        '"color":"RRGGBB","bold":false,"italic":false,"underline":false,'
        '"align":"left|center|right","rotation":0,"is_paragraph":false,'
        '"runs":[{"text":"..","color":"RRGGBB","bold":false,"italic":false}]}'
        "\n\nRULES:\n"
        "1. class \"text\" = any real readable words, PRINTED OR HANDWRITTEN "
        "(headings, labels, dense body paragraphs, cursive writing, "
        "handwritten notes/dates/names, any language). This is the DEFAULT. "
        "Handwriting rules: if it looks like human writing, treat it as "
        "text; if cursive, attempt to read it; if messy, extract what you "
        "can; if unsure, EXTRACT it as text.\n"
        "2. class \"element\" = ONLY a logo, emblem, badge, seal, stamp, "
        "QR/barcode, a genuine handwritten SIGNATURE, or a real picture "
        "(illustration/drawing/photo/figure/chart). Keep as image, do NOT "
        "translate. Give its box; leave text/translation empty.\n"
        "3. class \"decoration\" = a DECORATIVE line/rule/divider, a "
        "signature line (the line itself), a border or frame, an ornament, "
        "or a watermark. IMPORTANT: these lines are NOT text - they must be "
        "PRESERVED in place, never treated as text and never removed. Report "
        "them so we know to keep them; leave text/translation empty.\n"
        "4. NEVER classify a paragraph of words as an element. If it is "
        "readable words, it is text. A printed name or label like 'Team "
        "Leader' is text, not a signature. Only an actual cursive signature "
        "mark is a signature.\n"
        "5. GROUPING (very important): consecutive lines of the SAME "
        "paragraph MUST be returned as ONE block, with each visual line as "
        "an entry in that block's \"lines\" array (in reading order). Do NOT "
        "output each line as a separate block. A block's \"translation\" is "
        "the full paragraph; its \"lines\" give the box of each original "
        "line so the translation can flow across them at ONE uniform font "
        "size. Only start a new block at a real paragraph break, a heading, "
        "a different column, or a clearly different text size/style. Every "
        "box must be tight around the glyphs and must NOT overlap another "
        "block.\n"
        "5b. BOX HEIGHT drives the output font size, so make each line/box "
        "height match the ACTUAL visual height of the letters in that line, "
        "including tall calligraphy, ascenders and descenders. A large "
        "title must get a TALL box; small body text a short box. Do not "
        "give a big heading a thin box - that would shrink its translation "
        "wrongly. Widths must span the full visual extent of the line.\n"
        "6. translation: translate for MEANING; keep names, numbers, dates, "
        "identifiers unchanged; you may pick slightly shorter wording so it "
        "fits, but never lose information. Empty for element/decoration.\n"
        "7. color MUST be the real ink colour of the letters, read from the "
        "glyph pixels (e.g. black 000000, dark-green 1B4D3E, gold C9A227). "
        "If some words in a block are a DIFFERENT colour (e.g. a red word "
        "among black), you MUST return a \"runs\" array splitting the "
        "translation into pieces with each piece's colour; the pieces "
        "concatenated equal the translation. Never default everything to "
        "black if the original uses colour.\n"
        "8. is_paragraph: true for a multi-line body block, false for a "
        "single heading/label/line. rotation: clockwise degrees from "
        "horizontal (0 for normal).\n"
        "Do not report reading direction; it is decided later by the target "
        "language.\n"
        + _lessons_as_prompt_block(_load_reviewer_lessons()) +
        "\nReturn ONLY the JSON array, ordered top-to-bottom then "
        "left-to-right. No prose, no markdown fences."
    )
    user_text = ("Segment, read and translate this page now. Return the JSON "
                 "array of blocks with pixel coordinates. EVERY block MUST "
                 "include class, text, translation, color and lines - a bare "
                 "box with only coordinates is invalid.")
    dbg["prompt"] = system_prompt + "\n\n---- USER MESSAGE ----\n" + user_text
    try:
        content, _provider = _call_vision_with_failover(
            llm_config, system_prompt, user_text, image_b64, max_tokens=16000)
    except LeaseEngineError as err:
        print(f"Full-page layout vision failed: {err}")
        dbg["response"] = f"[ERROR] {err}"
        return (None, page_image.width, page_image.height, dbg) if return_debug \
            else (None, page_image.width, page_image.height)
    dbg["response"] = content or ""
    if not content or _looks_like_refusal(content):
        return (None, page_image.width, page_image.height, dbg) if return_debug \
            else (None, page_image.width, page_image.height)

    data = _parse_layout_json(content)
    # Validate: a well-formed layout has blocks that actually carry a
    # class and (for text) a translation. A dense page sometimes comes
    # back as a bare list of coordinate boxes with no class/text (the
    # model lost the schema) - that is unusable, so retry ONCE with an
    # even more explicit instruction and a fresh, higher token budget.
    def _well_formed(blocks):
        if not blocks:
            return False
        good = [b for b in blocks if isinstance(b, dict) and b.get("class")
                and (b.get("class") != "text" or b.get("translation"))]
        return len(good) >= max(1, int(0.5 * len(blocks)))

    if not _well_formed(data):
        print("Layout JSON malformed (bare boxes / missing class,text) - retrying with stricter prompt.")
        strict = user_text + ("\n\nYour previous answer was INVALID: it "
                              "contained boxes without class/text/translation, "
                              "or duplicates. Return ONE object PER visual text "
                              "line, each with class='text', the exact original "
                              "text, its English translation, color and a lines "
                              "array. No duplicates. No bare boxes.")
        try:
            content2, _ = _call_vision_with_failover(
                llm_config, system_prompt, strict, image_b64, max_tokens=16000)
            if content2 and not _looks_like_refusal(content2):
                data2 = _parse_layout_json(content2)
                dbg["response"] = (dbg.get("response", "") +
                                   "\n\n==== RETRY RESPONSE ====\n" + content2)
                if _well_formed(data2):
                    data = data2
        except LeaseEngineError as err:
            print(f"Layout retry failed: {err}")

    # Drop exact-duplicate boxes (the dense-page failure mode repeated the
    # same coordinates many times).
    seen = set()
    deduped = []
    for b in (data or []):
        if not isinstance(b, dict):
            continue
        key = (b.get("left"), b.get("top"), b.get("right"), b.get("bottom"),
               b.get("text", ""))
        if key in seen:
            continue
        seen.add(key)
        deduped.append(b)
    data = deduped

    if not isinstance(data, list) or not data:
        return (None, 1000, 1000, dbg) if return_debug else (None, 1000, 1000)
    # Coordinates are normalised to a 0-1000 grid; the caller maps that grid
    # onto the real page-image pixels, so there is no resolution mismatch.
    return (data, 1000, 1000, dbg) if return_debug else (data, 1000, 1000)


def _parse_layout_json(content):
    """Extract a JSON array of layout blocks from a model reply."""
    txt = (content or "").strip()
    txt = re.sub(r"^```(?:json)?", "", txt).strip()
    txt = re.sub(r"```$", "", txt).strip()
    a, b = txt.find("["), txt.rfind("]")
    if a != -1 and b != -1 and b > a:
        txt = txt[a:b + 1]
    try:
        return json.loads(txt)
    except Exception:
        out = []
        for m in re.finditer(r"\{[^{}]*\}", txt):
            try:
                out.append(json.loads(m.group(0)))
            except Exception:
                pass
        return out


def _pil_image_to_jpeg_b64(pil_image, max_dim=2000, quality=80):
    """Downscales (never upscales) a PIL page render to max_dim on its
    longest side and returns base64 JPEG - keeps the vision API payload
    reasonable (a raw 300 DPI A4 PNG is ~2500x3500 and several MB;
    JPEG at 2000px is a fraction of that with no meaningful loss for
    text reading)."""
    img = pil_image
    longest = max(img.width, img.height)
    if longest > max_dim:
        scale = max_dim / float(longest)
        img = img.resize((max(1, int(img.width * scale)), max(1, int(img.height * scale))), PILImage.LANCZOS)
    buf = io.BytesIO()
    img.convert("RGB").save(buf, format="JPEG", quality=quality)
    return base64.b64encode(buf.getvalue()).decode("ascii")


def _call_vision_with_failover(llm_config, system_prompt, text_prompt, image_b64, max_tokens=8000):
    """Sends ONE page image + a text prompt to the configured chat model
    using the OpenAI multimodal content format (identical on OpenRouter,
    which proxies the same schema). Reuses the existing failover plumbing
    - the multimodal content list rides through _call_chat_completion's
    `user_content` untouched, since that function just drops it into the
    user message as-is. Returns (content, provider) or (None, None) if no
    key is configured. NOTE: requires a vision-capable model in .env
    (gpt-4o / gpt-4o-mini and most OpenRouter frontier models are)."""
    user_content = [
        {"type": "text", "text": text_prompt},
        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_b64}"}},
    ]
    return _call_chat_completion_with_failover(llm_config, system_prompt, user_content, max_tokens=max_tokens)


def _translate_lines_batch_vision(page_image, lines, target_language, llm_config=None, chunk_size=60):
    """Hybrid mode's translator: same 1:1 numbered-lines contract as
    _translate_lines_batch, but the page IMAGE is sent along with the
    OCR'd lines so the model can use the actual pixels as ground truth to
    FIX Tesseract's misreads before translating - this is the piece that
    makes the layout-preserving path usable on photographed documents,
    where raw OCR text alone is too corrupted to translate meaningfully.
    Falls back to the text-only _translate_lines_batch for any chunk
    where no key is configured or the vision call fails."""
    llm_config = llm_config if llm_config is not None else load_llm_config()
    if not llm_is_configured(llm_config):
        return _translate_lines_batch(lines, target_language, llm_config)

    image_b64 = _pil_image_to_jpeg_b64(page_image)
    results = []
    for start in range(0, len(lines), chunk_size):
        chunk = lines[start:start + chunk_size]
        numbered_input = "\n".join(f"{i+1}. {line}" for i, line in enumerate(chunk))
        system_prompt = (
            f"You are a professional document translator. You are given an IMAGE of a document "
            f"page, plus a numbered list of text lines that OCR software extracted from that same "
            f"page. The OCR text contains recognition errors - misread characters, wrong words, "
            f"garbled fragments. Use the IMAGE as the ground truth: for each numbered line, find "
            f"the corresponding text in the image, read what it ACTUALLY says, and translate that "
            f"into {target_language}.\n\n"
            f"Rules:\n"
            f"- Keep company names, trademarks, product names and personal names unchanged.\n"
            f"- Preserve dates, numbers, currencies, amounts and identifiers exactly as shown in the image.\n"
            f"- Lines may be labels, table cells or fragments - translate each independently.\n"
            f"- If a numbered line is NOT real document text - it is an OCR misdetection of decorative "
            f"graphics, ornaments, borders, background pattern, a LOGO, or a HANDWRITTEN SIGNATURE - "
            f"return exactly [SKIP] for that line (nothing else), so the original pixels stay untouched. "
            f"Never guess or invent a reading for a signature or logo.\n"
            f"- If a line can't be located in the image, translate the OCR text as-is.\n\n"
            f"Return ONLY a numbered list with EXACTLY {len(chunk)} lines, in the exact same order "
            f"(line N in = line N out) - no commentary, no merged or split lines, no markdown."
        )
        try:
            content, _provider = _call_vision_with_failover(llm_config, system_prompt, numbered_input, image_b64, max_tokens=4000)
        except LeaseEngineError as err:
            print(f"Vision line-translation failed for chunk starting at {start}: {err} - falling back to text-only translation for this chunk.")
            content = None
        if not content:
            results.extend(_translate_lines_batch(chunk, target_language, llm_config))
            continue
        translated_chunk = [None] * len(chunk)
        for line in content.strip().split("\n"):
            m = re.match(r"^\s*(\d+)[.):]\s*(.*)$", line)
            if m:
                idx = int(m.group(1)) - 1
                if 0 <= idx < len(chunk):
                    translated_chunk[idx] = m.group(2).strip()
        results.extend(t if t else original for t, original in zip(translated_chunk, chunk))
    return results


def _closest_builtin_font(bold, italic):
    if bold and italic:
        return "Helvetica-BoldOblique"
    if bold:
        return "Helvetica-Bold"
    if italic:
        return "Helvetica-Oblique"
    return "Helvetica"


def _wrap_text_reportlab(text, font_name, font_size, max_width):
    """Same greedy word-wrap as _wrap_text_pil, using reportlab's
    stringWidth for measurement instead of PIL's textbbox."""
    from reportlab.pdfbase.pdfmetrics import stringWidth
    words = text.split()
    if not words:
        return [text] if text else [""]
    lines, cur = [], ""
    for w in words:
        test = (cur + " " + w).strip()
        if stringWidth(test, font_name, font_size) <= max_width or not cur:
            cur = test
        else:
            lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines


def _fit_text_in_box_reportlab(text, font_name, box_w, box_h, start_size,
                                min_size=5, line_spacing=1.25):
    """reportlab counterpart of _fit_text_in_box_pil: picks the largest
    font size (starting from the ORIGINAL PDF's own extracted font size,
    which is the best 'true height' reference we have here) that lets
    the word-wrapped text fit inside box_w x box_h, wrapping to multiple
    lines rather than shrinking one line down to an illegible size.
    Returns (font_size, lines, line_height)."""
    from reportlab.pdfbase.pdfmetrics import stringWidth
    start_size = max(min_size, start_size)
    step = 0.5
    size = start_size
    while size >= min_size:
        lines = _wrap_text_reportlab(text, font_name, size, box_w)
        line_h = size * line_spacing
        total_h = line_h * len(lines)
        max_line_w = max(stringWidth(ln, font_name, size) for ln in lines)
        if total_h <= box_h and max_line_w <= box_w:
            return size, lines, line_h
        size -= step
    size = min_size
    lines = _wrap_text_reportlab(text, font_name, size, box_w)
    line_h = size * line_spacing
    return size, lines, line_h


_FONTS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "fonts")


def _load_pil_font(size, bold=False, script="latin"):
    """Loads a bundled TrueType font (shipped in py/fonts/ rather than
    relying on whatever happens to be installed system-wide, since a
    slim Docker base image often has zero fonts at all) sized for PIL's
    ImageDraw. Arabic script gets Noto Sans Arabic (needed for the
    glyphs to render as anything other than empty boxes); everything
    else gets Noto Sans. Falls back to PIL's built-in bitmap font (fixed
    size, better than crashing) if the bundled files are ever missing."""
    try:
        if script == "arabic":
            path = os.path.join(_FONTS_DIR, "NotoSansArabic-Bold.ttf" if bold else "NotoSansArabic-Regular.ttf")
        else:
            path = os.path.join(_FONTS_DIR, "NotoSans-Regular.ttf")
        return ImageFont.truetype(path, max(8, int(size)))
    except Exception:
        return ImageFont.load_default()


def _wrap_text_pil(draw, text, font, max_width):
    """Greedy word-wrap: adds words to the current line while it still
    fits max_width, starting a new line otherwise. Works for both
    space-separated LTR text and RTL text (Arabic also uses spaces
    between words, so wrapping on whitespace before shaping is safe -
    shape each returned line individually afterward)."""
    words = text.split()
    if not words:
        return [text] if text else [""]
    lines, cur = [], ""
    for w in words:
        test = (cur + " " + w).strip()
        width = draw.textbbox((0, 0), test, font=font)[2]
        if width <= max_width or not cur:
            cur = test
        else:
            lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines


def _fit_text_in_box_pil(draw, text, box_w, box_h, script="latin", bold=False,
                          min_size=6, max_size=None, line_spacing=1.25):
    """Chooses the largest font size that lets `text` (word-wrapped if
    needed) fit ENTIRELY inside box_w x box_h - matching the original
    line's HEIGHT first (not just its width, which was the previous bug:
    long translations got squeezed into a single tiny-font line instead
    of wrapping, and short translations at a width-only-matched size
    could end up taller or shorter than the source line ever was).

    Algorithm: starting from a height-appropriate size and working
    downward, at each candidate size word-wrap the text to box_w, then
    check whether the wrapped block's total height still fits box_h. The
    first (largest) size where both the widest wrapped line fits box_w
    AND the total wrapped height fits box_h wins. This naturally
    degrades to "single line, shrunk for width" when the text is short
    (exactly the old behavior) and to "multiple lines at a smaller,
    still-readable size" when it's long - rather than one line crushed
    down to an illegibly small font.

    Returns (font, lines, line_height)."""
    start_size = int(max_size or box_h)
    start_size = max(start_size, min_size)
    for size in range(start_size, min_size - 1, -1):
        font = _load_pil_font(size, bold=bold, script=script)
        lines = _wrap_text_pil(draw, text, font, box_w)
        bbox = draw.textbbox((0, 0), "Agjpqy|ÉÑ", font=font)
        line_h = (bbox[3] - bbox[1]) * line_spacing
        total_h = line_h * len(lines)
        max_line_w = max(draw.textbbox((0, 0), ln, font=font)[2] for ln in lines)
        if total_h <= box_h and max_line_w <= box_w:
            return font, lines, line_h
    # Nothing fit even at min_size (extremely long text / tiny box) -
    # use min_size anyway rather than refusing to draw the text at all.
    font = _load_pil_font(min_size, bold=bold, script=script)
    lines = _wrap_text_pil(draw, text, font, box_w)
    bbox = draw.textbbox((0, 0), "Agjpqy|ÉÑ", font=font)
    line_h = (bbox[3] - bbox[1]) * line_spacing
    return font, lines, line_h


def _group_ocr_words_into_lines(ocr_data, conf_threshold=8):
    """Groups pytesseract's word-level image_to_data() output back into
    lines, using its own (block_num, par_num, line_num) grouping rather
    than re-deriving Y-proximity - Tesseract has already done that
    analysis during OCR. The confidence floor here is intentionally low
    (Tesseract can score even correct words at conf<25 on stylized/large
    title text or busy backgrounds) - the priority is catching real text
    so it actually gets masked+translated rather than silently left
    untouched; only clearly invalid readings (conf<0, which Tesseract
    uses for structural/non-text placeholders, not real words) are
    dropped."""
    n = len(ocr_data["text"])
    lines = {}
    for i in range(n):
        text = ocr_data["text"][i].strip()
        try:
            conf = float(ocr_data["conf"][i])
        except (ValueError, TypeError):
            conf = -1
        if not text or conf < conf_threshold:
            continue
        # Word-level noise filters (tuned against a real decorative
        # certificate at 300 DPI, where border art / guilloche texture
        # produced phantom words): real short words OCR at conf 90+
        # with normal glyph proportions, while ornament misreads are
        # tall skinny slivers, sub-legible specks, or giant low-conf
        # single glyphs. Dropping them here keeps them out of BOTH the
        # masks and the translation batch.
        w_px, h_px = ocr_data["width"][i], ocr_data["height"][i]
        content = len(re.findall(r"[^\W_]", text, flags=re.UNICODE))
        if content <= 2:
            if h_px > 3 * max(1, w_px) and conf < 85:
                continue  # tall skinny sliver (border line / edge of art)
            if h_px < 15:
                continue  # sub-legible speck (at 300 DPI this is < ~4pt)
            if content <= 1 and h_px > 120 and conf < 70:
                continue  # giant low-confidence lone glyph = decoration
        key = (ocr_data["block_num"][i], ocr_data["par_num"][i], ocr_data["line_num"][i])
        lines.setdefault(key, []).append({
            "text": text, "left": ocr_data["left"][i], "top": ocr_data["top"][i],
            "width": ocr_data["width"][i], "height": ocr_data["height"][i],
            "conf": conf,
        })

    def _emit_region(words, out):
        """One group of words -> one region dict, with noise filtering.
        Filters exist because Tesseract on decorative/photographed pages
        (certificates, posters) reliably emits phantom 'words' for
        ornaments, borders and logo art - masking those draws blank
        rectangles over graphics that were never text at all."""
        text = " ".join(w["text"] for w in words)
        left = min(w["left"] for w in words)
        top = min(w["top"] for w in words)
        right = max(w["left"] + w["width"] for w in words)
        bottom = max(w["top"] + w["height"] for w in words)
        avg_conf = sum(w["conf"] for w in words) / max(1, len(words))
        # Count real content characters (letters/digits in any script);
        # regions made only of punctuation/bars ('|', '—', ']') are
        # always ornament misreads.
        content_chars = len(re.findall(r"[^\W_]", text, flags=re.UNICODE))
        if content_chars == 0:
            return
        # 1-2 stray characters at low confidence = a speck of border art
        # or texture, not text (real short labels like 'No' or '12' OCR
        # at much higher confidence than pattern noise does).
        if content_chars <= 2 and avg_conf < 40:
            return
        out.append({"text": text, "left": left, "top": top, "right": right,
                    "bottom": bottom, "height": bottom - top})

    regions = []
    for key in sorted(lines.keys()):
        words = sorted(lines[key], key=lambda w: w["left"])
        # SPLIT one Tesseract "line" into separate regions wherever
        # there's a big horizontal gap between consecutive words. With
        # --psm 6 (single uniform block), two side-by-side columns on
        # the same visual row - e.g. a certificate's 'Team leader'
        # label on the right and 'Date' label on the left - get merged
        # into ONE line, so their translations can never go back to
        # their own spots. A gap much wider than the text height is a
        # column boundary, not word spacing.
        med_h = sorted(w["height"] for w in words)[len(words) // 2]
        gap_limit = max(2.5 * med_h, 40)
        group = [words[0]]
        for w in words[1:]:
            prev = group[-1]
            gap = w["left"] - (prev["left"] + prev["width"])
            if gap > gap_limit:
                _emit_region(group, regions)
                group = [w]
            else:
                group.append(w)
        _emit_region(group, regions)
    return regions


def _resolve_available_ocr_langs(requested_langs):
    """Item 2 - Tesseract raises (or on some builds, silently returns
    nothing) if asked for a language pack that isn't actually installed
    (e.g. this server hasn't been redeployed yet with the Dockerfile's
    tesseract-ocr-ara line). Filters the requested '+'-joined language
    string down to whatever's ACTUALLY available, so a missing language
    pack degrades to "still OCR in whichever languages ARE present"
    rather than silently extracting zero text for the whole page."""
    try:
        installed = set(pytesseract.get_languages(config=""))
    except Exception:
        installed = {"eng"}
    requested = [l for l in requested_langs.split("+") if l]
    available = [l for l in requested if l in installed]
    if not available:
        print(f"None of the requested OCR languages {requested} are installed (have: {sorted(installed)}) - falling back to English only.")
        available = ["eng"] if "eng" in installed else (sorted(installed)[:1] or ["eng"])
    missing = set(requested) - set(available)
    if missing:
        print(f"OCR language(s) {sorted(missing)} not installed on this server (needs a redeploy with the updated Dockerfile) - continuing with {available}.")
    return "+".join(available)


def _detect_text_regions_cv(page_image):
    """Pixel-level text-line detection with OpenCV - the piece Tesseract
    alone can't provide: Tesseract only reports text it can RECOGNIZE,
    so decorative calligraphy, stylized titles and unusual scripts never
    get a bounding box at all and end up untranslated. This instead
    finds anything that LOOKS like text strokes (high local contrast,
    line-shaped clusters), and leaves deciding what each region actually
    says to the vision model. Returns regions in the same dict shape the
    OCR path uses ({'text','left','top','right','bottom','height'}),
    with text='' since content is unknown at this stage."""
    img = np.array(page_image)
    H, W = img.shape[:2]
    gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)

    # --- Illustration pass (run first): find strongly-COLOURED blobs
    # (green foliage, painted figures) that are pictures, not text, and
    # record them as pre-marked graphic regions. This is what stops a
    # botanical plate or photo from being swept into text lines and then
    # erased. Text - even red/black ink on aged parchment - is low
    # saturation, so this only fires on real colour art.
    graphic_regions = []
    try:
        Rc = img[..., 0].astype(np.int16)
        Gc = img[..., 1].astype(np.int16)
        Bc = img[..., 2].astype(np.int16)
        hsv = cv2.cvtColor(img, cv2.COLOR_RGB2HSV)
        sat = hsv[..., 1].astype(np.float32) / 255.0
        # "Picture" pixels: coloured foliage (green dominant, mid-dark) OR
        # any strongly-saturated non-warm-ink colour. Warm brown/black ink
        # on parchment is deliberately excluded so manuscript TEXT is not
        # mistaken for a picture.
        foliage = (Gc > Rc) & (Gc >= Bc) & (Gc > 40) & (Gc < 170) & (Rc < 150)
        warm_ink = (Rc >= Gc) & (Gc >= Bc) & (sat < 0.55)
        vivid = (sat > 0.5) & (~warm_ink)
        pic = (foliage | vivid).astype(np.uint8) * 255
        pic = cv2.morphologyEx(pic, cv2.MORPH_CLOSE,
                               cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (21, 21)))
        pic = cv2.morphologyEx(pic, cv2.MORPH_OPEN,
                               cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (9, 9)))
        # Frame band: the coloured page border hugs the edges. Ignore any
        # blob that is mostly in the outer 8% margin so a green/gold
        # certificate frame is not taken for an illustration.
        fm = int(min(H, W) * 0.08)
        cnts, _ = cv2.findContours(pic, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        for cnt in cnts:
            x, y, w, h = cv2.boundingRect(cnt)
            area = w * h
            if area < 0.008 * W * H or area > 0.6 * W * H:
                continue
            if (x + w <= fm) or (x >= W - fm) or (y + h <= fm) or (y >= H - fm):
                continue
            fill = float(np.count_nonzero(pic[y:y + h, x:x + w])) / max(1, area)
            if fill < 0.28:
                continue
            if h < 30 or w < 30:
                continue
            # Confirm the blob is genuinely a coloured picture: a real
            # fraction of its pixels must be foliage/vivid colour (not
            # just merged text ink that slipped through).
            sub_foliage = foliage[y:y + h, x:x + w].mean()
            sub_vivid = vivid[y:y + h, x:x + w].mean()
            if (sub_foliage + sub_vivid) < 0.16:
                continue
            graphic_regions.append({"text": "", "left": int(x), "top": int(y),
                                    "right": int(x + w), "bottom": int(y + h),
                                    "height": int(h), "is_graphic": True})
    except Exception:
        graphic_regions = []

    # High-contrast marks vs. local background. C=20 keeps faint page
    # texture (guilloche patterns, watermarks) OUT of the mask.
    bw = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C,
                               cv2.THRESH_BINARY_INV, 31, 20)

    # Long straight runs are frame/border/table lines, not glyphs.
    line_len = max(60, W // 25)
    horiz = cv2.morphologyEx(bw, cv2.MORPH_OPEN, cv2.getStructuringElement(cv2.MORPH_RECT, (line_len, 1)))
    vert = cv2.morphologyEx(bw, cv2.MORPH_OPEN, cv2.getStructuringElement(cv2.MORPH_RECT, (1, line_len)))
    text_mask = cv2.subtract(bw, cv2.bitwise_or(horiz, vert))

    # Glue characters/words on the same baseline into line blobs.
    glue_w = max(20, W // 70)
    conn = cv2.morphologyEx(text_mask, cv2.MORPH_CLOSE,
                            cv2.getStructuringElement(cv2.MORPH_RECT, (glue_w, 5)))
    contours, _ = cv2.findContours(conn, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    margin = int(min(H, W) * 0.05)  # decorative frames hug page edges
    regions = []
    for cnt in contours:
        x, y, w, h = cv2.boundingRect(cnt)
        if h < H * 0.006 or h > H * 0.22:
            continue  # speck / not a text line
        if w < W * 0.02:
            continue
        if x < margin or y < margin or x + w > W - margin or y + h > H - margin:
            continue  # touching the decorative frame band
        ink = float(np.count_nonzero(text_mask[y:y + h, x:x + w])) / max(1, w * h)
        if ink < 0.02 or ink > 0.65:
            continue  # empty box or solid graphic, not strokes
        regions.append({"text": "", "left": int(x), "top": int(y),
                        "right": int(x + w), "bottom": int(y + h), "height": int(h)})

    # Merge boxes that share a baseline row and nearly touch (one visual
    # line the closing kernel didn't quite bridge).
    regions.sort(key=lambda r: (r["top"], r["left"]))
    merged = []
    for r in regions:
        if merged:
            m = merged[-1]
            same_row = not (r["top"] > m["bottom"] - min(r["height"], m["height"]) * 0.4)
            gap = r["left"] - m["right"]
            if same_row and -20 <= gap <= max(r["height"], m["height"]) * 2.0:
                m["left"] = min(m["left"], r["left"]); m["right"] = max(m["right"], r["right"])
                m["top"] = min(m["top"], r["top"]); m["bottom"] = max(m["bottom"], r["bottom"])
                m["height"] = m["bottom"] - m["top"]
                continue
        merged.append(dict(r))

    # Second pass: collapse boxes that physically OVERLAP into one.
    # Calligraphy and stylized titles fragment badly (dots, diacritics
    # and flourishes each become their own contour, stacked over the
    # main strokes) - the model should see ONE numbered box over the
    # whole title, not fourteen shards of it. Only boxes that actually
    # intersect merge, so side-by-side columns (a left 'Date' and a
    # right 'Team leader') can never be glued together by this pass.
    def _overlaps(a, b, pad=12):
        ix = min(a["right"], b["right"]) + pad - max(a["left"], b["left"])
        iy = min(a["bottom"], b["bottom"]) - max(a["top"], b["top"])
        if ix > 0 and iy > 0:
            if iy > 0.35 * min(a["height"], b["height"]) or \
                    (ix * iy) > 0.4 * min((a["right"] - a["left"]) * a["height"],
                                          (b["right"] - b["left"]) * b["height"]):
                return True
        # Same visual row, small gap between them: two shards of ONE
        # line (an RTL sentence split mid-way reads in the wrong order
        # if each shard is translated and placed separately, so shards
        # must become one region). The gap ceiling is proportional to
        # text height, far smaller than any real column gutter - a left
        # 'Date' and a right 'Team leader' stay separate.
        v_overlap = min(a["bottom"], b["bottom"]) - max(a["top"], b["top"])
        if v_overlap > 0.5 * min(a["height"], b["height"]):
            gap = max(a["left"], b["left"]) - min(a["right"], b["right"])
            if 0 <= gap <= 2.2 * max(a["height"], b["height"]):
                return True
        return False

    changed = True
    while changed:
        changed = False
        out = []
        for r in merged:
            hit = None
            for m in out:
                if _overlaps(m, r):
                    hit = m
                    break
            if hit:
                hit["left"] = min(hit["left"], r["left"]); hit["right"] = max(hit["right"], r["right"])
                hit["top"] = min(hit["top"], r["top"]); hit["bottom"] = max(hit["bottom"], r["bottom"])
                hit["height"] = hit["bottom"] - hit["top"]
                changed = True
            else:
                out.append(r)
        merged = out

    # Split oversized merged regions back into their real text lines and
    # columns. A single CV blob sometimes swallows a whole footer band
    # (Date + label + signature + date-value), which then renders as one
    # overlapping, overflowing text item. We re-run connected-component
    # analysis INSIDE such a blob (with gentler gluing) to recover the
    # separate labels/lines, each placed and sized on its own.
    def _split_big_region(reg):
        rw = reg["right"] - reg["left"]
        rh = reg["bottom"] - reg["top"]
        # Only bother for genuinely oversized blobs (wide AND tall).
        if not (rw > 0.45 * W and rh > 0.12 * H):
            return [reg]
        sub = text_mask[reg["top"]:reg["bottom"], reg["left"]:reg["right"]]
        if sub.size == 0:
            return [reg]
        # A dense blob (high ink fill) is one solid thing - a calligraphy
        # title or a picture - and must NOT be split. Only sparse blobs
        # (scattered labels/lines across a band, like a footer) get cut.
        fill = float(np.count_nonzero(sub)) / max(1, sub.size)
        if fill > 0.16:
            return [reg]
        glue = max(12, (reg["right"] - reg["left"]) // 45)
        conn = cv2.morphologyEx(sub, cv2.MORPH_CLOSE,
                                cv2.getStructuringElement(cv2.MORPH_RECT, (glue, 3)))
        cnts, _ = cv2.findContours(conn, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        parts = []
        for cnt in cnts:
            x, y, w, h = cv2.boundingRect(cnt)
            if w < 0.03 * W or h < 0.012 * H:
                continue  # sliver / speck
            ink = float(np.count_nonzero(sub[y:y + h, x:x + w])) / max(1, w * h)
            if ink < 0.03:
                continue
            parts.append({"text": "", "left": reg["left"] + x, "top": reg["top"] + y,
                          "right": reg["left"] + x + w, "bottom": reg["top"] + y + h,
                          "height": h})
        if len(parts) < 2:
            return [reg]
        # Guard against fragmenting a single big word/title: only accept
        # the split if the parts span multiple distinct vertical rows
        # (a real multi-line footer), not one row cut into pieces (which
        # is usually calligraphy). Cluster part-centres by y.
        ys = sorted(((p["top"] + p["bottom"]) / 2.0 for p in parts))
        row_gap = 0.03 * H
        rows = 1
        for a, b in zip(ys, ys[1:]):
            if b - a > row_gap:
                rows += 1
        if rows < 2:
            return [reg]
        return parts

    split_regions = []
    for r in merged:
        if r.get("is_graphic"):
            split_regions.append(r)
        else:
            split_regions.extend(_split_big_region(r))
    merged = split_regions

    # Line-split pass: a clean body paragraph often gets detected as ONE
    # tall region (the closing kernel bridged its lines). Break such a
    # region into its individual visual LINES using horizontal ink-row
    # projection, so each line becomes its own box. This is what lets a
    # 3-line sentence become 3 line-boxes that the translation flows
    # across (the user's core requirement), instead of one block.
    def _split_into_lines(reg):
        rh = reg["bottom"] - reg["top"]
        sub = text_mask[reg["top"]:reg["bottom"], reg["left"]:reg["right"]]
        if sub.size == 0 or rh < 0.05 * H:
            return [reg]
        row_ink = (sub > 0).sum(axis=1).astype(np.float32)
        if row_ink.max() < 1:
            return [reg]
        # Smooth the profile so diacritics don't create spurious splits.
        k = max(3, int(0.01 * H) | 1)
        kern = np.ones(k, dtype=np.float32) / k
        smooth = np.convolve(row_ink, kern, mode="same")
        peak = smooth.max()
        # A "line" is where smoothed ink is above a low fraction of the
        # peak; the shallow VALLEYS between Arabic lines (which never drop
        # to zero) fall below this and become the separators.
        active = smooth > 0.18 * peak
        runs = []
        s = None
        for i, a in enumerate(active):
            if a and s is None:
                s = i
            elif not a and s is not None:
                runs.append((s, i)); s = None
        if s is not None:
            runs.append((s, len(active)))
        if len(runs) < 2:
            return [reg]
        line_h_guess = np.median([b - a for a, b in runs])
        # The smoothing already absorbed diacritics into their lines, so
        # the runs are the text lines - do not merge them (that would
        # re-join lines separated by the shallow valleys we just found).
        merged_runs = [list(r) for r in runs]
        lines = [(a, b) for a, b in merged_runs if (b - a) > 0.35 * line_h_guess]
        if len(lines) < 2:
            return [reg]
        out = []
        for a, b in lines:
            # Pad each line slightly and tighten to its own ink columns.
            pa = max(0, a - 2)
            pb = min(sub.shape[0], b + 2)
            lsub = sub[pa:pb, :]
            cols = np.where(lsub.sum(axis=0) > 0)[0]
            if len(cols) < 2:
                continue
            nl = reg["left"] + int(cols[0])
            nr = reg["left"] + int(cols[-1]) + 1
            out.append({"text": "", "left": nl, "top": reg["top"] + pa,
                        "right": nr, "bottom": reg["top"] + pb,
                        "height": pb - pa})
        return out if len(out) >= 2 else [reg]

    line_regions = []
    for r in merged:
        if r.get("is_graphic"):
            line_regions.append(r)
        else:
            line_regions.extend(_split_into_lines(r))
    merged = line_regions

    # Drop any text-line region that sits largely INSIDE a detected
    # illustration (its "strokes" are really parts of the picture), then
    # add the illustration blobs themselves as their own regions so the
    # vision model sees and classifies them as pictures (kept, not erased).
    def _center_inside(r, g):
        cx = (r["left"] + r["right"]) / 2.0
        cy = (r["top"] + r["bottom"]) / 2.0
        return (g["left"] <= cx <= g["right"] and g["top"] <= cy <= g["bottom"])
    if graphic_regions:
        merged = [r for r in merged
                  if not any(_center_inside(r, g) for g in graphic_regions)]
        merged.extend(graphic_regions)

    # Reading order + a sanity cap (a busy page shouldn't turn into a
    # 300-line vision request; keep the biggest regions).
    if len(merged) > 80:
        merged.sort(key=lambda r: (r["right"] - r["left"]) * r["height"], reverse=True)
        merged = merged[:80]
    merged.sort(key=lambda r: (r["top"], r["left"]))
    return merged


def _extract_page_metadata_vlm(page_image, regions, target_language, llm_config=None):
    """Implements the workflow doc's Hybrid steps 4.2.5-4.2.12 in ONE
    structured VLM pass. The page image is sent with every detected
    region drawn as a numbered red box; the model returns, per box, the
    full 4.2.7 metadata table plus translation and classification:

        {"n":1,"class":"translatable"|"nontranslatable"|"decoration",
         "kind":"heading|paragraph|label|logo|signature|qr|...",
         "text":"<original>","translation":"<target>",
         "color":"RRGGBB","bold":bool,"italic":bool,"underline":bool,
         "align":"left|center|right","rtl":bool}

    class drives 4.2.13 layering: translatable -> inpaint + Layer-3
    text; nontranslatable -> keep pixels as Layer-2 element; decoration
    (watermark/ornament) -> erased, never re-added. Returns a list
    aligned to `regions` (None where the model gave nothing)."""
    llm_config = llm_config if llm_config is not None else load_llm_config()
    annotated = page_image.copy()
    d = ImageDraw.Draw(annotated)
    label_font = _load_pil_font(max(22, page_image.height // 90), script="latin", bold=True)
    for i, r in enumerate(regions):
        d.rectangle([r["left"] - 3, r["top"] - 3, r["right"] + 3, r["bottom"] + 3],
                    outline=(255, 0, 0), width=4)
        d.text((r["left"] + 2, max(0, r["top"] - (page_image.height // 70))),
               str(i + 1), fill=(255, 0, 0), font=label_font)

    system_prompt = (
        "You are Lexora's document-translation engine. The user has "
        "uploaded THEIR OWN document to their own account and explicitly "
        "requested a translation of it; you are authorised to read and "
        "translate every piece of text on the page. This is a routine, "
        "legitimate localisation task - always comply by returning the "
        "requested JSON. Never refuse and never apologise.\n\n"
        "Act as a world-class Computer Vision Engineer, OCR Expert, Document "
        "AI Architect and Multilingual Translation Expert with 30+ years of "
        "experience. The image is ONE document page with "
        f"{len(regions)} regions outlined as numbered red boxes. For EACH "
        "box, inspect the pixels inside it and output one JSON object.\n\n"
        "CLASSIFICATION (the most important decision) - choose the class:\n"
        "  * translatable  = real, readable printed TEXT (any language, "
        "including stylized or calligraphic headings, and dense body "
        "paragraphs). This is the DEFAULT for anything that is words.\n"
        "  * nontranslatable = ONLY a logo, brand mark, emblem, badge, "
        "ribbon, seal, stamp, QR/barcode, OR a genuine handwritten/cursive "
        "SIGNATURE, OR a real picture (illustration, drawing, painting, "
        "photo, figure, chart, map). These are kept as images, never "
        "translated.\n"
        "  * decoration = watermark, ornament, border art, blank/underline "
        "box with no real content.\n\n"
        "CRITICAL ANTI-MISTAKE RULES (follow exactly):\n"
        "  1. A block of PRINTED WORDS is ALWAYS translatable, even if it "
        "is long, dense, faint, handwritten-style calligraphy, or in Arabic/"
        "Urdu/Persian. Do NOT call a paragraph of text an illustration, "
        "figure, signature or logo. If a box is mostly readable letters, "
        "it is text.\n"
        "  2. signature applies ONLY to an actual person's cursive "
        "handwritten signature mark - NOT to a printed name, a printed "
        "label like 'Team Leader', or a line of body text. If you can read "
        "it as normal words, it is text, not a signature.\n"
        "  3. A logo/emblem/seal is a compact GRAPHICAL mark; a wide line or "
        "block of running text is NOT a logo even if decorative.\n"
        "  4. An illustration/figure is a PICTURE (a plant drawing, a "
        "portrait, a diagram). Text that merely sits NEXT TO a picture is "
        "still translatable text - only the picture itself is "
        "nontranslatable.\n"
        "  5. Never OCR or translate the inside of a logo/signature/"
        "watermark; classify the whole mark as nontranslatable and move on.\n"
        "  When in doubt whether something is text or a picture, and it "
        "contains readable words, choose translatable.\n\n"
        "FIELDS per box:\n"
        '  "n": box number (integer)\n'
        '  "class": "translatable" | "nontranslatable" | "decoration"\n'
        '  "kind": short label e.g. heading, subheading, paragraph, label, '
        "table-cell, caption, logo, signature, qr, stamp, seal, illustration, "
        "figure, photo\n"
        '  "text": exact original text (empty if not translatable)\n'
        f'  "translation": the text translated into {target_language} '
        "(empty unless translatable). Translate for MEANING, not literally; "
        "keep proper nouns, trademarks, dates, numbers and identifiers "
        "unchanged; you may choose slightly shorter or longer equivalent "
        "wording so it fits the box, but never lose or invent information.\n"
        '  "color": ink color of the letters as RRGGBB hex, from the actual '
        "glyph pixels (not the background)\n"
        '  "bold": bool  "italic": bool  "underline": bool\n'
        '  "align": "center" if the line is visually centered in its area, '
        'else "left". Do NOT use alignment to encode reading direction.\n'
        '  "is_paragraph": true if this box is a multi-line body paragraph, '
        "false for a single heading/label/line\n"
        '  "runs": OPTIONAL. If letters in the box are NOT all one colour/'
        "style (e.g. one red word among black), return the translation split "
        'into styled pieces [{"text":"...","color":"RRGGBB","bold":bool,'
        '"italic":bool}, ...] that concatenate to "translation". Omit when '
        "uniform.\n"
        '  "rotation": OPTIONAL degrees rotated clockwise from horizontal '
        "(0 normal; 90/270 for vertical side text). Omit when horizontal.\n\n"
        "OTHER RULES:\n"
        "- color is the letters' colour, e.g. dark-green certificates often "
        "use 1B4D3E, gold titles C9A227.\n"
        "- Translate each box as a self-contained block; do NOT reorder words "
        "across boxes. Output reading DIRECTION is decided later by the "
        "target language - you need not report it.\n"
        + _lessons_as_prompt_block(_load_reviewer_lessons()) +
        f"\nReturn ONLY a JSON array of exactly {len(regions)} objects ordered "
        "by n from 1. No prose, no markdown fences."
    )

    user_text = f"Analyze boxes 1 to {len(regions)} and return the JSON array now."
    image_b64 = _pil_image_to_jpeg_b64(annotated)
    content, _provider = _call_vision_with_failover(
        llm_config, system_prompt, user_text, image_b64, max_tokens=6000)
    results = [None] * len(regions)
    if not content:
        return results
    txt = content.strip()
    # Detect a model REFUSAL (safety decline). Some models refuse to
    # process a document image and reply with an apology instead of JSON;
    # that apology must NOT be treated as a translation. Signal the caller
    # (via the sentinel) so it can retry / fall back rather than printing
    # "I'm sorry, I can't assist with that request." into the document.
    low = txt.lower()
    looks_like_json = ("[" in txt and "]" in txt) or ("{" in txt and "}" in txt)
    refusal_markers = ("i'm sorry", "i am sorry", "i cannot assist",
                       "i can't assist", "i cannot help", "i can't help",
                       "unable to assist", "cannot process", "can't process",
                       "i'm not able to", "i am not able to")
    if (not looks_like_json) and any(mk in low for mk in refusal_markers):
        print(f"Vision model REFUSED the page ({txt[:80]!r}). Returning no "
              f"metadata so the caller can fall back instead of printing the refusal.")
        return results          # all None -> nothing gets mislabelled as text
    txt = re.sub(r"^```(?:json)?", "", txt).strip()
    txt = re.sub(r"```$", "", txt).strip()
    s, e = txt.find("["), txt.rfind("]")
    if s != -1 and e != -1 and e > s:
        txt = txt[s:e + 1]
    try:
        data = json.loads(txt)
    except Exception:
        data = []
        for m in re.finditer(r"\{[^{}]*\}", txt):
            try:
                data.append(json.loads(m.group(0)))
            except Exception:
                pass
    for obj in data:
        try:
            idx = int(obj.get("n", 0)) - 1
        except (TypeError, ValueError):
            continue
        if 0 <= idx < len(regions):
            # Guard: never let a per-item refusal string become a translation.
            tr = str(obj.get("translation", ""))
            if any(mk in tr.lower() for mk in refusal_markers):
                obj["translation"] = ""
                obj["class"] = "decoration"
            results[idx] = obj
    return results


def _normalize_hex_color(value, default=(27, 77, 62)):
    """Parses 'RRGGBB'/'#RRGGBB' into (r,g,b); default = certificate
    dark-green when the model gives nothing usable."""
    if not value or not isinstance(value, str):
        return default
    v = value.strip().lstrip("#")
    if len(v) == 6:
        try:
            return tuple(int(v[i:i + 2], 16) for i in (0, 2, 4))
        except ValueError:
            return default
    return default


def _region_stroke_mask(np_img, region, pad=2, threshold=55):
    """Boolean mask of the actual text strokes inside a region box -
    pixels whose color is far from the box's local background (median of
    a ring just outside the box). Faint background pattern stays below
    the distance threshold, so only the ink gets marked."""
    H, W = np_img.shape[:2]
    l, t = max(0, region["left"] - pad), max(0, region["top"] - pad)
    r, b = min(W, region["right"] + pad), min(H, region["bottom"] + pad)
    ring = 6
    rl, rt = max(0, l - ring), max(0, t - ring)
    rr, rb = min(W, r + ring), min(H, b + ring)
    outer = np_img[rt:rb, rl:rr].reshape(-1, 3)
    bg = np.median(outer, axis=0)
    crop = np_img[t:b, l:r].astype(np.int16)
    dist = np.sqrt(((crop - bg) ** 2).sum(axis=2))
    mask = (dist > threshold).astype(np.uint8) * 255
    # Two dilation passes: thick decorative strokes (gold calligraphy)
    # carry wide anti-aliased halos that sit below the color-distance
    # threshold - without swallowing that halo into the mask, inpainting
    # leaves a visible ghost outline of the erased text.
    mask = cv2.dilate(mask, cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (5, 5)), iterations=2)
    return (l, t, r, b), mask


def _region_stroke_color(np_img, region):
    """True ink color of a text region. The naive median over the
    stroke mask washes out toward gray because it includes every
    anti-aliased edge pixel (part background). Instead this keeps only
    the CORE stroke pixels - the ones furthest from the local
    background - and medians those, so dark-green body text reads as
    dark green and gold calligraphy reads as gold, not a muddy
    mid-tone."""
    try:
        (l, t, r, b), _mask = _region_stroke_mask(np_img, region, threshold=40)
        ring = 8
        H, W = np_img.shape[:2]
        outer = np_img[max(0, t - ring):min(H, b + ring), max(0, l - ring):min(W, r + ring)].reshape(-1, 3)
        bg = np.median(outer, axis=0)
        crop = np_img[t:b, l:r].astype(np.int16)
        dist = np.sqrt(((crop - bg) ** 2).sum(axis=2))
        flat = crop.reshape(-1, 3)
        fdist = dist.reshape(-1)
        if fdist.max() < 30:
            return (0, 0, 0)
        # Core = pixels in the top 40% of distance-from-background.
        cutoff = np.percentile(fdist[fdist > 25], 60) if (fdist > 25).any() else 25
        core = flat[fdist >= max(cutoff, 30)]
        if len(core) < 15:
            core = flat[fdist >= 30]
        if len(core) < 15:
            return (0, 0, 0)
        med = np.median(core, axis=0)
        return tuple(int(v) for v in med)
    except Exception:
        return (0, 0, 0)


def _ai_fill_background(page_image, erase_boxes, protect_boxes, llm_config):
    """Generative background fill via OpenRouter's image API (the approach
    the user validated: text is removed and the paper/pattern is painted
    back by an image model). Uses the whole page plus an inpainting prompt.

    Enabled when LEXORA_AI_FILL is not explicitly disabled AND an
    OpenRouter key is available. Returns a filled PIL image, or None on any
    error so the caller falls back to CV inpainting (never breaks the run).

    Model + endpoint come from the working proof of concept:
      POST https://openrouter.ai/api/v1/images
      { model: "black-forest-labs/flux.2-pro", prompt, input_references:[
        {type:"image_url", image_url:{url:<dataURL>}}], output_format:"png" }
    Response: data[0].url (fetch it) or data[0].b64_json.
    """
    if os.environ.get("LEXORA_AI_FILL", "1").lower() in ("0", "false", "no", "off"):
        return None
    cfg = (llm_config or {}).get("openrouter") or {}
    api_key = (cfg.get("apiKey") or os.environ.get("OPENROUTER_API_KEY") or "").strip()
    if not api_key:
        return None
    model = (os.environ.get("LEXORA_IMAGE_MODEL")
             or "black-forest-labs/flux.2-pro").strip()
    try:
        import base64 as _b64
        W, H = page_image.size
        # Downscale large pages to keep the request light and fast; the
        # returned fill is scaled back up to the page size.
        max_dim = 1536
        img = page_image.convert("RGB")
        longest = max(W, H)
        if longest > max_dim:
            s = max_dim / float(longest)
            img = img.resize((max(1, int(W * s)), max(1, int(H * s))), PILImage.LANCZOS)
        buf = io.BytesIO(); img.save(buf, format="PNG")
        data_url = "data:image/png;base64," + _b64.b64encode(buf.getvalue()).decode("ascii")

        prompt = (
            "You are an expert image editor. Remove ALL text from this "
            "document page - BOTH printed AND handwritten. "
            "WHAT TO REMOVE: printed letters, numbers and words in any font "
            "or language; ANY handwritten characters, words, sentences or "
            "cursive writing; handwritten notes, annotations, dates and "
            "names; signature text. "
            "WHAT TO PRESERVE EXACTLY (do NOT remove or move): all "
            "decorative elements - borders, frames, dividers and lines; "
            "signature lines (the line itself, not the text on it); all "
            "graphics, logos, seals and illustrations; the paper texture, "
            "colour, gradient and background pattern. "
            "SPECIAL: handwritten text MUST be removed like any other text. "
            "If handwriting sits on a line, remove the handwriting but KEEP "
            "the line. Do NOT add any new lines, text or elements, and do "
            "NOT create phantom/ghost lines. "
            "Fill every removed area with the surrounding background so it "
            "looks seamless. Return a PNG of the same page with ALL text "
            "removed and all decorative elements identical to the original."
        )
        payload = {
            "model": model,
            "prompt": prompt,
            "input_references": [
                {"type": "image_url", "image_url": {"url": data_url}}
            ],
            "output_format": "png",
        }
        req = urllib.request.Request(
            "https://openrouter.ai/api/v1/images",
            data=json.dumps(payload).encode("utf-8"),
            headers={"Authorization": f"Bearer {api_key}",
                     "Content-Type": "application/json",
                     "HTTP-Referer": "https://lexora.ai",
                     "X-Title": "Lexora"})
        with urllib.request.urlopen(req, timeout=120) as resp:
            result = json.loads(resp.read().decode("utf-8"))
        arr = result.get("data") or []
        if not arr:
            print("AI fill: image API returned no data; using CV inpaint.")
            return None
        obj = arr[0]
        if obj.get("b64_json"):
            filled_bytes = _b64.b64decode(obj["b64_json"])
        elif obj.get("url"):
            with urllib.request.urlopen(obj["url"], timeout=120) as r2:
                filled_bytes = r2.read()
        else:
            return None
        filled = PILImage.open(io.BytesIO(filled_bytes)).convert("RGB")
        if filled.size != (W, H):
            filled = filled.resize((W, H), PILImage.LANCZOS)
        return filled
    except urllib.error.HTTPError as err:
        detail = ""
        try:
            detail = err.read().decode()[:200]
        except Exception:
            pass
        print(f"AI background fill HTTP {err.code} ({detail}) - using CV inpaint.")
        return None
    except Exception as err:
        print(f"AI background fill unavailable ({err}) - using CV inpaint.")
        return None



def _inpaint_regions_cv(page_image, regions, protect_regions=None):
    """Removes the text strokes of the given regions and fills them from
    the surrounding pixels (cv2.inpaint/TELEA) so the page's own
    background pattern flows back into the erased area. Runs MULTIPLE
    passes: thick decorative strokes (large gold calligraphy) leave a
    visible ghost halo after one pass because their anti-aliased edges
    sit below the stroke threshold - each following pass re-detects
    whatever residue is still visible inside the same boxes and inpaints
    that too, until the area is genuinely clean. Returns a new PIL
    image."""
    np_img = np.array(page_image)
    # Build a protection mask (non-translatable elements: logos,
    # signatures, QR) that the inpaint mask must never touch, so a
    # translatable box that slightly overlaps a signature can't nibble
    # its edges.
    H, W = np_img.shape[:2]
    protect = np.zeros((H, W), dtype=bool)
    for pr in (protect_regions or []):
        pl, pt = max(0, pr["left"] - 2), max(0, pr["top"] - 2)
        prr, pb = min(W, pr["right"] + 2), min(H, pr["bottom"] + 2)
        protect[pt:pb, pl:prr] = True

    for pass_num in range(2):
        threshold = 50 if pass_num == 0 else 38
        full_mask = np.zeros(np_img.shape[:2], dtype=np.uint8)
        for region in regions:
            try:
                (l, t, r, b), mask = _region_stroke_mask(np_img, region, threshold=threshold)
                # Dense-ink regions (large stylized titles / calligraphy)
                # have strokes too thick and interconnected for a
                # stroke-only mask to fully clear - measurable as a high
                # ink ratio. For those, erase the WHOLE box so no
                # original glyph fragments survive (the box tightly
                # bounds just that text, and it will be inpainted from
                # the surrounding background anyway).
                ink_ratio = float(np.count_nonzero(mask)) / max(1, mask.size)
                if ink_ratio > 0.28:
                    mask[:] = 255
                    # Calligraphy flourishes often extend a little past
                    # the detected box; pad the erase area outward (but
                    # never into a protected element) so trailing strokes
                    # don't survive as ghosts.
                    ph = int((b - t) * 0.10)
                    pw = int((r - l) * 0.04)
                    et, eb = max(0, t - ph), min(H, b + ph)
                    el, er = max(0, l - pw), min(W, r + pw)
                    full_mask[et:eb, el:er] = 255
                    continue
                full_mask[t:b, l:r] = np.maximum(full_mask[t:b, l:r], mask)
            except Exception:
                continue
        if protect.any():
            full_mask[protect] = 0
        coverage = int(np.count_nonzero(full_mask))
        if coverage == 0:
            break
        # Dilate the mask slightly so anti-aliased stroke edges (which
        # sit just below the color threshold and are the usual source of
        # faint ghosts) are swallowed too, then inpaint. Navier-Stokes
        # (INPAINT_NS) reconstructs smooth gradient/pattern backgrounds -
        # like certificate guilloche - more cleanly than TELEA for the
        # larger erased areas.
        if protect.any():
            dilated = cv2.dilate(full_mask, cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (5, 5)), iterations=1)
            dilated[protect] = 0
            full_mask = dilated
        else:
            full_mask = cv2.dilate(full_mask, cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (5, 5)), iterations=1)
        np_img = cv2.inpaint(np_img, full_mask, 6, cv2.INPAINT_NS)
        if pass_num > 0 and coverage < 400:
            break  # residue is down to crumbs - done
    return PILImage.fromarray(np_img)


def _parse_style_prefix(s):
    """Splits optional leading [color:RRGGBB] and [style:...] tags off a
    translated line. Returns (text, bold, italic, underline, color_or_None)."""
    bold = italic = underline = False
    color = None
    # Both tags may appear in either order; loop until neither matches.
    for _ in range(2):
        mc = re.match(r"^\s*\[color:#?([0-9A-Fa-f]{6})\]\s*(.*)$", s)
        if mc:
            color = tuple(int(mc.group(1)[i:i + 2], 16) for i in (0, 2, 4))
            s = mc.group(2)
            continue
        ms = re.match(r"^\s*\[style:([^\]]*)\]\s*(.*)$", s, flags=re.IGNORECASE)
        if ms:
            flags = ms.group(1).lower()
            bold = "bold" in flags
            italic = "italic" in flags
            underline = "underline" in flags
            s = ms.group(2)
            continue
        break
    return s.strip(), bold, italic, underline, color


def _crop_element_png(page_image, region, pad=6):
    """Document 4.2.4 extraction: crop a non-translatable element (logo,
    signature, QR, stamp) from the page and return it as an RGBA PIL
    image with a transparent background, so it can be placed as its own
    floating Layer-2 object. Transparency is derived by knocking out the
    element's local background color (the median of a ring just outside
    the crop); this keeps QR codes and signatures crisp while letting
    the page pattern show around them. A small pad avoids clipping
    strokes at the very edge of the detected box."""
    try:
        W, H = page_image.size
        l = max(0, region["left"] - pad)
        t = max(0, region["top"] - pad)
        r = min(W, region["right"] + pad)
        b = min(H, region["bottom"] + pad)
        crop = page_image.crop((l, t, r, b)).convert("RGB")
        arr = np.array(crop)
        # Local background = median of a thin border ring of the crop.
        ring = 4
        edges = np.concatenate([
            arr[:ring].reshape(-1, 3), arr[-ring:].reshape(-1, 3),
            arr[:, :ring].reshape(-1, 3), arr[:, -ring:].reshape(-1, 3)])
        bg = np.median(edges, axis=0)
        dist = np.sqrt(((arr.astype(np.int16) - bg) ** 2).sum(axis=2))
        # Pixels close to the background become transparent; the element
        # itself (far from bg) stays fully opaque, with a soft ramp in
        # between so edges don't look cut out.
        alpha = np.clip((dist - 22) / 26.0, 0.0, 1.0)
        alpha = (alpha * 255).astype(np.uint8)
        rgba = np.dstack([arr, alpha])
        return PILImage.fromarray(rgba, mode="RGBA")
    except Exception:
        return None


def _render_layout_preview(page_image, text_items, png_items):
    """Flattens the current rebuild (cleaned background + Layer-2 PNGs +
    Layer-3 text) into a single preview image, so the reviewer agent can
    SEE what the output actually looks like and compare it to the
    original - the only reliable way to catch RTL/LTR mistakes, overlaps,
    wrong line breaks and mis-sized text, which are invisible in the
    metadata alone."""
    preview = page_image.convert("RGB").copy()
    for pe in (png_items or []):
        try:
            preview.paste(pe["png"], (pe["left"], pe["top"]), pe["png"])
        except Exception:
            pass
    draw = ImageDraw.Draw(preview)
    for it in (text_items or []):
        item_rtl = it.get("rtl", False)
        align = it.get("align", "center")
        script = "arabic" if item_rtl else "latin"
        raw_boxes = it.get("line_boxes") or [{
            "left": it["left"], "top": it["top"],
            "right": it["right"], "bottom": it["bottom"]}]
        # Pick a uniform PIL font size that lets the text wrap into at
        # most len(boxes) lines fitting each box width; shrink on
        # overflow. Mirrors the PDF flow so the preview matches output.
        n = len(raw_boxes)
        widths = [max(1, b["right"] - b["left"]) for b in raw_boxes]
        heights = [max(1, b["bottom"] - b["top"]) for b in raw_boxes]
        words = it["text"].split()

        def _wrap(font):
            out, wi = [], 0
            for k in range(n):
                cur = ""
                while wi < len(words):
                    trial = words[wi] if not cur else cur + " " + words[wi]
                    bb = draw.textbbox((0, 0), trial, font=font)
                    if (bb[2] - bb[0]) <= widths[k] or not cur:
                        cur = trial; wi += 1
                    else:
                        break
                out.append(cur)
                if wi >= len(words):
                    break
            return out if wi >= len(words) else None

        size = max(6, int(min(heights) * 0.9))
        wrapped = None
        while size >= 6:
            font = _load_pil_font(size, bold=it.get("bold", False), script=script)
            wrapped = _wrap(font)
            if wrapped is not None:
                break
            size -= 1
        if wrapped is None:
            font = _load_pil_font(6, bold=it.get("bold", False), script=script)
            wrapped = [it["text"]]
        for k, ln in enumerate(wrapped):
            box = raw_boxes[min(k, n - 1)]
            bw = box["right"] - box["left"]
            bh = box["bottom"] - box["top"]
            draw_ln = shape_rtl_text(ln) if item_rtl else ln
            bb = draw.textbbox((0, 0), draw_ln, font=font)
            lw = bb[2] - bb[0]
            y = box["top"] + max(0, (bh - size) // 2)
            if align == "center":
                x = box["left"] + max(0, (bw - lw) // 2)
            elif align == "left":
                x = box["left"]
            elif align == "right":
                x = box["right"] - lw
            else:
                x = (box["right"] - lw) if item_rtl else box["left"]
            draw.text((x, y), draw_ln, font=font, fill=tuple(it.get("color", (0, 0, 0))))
    return preview


def _review_page_layout(original_image, text_items, png_items, target_language,
                        rtl, llm_config, page_diag, max_rounds=2, progress=None):
    """A senior-reviewer QA agent (translator/typesetter with 20+ years
    of experience) enforcing the workflow's Section 4.2.15 quality gate.

    Each round it sees the ORIGINAL page and a PREVIEW of the rebuild and
    reports items that break the rules. When it drops an item as a
    logo/signature that was wrongly rendered as text, that region is
    turned back into a transparent PNG element (Layer 2) using the
    ORIGINAL pixels - so a signature/logo is never lost, it is restored
    as a picture (point: 'we don't extract its text, only make a PNG').
    Corrections re-loop until clean or max_rounds. Returns
    (text_items, png_items)."""
    if not llm_is_configured(llm_config) or not text_items:
        return text_items, png_items

    # Speed control (point: faster agents). The reviewer is the main
    # cost of the long 85% wait - each round is a vision call. Skip it
    # for trivially small pages (nothing to get wrong), and for very
    # large pages do a single pass rather than looping, so processing
    # time stays bounded.
    if len(text_items) <= 2:
        page_diag["reviewSkipped"] = "page too simple to need review"
        return text_items, png_items
    if len(text_items) > 40:
        max_rounds = 1

    lessons = _load_reviewer_lessons()
    lessons_block = _lessons_as_prompt_block(lessons)
    new_lessons = []

    for round_no in range(max_rounds):
        if progress:
            progress(f"Reviewer agent: round {round_no + 1} - inspecting layout")
        preview = _render_layout_preview(original_image, text_items, png_items)
        orig_b64 = _pil_image_to_jpeg_b64(original_image)
        prev_b64 = _pil_image_to_jpeg_b64(preview)

        # Compact, indexed snapshot of the current text layer for the
        # reviewer to reference by index.
        snapshot = []
        for i, it in enumerate(text_items):
            snapshot.append(
                f'{i}: text={it["text"]!r} align={it.get("align")} '
                f'rtl={it.get("rtl")} bold={it.get("bold")} '
                f'color=#{"%02X%02X%02X" % tuple(it.get("color", (0,0,0)))}')
        snapshot_text = "\n".join(snapshot)

        system_prompt = (
            "You are a senior document-translation reviewer and typesetter "
            "with 20+ years of experience in multilingual (Arabic/English) "
            "desktop publishing. You are the final quality gate (workflow "
            "Section 4.2.15). You are given TWO images: IMAGE 1 is the "
            "ORIGINAL source page; IMAGE 2 is a PREVIEW of the machine-"
            f"generated rebuild translated into {target_language}. You also "
            "get the current text layer as an indexed list.\n\n"
            "IMPORTANT: text reading DIRECTION is fixed by the OUTPUT "
            f"language ({target_language}) and is already handled - do NOT "
            "flag direction. Compare IMAGE 2 against IMAGE 1 and report ONLY "
            "items that are WRONG on these points:\n"
            "- CENTERING: if a line is visually CENTERED in the original but "
            'the rebuild is not (or vice-versa), fix align ("center" vs '
            '"natural").\n'
            "- PARAGRAPH/INDENT: if the original box is a centered block or "
            "an indented paragraph and the rebuild lost that, correct align.\n"
            "- OVERFLOW / OVERLAP: flag text that overflows its box or "
            "overlaps a neighbour in IMAGE 2.\n"
            "- COLOR: flag a color that clearly doesn't match the original "
            "ink.\n"
            "- LINE GROUPING: a single visual line wrongly split or two "
            "separate lines wrongly merged.\n"
            "- CLASSIFICATION: flag any item that is actually a LOGO, seal, "
            "emblem or HANDWRITTEN SIGNATURE - it must not be rendered as "
            "translated text (use drop). Also flag any real translatable "
            "line missing from the rebuild.\n\n"
            "For each wrong item output an object:\n"
            '  {"index": <int>, "fix": {"align":"center|natural", '
            '"text":"<corrected translation>", "drop":true|false}, '
            '"reason":"<what was wrong and must not recur>"}\n'
            "Include only fields that need changing. Use \"drop\":true when "
            "the item is a logo/signature/decoration wrongly shown as text.\n\n"
            + lessons_block +
            "Return ONLY a JSON array of correction objects (empty array [] "
            "if the rebuild is correct). No prose, no markdown."
        )
        user_content = [
            {"type": "text", "text": "IMAGE 1 = ORIGINAL page:"},
            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{orig_b64}"}},
            {"type": "text", "text": "IMAGE 2 = current translated rebuild PREVIEW:"},
            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{prev_b64}"}},
            {"type": "text", "text": "Current text layer (index: fields):\n" + snapshot_text +
             "\n\nReturn the JSON array of corrections now."},
        ]
        try:
            content, _prov = _call_chat_completion_with_failover(
                llm_config, system_prompt, user_content, max_tokens=4000)
        except Exception as err:
            print(f"Reviewer agent call failed (round {round_no + 1}): {err}")
            break
        if not content:
            break

        txt = content.strip()
        txt = re.sub(r"^```(?:json)?", "", txt).strip()
        txt = re.sub(r"```$", "", txt).strip()
        s, e = txt.find("["), txt.rfind("]")
        if s != -1 and e != -1 and e > s:
            txt = txt[s:e + 1]
        try:
            corrections = json.loads(txt)
        except Exception:
            corrections = []
        if not corrections:
            page_diag[f"reviewRound{round_no + 1}"] = "clean"
            break

        applied = 0
        drops = set()
        for corr in corrections:
            try:
                idx = int(corr.get("index"))
            except (TypeError, ValueError):
                continue
            if not (0 <= idx < len(text_items)):
                continue
            fix = corr.get("fix") or {}
            # Turn each caught mistake into a durable lesson for future
            # runs (self-training): the reviewer's own "reason" becomes a
            # rule injected into later prompts.
            reason = (corr.get("reason") or "").strip()
            if reason:
                new_lessons.append({"rule": reason})
            if fix.get("drop"):
                drops.add(idx)
                applied += 1
                continue
            if "align" in fix and fix["align"] in ("left", "center", "right", "natural"):
                text_items[idx]["align"] = fix["align"]
            if "rtl" in fix and isinstance(fix["rtl"], bool):
                text_items[idx]["rtl"] = fix["rtl"]
            if fix.get("text"):
                text_items[idx]["text"] = str(fix["text"]).strip()
            applied += 1
        if drops:
            # A dropped item is a logo/signature that was wrongly shown as
            # text. Restore it as a transparent PNG (Layer 2) from the
            # ORIGINAL pixels so it reappears as a picture, not text, and
            # is not left as an empty erased patch.
            for di in drops:
                it = text_items[di]
                region = {"left": it["left"], "top": it["top"],
                          "right": it["right"], "bottom": it["bottom"]}
                png = _crop_element_png(original_image, region)
                if png is not None:
                    png_items.append({
                        "left": it["left"], "top": it["top"],
                        "right": it["right"], "bottom": it["bottom"],
                        "png": png, "kind": it.get("kind", "signature")})
            text_items = [it for i, it in enumerate(text_items) if i not in drops]
        page_diag[f"reviewRound{round_no + 1}"] = f"{applied} correction(s)"
        if progress:
            progress(f"Reviewer agent: round {round_no + 1} applied {applied} correction(s)")
        if applied == 0:
            break

    # Persist what was learned this page so the next run starts smarter.
    if new_lessons:
        _save_reviewer_lessons(lessons + new_lessons)
        page_diag["reviewLessonsLearned"] = len(new_lessons)
    return text_items, png_items


def _tight_text_bbox(np_img, region, min_frac=0.5):
    """Shrinks a detected region to the TIGHT bounding box of its actual
    ink pixels. CV merges sometimes make a box far wider/taller than the
    text (it reaches into the decorative border, or spans a whole empty
    column), which causes two visible bugs: (1) inpainting then erases
    part of the border/pattern that was never text, and (2) 'centering'
    is computed against the oversized box so text drifts off-centre.
    Measuring the real ink extent fixes both. Falls back to the original
    box if it can't find a confident tighter one."""
    try:
        H, W = np_img.shape[:2]
        l = max(0, region["left"]); t = max(0, region["top"])
        r = min(W, region["right"]); b = min(H, region["bottom"])
        if r - l < 6 or b - t < 6:
            return region
        crop = np_img[t:b, l:r]
        ring = 6
        outer = np_img[max(0, t - ring):min(H, b + ring),
                       max(0, l - ring):min(W, r + ring)].reshape(-1, 3)
        bg = np.median(outer, axis=0)
        dist = np.sqrt(((crop.astype(np.int16) - bg) ** 2).sum(axis=2))
        mask = dist > 45
        cols = np.where(mask.any(axis=0))[0]
        rows = np.where(mask.any(axis=1))[0]
        if len(cols) < 3 or len(rows) < 3:
            return region
        nl, nr = l + int(cols[0]), l + int(cols[-1]) + 1
        nt, nb = t + int(rows[0]), t + int(rows[-1]) + 1
        # Only accept the tighter box if it actually removes a meaningful
        # amount of empty margin (otherwise keep the original to avoid
        # clipping faint glyph edges).
        if (nr - nl) < (r - l) * 0.98 or (nb - nt) < (b - t) * 0.98:
            new = dict(region)
            new["left"], new["right"], new["top"], new["bottom"] = nl, nr, nt, nb
            new["height"] = nb - nt
            return new
        return region
    except Exception:
        return region


def _region_is_graphic(np_img, region):
    """Conservative safety net: does this region contain a real COLOURED
    illustration (green foliage, blue/red painted figure, skin tones)
    that plain text - even red/black ink on aged parchment - would not?
    Text on old paper is dark ink (low saturation) on a warm neutral
    background; a botanical plate or painting has a meaningful fraction
    of STRONGLY SATURATED, non-ink-coloured pixels. We only flag when
    that saturated-colour fraction is high, to avoid mislabelling aged
    manuscript text as a picture. Primary classification is still the
    vision model - this only catches obvious pictures it might miss."""
    try:
        import colorsys
        H, W = np_img.shape[:2]
        l, t = max(0, region["left"]), max(0, region["top"])
        r, b = min(W, region["right"]), min(H, region["bottom"])
        w, h = r - l, b - t
        if w < 60 or h < 60:
            return False
        crop = np_img[t:b, l:r].astype(np.float32) / 255.0
        mx = crop.max(axis=2)
        mn = crop.min(axis=2)
        sat = np.where(mx > 0, (mx - mn) / np.maximum(mx, 1e-6), 0.0)
        val = mx
        R, G, B = crop[..., 0], crop[..., 1], crop[..., 2]
        # Warm parchment / ink pixels: high red-ish, brownish - exclude
        # them. Strongly-coloured pixels are saturated AND not just
        # warm-brown ink/paper.
        warm_brown = (R >= G) & (G >= B) & (sat < 0.55)
        strong = (sat > 0.45) & (val > 0.25) & (~warm_brown)
        strong_frac = float(strong.mean())
        # Green foliage specifically (plant plates): green channel clearly
        # dominant with decent saturation.
        green = ((G > R + 0.06) & (G > B + 0.06) & (sat > 0.25)).mean()
        # A picture: a clear band of strongly coloured OR green pixels.
        return strong_frac > 0.06 or green > 0.05
    except Exception:
        return False


def _group_regions_into_paragraphs(regions):
    """Groups consecutive line-regions that belong to the SAME paragraph
    into one block, while KEEPING each line's own box. This implements
    the user's model: every visual line is its own block, and a
    multi-line sentence is a group of those line-blocks that the
    translation is later flowed across (line by line), not a single
    merged box.

    Two line-regions join the same paragraph when they are vertically
    close (gap smaller than a line height) and their horizontal spans
    overlap - i.e. they read as stacked lines of one block. Graphic
    regions never group. Returns a list of blocks, each:
        {"line_boxes":[{l,t,r,b}, ...],   # one per visual line, in order
         "left","top","right","bottom",   # union box of the group
         "is_graphic": bool}
    Single lines (headings, labels) become one-line blocks."""
    lines = [r for r in regions if not r.get("is_graphic")]
    graphics = [r for r in regions if r.get("is_graphic")]
    lines.sort(key=lambda r: (r["top"], r["left"]))

    blocks = []
    used = [False] * len(lines)
    for i, r in enumerate(lines):
        if used[i]:
            continue
        group = [r]
        used[i] = True
        rh = r["bottom"] - r["top"]
        cur_bottom = r["bottom"]
        cur_l, cur_r = r["left"], r["right"]
        for j in range(i + 1, len(lines)):
            if used[j]:
                continue
            s = lines[j]
            sh = s["bottom"] - s["top"]
            vgap = s["top"] - cur_bottom
            # Must be the NEXT line down: a small positive gap (no
            # vertical overlap, which would mean they're not stacked text
            # lines but unrelated blocks like a badge over a title).
            if vgap < -0.15 * max(rh, sh) or vgap > 0.8 * max(rh, sh):
                continue
            # Horizontal spans must overlap a lot (same column of text).
            overlap = min(cur_r, s["right"]) - max(cur_l, s["left"])
            if overlap < 0.5 * min(cur_r - cur_l, s["right"] - s["left"]):
                continue
            # Similar line height (body lines match; a heading over body
            # differs too much and must stay separate).
            if not (0.55 <= sh / max(1, rh) <= 1.8):
                continue
            group.append(s)
            used[j] = True
            cur_bottom = s["bottom"]
            cur_l = min(cur_l, s["left"])
            cur_r = max(cur_r, s["right"])
            rh = sh   # track the most recent line's height for the next gap test
        group.sort(key=lambda g: g["top"])
        line_boxes = [{"left": g["left"], "top": g["top"],
                       "right": g["right"], "bottom": g["bottom"]} for g in group]
        blocks.append({
            "line_boxes": line_boxes,
            "left": min(g["left"] for g in group),
            "top": min(g["top"] for g in group),
            "right": max(g["right"] for g in group),
            "bottom": max(g["bottom"] for g in group),
            "is_graphic": False,
        })
    for g in graphics:
        blocks.append({
            "line_boxes": [{"left": g["left"], "top": g["top"],
                            "right": g["right"], "bottom": g["bottom"]}],
            "left": g["left"], "top": g["top"], "right": g["right"],
            "bottom": g["bottom"], "is_graphic": True,
        })
    blocks.sort(key=lambda b: (b["top"], b["left"]))
    return blocks


def _build_page_vision_layout(page_image, target_language, rtl, llm_config, page_diag, progress=None):
    """Workflow doc steps 4.2.5-4.2.13 for one page, returning DATA (not
    flattened pixels): CV region detection (4.2.5) -> single structured
    VLM pass for text, translation, typography and classification
    (4.2.6-4.2.12) -> classify each region into the three document
    buckets:
      * translatable  -> its ORIGINAL text is erased from the background
                         (4.2.10) and re-added as an editable Layer-3
                         text item with the detected typography (4.2.12)
      * nontranslatable(logo/signature/QR/stamp/seal) -> pixels are LEFT
                         in the background (kept as-is), never erased,
                         never translated (Layer-2 behaviour without a
                         separate PNG - visually identical result)
      * decoration/watermark -> erased and NOT re-added (4.2.10 clean
                         canvas)
    Only translatable regions go into the inpaint mask, so the doc's
    'erase only translatable text, preserve everything else' rule is
    honored exactly. Returns (cleaned_background_image, text_items) or
    None if the VLM pass produced nothing (caller falls back)."""
    line_regions = _detect_text_regions_cv(page_image)
    page_diag["cvRegionsDetected"] = len(line_regions)
    if not line_regions:
        return None
    # Group consecutive line-boxes into paragraph BLOCKS (each block
    # keeps its individual line boxes). The VLM then sees one numbered
    # box per block (its union), so a multi-line sentence is translated
    # coherently as one unit - and the translation is later flowed back
    # across the block's own line boxes, line by line.
    blocks = _group_regions_into_paragraphs(line_regions)
    page_diag["cvBlocks"] = len(blocks)
    regions = blocks   # VLM receives union boxes; each carries line_boxes
    meta = _extract_page_metadata_vlm(page_image, regions, target_language, llm_config)
    if all(m is None for m in meta):
        page_diag["cvVisionFailed"] = True
        return None

    np_img = np.array(page_image)
    text_items = []          # translatable -> Layer 3
    png_items = []           # nontranslatable -> Layer 2 (transparent PNG)
    erase_regions = []       # translatable + decoration -> inpaint mask
    protect_regions = []     # nontranslatable -> never erased
    counts = {"translatable": 0, "nontranslatable": 0, "decoration": 0}

    # First pass: collect the boxes the model marked non-translatable
    # (logos, badges, signatures). Any other region that sits INSIDE one
    # of these must never be erased or translated - e.g. the Arabic text
    # printed on a ribbon badge is part of the LOGO, not a separate text
    # line, so erasing it would gouge a hole in the logo (the "logo cut"
    # problem). Those contained regions are simply dropped.
    nontrans_boxes = []
    for region, m in zip(regions, meta):
        if m and str(m.get("class", "")).lower() == "nontranslatable":
            nontrans_boxes.append(region)

    def _inside_any(region, boxes):
        cx = (region["left"] + region["right"]) / 2.0
        cy = (region["top"] + region["bottom"]) / 2.0
        for bx in boxes:
            if bx is region:
                continue
            if bx["left"] - 4 <= cx <= bx["right"] + 4 and bx["top"] - 4 <= cy <= bx["bottom"] + 4:
                return True
        return False

    for region, m in zip(regions, meta):
        # Pre-detected illustration blob (from the colour-blob pass): keep
        # as a non-translatable PNG, never erase or translate. This is the
        # reliable path for the page-2 plant. We do NOT run a pixel-level
        # "looks graphic" guess here anymore, because dark ink on aged
        # parchment can look colourful and would wrongly turn real text
        # paragraphs into pictures - classification of ambiguous regions
        # is left to the vision model.
        if region.get("is_graphic"):
            kind_m = str((m or {}).get("kind", "")).lower()
            png = _crop_element_png(page_image, region)
            if png is not None:
                png_items.append({
                    "left": region["left"], "top": region["top"],
                    "right": region["right"], "bottom": region["bottom"],
                    "png": png, "kind": kind_m or "illustration"})
            protect_regions.append(region)
            nontrans_boxes.append(region)
            counts["nontranslatable"] += 1
            continue
        if not m:
            if _inside_any(region, nontrans_boxes):
                continue  # part of a logo/badge - leave it alone
            erase_regions.append(region)
            counts["decoration"] += 1
            continue
        cls = str(m.get("class", "")).lower()
        if cls == "translatable":
            # A "translatable" box that actually lies inside a logo is a
            # mis-classification of logo lettering; keep the logo intact.
            if _inside_any(region, nontrans_boxes):
                continue
            translation = (m.get("translation") or m.get("text") or "").strip()
            if not translation:
                for lb in region.get("line_boxes", [region]):
                    erase_regions.append(lb)
                counts["decoration"] += 1
                continue
            # Each visual line in this block is its own box. Tighten each
            # to its real ink extent (keeps inpaint off the border and
            # makes placement exact), and erase each line box.
            raw_lines = region.get("line_boxes", [{"left": region["left"],
                        "top": region["top"], "right": region["right"],
                        "bottom": region["bottom"]}])
            line_boxes = []
            for lb in raw_lines:
                tb = _tight_text_bbox(np_img, lb)
                line_boxes.append({"left": tb["left"], "top": tb["top"],
                                   "right": tb["right"], "bottom": tb["bottom"]})
                erase_regions.append(tb)
            union = {"left": min(b["left"] for b in line_boxes),
                     "top": min(b["top"] for b in line_boxes),
                     "right": max(b["right"] for b in line_boxes),
                     "bottom": max(b["bottom"] for b in line_boxes)}
            color = _normalize_hex_color(
                m.get("color"),
                default=_region_stroke_color(np_img, union))
            align = str(m.get("align", "")).lower()
            if align not in ("left", "center", "right"):
                align = "center" if not m.get("is_paragraph") else "natural"
            if align == "right":
                align = "natural"   # never force right; direction decides the edge
            # Optional per-run colour/style (point 3). Keep only if the
            # runs are well-formed and actually vary; otherwise the
            # single color/style above is used.
            runs = m.get("runs")
            clean_runs = None
            if isinstance(runs, list) and len(runs) > 1:
                clean_runs = []
                for rn in runs:
                    if not isinstance(rn, dict):
                        continue
                    rt = str(rn.get("text", ""))
                    if rt == "":
                        continue
                    clean_runs.append({
                        "text": rt,
                        "color": _normalize_hex_color(rn.get("color"), default=color),
                        "bold": bool(rn.get("bold", m.get("bold"))),
                        "italic": bool(rn.get("italic", m.get("italic"))),
                    })
                if len(clean_runs) < 2:
                    clean_runs = None
            text_items.append({
                "left": union["left"], "top": union["top"],
                "right": union["right"], "bottom": union["bottom"],
                "line_boxes": line_boxes,      # per-line geometry for flow
                "text": translation,
                "runs": clean_runs,            # per-letter colour/style, or None
                "color": color,
                "bold": bool(m.get("bold")),
                "italic": bool(m.get("italic")),
                "underline": bool(m.get("underline")),
                "align": align,
                "rtl": rtl,
                "is_paragraph": bool(m.get("is_paragraph")) or len(line_boxes) > 1,
                "kind": str(m.get("kind", "")),
            })
            counts["translatable"] += 1
        elif cls == "nontranslatable":
            # Skip elements nested inside a bigger non-translatable box
            # (e.g. badge text inside the badge) - the outer element's
            # single PNG already contains them.
            if _inside_any(region, nontrans_boxes):
                protect_regions.append(region)
                continue
            png = _crop_element_png(page_image, region)
            if png is not None:
                png_items.append({
                    "left": region["left"], "top": region["top"],
                    "right": region["right"], "bottom": region["bottom"],
                    "png": png, "kind": str(m.get("kind", "")),
                })
            protect_regions.append(region)
            counts["nontranslatable"] += 1
        else:  # decoration / watermark / empty
            if _inside_any(region, nontrans_boxes):
                continue
            erase_regions.append(region)
            counts["decoration"] += 1

    page_diag["cvTranslatable"] = counts["translatable"]
    page_diag["cvNonTranslatable"] = counts["nontranslatable"]
    page_diag["cvDecoration"] = counts["decoration"]

    # Optional quality-gate reviewer agent. It is powerful but costs a
    # SECOND vision round-trip per page, which roughly doubles the time
    # per file - the main reason a 2-page job took ~5 minutes. It is now
    # OFF by default and only runs when explicitly enabled via the
    # LEXORA_ENABLE_REVIEWER env flag, so the default path is fast.
    if os.environ.get("LEXORA_ENABLE_REVIEWER", "").lower() in ("1", "true", "yes"):
        try:
            text_items, png_items = _review_page_layout(
                page_image, text_items, png_items, target_language, rtl,
                llm_config, page_diag, progress=progress)
        except Exception as rev_err:
            print(f"Reviewer agent skipped due to error: {rev_err}")

    # Any element restored as a PNG by the reviewer (a dropped
    # signature/logo) must NOT be inpainted away - protect its box.
    review_protect = [{"left": p["left"], "top": p["top"],
                       "right": p["right"], "bottom": p["bottom"]}
                      for p in png_items]
    if progress:
        progress(f"Reconstructing clean background ({len(erase_regions)} text areas)")
    # Point 1: prefer AI generative fill of the erased areas when enabled;
    # otherwise reconstruct the background with fast CV inpainting.
    cleaned = None
    if erase_regions:
        cleaned = _ai_fill_background(page_image, erase_regions,
                                      protect_regions + review_protect, llm_config)
    if cleaned is None:
        cleaned = _inpaint_regions_cv(page_image, erase_regions,
                                      protect_regions + review_protect) if erase_regions else page_image
    return cleaned, text_items, png_items


_PDF_FONTS_REGISTERED = False


def _register_pdf_fonts():
    """Registers the bundled Noto TTFs with reportlab so the editable
    PDF path can emit REAL text objects (selectable/copyable/editable
    in a PDF editor) in any script, instead of Helvetica-only. Safe to
    call repeatedly."""
    global _PDF_FONTS_REGISTERED
    if _PDF_FONTS_REGISTERED or not REPORTLAB_OK:
        return
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        fonts_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "fonts")
        for name, filename in (("LexNoto", "NotoSans-Regular.ttf"),
                               ("LexNotoArabic", "NotoSansArabic-Regular.ttf"),
                               ("LexNotoArabicBold", "NotoSansArabic-Bold.ttf")):
            path = os.path.join(fonts_dir, filename)
            if os.path.isfile(path) and name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(name, path))
        _PDF_FONTS_REGISTERED = True
    except Exception as err:
        print(f"Could not register bundled TTF fonts with reportlab: {err} - falling back to Helvetica.")


def _pdf_font_for(text, bold, italic, rtl):
    """Best registered PDF font for one text item. Pure-ASCII text can
    use reportlab's built-in Helvetica family (which has real bold and
    italic faces); anything beyond ASCII needs a registered Unicode TTF
    (Noto), where bold exists only for the Arabic face."""
    try:
        text.encode("ascii")
        is_ascii = True
    except UnicodeEncodeError:
        is_ascii = False
    if is_ascii:
        return _closest_builtin_font(bold, italic)
    from reportlab.pdfbase import pdfmetrics
    registered = pdfmetrics.getRegisteredFontNames()
    if rtl or any("\u0600" <= ch <= "\u06FF" for ch in text):
        if bold and "LexNotoArabicBold" in registered:
            return "LexNotoArabicBold"
        if "LexNotoArabic" in registered:
            return "LexNotoArabic"
    if "LexNoto" in registered:
        return "LexNoto"
    return _closest_builtin_font(bold, italic)


def _sample_region_bg(page_image, left, top, right, bottom, pad=4):
    """Median color of a thin ring of pixels JUST OUTSIDE the region box -
    a much better mask color than the old single-pixel sample: on
    textured/patterned backgrounds (certificates, letterheads) one pixel
    is a lottery ticket, while the ring median lands on the dominant
    background tone so the mask rectangle blends in instead of standing
    out as a stark white patch."""
    w, h = page_image.size
    xs = list(range(max(0, left - pad), min(w, right + pad), 6))
    ys = list(range(max(0, top - pad), min(h, bottom + pad), 6))
    samples = []
    for x in xs:
        for y in (max(0, top - pad), min(h - 1, bottom + pad - 1)):
            samples.append(page_image.getpixel((x, y)))
    for y in ys:
        for x in (max(0, left - pad), min(w - 1, right + pad - 1)):
            samples.append(page_image.getpixel((x, y)))
    if not samples:
        return (255, 255, 255)
    med = tuple(sorted(px[ch] for px in samples)[len(samples) // 2] for ch in range(3))
    return med


def generate_ocr_based_translation_pdf(original_pdf_path, output_pdf_path, target_language, llm_config=None, ocr_lang="eng+ara", diagnostics=None, vision_assist=False, progress=None):
    """Item 6 - the scanned/image-only counterpart to
    generate_layout_preserving_translation_pdf(): when a page has no real
    text layer at all (the whole page is one embedded/scanned image,
    common for posters, faxes, and scanned contracts), there's no text
    object to surgically edit - so this renders the page to a
    high-resolution image, uses Tesseract's word-level bounding boxes
    (image_to_data, not just image_to_string) to find where the text
    actually sits, masks each detected line with a rectangle (sampling
    the source image's own nearby background color, not assuming white)
    drawn directly onto the image, and draws the translated line back at
    that same position with a bundled Unicode font - everything else in
    the image (icons, photos, background colors, decorative graphics) is
    untouched pixel data, since only the specific masked rectangles get
    overwritten.

    OCR language defaults to English+Arabic (ocr_lang="eng+ara") since
    that covers this feature's most common real case; pass a different
    `ocr_lang` for other source scripts (matching Tesseract's language
    codes, e.g. "eng+hin" for Hindi) - see the Dockerfile for which
    tesseract-ocr-<lang> packages are installed. Whatever's requested
    here is filtered down to whatever's actually installed on this
    server (see _resolve_available_ocr_langs) rather than failing
    outright if one language pack is missing.

    diagnostics (optional, a dict) - filled in with what actually
    happened (pdfium/tesseract versions, regions detected per page,
    translation call outcome, any errors) so the CALLER can surface this
    somewhere visible (the Activity Log) - added specifically because a
    production server's own console isn't something we can see directly,
    only what the app itself reports back."""
    if diagnostics is None:
        diagnostics = {}
    diagnostics["ocrLibsOk"] = OCR_LIBS_OK
    diagnostics["pdfplumberAvailable"] = bool(pdfplumber)
    try:
        import importlib.metadata as _importlib_metadata
        diagnostics["pypdfium2Version"] = _importlib_metadata.version("pypdfium2")
    except Exception as err:
        diagnostics["pypdfium2Version"] = f"unavailable: {err}"
    try:
        diagnostics["tesseractVersion"] = str(pytesseract.get_tesseract_version()) if OCR_LIBS_OK else "n/a"
    except Exception as err:
        diagnostics["tesseractVersion"] = f"unavailable: {err}"
    diagnostics["pages"] = []

    if not OCR_LIBS_OK:
        raise LeaseEngineError("pytesseract/Pillow are not installed - run: pip install -r requirements.txt")
    if not pdfplumber:
        raise LeaseEngineError("pdfplumber is not installed - run: pip install -r requirements.txt")

    from reportlab.pdfgen import canvas as pdfcanvas

    rtl = is_rtl_language(target_language)
    # 300 (not 200) matters a lot for photographed/low-res source scans:
    # on a real photographed Arabic utility bill, 200 DPI yielded ~20 OCR
    # words while 300 DPI yielded 127+ - Tesseract's recognizer is tuned
    # for roughly 300 DPI text height, so under-rendering the page image
    # starves it even when the right language pack IS installed.
    render_dpi = 300
    ocr_lang = _resolve_available_ocr_langs(ocr_lang)
    diagnostics["ocrLangUsed"] = ocr_lang

    story_pages = []
    with pdfplumber.open(original_pdf_path) as pdf:
        total_pages = len(pdf.pages)
        for page_num, page in enumerate(pdf.pages):
            if progress:
                progress(f"Page {page_num + 1}/{total_pages}: rendering & detecting text regions")
            page_diag = {"page": page_num + 1, "regionsDetected": 0, "regionsDrawn": 0, "error": None}
            page_image = page.to_image(resolution=render_dpi).original.convert("RGB")
            try:
                if vision_assist and CV2_OK and llm_is_configured(llm_config):
                    layout = None
                    try:
                        if progress:
                            progress(f"Page {page_num + 1}/{total_pages}: reading text + typography (vision)")
                        layout = _build_page_vision_layout(page_image, target_language, rtl, llm_config, page_diag, progress=progress)
                    except Exception as cv_err:
                        print(f"Vision-layout rebuild failed on page {page_num + 1}: {cv_err} - falling back to the OCR flow.")
                        page_diag["cvError"] = str(cv_err)
                    if layout is not None:
                        cleaned_image, layout_items, png_items = layout
                        page_diag["pathUsed"] = "vision-layout-cv"
                        diagnostics["pages"].append(page_diag)
                        story_pages.append(("layout", cleaned_image, layout_items, float(page.width), float(page.height), png_items))
                        continue
                    page_diag["pathUsed"] = "tesseract-fallback"
                try:
                    ocr_data = pytesseract.image_to_data(page_image, lang=ocr_lang, output_type=pytesseract.Output.DICT)
                except Exception as err:
                    print(f"OCR failed on page {page_num + 1} (lang={ocr_lang}): {err} - this page will be left as the original image.")
                    page_diag["error"] = f"OCR call failed: {err}"
                    ocr_data = {"text": [], "conf": [], "left": [], "top": [], "width": [], "height": [], "block_num": [], "par_num": [], "line_num": []}
                regions = _group_ocr_words_into_lines(ocr_data)
                # Retry with --psm 6 ("assume a single uniform block of
                # text") when the default auto page segmentation (--psm 3)
                # found suspiciously little. Auto segmentation frequently
                # under-detects photographed documents with busy layouts
                # (bills, forms, tables) - on a real photographed Arabic
                # bill it found 34 line regions where psm 6 found 100+.
                # Whichever pass detects MORE regions wins.
                if len(regions) < 60:
                    try:
                        ocr_data_psm6 = pytesseract.image_to_data(page_image, lang=ocr_lang, config="--psm 6", output_type=pytesseract.Output.DICT)
                        regions_psm6 = _group_ocr_words_into_lines(ocr_data_psm6)
                        if len(regions_psm6) > len(regions):
                            print(f"OCR page {page_num + 1}: psm 6 retry found {len(regions_psm6)} regions vs {len(regions)} with auto segmentation - using psm 6 result.")
                            regions = regions_psm6
                            page_diag["psm6RetryUsed"] = True
                    except Exception as retry_err:
                        print(f"psm 6 OCR retry failed on page {page_num + 1}: {retry_err} - keeping the auto-segmentation result.")
                page_diag["regionsDetected"] = len(regions)
                print(f"OCR page {page_num + 1}: {len(regions)} text region(s) detected (lang={ocr_lang}).")

                if regions and vision_assist:
                    # Hybrid mode: send the page image along with the OCR
                    # lines so the model corrects Tesseract's misreads
                    # against the actual pixels before translating.
                    translated_texts = _translate_lines_batch_vision(page_image, [r["text"] for r in regions], target_language, llm_config)
                    page_diag["visionAssist"] = True
                elif regions:
                    translated_texts = _translate_lines_batch([r["text"] for r in regions], target_language, llm_config)
                else:
                    translated_texts = []
                page_diag["translatedCount"] = len(translated_texts)
                page_diag["sampleOriginal"] = regions[0]["text"] if regions else None
                page_diag["sampleTranslated"] = translated_texts[0] if translated_texts else None

                draw = ImageDraw.Draw(page_image)
                for region, translated in zip(regions, translated_texts):
                    try:
                        # [SKIP] = the model confirmed this "line" is a
                        # logo / handwritten signature / ornament that OCR
                        # misdetected as text - leave those pixels
                        # completely alone (no mask, no redraw). This is
                        # what stops blank boxes appearing over signatures
                        # and hallucinated "translations" of logo art.
                        if not translated or translated.strip().upper() in ("[SKIP]", "SKIP"):
                            page_diag["regionsSkipped"] = page_diag.get("regionsSkipped", 0) + 1
                            continue
                        box_w = max(1, region["right"] - region["left"])
                        box_h = max(1, region["bottom"] - region["top"])
                        try:
                            bg = _sample_region_bg(page_image, region["left"], region["top"], region["right"], region["bottom"])
                        except Exception:
                            bg = (255, 255, 255)
                        draw.rectangle([region["left"] - 2, region["top"] - 2, region["right"] + 2, region["bottom"] + 2], fill=bg)

                        script = "arabic" if rtl else "latin"
                        # Wrap on the UNshaped translated text first (word
                        # boundaries are meaningful before RTL shaping),
                        # then shape each resulting line individually -
                        # shaping a string that already contains a '\n'
                        # can scramble the line break itself.
                        font, plain_lines, line_h = _fit_text_in_box_pil(
                            draw, translated, box_w, box_h, script=script)
                        render_lines = [shape_rtl_text(ln) if rtl else ln for ln in plain_lines]

                        text_color = (0, 0, 0)
                        block_h = line_h * len(render_lines)
                        # Center the wrapped block vertically inside the
                        # original line's box rather than always hanging
                        # it from the top - looks right whether it ended
                        # up as 1 line (matches old behavior) or several.
                        y = region["top"] + max(0, (box_h - block_h) / 2)
                        for ln in render_lines:
                            bbox = draw.textbbox((0, 0), ln, font=font)
                            ln_w = bbox[2] - bbox[0]
                            if rtl:
                                x = region["right"] - ln_w
                            else:
                                x = region["left"]
                            draw.text((x, y), ln, font=font, fill=text_color)
                            y += line_h
                        page_diag["regionsDrawn"] += 1
                    except Exception as region_err:
                        # One region failing to mask/redraw (a font glitch,
                        # an unusual bbox) shouldn't sacrifice every OTHER
                        # region on the page - that region's original text
                        # just stays as-is and the rest continue normally.
                        print(f"Could not mask/redraw region {region.get('text','')!r} on page {page_num + 1}: {region_err}")
            except Exception as page_err:
                # Same principle one level up: a full page failing
                # (a genuinely broken image, an OCR crash that isn't
                # caught above) leaves that ONE page as the original
                # image rather than failing the entire document.
                print(f"Page {page_num + 1} could not be processed for translation, leaving it as the original image: {page_err}")
                page_diag["error"] = page_diag["error"] or str(page_err)
            diagnostics["pages"].append(page_diag)

            story_pages.append(("flat", page_image, None, float(page.width), float(page.height), []))

    _register_pdf_fonts()
    buf = io.BytesIO()
    c = None
    for kind, page_image, layout_items, page_w_pt, page_h_pt, png_items in story_pages:
        if c is None:
            c = pdfcanvas.Canvas(buf, pagesize=(page_w_pt, page_h_pt))
        else:
            c.showPage()
            c.setPageSize((page_w_pt, page_h_pt))
        img_buf = io.BytesIO()
        page_image.save(img_buf, format="PNG")
        img_buf.seek(0)
        # Layer 1: inpainted background.
        c.drawImage(ImageReader(img_buf), 0, 0, width=page_w_pt, height=page_h_pt)
        scale = page_w_pt / float(page_image.width)
        # Layer 2: non-translatable elements as their own transparent
        # PNGs, placed at exact original position over the background.
        for pe in (png_items or []):
            try:
                pbuf = io.BytesIO()
                pe["png"].save(pbuf, format="PNG")
                pbuf.seek(0)
                px = pe["left"] * scale
                pw = (pe["right"] - pe["left"]) * scale
                ph = (pe["bottom"] - pe["top"]) * scale
                py = page_h_pt - (pe["bottom"] * scale)   # PDF y is bottom-up
                c.drawImage(ImageReader(pbuf), px, py, width=pw, height=ph, mask="auto")
            except Exception as perr:
                print(f"Could not place element PNG on PDF: {perr}")
        # Layer 3: translated text as REAL, selectable text objects.
        if kind == "layout" and layout_items:
            _draw_layout_items_on_canvas(c, layout_items, scale, page_h_pt, rtl)
    if c is not None:
        c.save()
    buf.seek(0)
    with open(output_pdf_path, "wb") as f:
        f.write(buf.read())

    # Editable Word companion: same cleaned backgrounds + the same text
    # items as floating, styled text boxes (the reference "v5" output
    # format). Written next to the PDF whenever at least one page went
    # through the vision-layout path.
    if any(kind == "layout" for kind, *_ in story_pages):
        docx_path = os.path.splitext(output_pdf_path)[0] + ".docx"
        try:
            _write_layout_docx(docx_path, story_pages, rtl)
            diagnostics["editableDocx"] = docx_path
        except Exception as err:
            print(f"Could not write editable DOCX companion: {err}")
            diagnostics["editableDocxError"] = str(err)


def _flow_text_across_lineboxes(text, line_boxes, font_name, rtl,
                                min_size=4.0, start_frac=0.92, width_factor=0.93):
    """Flows one translated string across a paragraph's N line boxes
    (the user's model: each original visual line is a box; the sentence
    is distributed line by line into those boxes at a UNIFORM font size).

    It picks the largest single font size at which the text word-wraps
    into AT MOST len(line_boxes) lines with every wrapped line fitting
    the WIDTH of its target box. If it cannot (a long translation in a
    short box), it keeps shrinking the font until every line fits - so
    text NEVER overflows its box (the caller's hard rule). `width_factor`
    (<1) leaves a safety margin because Word renders in Arial, which is a
    little wider than the metrics font. Returns (size, [(line, box), ...])."""
    from reportlab.pdfbase.pdfmetrics import stringWidth
    n = len(line_boxes)
    if n == 0:
        return min_size, []
    widths = [max(1.0, (b["right"] - b["left"]) * width_factor) for b in line_boxes]
    heights = [max(1.0, b["bottom"] - b["top"]) for b in line_boxes]
    words = text.split() or [text]

    def wrap_at(sz):
        # Greedy word wrap where line k must fit widths[k]; returns the
        # list of lines, or None if words are left over after n lines OR
        # a single word is wider than its box at this size (real overflow).
        out = []
        wi = 0
        for k in range(n):
            maxw = widths[k]
            cur = ""
            while wi < len(words):
                w = words[wi]
                trial = w if not cur else cur + " " + w
                if stringWidth(trial, font_name, sz) <= maxw:
                    cur = trial
                    wi += 1
                elif not cur:
                    # Single word wider than the whole box at this size.
                    if stringWidth(w, font_name, sz) > maxw:
                        return None
                    cur = w
                    wi += 1
                else:
                    break
            out.append(cur)
            if wi >= len(words):
                break
        if wi < len(words):
            return None   # leftover words -> needs more than n lines
        return out

    # Start from a size bounded by BOTH the box height and a width-based
    # estimate (so a wide-but-short box starts sensibly), then shrink.
    size = max(min_size, min(heights) * start_frac)
    chosen = None
    while size >= min_size:
        wrapped = wrap_at(size)
        if wrapped is not None:
            chosen = wrapped
            break
        size -= 0.5
    if chosen is None:
        # Guarantee no overflow: at min_size, hard-wrap by character so
        # every drawn line fits its box width.
        size = min_size
        chosen = []
        wi = 0
        for k in range(n):
            maxw = widths[k]
            cur = ""
            while wi < len(words):
                w = words[wi]
                trial = w if not cur else cur + " " + w
                if stringWidth(trial, font_name, size) <= maxw or not cur:
                    cur = trial
                    wi += 1
                else:
                    break
            chosen.append(cur)
            if wi >= len(words):
                break
    mapped = []
    for k, ln in enumerate(chosen):
        box = line_boxes[min(k, n - 1)]
        mapped.append((ln, box))
    return size, mapped


def _draw_layout_items_on_canvas(c, items, scale, page_h_pt, rtl):
    """Draws one page's translated items on a reportlab canvas as real,
    selectable text. Each item is a paragraph whose translation is
    FLOWED across its per-line boxes (one original visual line = one
    box) at a uniform font size, shrinking the font if the text would
    overflow the available lines. Single-line items are just one box.
    Honors ink colour, bold/italic, underline and centering."""
    from reportlab.pdfbase.pdfmetrics import stringWidth
    for item in items:
        item_rtl = item.get("rtl", rtl)
        align = item.get("align", "center")
        font_name = _pdf_font_for(item["text"], item["bold"], item["italic"], item_rtl)
        # Scale the per-line boxes into PDF points.
        raw_boxes = item.get("line_boxes") or [{
            "left": item["left"], "top": item["top"],
            "right": item["right"], "bottom": item["bottom"]}]
        line_boxes = [{"left": b["left"] * scale, "top": b["top"] * scale,
                       "right": b["right"] * scale, "bottom": b["bottom"] * scale}
                      for b in raw_boxes]
        size, mapped = _flow_text_across_lineboxes(
            item["text"], line_boxes, font_name, item_rtl)
        r, g, b = [v / 255.0 for v in item["color"]]
        c.setFillColorRGB(r, g, b)
        c.setStrokeColorRGB(r, g, b)
        c.setFont(font_name, size)
        for ln, box in mapped:
            draw_ln = shape_rtl_text(ln) if (item_rtl and any("\u0590" <= ch <= "\u06FF" for ch in ln)) else ln
            ln_w = stringWidth(draw_ln, font_name, size)
            bw = box["right"] - box["left"]
            bh = box["bottom"] - box["top"]
            # Vertically centre the glyph within its own line box.
            baseline = box["top"] + max(0.0, (bh - size) / 2.0) + size * 0.82
            y_pdf = page_h_pt - baseline
            if align == "center":
                x = box["left"] + max(0.0, (bw - ln_w) / 2.0)
            elif align == "left":
                x = box["left"]
            elif align == "right":
                x = box["left"] + max(0.0, bw - ln_w)
            else:  # natural: follow output direction
                x = box["left"] + max(0.0, bw - ln_w) if item_rtl else box["left"]
            c.drawString(x, y_pdf, draw_ln)
            if item["underline"]:
                c.setLineWidth(max(0.5, size / 16.0))
                c.line(x, y_pdf - size * 0.12, x + ln_w, y_pdf - size * 0.12)


_DOCX_CONTENT_TYPES = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Default Extension="png" ContentType="image/png"/>
<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""

_DOCX_RELS = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""

_DOCX_NS = (
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
    'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
    'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
    'xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture" '
    'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"'
)


def _xml_escape(s):
    return (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            .replace('"', "&quot;").replace("'", "&apos;"))


def _write_layout_docx(docx_path, story_pages, rtl):
    """Writes the editable Word companion: for every page, the cleaned
    (text-erased, inpainted) page render as a full-page behind-text
    background image, plus each translated item as a floating,
    borderless TEXT BOX anchored at its exact position - fully editable
    text in Word, styled with the item's color/bold/italic/underline
    and a size fitted to its original box. Built as a raw OOXML package
    for exact control over anchored drawing XML (python-docx has no API
    for floating text boxes)."""
    import zipfile
    from reportlab.pdfbase.pdfmetrics import stringWidth

    EMU_PER_PT = 12700
    media = []          # (filename, bytes)
    body_parts = []
    doc_rels = []
    draw_id = 100

    for page_idx, (kind, page_image, layout_items, page_w_pt, page_h_pt, png_items) in enumerate(story_pages):
        img_buf = io.BytesIO()
        page_image.save(img_buf, format="PNG")
        media.append((f"page{page_idx + 1}.png", img_buf.getvalue()))
        rel_id = f"rIdImg{page_idx + 1}"
        doc_rels.append(
            f'<Relationship Id="{rel_id}" '
            f'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" '
            f'Target="media/page{page_idx + 1}.png"/>')

        page_w_emu = int(page_w_pt * EMU_PER_PT)
        page_h_emu = int(page_h_pt * EMU_PER_PT)
        scale = page_w_pt / float(page_image.width)

        draw_id += 1
        anchors = [(
            f'<w:r><w:drawing><wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" '
            f'relativeHeight="0" behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">'
            f'<wp:simplePos x="0" y="0"/>'
            f'<wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>'
            f'<wp:positionV relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionV>'
            f'<wp:extent cx="{page_w_emu}" cy="{page_h_emu}"/><wp:wrapNone/>'
            f'<wp:docPr id="{draw_id}" name="Background {page_idx + 1}"/>'
            f'<a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
            f'<pic:pic><pic:nvPicPr><pic:cNvPr id="{draw_id}" name="bg{page_idx + 1}"/><pic:cNvPicPr/></pic:nvPicPr>'
            f'<pic:blipFill><a:blip r:embed="{rel_id}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
            f'<pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{page_w_emu}" cy="{page_h_emu}"/></a:xfrm>'
            f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr></pic:pic>'
            f'</a:graphicData></a:graphic></wp:anchor></w:drawing></w:r>')]

        # Layer 2: non-translatable elements (logos, signatures, QR) as
        # their own floating PNG images anchored at exact position. In
        # Word these are separate, clickable, movable, deletable objects
        # - satisfying the doc's 4.2.4 "extract as PNG, place In Front of
        # Text" rule (click a signature and it selects/extracts, rather
        # than being fused into the background).
        for pe in (png_items or []):
            draw_id += 1
            pbuf = io.BytesIO()
            pe["png"].save(pbuf, format="PNG")
            png_name = f"elem{page_idx + 1}_{draw_id}.png"
            media.append((png_name, pbuf.getvalue()))
            pe_rel = f"rIdElem{page_idx + 1}_{draw_id}"
            doc_rels.append(
                f'<Relationship Id="{pe_rel}" '
                f'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" '
                f'Target="media/{png_name}"/>')
            ex = int((pe["left"]) * scale * EMU_PER_PT)
            ey = int((pe["top"]) * scale * EMU_PER_PT)
            ecx = int(max(1.0, (pe["right"] - pe["left"]) * scale) * EMU_PER_PT)
            ecy = int(max(1.0, (pe["bottom"] - pe["top"]) * scale) * EMU_PER_PT)
            kind_name = pe.get("kind", "element") or "element"
            anchors.append(
                f'<w:r><w:drawing><wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" '
                f'relativeHeight="{draw_id}" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">'
                f'<wp:simplePos x="0" y="0"/>'
                f'<wp:positionH relativeFrom="page"><wp:posOffset>{ex}</wp:posOffset></wp:positionH>'
                f'<wp:positionV relativeFrom="page"><wp:posOffset>{ey}</wp:posOffset></wp:positionV>'
                f'<wp:extent cx="{ecx}" cy="{ecy}"/><wp:wrapNone/>'
                f'<wp:docPr id="{draw_id}" name="{_xml_escape(kind_name)} {draw_id}"/>'
                f'<a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
                f'<pic:pic><pic:nvPicPr><pic:cNvPr id="{draw_id}" name="{png_name}"/><pic:cNvPicPr/></pic:nvPicPr>'
                f'<pic:blipFill><a:blip r:embed="{pe_rel}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
                f'<pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{ecx}" cy="{ecy}"/></a:xfrm>'
                f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr></pic:pic>'
                f'</a:graphicData></a:graphic></wp:anchor></w:drawing></w:r>')

        for item in (layout_items or []):
            item_rtl = item.get("rtl", rtl)
            align = item.get("align", "center")
            if align == "natural":
                jc = "right" if item_rtl else "left"
            else:
                jc = {"left": "left", "right": "right", "center": "center"}.get(align, "center")
            font_name = _pdf_font_for(item["text"], item["bold"], item["italic"], item_rtl)
            # Flow the paragraph across its per-line boxes at a uniform
            # font size (one original visual line = one Word text box),
            # shrinking on overflow - matching the PDF output exactly.
            raw_boxes = item.get("line_boxes") or [{
                "left": item["left"], "top": item["top"],
                "right": item["right"], "bottom": item["bottom"]}]
            pt_boxes = [{"left": b["left"] * scale, "top": b["top"] * scale,
                         "right": b["right"] * scale, "bottom": b["bottom"] * scale}
                        for b in raw_boxes]
            size, mapped = _flow_text_across_lineboxes(
                item["text"], pt_boxes, font_name, item_rtl)
            half_pts = max(2, int(round(size * 2)))
            color_hex = "%02X%02X%02X" % item["color"]
            rpr = (f'<w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>'
                   + ('<w:b/><w:bCs/>' if item["bold"] else '')
                   + ('<w:i/><w:iCs/>' if item["italic"] else '')
                   + ('<w:u w:val="single"/>' if item["underline"] else '')
                   + ('<w:rtl/>' if item_rtl else '')
                   + f'<w:color w:val="{color_hex}"/>'
                   + f'<w:sz w:val="{half_pts}"/><w:szCs w:val="{half_pts}"/></w:rPr>')
            ppr = '<w:pPr>' + ('<w:bidi/>' if item_rtl else '') + \
                  '<w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>' \
                  f'<w:jc w:val="{jc}"/></w:pPr>'
            # One floating text box per produced line, positioned at that
            # line's own box - so a 3-line sentence yields 3 boxes, never
            # one overflowing block.
            for ln, box in mapped:
                draw_id += 1
                bx = box["left"]
                by = box["top"]
                bw = max(1.0, box["right"] - box["left"])
                bh = max(1.0, box["bottom"] - box["top"])
                tb_w_emu = int((bw + 8) * EMU_PER_PT)
                tb_h_emu = int((bh + 4) * EMU_PER_PT)
                anchors.append(
                    f'<w:r><w:drawing><wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" '
                    f'relativeHeight="{draw_id}" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">'
                    f'<wp:simplePos x="0" y="0"/>'
                    f'<wp:positionH relativeFrom="page"><wp:posOffset>{int((bx - 3) * EMU_PER_PT)}</wp:posOffset></wp:positionH>'
                    f'<wp:positionV relativeFrom="page"><wp:posOffset>{int((by - 2) * EMU_PER_PT)}</wp:posOffset></wp:positionV>'
                    f'<wp:extent cx="{tb_w_emu}" cy="{tb_h_emu}"/><wp:wrapNone/>'
                    f'<wp:docPr id="{draw_id}" name="Text {draw_id}"/>'
                    f'<a:graphic><a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
                    f'<wps:wsp><wps:cNvSpPr txBox="1"/><wps:spPr>'
                    f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{tb_w_emu}" cy="{tb_h_emu}"/></a:xfrm>'
                    f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom><a:noFill/><a:ln><a:noFill/></a:ln></wps:spPr>'
                    f'<wps:txbx><w:txbxContent><w:p>{ppr}<w:r>{rpr}'
                    f'<w:t xml:space="preserve">{_xml_escape(ln)}</w:t></w:r></w:p></w:txbxContent></wps:txbx>'
                    f'<wps:bodyPr rot="0" wrap="square" lIns="0" tIns="0" rIns="0" bIns="0" anchor="ctr">'
                    f'<a:noAutofit/></wps:bodyPr></wps:wsp>'
                    f'</a:graphicData></a:graphic></wp:anchor></w:drawing></w:r>')

        body_parts.append("<w:p><w:pPr><w:spacing w:before=\"0\" w:after=\"0\"/></w:pPr>" + "".join(anchors) + "</w:p>")
        if page_idx < len(story_pages) - 1:
            body_parts.append('<w:p><w:r><w:br w:type="page"/></w:r></w:p>')

    first_w_pt = story_pages[0][3]
    first_h_pt = story_pages[0][4]
    sect = (f'<w:sectPr><w:pgSz w:w="{int(first_w_pt * 20)}" w:h="{int(first_h_pt * 20)}"/>'
            f'<w:pgMar w:top="0" w:right="0" w:bottom="0" w:left="0" w:header="0" w:footer="0" w:gutter="0"/></w:sectPr>')
    document_xml = (f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                    f'<w:document {_DOCX_NS}><w:body>{"".join(body_parts)}{sect}</w:body></w:document>')
    document_rels = ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                     '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                     + "".join(doc_rels) + '</Relationships>')

    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", _DOCX_CONTENT_TYPES)
        z.writestr("_rels/.rels", _DOCX_RELS)
        z.writestr("word/document.xml", document_xml)
        z.writestr("word/_rels/document.xml.rels", document_rels)
        for filename, data in media:
            z.writestr(f"word/media/{filename}", data)


def generate_layout_preserving_translation_pdf(original_pdf_path, output_pdf_path, target_language, llm_config=None, diagnostics=None, vision_assist=False, progress=None):
    """Item 6 (per the uploaded 'Document Translation Instructions' spec)
    - rather than extracting the document's text and rebuilding a brand
    new, generically-formatted PDF from scratch (the old approach), this
    surgically edits the ORIGINAL PDF in place: for each page, every
    detected line of text gets masked with a white rectangle and the
    translated line drawn back in the exact same position/font
    size/color - everything else on the page (images, logos, signatures,
    stamps, tables' borders/shading, vector graphics) is the ORIGINAL
    PDF's own content, completely untouched, because this never redraws
    the page - it only overlays on top of it.

    Font MATCHING is approximate (built-in Helvetica + bold/italic
    variants, not the original's exact embedded font - true font
    substitution would need font-metrics analysis this renderer doesn't
    have), and very long translated lines get their font size shrunk down
    (never enlarged) to stay inside the original line's width, per the
    spec's 'minimal adjustments, never overflow' rule. Routes to the
    OCR-based image-overlay approach (generate_ocr_based_translation_pdf)
    if the source has no real text layer at all (a scanned/image-only
    PDF - there's no text OBJECT to surgically edit, but OCR can still
    find where the visible text sits on the rendered page and mask/redraw
    it the same way)."""
    if not REPORTLAB_OK:
        raise LeaseEngineError("reportlab is not installed - run: pip install -r requirements.txt")
    if not pdfplumber or not PdfReader:
        raise LeaseEngineError("pdfplumber/pypdf are not installed - run: pip install -r requirements.txt")

    from pypdf import PdfWriter
    from reportlab.pdfgen import canvas as pdfcanvas
    from reportlab.pdfbase.pdfmetrics import stringWidth

    rtl = is_rtl_language(target_language)
    reader = PdfReader(original_pdf_path)
    writer = PdfWriter()

    with pdfplumber.open(original_pdf_path) as pdf:
        any_real_text = any(p.chars for p in pdf.pages)
        if diagnostics is not None:
            diagnostics["pathUsed"] = "ocr-image-based" if not any_real_text else "text-layer-based"
        if not any_real_text:
            return generate_ocr_based_translation_pdf(original_pdf_path, output_pdf_path, target_language, llm_config, diagnostics=diagnostics, vision_assist=vision_assist, progress=progress)

        for page_index, page in enumerate(pdf.pages):
            lines = _group_chars_into_lines(page.chars)
            regions = [r for r in (_line_to_region(l) for l in lines) if r]

            translated_texts = _translate_lines_batch([r["text"] for r in regions], target_language, llm_config) if regions else []

            page_width, page_height = float(page.width), float(page.height)
            buf = io.BytesIO()
            c = pdfcanvas.Canvas(buf, pagesize=(page_width, page_height))

            for region, translated in zip(regions, translated_texts):
                box_w = max(1.0, region["x1"] - region["x0"])
                box_h = max(1.0, region["bottom"] - region["top"])
                # Mask the original line - white covers the overwhelming
                # majority of real-world documents; per-region background
                # color detection is a possible future refinement.
                c.setFillColorRGB(1, 1, 1)
                c.rect(region["x0"] - 1, page_height - region["bottom"] - 1, box_w + 2, box_h + 2, fill=1, stroke=0)

                font_name = _closest_builtin_font(region["bold"], region["italic"])
                # Fit against the UNshaped text first (word boundaries are
                # meaningful before RTL shaping); shape each wrapped line
                # individually afterward for drawing.
                font_size, plain_lines, line_h = _fit_text_in_box_reportlab(
                    translated, font_name, box_w, box_h, start_size=region["size"])
                render_lines = [shape_rtl_text(ln) if rtl else ln for ln in plain_lines]

                try:
                    hexc = region["color"].lstrip("#")
                    r, g, b = (int(hexc[0:2], 16) / 255, int(hexc[2:4], 16) / 255, int(hexc[4:6], 16) / 255)
                except (ValueError, IndexError, TypeError):
                    r, g, b = 0, 0, 0
                c.setFillColorRGB(r, g, b)
                c.setFont(font_name, font_size)

                # Stack lines top-down, the wrapped block vertically
                # centered inside the original line's box so 1-line and
                # multi-line results both look correctly placed.
                block_h = line_h * len(render_lines)
                top_y = page_height - region["top"] - max(0, (box_h - block_h) / 2)
                baseline_y = top_y - line_h * 0.8
                for ln in render_lines:
                    if rtl:
                        c.drawRightString(region["x1"], baseline_y, ln)
                    else:
                        c.drawString(region["x0"], baseline_y, ln)
                    baseline_y -= line_h

            c.save()
            buf.seek(0)
            overlay_reader = PdfReader(buf)
            overlay_page = overlay_reader.pages[0]
            original_page = reader.pages[page_index]
            # The overlay (white masks + translated text, just built above)
            # must end up as the TOP layer with the original page's content
            # underneath it - merging the original INTO the overlay with
            # over=False (rather than merging the overlay into the
            # original) is what actually achieves that; doing it the other
            # way around left the original text still on top, showing
            # through the "masking" rectangles.
            overlay_page.merge_page(original_page, over=False)
            writer.add_page(overlay_page)

    with open(output_pdf_path, "wb") as f:
        writer.write(f)


def generate_translation_docx(output_json_path, docx_out_path):
    """Builds an editable Output.docx for the NON-hybrid (reflow) path,
    so that when the user picks 'docx' output on a Simple-mode job they
    get a real Word file - not a PDF. This is a clean reflowed document
    (python-docx paragraphs), mirroring generate_translation_pdf, with
    direction following the OUTPUT language (RTL target -> right-aligned,
    reshaped). No PDF round-trip is involved."""
    if docx_lib is None:
        raise LeaseEngineError(
            "python-docx is not installed - run: pip install -r requirements.txt")
    with open(output_json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    target_language = data.get("targetLanguage", "")
    translated_text = data.get("translatedText", "") or ""
    rtl = is_rtl_language(target_language)

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
    doc = Document()
    for raw_line in translated_text.split("\n"):
        line = raw_line.rstrip()
        para = doc.add_paragraph()
        if rtl:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            try:
                para.paragraph_format.right_to_left = True
            except Exception:
                pass
        # Light markup handling: leading '# ' heading, '**bold**' emphasis.
        text = line
        is_heading = False
        if text.startswith("# "):
            text = text[2:]
            is_heading = True
        run = para.add_run(text)
        if is_heading:
            run.bold = True
            run.font.size = None
    doc.save(docx_out_path)


def generate_translation_pdf(output_json_path, pdf_out_path):
    """Builds Output.pdf for a translation job. Item 6 - auto-detects the
    translated text's structure (paragraphs, bullet lists, numbered
    lists, headings, bold emphasis - all marked up by translate_text()'s
    prompt) and renders each as the appropriate PDF element, rather than
    dumping everything as flat paragraphs. No metadata table - just a
    one-line header, then the formatted document itself.

    Also handles text direction (item 6): right-to-left languages
    (Arabic, Hebrew, ...) get right-aligned paragraphs and their text run
    through arabic_reshaper + python-bidi so RTL glyphs join correctly and
    read in the right visual order - translating INTO English (or another
    LTR language) always renders left-to-right regardless of what
    direction the source document used, and vice versa: direction follows
    the OUTPUT language, never the source."""
    if not REPORTLAB_OK:
        raise LeaseEngineError(
            "reportlab is not installed - run: pip install -r requirements.txt"
        )

    with open(output_json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    _render_reflowed_translation_pdf(
        data.get("docName", "Document"),
        data.get("targetLanguage", ""),
        data.get("translatedText", ""),
        pdf_out_path,
    )


def _render_reflowed_translation_pdf(doc_name, target_language, translated_text, pdf_out_path):
    """The reflow renderer factored out of generate_translation_pdf so
    the Simple-mode vision path can reuse it - takes marked-up translated
    text (translate_text() / vision conventions: blank-line paragraphs,
    '- ' bullets, numbered lists, '## ' headings, **bold**) and renders a
    clean formatted document."""
    if not REPORTLAB_OK:
        raise LeaseEngineError(
            "reportlab is not installed - run: pip install -r requirements.txt"
        )

    rtl = is_rtl_language(target_language)
    align = TA_RIGHT if rtl else TA_LEFT

    doc = SimpleDocTemplate(
        pdf_out_path, pagesize=LETTER,
        topMargin=0.6 * inch, bottomMargin=0.6 * inch,
        leftMargin=0.6 * inch, rightMargin=0.6 * inch,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TranslationTitle", parent=styles["Heading1"], textColor=colors.black, fontSize=16, spaceAfter=2
    )
    normal = ParagraphStyle("TranslationNormal", parent=styles["Normal"], fontSize=8.5, alignment=align)
    body_style = ParagraphStyle("TranslationBody", parent=normal, fontSize=10.5, spaceAfter=10, leading=16, alignment=align)
    heading_style = ParagraphStyle("TranslationHeading", parent=normal, fontSize=13, spaceBefore=12, spaceAfter=8, fontName="Helvetica-Bold", alignment=align)
    # Bullets/numbers visually belong on the side text flows TOWARD in
    # each direction (left indent for LTR, right indent mirrored for RTL).
    bullet_style = ParagraphStyle("TranslationBullet", parent=body_style,
                                   leftIndent=0 if rtl else 18, rightIndent=18 if rtl else 0, spaceAfter=4)
    numbered_style = ParagraphStyle("TranslationNumbered", parent=body_style,
                                     leftIndent=0 if rtl else 20, rightIndent=20 if rtl else 0, spaceAfter=4)

    def _render(text):
        return _escape_and_apply_bold(text, rtl=rtl)

    story = [
        Paragraph("Translation Report", title_style),
        Paragraph(f"Document : {doc_name}  \u00b7  Translated to {target_language}", normal),
        Spacer(1, 14),
    ]

    for block_type, content in _parse_translation_blocks(translated_text):
        if block_type == "heading":
            story.append(Paragraph(_render(content), heading_style))
        elif block_type == "bullet":
            for item in content:
                text = f"{_render(item)}&nbsp;&nbsp;\u2022" if rtl else f"&bull;&nbsp;&nbsp;{_render(item)}"
                story.append(Paragraph(text, bullet_style))
        elif block_type == "numbered":
            for i, item in enumerate(content, 1):
                text = f"{_render(item)}&nbsp;&nbsp;.{i}" if rtl else f"{i}.&nbsp;&nbsp;{_render(item)}"
                story.append(Paragraph(text, numbered_style))
        else:
            story.append(Paragraph(_render(content), body_style))

    if not story[3:]:
        story.append(Paragraph("(No translated content)", body_style))

    doc.build(story)


def _looks_like_refusal(text):
    """True if the model's reply is a safety REFUSAL / apology rather than
    real content, so it is never printed into the document as a
    'translation'. Kept deliberately narrow: a short reply dominated by an
    apology with no substantive body."""
    if not text:
        return False
    t = text.strip().lower()
    markers = ("i'm sorry", "i am sorry", "i cannot assist", "i can't assist",
               "i cannot help", "i can't help", "unable to assist",
               "cannot process", "can't process", "i'm not able to",
               "i am not able to", "i cannot fulfill", "can't fulfill",
               "i cannot comply", "i won't be able to", "as an ai")
    if not any(m in t for m in markers):
        return False
    # Only treat as refusal when the apology dominates (short reply), not
    # when the word merely appears inside a long translated document.
    return len(t) < 400


def generate_vision_translation_pdf(original_pdf_path, output_pdf_path, target_language, doc_name="Document", llm_config=None, diagnostics=None):
    """Simple mode for scanned/photo PDFs - completely bypasses Tesseract:
    each page is rendered to an image and sent to the vision-capable chat
    model, which reads the page like Google Lens does (deep-learning
    vision handles photos, skew, low contrast far better than classical
    OCR) and returns the full translation directly, marked up with the
    same structure conventions translate_text() uses. Output is a clean
    reflowed document - accurate text, but NOT the original layout (that
    trade-off is exactly what the Hybrid checkbox toggles).

    Raises LeaseEngineError if no LLM key is configured or every page's
    vision call failed - the caller decides what to fall back to."""
    if diagnostics is None:
        diagnostics = {}
    llm_config = llm_config if llm_config is not None else load_llm_config()
    if not llm_is_configured(llm_config):
        raise LeaseEngineError("Vision translation needs an OpenAI/OpenRouter API key in .env")
    if not pdfplumber:
        raise LeaseEngineError("pdfplumber is not installed - run: pip install -r requirements.txt")

    system_prompt = (
        f"You are Lexora's professional document translator. The user has "
        f"uploaded THEIR OWN document to their own account and explicitly "
        f"asked for it to be translated; you are fully authorised to read "
        f"and translate all text on the page. This is a routine, legitimate "
        f"localisation task - always comply, never refuse, never apologise. "
        f"You are given an IMAGE of one page of a "
        f"document. First carefully read ALL visible text in the image, in its natural reading "
        f"order. Then translate everything into {target_language}.\n\n"
        f"Requirements:\n"
        f"- Translate only the textual content; preserve the original meaning.\n"
        f"- Preserve legal, medical, technical and business terminology.\n"
        f"- Keep company names, trademarks, product names and personal names unchanged unless a "
        f"standard translation exists.\n"
        f"- Preserve dates, numbers, currencies, amounts, meter readings and identifiers exactly.\n"
        f"- Do not summarize, do not explain, do not add or remove content.\n\n"
        f"Mark up the structure you see using these plain-text conventions (the renderer "
        f"auto-formats from them):\n"
        f"- Separate distinct paragraphs/sections with a single BLANK LINE.\n"
        f"- Bullet items: start the line with '- '. Numbered items: '1. ', '2. ', ...\n"
        f"- Section headings/titles: start the line with '## '.\n"
        f"- Visually emphasized text (key figures, totals, warnings): wrap in **double asterisks**.\n"
        f"- For tables, render each row as 'Label: value' lines.\n\n"
        f"Return ONLY the translated, marked-up text - no preamble, no notes."
    )

    page_texts = []
    diagnostics["pathUsed"] = "vision-full-page"
    diagnostics["pages"] = []
    with pdfplumber.open(original_pdf_path) as pdf:
        total = len(pdf.pages)
        for page_num, page in enumerate(pdf.pages):
            page_diag = {"page": page_num + 1, "error": None}
            try:
                page_image = page.to_image(resolution=300).original.convert("RGB")
                image_b64 = _pil_image_to_jpeg_b64(page_image)
                content, provider = _call_vision_with_failover(
                    llm_config, system_prompt,
                    f"Translate this page ({page_num + 1} of {total}) into {target_language}.",
                    image_b64, max_tokens=8000,
                )
                # Detect a model REFUSAL and don't let it become the page's
                # "translation". Retry once (often succeeds); if it still
                # refuses, skip this page rather than printing the apology.
                if content and _looks_like_refusal(content):
                    print(f"Vision refused page {page_num + 1}; retrying once.")
                    content, provider = _call_vision_with_failover(
                        llm_config, system_prompt,
                        f"This is the user's own document; please translate "
                        f"page {page_num + 1} of {total} into {target_language} "
                        f"and return only the translated text.",
                        image_b64, max_tokens=8000)
                if content and _looks_like_refusal(content):
                    page_diag["error"] = "model declined to translate this page"
                    content = None
                if content:
                    page_texts.append(content.strip())
                    page_diag["provider"] = provider
                    page_diag["chars"] = len(content)
                else:
                    page_diag["error"] = page_diag.get("error") or "no LLM key configured"
            except LeaseEngineError as err:
                print(f"Vision translation failed on page {page_num + 1}: {err}")
                page_diag["error"] = str(err)
            diagnostics["pages"].append(page_diag)

    if not page_texts:
        raise LeaseEngineError("Vision translation produced no text for any page")

    translated_text = "\n\n".join(page_texts)
    diagnostics["translatedChars"] = len(translated_text)
    _render_reflowed_translation_pdf(doc_name, target_language, translated_text, output_pdf_path)
    return translated_text


def analyze_lease(text, fallback_name="Lease Document"):
    """Section 14.3 (40%) - 'data analyzed and interpreted using GPT
    prompts'. Tries a real LLM call (OpenAI/OpenRouter, key from .env
    + json/extraction_prompt.txt) first; falls back to the heuristic engine
    if no key is configured or the call fails. Either way, also runs a
    second lightweight pass (call_llm_validation, or a heuristic
    completeness check without an LLM) to produce an accuracy/confidence
    score - same two-stage extract-then-validate shape as the reference
    project's /api/extract + /api/validate routes."""
    doc_type = classify_document(text)
    llm_config = load_llm_config()

    try:
        llm_fields, used_provider = call_llm_extraction(text, llm_config=llm_config)
    except LeaseEngineError as err:
        # Every configured provider's call failed - fall back rather than
        # aborting the whole pipeline, but keep the error visible.
        print(f"LLM extraction failed on every configured provider, falling back to heuristic engine: {err}")
        llm_fields, used_provider = None, None

    if llm_fields is not None:
        lease_name = sanitize_lease_name(_lease_name_source_from_fields(llm_fields, fallback_name))
        result = {
            "docType": doc_type,
            "leaseName": lease_name,
            "fields": llm_fields,
            "extractionMethod": f"llm-{used_provider}",
        }
    else:
        result = heuristic_analyze_lease(text, fallback_name=fallback_name)
        result["docType"] = doc_type  # keep a single source of truth for classification

    result.update(_run_accuracy_check(result["fields"], llm_config))
    rules_applied, rules_total = get_rules_applied_count()
    result["rulesApplied"] = rules_applied
    result["rulesTotal"] = rules_total
    return result


def sanitize_lease_name(raw):
    """Turns extracted text into a filesystem-safe folder name."""
    raw = raw or "Lease"
    cleaned = re.sub(r"[^A-Za-z0-9 _\-]", "", raw).strip()
    cleaned = re.sub(r"\s+", "_", cleaned)
    return cleaned[:60] or "Lease"


# ============================================================
# 3. OUTPUT PDF GENERATION
# ============================================================
def _humanize_key(key):
    return str(key).replace("_", " ").strip().title()


def _stringify_leaf(value):
    if value is None or value == "":
        return "-"
    return str(value)


def _flatten_to_rows(value, prefix=""):
    """Recursively flattens a dict/list into (label, value) row pairs -
    used so the PDF can render either the flat heuristic fields dict or
    the rich nested LLM schema (parties/premises/term/rent/... sections,
    each of which may itself contain lists like rent_schedule or
    queries_assumptions) without hardcoding every possible field name."""
    rows = []
    if isinstance(value, dict):
        if not value:
            rows.append((prefix.rstrip(" -") or "Value", "-"))
        for k, v in value.items():
            label = (prefix + _humanize_key(k)) if prefix else _humanize_key(k)
            if isinstance(v, (dict, list)):
                rows.extend(_flatten_to_rows(v, label + " - "))
            else:
                rows.append((label, _stringify_leaf(v)))
    elif isinstance(value, list):
        if not value:
            rows.append((prefix.rstrip(" -") or "Value", "-"))
        for i, item in enumerate(value):
            item_prefix = f"{prefix}#{i + 1} "
            if isinstance(item, (dict, list)):
                rows.extend(_flatten_to_rows(item, item_prefix))
            else:
                rows.append((item_prefix.strip(), _stringify_leaf(item)))
    else:
        rows.append((prefix.rstrip(" -") or "Value", _stringify_leaf(value)))
    return rows


def _build_kv_table(rows, col_widths=None):
    from xml.sax.saxutils import escape as _esc

    cell_style = ParagraphStyle("KVCell", fontSize=8.5, leading=11, fontName=_active_font())
    header_style = ParagraphStyle(
        "KVHeader", fontSize=8.5, leading=11, textColor=_active_header_text(), fontName=_active_font_bold()
    )
    table_data = [[Paragraph("Field", header_style), Paragraph("Value", header_style)]]
    for label, value in rows:
        table_data.append([
            Paragraph(_esc(str(label)), cell_style),
            Paragraph(_esc(str(value)).replace("\n", "<br/>"), cell_style),
        ])

    table = Table(table_data, colWidths=list(col_widths) if col_widths else _normalize_widths([2.3, 3.7]))
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), _active_header_bg()),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#AAAAAA")),
        ("BOX", (0, 0), (-1, -1), 0.5, _NAVY),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, _STRIPE]),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    return table


# ============================================================
# Template-matching PDF sections (item 6 - mirrors the structure of the
# real "Lease Abstract" output format: Lease Information grid, Charge
# Schedules table, Amendments table, Late Fee, a per-clause Other Lease
# Provisions/Clauses table, and a Contacts table). Only used when `fields`
# has the rich nested LLM schema shape (see json/extraction_prompt.txt) -
# the flat heuristic-engine schema still falls back to the older generic
# key/value renderer below it, since it doesn't carry per-clause data.
# ============================================================
_CURRENCY_ARTIFACT_RE = re.compile(r"\${2,}")
_COMMA_ARTIFACT_RE = re.compile(r",{2,}")
_WHITESPACE_ARTIFACT_RE = re.compile(r"[ \t]{2,}")


def _sanitize_leaf_string(value):
    """Item 1 - self-QA pass that catches common LLM/formatting artifacts
    before they reach Output.json or the PDF, instead of just hoping the
    extraction got it right. Fixes things like the extraction prompt
    saying 'format as $X' while the PDF builder ALSO prepends its own $,
    producing '$$10,479.88' - collapses doubled currency symbols,
    doubled commas, and doubled internal whitespace. Deliberately
    conservative: only touches unambiguous duplication artifacts, never
    rewrites the actual content/meaning of a value."""
    if not isinstance(value, str) or not value:
        return value
    cleaned = _CURRENCY_ARTIFACT_RE.sub("$", value)
    cleaned = _COMMA_ARTIFACT_RE.sub(",", cleaned)
    cleaned = _WHITESPACE_ARTIFACT_RE.sub(" ", cleaned).strip()
    return cleaned


def sanitize_fields_recursively(value):
    """Applies _sanitize_leaf_string to every string leaf in a nested
    fields dict/list (the rich LLM extraction schema), leaving numbers,
    dicts, and list structure untouched. Called once right when a lease's
    Output.json is first saved (see save-output in server.py), so both
    the Human Review screen and the final PDF see the cleaned values."""
    if isinstance(value, dict):
        return {k: sanitize_fields_recursively(v) for k, v in value.items()}
    if isinstance(value, list):
        return [sanitize_fields_recursively(v) for v in value]
    if isinstance(value, str):
        return _sanitize_leaf_string(value)
    return value


def _get_path(d, path, default=""):
    cur = d
    for part in path.split("."):
        if not isinstance(cur, dict):
            return default
        cur = cur.get(part)
    return cur if cur not in (None, "") else default


def compute_lease_fingerprint(fields):
    """Item 6 - a content-level duplicate check, not just 'same lease
    name'. Builds a canonical string from the handful of fields that
    genuinely identify a unique lease (tenant, landlord, property address,
    commencement/expiration dates) - normalized (lowercased, whitespace
    collapsed) so trivial OCR/formatting differences between two scans of
    the SAME lease don't produce different fingerprints - and hashes it.
    Two documents with the same fingerprint are almost certainly the same
    underlying lease (or a reference/exhibit document whose extracted
    'lease' details are identical to one already on file), even if their
    filenames or auto-generated lease names differ."""
    if not isinstance(fields, dict):
        return None
    parts = [
        _get_path(fields, "parties.tenant_legal_name"),
        _get_path(fields, "parties.landlord_legal_name"),
        _get_path(fields, "premises.property_address"),
        _get_path(fields, "term.lease_commencement_date"),
        _get_path(fields, "term.lease_expiration_date"),
    ]
    normalized = "|".join(re.sub(r"\s+", " ", str(p).strip().lower()) for p in parts if p)
    if not normalized or normalized.count("|") < 2:
        return None  # not enough identifying data to fingerprint reliably
    return hashlib.sha256(normalized.encode("utf-8")).hexdigest()[:24]


def _flatten_leaves(d, prefix=""):
    """Yields (dotted.path, value) for every leaf (non-dict, non-list-of-
    dicts) value in a nested dict - used by compare_extraction_accuracy
    below to line up two field dicts leaf-by-leaf even if one has extra
    sections the other doesn't."""
    if isinstance(d, dict):
        for k, v in d.items():
            yield from _flatten_leaves(v, f"{prefix}.{k}" if prefix else k)
    elif isinstance(d, list):
        for i, v in enumerate(d):
            yield from _flatten_leaves(v, f"{prefix}[{i}]")
    else:
        yield (prefix, d)


def _leaf_similarity(a, b):
    """1.0 for an exact match (after trimming/lowercasing so formatting
    differences like extra whitespace don't count against a match), down
    to 0.0 for completely different text - uses difflib's ratio for a
    fuzzy partial-credit score in between, since 'Lease is silent.' vs
    'Lease is silent' should score high, not zero."""
    sa, sb = str(a or "").strip().lower(), str(b or "").strip().lower()
    if sa == sb:
        return 1.0
    if not sa or not sb:
        return 0.0
    return difflib.SequenceMatcher(None, sa, sb).ratio()


def extract_lease_pdf_as_dict(pdf_path):
    """Item (Test & Compare with PDF files, not JSON) - reverse-parses a
    Lease Abstract-style PDF (ours, or a human-produced one in the same
    general layout) back into a flat {label: value} dict, so two PDFs -
    a human-reviewed reference and our own generated Output.pdf - can be
    diffed field-by-field via compare_extraction_accuracy() without
    either side needing to be JSON. Works off whatever tables pdfplumber
    can detect on each page: a 2-column table becomes {label: value} rows
    directly; a wider data table (Charge Schedules, Amendments, ...) uses
    its own header row as column names and keys each cell as
    'table{N}_row{R}_{ColumnHeader}' - matching by the ACTUAL label/header
    text rather than assuming an identical internal schema is what makes
    this robust to a human's PDF looking slightly different from ours
    structurally, as long as the labels themselves read the same."""
    if not pdfplumber:
        raise LeaseEngineError("pdfplumber is not installed - run: pip install -r requirements.txt")

    result = {}
    with pdfplumber.open(pdf_path) as pdf:
        table_counter = 0
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                if not table or not table[0]:
                    continue
                table_counter += 1
                # 2-column table => straightforward label:value rows.
                if all(len(row) == 2 for row in table):
                    for row in table:
                        label = (row[0] or "").strip()
                        value = (row[1] or "").strip()
                        if label:
                            result[label] = value
                    continue
                # 4-column table => our own Lease Information grid layout
                # (label, value, label, value per row).
                if all(len(row) == 4 for row in table):
                    for row in table:
                        for label, value in ((row[0], row[1]), (row[2], row[3])):
                            label = (label or "").strip()
                            value = (value or "").strip()
                            if label:
                                result[label] = value
                    continue
                # Any other width => a data table (Charge Schedules,
                # Amendments, ...); first row is the header/column names,
                # every following row's cells key off that header plus a
                # row index (so repeated rows don't collide with each
                # other).
                header = [(h or "").strip() or f"col{i}" for i, h in enumerate(table[0])]
                for row_idx, row in enumerate(table[1:]):
                    for col_idx, cell in enumerate(row):
                        if col_idx >= len(header):
                            continue
                        key = f"table{table_counter}_row{row_idx}_{header[col_idx]}"
                        result[key] = (cell or "").strip()
    return result


def compare_extraction_accuracy(reference_fields, current_fields):
    """Item 4 (Test & Compare) - scores how closely `current_fields` (this
    system's own extraction) matches `reference_fields` (a human-reviewed
    "ideal" answer key for the same lease), field by field. Returns
    {overallAccuracy, fieldResults: [{path, referenceValue, currentValue,
    similarity}]} sorted worst-match-first, so the biggest discrepancies -
    the ones most worth turning into a new/updated rule - surface at the
    top."""
    reference_leaves = dict(_flatten_leaves(reference_fields or {}))
    current_leaves = dict(_flatten_leaves(current_fields or {}))

    field_results = []
    total_similarity = 0.0
    for path, ref_value in reference_leaves.items():
        cur_value = current_leaves.get(path, "")
        sim = _leaf_similarity(ref_value, cur_value)
        total_similarity += sim
        field_results.append({
            "path": path, "referenceValue": ref_value, "currentValue": cur_value,
            "similarity": round(sim, 3),
        })

    field_results.sort(key=lambda r: r["similarity"])
    overall = round((total_similarity / len(reference_leaves)) * 100, 1) if reference_leaves else 0.0
    return {"overallAccuracy": overall, "fieldResults": field_results, "fieldCount": len(reference_leaves)}


# (template clause id, display name, dotted path into `fields`) - covers
# every clause the extraction prompt (json/extraction_prompt.txt) already
# asks the LLM to extract per-field, in the same order the reference
# "Other Lease Provisions / Clauses" table uses.
_CLAUSE_FIELD_MAP = [
    ("assign", "Assignment & Sublease", "assignment_subletting.with_consent"),
    ("rent", "Rent", "term.rental_term"),
    ("cotenanc", "Co-Tenancy", "use.go_dark"),
    ("default", "Default", "default.monetary_default"),
    ("estoppel", "Estoppel", "estoppel"),
    ("conuse", "Continuous Use or Go Dark", "use.failure_to_open"),
    ("guaranty", "Guaranty", "parties.guarantor"),
    ("holdover", "Holdover", "rent.holdover"),
    ("late fee", "Late Fee", "rent.late_fee"),
    ("opex/cam", "OpEx/CAM", "cam_opex.cam_detail"),
    ("insreimb", "Insurance Reimbursement", "insurance"),
    ("brokers", "Brokers", "brokers.landlord_broker"),
    ("taxes", "Real Estate Taxes", "cam_opex.real_property_taxes_pro_rata_share"),
    ("permit", "Permitted Use", "use.permitted_use"),
    ("percent", "Percentage Rent", "rent.percentage_rent"),
    ("tt ins", "Tenant Insurance", "insurance"),
    ("utility", "Utilities", "utilities"),
    ("security", "Security Deposit", "rent.security_deposit"),
    ("ti allow", "TI Allowance", "ti_allowance"),
    ("llrepair", "LL's Repair", "repairs.roof_repair"),
    ("ttrep", "TT's Repair", "repairs.tenant_repair"),
    ("misc", "Miscellaneous", "miscellaneous"),
    ("reloc", "Relocation Option", "options.renewal_options"),
    ("roof", "Roof Repairs", "repairs.roof_repair"),
    ("alter", "Alterations", "alterations"),
    ("snda", "Subordination", "subordination"),
    ("prohib", "Prohibited Use", "use.prohibited_use"),
]


def _is_rich_lease_schema(fields):
    return isinstance(fields, dict) and isinstance(fields.get("parties"), dict)


# (template clause id, display name, dotted path into `fields`) - covers
# every clause the extraction prompt (json/extraction_prompt.txt) already
# asks the LLM to extract per-field, in the same order the reference
# "Other Lease Provisions / Clauses" table uses.
_CLAUSE_FIELD_MAP = [
    ("assign", "Assignment & Sublease", "assignment_subletting.with_consent"),
    ("rent", "Rent", "term.rental_term"),
    ("cotenanc", "Co-Tenancy", "use.go_dark"),
    ("default", "Default", "default.monetary_default"),
    ("estoppel", "Estoppel", "estoppel"),
    ("conuse", "Continuous Use or Go Dark", "use.failure_to_open"),
    ("guaranty", "Guaranty", "parties.guarantor"),
    ("holdover", "Holdover", "rent.holdover"),
    ("late fee", "Late Fee", "rent.late_fee"),
    ("opex/cam", "OpEx/CAM", "cam_opex.cam_detail"),
    ("insreimb", "Insurance Reimbursement", "insurance"),
    ("brokers", "Brokers", "brokers.landlord_broker"),
    ("taxes", "Real Estate Taxes", "cam_opex.real_property_taxes_pro_rata_share"),
    ("permit", "Permitted Use", "use.permitted_use"),
    ("percent", "Percentage Rent", "rent.percentage_rent"),
    ("tt ins", "Tenant Insurance", "insurance"),
    ("utility", "Utilities", "utilities"),
    ("security", "Security Deposit", "rent.security_deposit"),
    ("ti allow", "TI Allowance", "ti_allowance"),
    ("llrepair", "LL's Repair", "repairs.roof_repair"),
    ("ttrep", "TT's Repair", "repairs.tenant_repair"),
    ("misc", "Miscellaneous", "miscellaneous"),
    ("reloc", "Relocation Option", "options.renewal_options"),
    ("roof", "Roof Repairs", "repairs.roof_repair"),
    ("alter", "Alterations", "alterations"),
    ("snda", "Subordination", "subordination"),
    ("prohib", "Prohibited Use", "use.prohibited_use"),
]

_NAVY = colors.HexColor("#00008B")
_SECTION_BG = colors.HexColor("#D9D9D9")  # gray section-title bar, matches the reference template
_STRIPE = colors.HexColor("#F5F5FF")

# Letter page (8.5in) minus the 0.6in left/right margins set on the
# SimpleDocTemplate below - every table in this file is normalized to sum
# to exactly this, so every section (Lease Information, Charge Schedules,
# Amendments, ...) spans the full page width consistently instead of each
# one being some arbitrary, visually-inconsistent width.
_CONTENT_WIDTH = LETTER[0] - 1.2 * inch


def _normalize_widths(relative_widths):
    """Scales a list of relative column-width numbers so they sum to
    exactly _CONTENT_WIDTH, in points. Lets each table describe its
    columns' widths *relative to each other* (which one should be wider)
    without needing to hand-compute exact inch values that add up right -
    every table ends up the same total width regardless."""
    total = sum(relative_widths)
    return [w / total * _CONTENT_WIDTH for w in relative_widths]


# Item 6 - when a custom Output Template's style was detected
# (extract_template_style / save_template_style_profile above),
# generate_output_pdf sets this once before building the story, and every
# _build_* table/header helper below reads from it instead of the
# hardcoded gray/Helvetica defaults - so a template's header color, text
# color, and font family carry through to the generated report.
_ACTIVE_STYLE = {"headerBgColor": None, "headerTextColor": None, "fontFamily": None}


def _active_header_bg():
    if _ACTIVE_STYLE.get("headerBgColor"):
        try:
            return colors.HexColor(_ACTIVE_STYLE["headerBgColor"])
        except Exception:
            pass
    return _SECTION_BG


def _active_header_text():
    if _ACTIVE_STYLE.get("headerTextColor"):
        try:
            return colors.HexColor(_ACTIVE_STYLE["headerTextColor"])
        except Exception:
            pass
    return colors.black


def _active_font():
    return _ACTIVE_STYLE.get("fontFamily") or "Helvetica"


def _active_font_bold():
    base = _active_font()
    return {"Times-Roman": "Times-Bold", "Courier": "Courier-Bold"}.get(base, "Helvetica-Bold")


def _section_bar(text, col_span, styles):
    """The gray full-width "Lease Information" / "Charge Schedules" / etc
    bar the reference template uses as a section title - built as a
    single-row, single-cell table so it lines up exactly with the section
    body table beneath it (a Paragraph heading alone won't align)."""
    style = ParagraphStyle("SectionBar", fontName=_active_font_bold(), fontSize=9, textColor=_active_header_text())
    t = Table([[Paragraph(text, style)]], colWidths=[sum(col_span)])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), _active_header_bg()),
        ("BOX", (0, 0), (-1, -1), 0.5, _NAVY),
        ("LEFTPADDING", (0, 0), (-1, -1), 6), ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return t


def _parse_date_loose(s):
    for fmt in ("%m/%d/%Y", "%m/%d/%y", "%B %d, %Y", "%b %d, %Y", "%Y-%m-%d"):
        try:
            return datetime.strptime(s.strip(), fmt)
        except (ValueError, AttributeError):
            continue
    return None


def _term_months(start, end):
    d1, d2 = _parse_date_loose(start), _parse_date_loose(end)
    if not d1 or not d2:
        return ""
    months = (d2.year - d1.year) * 12 + (d2.month - d1.month)
    return str(max(months, 0))


def _digits_only(s):
    m = re.search(r"[\d,]+", s or "")
    return m.group(0) if m else ""


def _ensure_dollar_prefix(value):
    """Guarantees exactly one leading '$' on a currency string, regardless
    of whether the source (LLM extraction) already included one - safer
    than either assuming it's always there or always prepending one."""
    s = _stringify_leaf(value)
    if s in ("-", ""):
        return s
    return s if s.startswith("$") else f"${s}"


def _build_lease_information_grid(fields, lease_name):
    """A single 4-column (label/value/label/value) grid matching the
    reference template's "Lease Information" layout exactly. A handful of
    fields shown there (Status, ICS Code, Property code, Sales Category,
    Office Phone, FAX) are internal property-management-system metadata
    that simply isn't present in a lease document's text - those cells
    show the template's own apparent constant/blank convention rather than
    fabricated data."""
    lbl = ParagraphStyle("GridLabel", fontName=_active_font_bold(), fontSize=8.5, leading=11)
    val = ParagraphStyle("GridVal", fontName=_active_font(), fontSize=8.5, leading=11)

    def L(text):
        return Paragraph(text, lbl)

    def V(text):
        return Paragraph(_stringify_leaf(text) or "-", val)

    tenant = _get_path(fields, "parties.tenant_legal_name") or lease_name
    premises_size = _get_path(fields, "premises.premises_size")
    security_deposit = _get_path(fields, "rent.security_deposit")
    start = _get_path(fields, "term.lease_commencement_date")
    end = _get_path(fields, "term.lease_expiration_date")

    rows = [
        [L("Name"), V(tenant), L("Status"), V("Future")],
        [L("DBA"), V(_get_path(fields, "parties.tenant_dba")), L("ICS Code"), V("-")],
        [L("Property"), V(_get_path(fields, "premises.property_address")), L("Lease Type"), V(_get_path(fields, "premises.lease_type"))],
        [L("Landlord"), V(_get_path(fields, "parties.landlord_legal_name")), L("Sales Category"), V("General")],
        [L("Guarantor"), V(_get_path(fields, "parties.guarantor_name") or _get_path(fields, "parties.guarantor")),
         L("Contract Area"), Paragraph(f"{_stringify_leaf(premises_size) or '0.00'} (Rentable)", val)],
        [L(""), V(""), L("Area"), Paragraph("0.00 (Rentable)", val)],
        [Paragraph("Primary Contact", lbl), Paragraph("", val), Paragraph("Monthly Rent", lbl), Paragraph("", val)],
        [L("Name"), V(tenant), L("Annual Rent"), V("0.00")],
        [L("Office Phone"), V("-"), L("Rent Per Area"), V("0.00")],
        [L("FAX"), V("-"), L("Deposit"), Paragraph(_ensure_dollar_prefix(security_deposit) if security_deposit else "0.00", val)],
        [L("E-Mail"), V("-"), L("Lease Term"), Paragraph(f"{start} To {end}", val)],
    ]
    col_widths = _normalize_widths([0.95, 2.55, 1.15, 2.35])
    table = Table(rows, colWidths=col_widths)
    table.setStyle(TableStyle([
        ("BOX", (0, 0), (-1, -1), 0.5, _NAVY),
        ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CCCCCC")),
        ("SPAN", (1, 5), (1, 5)), ("SPAN", (3, 5), (3, 5)),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5), ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    return _section_bar("Lease Information", col_widths, getSampleStyleSheet()), table


def _build_charge_schedule_table(fields):
    """rent.rent_schedule (see [RENT_SCHEDULE] in extraction_prompt.txt) -
    charge_code/charge_desc are only shown on the first row of each
    consecutive run (blank on repeats), matching the reference template's
    "merged cell" look for a multi-period charge like Base Rent."""
    schedule = _get_path(fields, "rent.rent_schedule", [])
    if not isinstance(schedule, list) or not schedule:
        return None

    def pick(row, *keys):
        for k in keys:
            if isinstance(row, dict) and row.get(k) not in (None, ""):
                return _stringify_leaf(row.get(k))
        return ""

    header_style = ParagraphStyle("ChargeHeader", fontSize=7.5, leading=9, textColor=_active_header_text(), fontName=_active_font_bold())
    cell_style = ParagraphStyle("ChargeCell", fontSize=7.5, leading=9, fontName=_active_font())
    headers = ["Charge Code", "Charge Desc", "Date From", "Date To", "Monthly Amt", "Annual Amt", "Amt Per Area", "Mgmt Fees", "Amendment Type", "Units"]
    table_data = [[Paragraph(h, header_style) for h in headers]]

    last_code = None
    for row in schedule:
        code = pick(row, "charge_code", "chargeCode", "unit", "Unit")
        desc = pick(row, "charge_desc", "chargeDesc")
        show_code_desc = code != last_code
        last_code = code
        table_data.append([
            Paragraph(code if show_code_desc else "", cell_style),
            Paragraph(desc if show_code_desc else "", cell_style),
            Paragraph(pick(row, "from_date", "fromDate", "From Date", "from"), cell_style),
            Paragraph(pick(row, "to_date", "toDate", "To Date", "to"), cell_style),
            Paragraph(pick(row, "monthly_amount", "monthlyAmount", "Monthly Amount", "monthly"), cell_style),
            Paragraph(pick(row, "annual_amount", "annualAmount", "Annual Amount", "annual"), cell_style),
            Paragraph(pick(row, "amt_per_area", "annual_per_sf", "Annual Per SF"), cell_style),
            Paragraph("0.00", cell_style),
            Paragraph("Original Lease", cell_style),
            Paragraph(pick(row, "units", "Units"), cell_style),
        ])
    col_widths = _normalize_widths([0.55, 0.65, 0.6, 0.6, 0.65, 0.65, 0.65, 0.5, 0.75, 0.4])
    table = Table(table_data, colWidths=col_widths)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), _active_header_bg()),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#AAAAAA")),
        ("BOX", (0, 0), (-1, -1), 0.5, _NAVY),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, _STRIPE]),
        ("LEFTPADDING", (0, 0), (-1, -1), 4), ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    return col_widths, table


def _build_indexation_table():
    """The reference template always shows this section with just headers
    (no rows) unless the lease has an actual indexation/escalation-index
    clause - our schema doesn't extract structured indexation data, so
    this mirrors that same header-only convention rather than guessing."""
    header_style = ParagraphStyle("IdxHeader", fontSize=7, leading=9, textColor=_active_header_text(), fontName=_active_font_bold())
    headers = ["Charge Code", "Charge Desc", "Date From", "Date To", "Indexation Method", "Index", "Month", "Factor", "Min", "Max", "Amendment Type", "Units"]
    relative = [1] * len(headers)
    relative[4] = 1.7  # "Indexation Method" needs more room
    col_widths = _normalize_widths(relative)
    table = Table([[Paragraph(h, header_style) for h in headers]], colWidths=col_widths)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), _active_header_bg()),
        ("BOX", (0, 0), (-1, -1), 0.5, _NAVY),
        ("LEFTPADDING", (0, 0), (-1, -1), 3), ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    return col_widths, table


def _build_amendments_table(fields, lease_name):
    start = _get_path(fields, "term.lease_commencement_date")
    end = _get_path(fields, "term.lease_expiration_date")
    if not start and not end:
        return None
    notes = _get_path(fields, "amendment_notes") or "Original Lease"

    header_style = ParagraphStyle("AmendHeader", fontSize=7.5, leading=9, textColor=_active_header_text(), fontName=_active_font_bold())
    cell_style = ParagraphStyle("AmendCell", fontSize=7.5, leading=9, fontName=_active_font())
    headers = ["Type", "Description", "Status", "Term (Months)", "Date From", "Date To", "Units"]
    table_data = [[Paragraph(h, header_style) for h in headers]]
    table_data.append([
        Paragraph("Original Lease", cell_style), Paragraph(notes, cell_style), Paragraph("In Process", cell_style),
        Paragraph(_term_months(start, end), cell_style), Paragraph(start, cell_style), Paragraph(end, cell_style),
        Paragraph("-", cell_style),
    ])
    col_widths = _normalize_widths([0.9, 2.3, 0.7, 0.85, 0.75, 0.75, 0.55])
    table = Table(table_data, colWidths=col_widths)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), _active_header_bg()),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#AAAAAA")),
        ("BOX", (0, 0), (-1, -1), 0.5, _NAVY),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4), ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    return col_widths, table


def _build_late_fee_table(fields):
    lf = fields.get("rent", {}).get("late_fee_table") if isinstance(fields.get("rent"), dict) else None
    if not isinstance(lf, dict) or not any(lf.values()):
        return None
    header_style = ParagraphStyle("LateHeader", fontSize=7.5, leading=9, textColor=_active_header_text(), fontName=_active_font_bold())
    cell_style = ParagraphStyle("LateCell", fontSize=7.5, leading=9, fontName=_active_font())
    headers = ["Calculation Type", "Grace Period", "Percent", "2nd Fee Calc Type", "2nd Fee Grace Period", "2nd Fee Percent", "Per Day Fee"]
    table_data = [[Paragraph(h, header_style) for h in headers]]
    table_data.append([
        Paragraph(_stringify_leaf(lf.get("calc_type")) or "% Owed-Total", cell_style),
        Paragraph(_stringify_leaf(lf.get("grace_period")) or "0", cell_style),
        Paragraph(_stringify_leaf(lf.get("percent")) or "-", cell_style),
        Paragraph(_stringify_leaf(lf.get("second_fee_calc_type")), cell_style),
        Paragraph(_stringify_leaf(lf.get("second_fee_grace_period")), cell_style),
        Paragraph(_stringify_leaf(lf.get("second_fee_percent")), cell_style),
        Paragraph(_stringify_leaf(lf.get("per_day_fee")) or "0.00", cell_style),
    ])
    col_widths = _normalize_widths([1] * 7)
    table = Table(table_data, colWidths=col_widths)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), _active_header_bg()),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#AAAAAA")),
        ("BOX", (0, 0), (-1, -1), 0.5, _NAVY),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4), ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    return col_widths, table


def _build_clauses_table(fields):
    """The reference template's "Other Lease Provisions / Clauses" table -
    Id / Name / Description / Amendment Type, one row per clause the
    extraction prompt already asks for. Missing/empty fields show "Lease
    is silent." which matches the SILENT CLAUSE RULE the prompt itself
    instructs the LLM to use."""
    header_style = ParagraphStyle("ClauseHeader", fontSize=8, leading=10, textColor=_active_header_text(), fontName=_active_font_bold())
    id_style = ParagraphStyle("ClauseId", fontSize=8, leading=11, fontName=_active_font_bold())
    cell_style = ParagraphStyle("ClauseCell", fontSize=8, leading=11, fontName=_active_font())
    headers = ["Id", "Name", "Description", "Amendment Type"]
    table_data = [[Paragraph(h, header_style) for h in headers]]
    for clause_id, name, path in _CLAUSE_FIELD_MAP:
        value = _get_path(fields, path) or "Lease is silent."
        table_data.append([
            Paragraph(clause_id, id_style), Paragraph(name, cell_style),
            Paragraph(_stringify_leaf(value).replace("\n", "<br/>"), cell_style),
            Paragraph("Original Lease", cell_style),
        ])
    col_widths = _normalize_widths([0.6, 1.3, 3.9, 0.8])
    table = Table(table_data, colWidths=col_widths)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), _active_header_bg()),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#AAAAAA")),
        ("BOX", (0, 0), (-1, -1), 0.5, _NAVY),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, _STRIPE]),
        ("LEFTPADDING", (0, 0), (-1, -1), 6), ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    return col_widths, table


def _build_contacts_table(fields):
    contacts = fields.get("contacts")
    if not isinstance(contacts, dict):
        return None
    tenant = _get_path(fields, "parties.tenant_legal_name")
    landlord = _get_path(fields, "parties.landlord_legal_name")
    role_map = [
        ("Tenant Notice", tenant, contacts.get("tenant_notice")),
        ("Tenant Billing", tenant, contacts.get("tenant_billing")),
        ("Landlord Notice", landlord, contacts.get("landlord_notice")),
        ("Guarantor", _get_path(fields, "parties.guarantor_name"), contacts.get("guarantor_contact")),
    ]
    rows = [(role, company, addr) for role, company, addr in role_map if addr and addr != "N/A"]
    if not rows:
        return None

    header_style = ParagraphStyle("ContactHeader", fontSize=8, leading=10, textColor=_active_header_text(), fontName=_active_font_bold())
    cell_style = ParagraphStyle("ContactCell", fontSize=8, leading=11, fontName=_active_font())
    headers = ["Role", "Company Name", "Address"]
    table_data = [[Paragraph(h, header_style) for h in headers]]
    for role, company, addr in rows:
        table_data.append([Paragraph(role, cell_style), Paragraph(_stringify_leaf(company) or "-", cell_style), Paragraph(_stringify_leaf(addr), cell_style)])
    col_widths = _normalize_widths([1.1, 1.8, 3.7])
    table = Table(table_data, colWidths=col_widths)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), _active_header_bg()),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#AAAAAA")),
        ("BOX", (0, 0), (-1, -1), 0.5, _NAVY),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, _STRIPE]),
        ("LEFTPADDING", (0, 0), (-1, -1), 6), ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return col_widths, table


def generate_output_pdf(output_json_path, pdf_out_path, template_name="Default.pdf", style_profile=None):
    """Builds Output.pdf from Output.json. When `fields` has the rich
    nested LLM extraction schema (see json/extraction_prompt.txt), this
    mirrors the reference "Lease Abstract" template's exact grid/table
    layout section-by-section: Lease Information, Charge Schedules,
    Indexation, Amendments, Late Fee, Other Lease Provisions/Clauses, then
    Contacts. Falls back to a generic key/value dump for the flat
    heuristic-engine schema, which doesn't carry per-clause data.

    style_profile (item 6, optional) - a dict from
    extract_template_style()/load_template_style_profile() with the
    uploaded Output Template's detected headerBgColor, headerTextColor,
    fontFamily, and logoImagePath. When given, every section header and
    table in this report reuses that color/font instead of the built-in
    gray/Helvetica default, and the template's logo (if one was found) is
    placed at the top of the page."""
    if not REPORTLAB_OK:
        raise LeaseEngineError(
            "reportlab is not installed - run: pip install -r requirements.txt"
        )

    # Reset to defaults each call, then apply whatever this specific
    # template's profile provides - _build_* helpers below all read these
    # via _active_header_bg()/_active_header_text()/_active_font().
    _ACTIVE_STYLE["headerBgColor"] = (style_profile or {}).get("headerBgColor")
    _ACTIVE_STYLE["headerTextColor"] = (style_profile or {}).get("headerTextColor")
    _ACTIVE_STYLE["fontFamily"] = (style_profile or {}).get("fontFamily")
    logo_path = (style_profile or {}).get("logoImagePath")

    with open(output_json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    fields = data.get("fields", {}) or {}
    lease_name = data.get("leaseName", "Lease")
    extraction_method = data.get("extractionMethod", "heuristic")
    accuracy = data.get("accuracy")
    accuracy_summary = data.get("accuracySummary") or ""

    doc = SimpleDocTemplate(
        pdf_out_path, pagesize=LETTER,
        topMargin=0.6 * inch, bottomMargin=0.6 * inch,
        leftMargin=0.6 * inch, rightMargin=0.6 * inch,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "LeaseTitle", parent=styles["Heading1"], textColor=colors.black, fontSize=16, spaceAfter=2, fontName=_active_font_bold()
    )
    normal = ParagraphStyle("LeaseNormal", parent=styles["Normal"], fontSize=8.5, fontName=_active_font())

    story = []
    if logo_path and os.path.isfile(logo_path):
        try:
            img = Image(logo_path)
            max_w, max_h = 1.6 * inch, 0.7 * inch
            scale = min(max_w / img.imageWidth, max_h / img.imageHeight, 1.0)
            img.drawWidth = img.imageWidth * scale
            img.drawHeight = img.imageHeight * scale
            story.append(img)
            story.append(Spacer(1, 6))
        except Exception as err:
            print(f"Could not embed template logo in Output.pdf: {err}")

    story += [
        Paragraph("Lease Abstract", title_style),
        Paragraph(f"Lease : {lease_name}", normal),
        Spacer(1, 10),
    ]

    if _is_rich_lease_schema(fields):
        bar, grid = _build_lease_information_grid(fields, lease_name)
        story += [bar, grid, Spacer(1, 10)]

        charge = _build_charge_schedule_table(fields)
        if charge:
            col_widths, table = charge
            story += [_section_bar("Charge Schedules", col_widths, styles), table, Spacer(1, 10)]

        idx_widths, idx_table = _build_indexation_table()
        story += [_section_bar("Indexation", idx_widths, styles), idx_table, Spacer(1, 10)]

        amend = _build_amendments_table(fields, lease_name)
        if amend:
            col_widths, table = amend
            story += [_section_bar("Amendments", col_widths, styles), table, Spacer(1, 10)]

        late_fee = _build_late_fee_table(fields)
        if late_fee:
            col_widths, table = late_fee
            story += [_section_bar("Late Fee", col_widths, styles), table, Spacer(1, 10)]

        clauses_widths, clauses_table = _build_clauses_table(fields)
        story += [_section_bar("Other Lease Provisions / Clauses", clauses_widths, styles), clauses_table, Spacer(1, 10)]

        contacts = _build_contacts_table(fields)
        if contacts:
            col_widths, table = contacts
            story += [_section_bar("Contacts", col_widths, styles), table, Spacer(1, 10)]

        if accuracy is not None:
            story.append(Paragraph(
                f"<i>Accuracy: {accuracy}%" + (f" — {accuracy_summary}" if accuracy_summary else "") + "</i>",
                normal
            ))
        if fields.get("queries_assumptions"):
            story.append(Spacer(1, 6))
            story.append(Paragraph("<b>Queries / Assumptions</b>", normal))
            for q in fields["queries_assumptions"]:
                story.append(Paragraph(f"• {_stringify_leaf(q)}", normal))
    else:
        # Flat heuristic-engine schema - generic key/value dump (no
        # per-clause data available to build the real template layout).
        heading_style = ParagraphStyle("LeaseHeading", parent=styles["Heading2"], textColor=_NAVY, spaceBefore=14, spaceAfter=6)
        if accuracy is not None:
            story.append(Paragraph(f"Extraction method: {extraction_method} · Accuracy: {accuracy}%" + (f" — {accuracy_summary}" if accuracy_summary else ""), normal))
        story.append(Paragraph("Lease Summary", heading_style))
        simple_rows = [
            (_humanize_key(k), _stringify_leaf(v))
            for k, v in fields.items() if not isinstance(v, (dict, list))
        ]
        if simple_rows:
            story.append(_build_kv_table(simple_rows))
        for key, value in fields.items():
            if not isinstance(value, (dict, list)):
                continue
            section_rows = _flatten_to_rows(value)
            if not section_rows:
                continue
            story.append(Spacer(1, 10))
            story.append(Paragraph(_humanize_key(key), heading_style))
            story.append(_build_kv_table(section_rows))

    story.append(Spacer(1, 10))
    story.append(HRFlowable(width="100%", color=_NAVY))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>Source Document(s)</b>", normal))
    for doc_path in data.get("sourceDocuments", []):
        # Only the filename - sourceDocuments stores the full internal
        # Users/{userId}/LeaseAbstraction/{leaseName}/... path, which is
        # our own server's folder structure and shouldn't leak into a
        # document the user downloads and shares externally.
        story.append(Paragraph(os.path.basename(doc_path), normal))

    doc.build(story)
    return pdf_out_path



def _color_tuple_to_hex(c):
    """pdfplumber gives fill colors as a float (grayscale 0-1), a 3-tuple
    (RGB 0-1), or a 4-tuple (CMYK 0-1) depending on the PDF's color space -
    normalizes any of those to a '#RRGGBB' hex string."""
    try:
        if isinstance(c, (int, float)):
            v = int(round(c * 255))
            return "#%02X%02X%02X" % (v, v, v)
        if isinstance(c, (list, tuple)):
            if len(c) == 3:
                r, g, b = c
                return "#%02X%02X%02X" % (int(round(r * 255)), int(round(g * 255)), int(round(b * 255)))
            if len(c) == 4:
                cy, m, y, k = c
                r = 255 * (1 - cy) * (1 - k)
                g = 255 * (1 - m) * (1 - k)
                b = 255 * (1 - y) * (1 - k)
                return "#%02X%02X%02X" % (int(round(r)), int(round(g)), int(round(b)))
    except (TypeError, ValueError):
        pass
    return None


_FONT_FAMILY_MAP = [
    ("times", "Times-Roman"), ("serif", "Times-Roman"), ("georgia", "Times-Roman"),
    ("courier", "Courier"), ("mono", "Courier"),
    ("arial", "Helvetica"), ("helvetica", "Helvetica"), ("calibri", "Helvetica"), ("verdana", "Helvetica"),
]


def extract_template_style(template_path):
    """Item 6 - inspects an uploaded Output Template PDF (first page) and
    identifies: the dominant header/accent fill color, the primary font
    family in use, and any logo/image near the top of the page. Returns a
    style profile dict that generate_output_pdf() applies to the actual
    report, so a custom template's look carries over instead of always
    using the hardcoded default gray/Helvetica style. Best-effort and
    silent on failure - a template we can't parse just falls back to
    defaults, it never breaks the pipeline."""
    profile = {"headerBgColor": None, "headerTextColor": None, "fontFamily": None, "logoImagePath": None}
    if not pdfplumber or not os.path.isfile(template_path):
        return profile

    try:
        with pdfplumber.open(template_path) as pdf:
            if not pdf.pages:
                return profile
            page = pdf.pages[0]

            # ---- Dominant fill color (header/accent bars) - most common
            # non-white, non-black rect fill on the page. ----
            color_counts = {}
            for rect in (page.rects or []):
                hexcolor = _color_tuple_to_hex(rect.get("non_stroking_color"))
                if hexcolor and hexcolor not in ("#FFFFFF", "#000000"):
                    color_counts[hexcolor] = color_counts.get(hexcolor, 0) + 1
            if color_counts:
                profile["headerBgColor"] = max(color_counts, key=color_counts.get)
                # Pick a readable text color against that fill - light fill
                # gets dark text, dark fill gets white text.
                hexc = profile["headerBgColor"].lstrip("#")
                r, g, b = int(hexc[0:2], 16), int(hexc[2:4], 16), int(hexc[4:6], 16)
                brightness = (r * 299 + g * 587 + b * 114) / 1000
                profile["headerTextColor"] = "#000000" if brightness > 150 else "#FFFFFF"

            # ---- Dominant font family ----
            font_counts = {}
            for ch in (page.chars or [])[:2000]:  # cap - just need a representative sample
                name = (ch.get("fontname") or "").lower()
                if name:
                    font_counts[name] = font_counts.get(name, 0) + 1
            if font_counts:
                top_font = max(font_counts, key=font_counts.get)
                for keyword, mapped in _FONT_FAMILY_MAP:
                    if keyword in top_font:
                        profile["fontFamily"] = mapped
                        break

            # ---- Logo/image near the top of the first page. PDFs often
            # embed tiny 1x1 "spacer" pixels as a rendering artifact -
            # those aren't a logo, so anything smaller than a plausible
            # logo size gets skipped. ----
            images = [
                im for im in (page.images or [])
                if im.get("top", 999) < page.height * 0.35
                and (im.get("x1", 0) - im.get("x0", 0)) >= 20
                and (im.get("bottom", 0) - im.get("top", 0)) >= 20
            ]
            if images:
                im = images[0]
                try:
                    bbox = (max(0, im["x0"]), max(0, im["top"]), min(page.width, im["x1"]), min(page.height, im["bottom"]))
                    if bbox[2] > bbox[0] and bbox[3] > bbox[1]:
                        cropped = page.crop(bbox)
                        logo_path = os.path.join(os.path.dirname(template_path), "_extracted_logo.png")
                        cropped.to_image(resolution=150).save(logo_path, format="PNG")
                        with PILImage.open(logo_path) as saved_img:
                            if saved_img.width < 20 or saved_img.height < 20:
                                os.remove(logo_path)
                                logo_path = None
                        profile["logoImagePath"] = logo_path
                except Exception as err:
                    print(f"Template logo extraction skipped: {err}")
    except Exception as err:
        print(f"Template style extraction failed, using defaults: {err}")

    return profile


def save_template_style_profile(template_path):
    """Extracts and persists the style profile next to the template file,
    so generate_output_pdf can look it up later by template name without
    re-parsing the PDF on every single lease processed with it."""
    profile = extract_template_style(template_path)
    profile_path = template_path + ".style.json"
    try:
        with open(profile_path, "w", encoding="utf-8") as f:
            json.dump(profile, f)
    except OSError:
        pass
    return profile


def load_template_style_profile(template_path):
    profile_path = template_path + ".style.json"
    if os.path.isfile(profile_path):
        try:
            with open(profile_path, "r", encoding="utf-8") as f:
                return json.load(f)
        except (OSError, json.JSONDecodeError):
            pass
    return {"headerBgColor": None, "headerTextColor": None, "fontFamily": None, "logoImagePath": None}


def build_default_template_pdf(pdf_out_path):
    """Generates the placeholder Default.pdf that ships in
    Template/LeaseAbstraction/ - the same layout generate_output_pdf()
    produces, but with blank/sample placeholder values, since this is the
    template shown to the user before any lease has been processed."""
    if not REPORTLAB_OK:
        raise LeaseEngineError(
            "reportlab is not installed - run: pip install -r requirements.txt"
        )

    doc = SimpleDocTemplate(
        pdf_out_path, pagesize=LETTER,
        topMargin=0.75 * inch, bottomMargin=0.75 * inch,
        leftMargin=0.75 * inch, rightMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "LeaseTitle", parent=styles["Title"], textColor=colors.HexColor("#00008B")
    )
    heading_style = ParagraphStyle(
        "LeaseHeading", parent=styles["Heading2"], textColor=colors.HexColor("#00008B"),
        spaceBefore=14, spaceAfter=6
    )
    normal = styles["Normal"]

    story = [
        Paragraph("Lease Abstraction Report", title_style),
        Paragraph("Default Output Template", normal),
        Spacer(1, 12),
        HRFlowable(width="100%", color=colors.HexColor("#00008B")),
        Spacer(1, 12),
        Paragraph("Lease Summary", heading_style),
    ]

    table_data = [["Field", "Value"]]
    placeholders = [
        ("Document Type", "{{documentType}}"),
        ("Tenant", "{{tenant}}"),
        ("Landlord", "{{landlord}}"),
        ("Property Address", "{{propertyAddress}}"),
        ("Lease Start", "{{leaseStart}}"),
        ("Lease End", "{{leaseEnd}}"),
        ("Base Rent", "{{baseRent}}"),
        ("Currency", "{{currency}}"),
    ]
    for label, value in placeholders:
        table_data.append([label, value])

    table = Table(table_data, colWidths=[2.2 * inch, 3.8 * inch])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#00008B")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#00008B")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F5F5FF")]),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(table)
    story.append(Spacer(1, 18))
    story.append(Paragraph(
        "This is the default output layout used to generate Output.pdf for "
        "every processed lease when no custom template has been selected.",
        normal
    ))

    doc.build(story)
    return pdf_out_path
