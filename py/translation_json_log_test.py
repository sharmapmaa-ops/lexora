"""
Test for the translation-batch JSON request/response logging feature,
per explicit direction: "jab hum openrouter ko json data bhejte hain
tab sabse pehle json file download honi chahiye, fir jab openrouter se
translation ka data json me mile to wo sabse pehle download ho jana
chahiye". _translate_docx_segments_in_place now returns a 6th value -
a list of {request, response} dicts, one per real LLM batch call - and
translate_existing_docx propagates it through as
"translation_request_response_log" for the API layer to hand to the
client for immediate per-batch downloads.

Run: python3 translation_json_log_test.py
"""
import os
import sys
import json

sys.path.insert(0, os.path.dirname(__file__))
from docx import Document
import aspose_test_pipeline as asp
import lease_engine as le


def assert_(cond, label):
    if not cond:
        print("FAIL:", label)
        sys.exit(1)
    print("PASS:", label)


def run():
    print("=== Test 1: request/response log structure, single batch ===")
    def mock_call_ok(llm_config, system_prompt, user_content, max_tokens=12000):
        req_items = json.loads(user_content)
        resp_items = [{"id": item["id"], "text": "TR: " + item["text"]} for item in req_items]
        return json.dumps(resp_items), "mock-provider"

    original = le._call_chat_completion_with_failover
    le._call_chat_completion_with_failover = mock_call_ok
    try:
        doc = Document()
        doc.add_paragraph("Hello world")
        doc.add_paragraph("Second paragraph")
        result = asp._translate_docx_segments_in_place(doc, "French", {})
        assert_(len(result) == 6, "Function now returns a 6-tuple (5 original values + the new log)")
        translated_count, skipped_count, failed_batches, total_batches, llm_calls_by_provider, log = result
        assert_(translated_count == 2, "Both paragraphs translated")
        assert_(len(log) == 1, "One batch, one log entry")
        assert_(log[0]["request"] is not None, "The request JSON was captured")
        assert_(log[0]["response"] is not None, "The response JSON was captured")
        req_parsed = json.loads(log[0]["request"])
        assert_(req_parsed == [{"id": 1, "text": "Hello world"}, {"id": 2, "text": "Second paragraph"}], "Request JSON matches exactly what was sent")
        resp_parsed = json.loads(log[0]["response"])
        assert_(resp_parsed[0]["text"] == "TR: Hello world", "Response JSON matches exactly what was received")
    finally:
        le._call_chat_completion_with_failover = original

    print("\n=== Test 2: a FAILED batch still logs its request, with response left None ===")
    def mock_call_fail(llm_config, system_prompt, user_content, max_tokens=12000):
        return None, None  # simulates a failed call, matching the real failure path

    le._call_chat_completion_with_failover = mock_call_fail
    try:
        doc2 = Document()
        doc2.add_paragraph("Some text")
        result2 = asp._translate_docx_segments_in_place(doc2, "French", {})
        _, skipped_count2, failed_batches2, _, _, log2 = result2
        assert_(failed_batches2 == 1, "One failed batch")
        assert_(skipped_count2 == 1, "The paragraph is left untranslated (skipped)")
        assert_(len(log2) == 1, "The failed batch's own request is still logged")
        assert_(log2[0]["request"] is not None, "Request JSON captured even though the call failed")
        assert_(log2[0]["response"] is None, "Response stays None for a failed batch - nothing was received")
    finally:
        le._call_chat_completion_with_failover = original

    print("\n=== Test 3: translate_existing_docx propagates the log through its own return dict ===")
    real_docx = "/home/claude/work/two_docs_review/ocr_trace/ocr_step1.docx"
    if os.path.exists(real_docx):
        le._call_chat_completion_with_failover = mock_call_ok
        try:
            import tempfile
            output_path = os.path.join(tempfile.mkdtemp(), "output.docx")
            result3 = asp.translate_existing_docx(real_docx, "English", output_path, llm_config={})
            assert_("translation_request_response_log" in result3, "The top-level pipeline result includes the new log key")
            assert_(isinstance(result3["translation_request_response_log"], list), "It's a real list of batch entries")
            assert_(len(result3["translation_request_response_log"]) >= 1, "At least one real batch was logged for this real document")
        finally:
            le._call_chat_completion_with_failover = original
    else:
        print("  (real document not present in this environment - skipped)")

    print("\nALL TESTS PASSED")


if __name__ == "__main__":
    run()
