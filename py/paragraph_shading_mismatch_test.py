"""
Regression test for _fix_paragraph_shading_mismatch_within_cell, per a
real user-provided, zoomed MS Word screenshot: header cells like "Rent
Value"/"Total Value" showed a visible white horizontal line cutting
through their otherwise-solid dark background. Root cause confirmed
directly in the real document's XML: each such cell is genuinely TWO
separate paragraphs (not one paragraph with concatenated text) - the
first paragraph correctly carries its own w:pPr/w:shd matching the
cell's dark fill, but the second paragraph has no shading at all,
rendering as a visible white gap. A systematic whole-document scan
found 196 real instances of this exact mismatch.

Run: python3 paragraph_shading_mismatch_test.py
"""
import os
import sys

sys.path.insert(0, os.path.dirname(__file__))
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import aspose_test_pipeline as asp


def assert_(cond, label):
    if not cond:
        print("FAIL:", label)
        sys.exit(1)
    print("PASS:", label)


def run():
    print("=== Test 1 (THE REAL REPORTED DOCUMENT): comprehensive whole-document scan ===")
    real_docx = "/home/claude/work/two_docs_review/doc1.docx"
    if os.path.exists(real_docx):
        doc = Document(real_docx)
        fixed = asp._fix_paragraph_shading_mismatch_within_cell(doc)
        assert_(fixed == 196, "Exactly 196 real paragraphs fixed (matches the real confirmed whole-document count)")

        remaining = 0
        for tc in doc.element.body.iter(qn("w:tc")):
            tcPr = tc.find(qn("w:tcPr"))
            cell_shd = tcPr.find(qn("w:shd")) if tcPr is not None else None
            cell_fill = cell_shd.get(qn("w:fill")) if cell_shd is not None else None
            if not cell_fill or cell_fill in ("auto", "FFFFFF"):
                continue
            for p in tc.findall(qn("w:p")):
                pPr = p.find(qn("w:pPr"))
                para_shd = pPr.find(qn("w:shd")) if pPr is not None else None
                if para_shd is None or para_shd.get(qn("w:fill")) != cell_fill:
                    remaining += 1
        assert_(remaining == 0, "No paragraph anywhere retains a shading mismatch vs its own cell's fill")
    else:
        print("  (real document not present in this environment - skipped)")

    print("\n=== Test 2 (synthetic, exact reported case): a 2-paragraph cell with one paragraph missing shading entirely ===")
    doc2 = Document()
    t2 = doc2.add_table(rows=1, cols=1)
    cell2 = t2.rows[0].cells[0]
    tcPr2 = cell2._tc.get_or_add_tcPr()
    cell_shd2 = OxmlElement("w:shd")
    cell_shd2.set(qn("w:fill"), "666666")
    tcPr2.append(cell_shd2)

    p1 = cell2.paragraphs[0]
    p1.text = "Rent Value"
    pPr1 = p1._p.get_or_add_pPr()
    shd1 = OxmlElement("w:shd")
    shd1.set(qn("w:fill"), "666666")
    pPr1.append(shd1)

    p2 = cell2.add_paragraph("Rent value")  # NO shading at all - the real bug

    fixed2 = asp._fix_paragraph_shading_mismatch_within_cell(doc2)
    assert_(fixed2 == 1, "Exactly 1 paragraph fixed (the second one, missing shading)")
    p2_shd = p2._p.find(qn("w:pPr")).find(qn("w:shd"))
    assert_(p2_shd is not None and p2_shd.get(qn("w:fill")) == "666666", "The second paragraph now matches the cell's own dark fill")
    p1_shd_after = p1._p.find(qn("w:pPr")).find(qn("w:shd"))
    assert_(p1_shd_after.get(qn("w:fill")) == "666666", "The first (already-correct) paragraph is unaffected")

    print("\n=== Test 3: a paragraph with a WRONG (mismatched) shading color is corrected, not just a missing one ===")
    doc3 = Document()
    t3 = doc3.add_table(rows=1, cols=1)
    cell3 = t3.rows[0].cells[0]
    tcPr3 = cell3._tc.get_or_add_tcPr()
    cell_shd3 = OxmlElement("w:shd")
    cell_shd3.set(qn("w:fill"), "666666")
    tcPr3.append(cell_shd3)
    p3 = cell3.paragraphs[0]
    p3.text = "Some header"
    pPr3 = p3._p.get_or_add_pPr()
    wrong_shd = OxmlElement("w:shd")
    wrong_shd.set(qn("w:fill"), "EEEEEE")  # wrong color, not matching the cell
    pPr3.append(wrong_shd)
    fixed3 = asp._fix_paragraph_shading_mismatch_within_cell(doc3)
    assert_(fixed3 == 1, "One correction made")
    assert_(wrong_shd.get(qn("w:fill")) == "666666", "The mismatched color is corrected to match the cell's own fill")

    print("\n=== Test 4: an ordinary WHITE (no real background) cell is never touched ===")
    doc4 = Document()
    t4 = doc4.add_table(rows=1, cols=1)
    cell4 = t4.rows[0].cells[0]
    cell4.text = "Ordinary cell"
    cell4.add_paragraph("Second paragraph, no shading needed")
    fixed4 = asp._fix_paragraph_shading_mismatch_within_cell(doc4)
    assert_(fixed4 == 0, "Zero fixes - a cell with no real background fill has no white-line gap to create")

    print("\nALL TESTS PASSED")


if __name__ == "__main__":
    run()
