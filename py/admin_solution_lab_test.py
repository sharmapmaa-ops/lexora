"""
Test for the admin-panel Solution Lab feature, per explicit direction:
"admin panel-claude me mujhe har issue ke label aur select box add
karo solution ke liye aur mujhe sabhi solution usi me dedo taki main
one by one koi bhi option select karke check kar saku". Covers the
registry structure, that every registered variant function exists and
runs cleanly, and that each issue's variants are genuinely different
implementations (not the same function registered twice).

Run: python3 admin_solution_lab_test.py
"""
import os
import sys

sys.path.insert(0, os.path.dirname(__file__))
from docx import Document
import aspose_test_pipeline as asp


def assert_(cond, label):
    if not cond:
        print("FAIL:", label)
        sys.exit(1)
    print("PASS:", label)


def run():
    print("=== Test 1: registry structure is well-formed ===")
    registry = asp.ADMIN_SOLUTION_LAB_REGISTRY
    assert_(len(registry) == 3, "3 issues registered (white-line, character-split, table-overflow)")
    for issue_id, issue in registry.items():
        assert_("label" in issue and issue["label"], f"{issue_id} has a real label")
        assert_(len(issue["variants"]) >= 2, f"{issue_id} has at least 2 selectable variants")
        fn_names = [v["fn"] for v in issue["variants"].values()]
        assert_(len(fn_names) == len(set(fn_names)), f"{issue_id}'s variants all point to genuinely DIFFERENT functions")

    print("\n=== Test 2: every registered variant function exists and is callable ===")
    for issue_id, issue in registry.items():
        for variant_id, variant in issue["variants"].items():
            fn = getattr(asp, variant["fn"], None)
            assert_(fn is not None and callable(fn), f"{issue_id}/{variant_id} -> {variant['fn']} exists and is callable")

    print("\n=== Test 3 (THE REAL REPORTED DOCUMENT): every variant runs cleanly, no exceptions, saves validly ===")
    real_docx = "/home/claude/work/regression_check/real_output.docx"
    if os.path.exists(real_docx):
        for issue_id, issue in registry.items():
            for variant_id, variant in issue["variants"].items():
                doc = Document(real_docx)
                fn = getattr(asp, variant["fn"])
                try:
                    result = fn(doc)
                    doc.save("/tmp/admin_lab_test_output.docx")
                    reopened = Document("/tmp/admin_lab_test_output.docx")
                    assert_(True, f"{issue_id}/{variant_id}: ran (result={result}), saved, and reopened without error")
                except Exception as e:
                    assert_(False, f"{issue_id}/{variant_id}: threw an exception - {e}")
    else:
        print("  (real document not present in this environment - skipped)")

    print("\n=== Test 4: white-line Variant B (merge) genuinely produces ONE paragraph from TWO ===")
    doc4 = Document()
    t4 = doc4.add_table(rows=1, cols=1)
    cell4 = t4.rows[0].cells[0]
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr4 = cell4._tc.get_or_add_tcPr()
    shd4 = OxmlElement("w:shd")
    shd4.set(qn("w:fill"), "666666")
    tcPr4.append(shd4)
    cell4.text = "Line one"
    cell4.add_paragraph("Line two")
    assert_(len(cell4._tc.findall(qn("w:p"))) == 2, "Setup: 2 paragraphs before the fix")
    asp._fix_paragraph_shading_mismatch_within_cell_variant_b_merge(doc4)
    assert_(len(cell4._tc.findall(qn("w:p"))) == 1, "Exactly 1 paragraph remains after Variant B's merge")
    assert_(len(cell4._tc.findall(".//" + qn("w:br"))) == 1, "A line-break preserves the visual line-split")

    print("\n=== Test 5: table-overflow Variant B genuinely REMOVES the tblInd element (not just resets its value) ===")
    doc5 = Document()
    t5 = doc5.add_table(rows=1, cols=1)
    tblPr5 = t5._tbl.find(qn("w:tblPr"))
    tblInd5 = OxmlElement("w:tblInd")
    tblInd5.set(qn("w:w"), "771")
    tblInd5.set(qn("w:type"), "dxa")
    tblPr5.append(tblInd5)
    tr5 = t5.rows[0]._tr
    tblPrEx5 = OxmlElement("w:tblPrEx")
    row_ind5 = OxmlElement("w:tblInd")
    row_ind5.set(qn("w:w"), "1950")
    tblPrEx5.append(row_ind5)
    tr5.insert(0, tblPrEx5)
    fixed5 = asp._fix_row_level_table_indent_override_variant_b_remove(doc5)
    assert_(fixed5 == 1, "One row-level override removed")
    assert_(tr5.find(qn("w:tblPrEx")) is None or tr5.find(qn("w:tblPrEx")).find(qn("w:tblInd")) is None, "The tblInd element is genuinely gone, not just reset to a new number")

    print("\nALL TESTS PASSED")


if __name__ == "__main__":
    run()
