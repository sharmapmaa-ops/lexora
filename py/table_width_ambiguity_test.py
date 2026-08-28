"""
Regression test for the ambiguous <w:tblW w:type="auto" w:w="0"/> fix,
per a real reported issue: tables showed unexpected left-gap or
width-overflow in real MS Word screenshots, while the same document's
XML (indent, gridCol widths) measured completely consistent/correct
and LibreOffice rendered them correctly - the same class of
LibreOffice/real-Word OOXML rendering discrepancy already documented
elsewhere in this codebase. Explicitly setting w:tblW to the table's
own real column-width sum removes the auto-vs-fixed ambiguity.

Flagged explicitly (not glossed over): this is a defensive correction
based on removing a genuinely ambiguous OOXML combination, not a
confirmed root-cause fix, since real MS Word rendering could not be
tested directly in this environment.

Run: python3 table_width_ambiguity_test.py
"""
import os
import sys

sys.path.insert(0, os.path.dirname(__file__))
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import aspose_test_pipeline as asp


def assert_(cond, label):
    if not cond:
        print("FAIL:", label)
        sys.exit(1)
    print("PASS:", label)


def run():
    print("=== Test 1: real document - ambiguous tblW resolved to an explicit real value ===")
    real_docx = "/home/claude/work/table_indent_debug/final_output.docx"
    if os.path.exists(real_docx):
        doc = Document(real_docx)
        fixed = asp._fix_ambiguous_table_width(doc)
        assert_(fixed > 0, "A real number of tables had their ambiguous tblW resolved (got " + str(fixed) + ")")
        remaining_ambiguous = 0
        for t in doc.tables:
            tblPr = t._tbl.find(qn("w:tblPr"))
            tblW = tblPr.find(qn("w:tblW")) if tblPr is not None else None
            if tblW is not None and tblW.get(qn("w:type")) == "auto" and tblW.get(qn("w:w")) in ("0", None):
                remaining_ambiguous += 1
        assert_(remaining_ambiguous == 0, "No table retains the ambiguous auto/0 tblW combination after the fix")
    else:
        print("  (real document not present in this environment - skipped)")

    print("\n=== Test 2 (synthetic, exact mechanics): tblW set to the real column-width sum ===")
    doc2 = Document()
    t2 = doc2.add_table(rows=1, cols=3)
    tblPr2 = t2._tbl.find(qn("w:tblPr"))
    tblW2 = tblPr2.find(qn("w:tblW"))
    tblW2.set(qn("w:type"), "auto")
    tblW2.set(qn("w:w"), "0")
    grid2 = t2._tbl.find(qn("w:tblGrid"))
    cols2 = grid2.findall(qn("w:gridCol"))
    widths = [3467, 3226, 2730]
    for gc, w in zip(cols2, widths):
        gc.set(qn("w:w"), str(w))

    fixed2 = asp._fix_ambiguous_table_width(doc2)
    assert_(fixed2 == 1, "Exactly 1 table fixed")
    tblW2_after = tblPr2.find(qn("w:tblW"))
    assert_(tblW2_after.get(qn("w:type")) == "dxa", "tblW type is now 'dxa' (an explicit, unambiguous width)")
    assert_(tblW2_after.get(qn("w:w")) == str(sum(widths)), "tblW value is exactly the real sum of the column widths (" + str(sum(widths)) + ")")

    print("\n=== Test 3: a table that already has a real, non-ambiguous tblW is left untouched ===")
    doc3 = Document()
    t3 = doc3.add_table(rows=1, cols=2)
    tblPr3 = t3._tbl.find(qn("w:tblPr"))
    tblW3 = tblPr3.find(qn("w:tblW"))
    tblW3.set(qn("w:type"), "dxa")
    tblW3.set(qn("w:w"), "5000")
    fixed3 = asp._fix_ambiguous_table_width(doc3)
    assert_(fixed3 == 0, "Zero tables touched - already had a real, explicit width")
    tblW3_after = tblPr3.find(qn("w:tblW"))
    assert_(tblW3_after.get(qn("w:w")) == "5000", "The original explicit width value is completely unchanged")

    print("\nALL TESTS PASSED")


if __name__ == "__main__":
    run()
