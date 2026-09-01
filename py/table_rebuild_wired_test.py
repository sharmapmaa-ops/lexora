"""
Test confirming _rebuild_table_from_ocr_structure is GENUINELY WIRED
INTO the real pipeline (translate_existing_docx), not just a
standalone, tested-in-isolation function - per explicit, direct
feedback: "agar koi chiz bina wire kiye bheju to wo kis kaam ki
hogi". A pre-translation snapshot of every table is taken at the top
of translate_existing_docx (the SAME docx that's already Step-1/OCR
output, before translation touches it), and after every existing fix
runs, each table is rebuilt using that snapshot as its structure-
source and the fully-processed table as its content-source - no
separate uploaded OCR file needed, since Step 2 already has both
states internally in one execution.

Run: python3 table_rebuild_wired_test.py
"""
import os
import sys
import json
import tempfile

sys.path.insert(0, os.path.dirname(__file__))
from docx import Document
import aspose_test_pipeline as asp
import lease_engine as le


def assert_(cond, label):
    if not cond:
        print("FAIL:", label)
        sys.exit(1)
    print("PASS:", label)


def _mock_translate(prefix=""):
    def mock_call(llm_config, system_prompt, user_content, max_tokens=12000):
        req_items = json.loads(user_content)
        resp_items = [{"id": item["id"], "text": prefix + item["text"]} for item in req_items]
        return json.dumps(resp_items), "mock-provider"
    return mock_call


def run():
    print("=== Test 1 (THE REAL REPORTED DOCUMENT, FULL REAL PIPELINE): every real table gets rebuilt, none crash ===")
    real_ocr_docx = "/home/claude/work/two_docs_review/ocr_trace/ocr_step1.docx"
    if os.path.exists(real_ocr_docx):
        original = le._call_chat_completion_with_failover
        le._call_chat_completion_with_failover = _mock_translate("TR: ")
        try:
            output_path = tempfile.mktemp(suffix=".docx")
            result = asp.translate_existing_docx(real_ocr_docx, "English", output_path, llm_config={})
            assert_("tables_rebuilt" in result, "The real pipeline's own return dict reports rebuild stats")
            assert_(result["tables_rebuild_failed"] == 0, f"Zero tables failed to rebuild (got {result['tables_rebuild_failed']}) - the real gridSpan-less short-row bug this test also guards against")
            assert_(result["tables_rebuilt"] >= 30, f"The large majority of real tables were genuinely rebuilt (got {result['tables_rebuilt']})")

            reopened = Document(output_path)
            assert_(len(reopened.tables) >= 30, "The saved, reopened output document still has its real tables intact after the rebuild pass")

            found_translated = False
            for t in reopened.tables:
                for row in t.rows:
                    for cell in row.cells:
                        if cell.text.startswith("TR: "):
                            found_translated = True
                            break
            assert_(found_translated, "Rebuilt cells carry the ACTUAL translated content (not blank, not source-language leftovers)")
        finally:
            le._call_chat_completion_with_failover = original
    else:
        print("  (real document not present in this environment - skipped)")

    print("\n=== Test 2: a real irregular row (fewer physical cells than the table's nominal column count, no gridSpan) no longer crashes the whole table's rebuild ===")
    from docx.oxml.ns import qn
    from docx.table import Table
    ocr_doc = Document()
    ocr_t = ocr_doc.add_table(rows=2, cols=3)
    ocr_t.rows[0].cells[0].text = "A"
    ocr_t.rows[0].cells[1].text = "B"
    ocr_t.rows[0].cells[2].text = "C"
    # row 1: genuinely remove one cell, matching the real confirmed pattern (fewer <w:tc> than row 0, no gridSpan)
    row1_tr = ocr_t.rows[1]._tr
    tcs = row1_tr.findall(qn("w:tc"))
    row1_tr.remove(tcs[-1])

    output_doc = Document()
    output_t = output_doc.add_table(rows=2, cols=3)
    for c, txt in zip(output_t.rows[0].cells, ["X", "Y", "Z"]):
        c.text = txt

    result2 = asp._rebuild_table_from_ocr_structure(output_doc, output_t, ocr_t, "English")
    assert_(result2["rebuilt"] is True, "Rebuild succeeds even with a real irregular (short) row present, instead of throwing")

    print("\nALL TESTS PASSED")


if __name__ == "__main__":
    run()
