"""
Regression test for _fix_negative_usable_width_from_paragraph_indent,
per a real reported issue found during a comprehensive two-document
before/after review: a table cell rendered "Brokerage Fee (Not
included..." as one character per line for the entire remaining page
height. Root cause confirmed directly in the real document's XML: the
cell's own w:tcW was 1136 twips, but its paragraph carried
w:ind/@right=1677 twips - MORE than the cell's entire width. Combined
with the cell's own w:tcMar/@right=400, the real usable width was
1136 - 400 - 1677 = -941 twips, genuinely negative - Word had no space
to lay out even a single character normally, so it broke every
character onto its own line.

Run: python3 negative_usable_width_test.py
"""
import os
import sys

sys.path.insert(0, os.path.dirname(__file__))
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import aspose_test_pipeline as asp


def assert_(cond, label):
    if not cond:
        print("FAIL:", label)
        sys.exit(1)
    print("PASS:", label)


def run():
    print("=== Test 1 (THE REAL REPORTED CASE): real document - the exact negative-width cell is fixed ===")
    real_docx = "/home/claude/work/two_docs_review/doc1.docx"
    if os.path.exists(real_docx):
        doc = Document(real_docx)
        fixed = asp._fix_negative_usable_width_from_paragraph_indent(doc)
        assert_(fixed == 1, "Exactly 1 cell fixed (matches the real confirmed count)")

        # Verify the specific real cell no longer has a poisonous right-indent.
        remaining = 0
        for tc in doc.element.body.iter(qn("w:tc")):
            tcPr = tc.find(qn("w:tcPr"))
            tcW = tcPr.find(qn("w:tcW")) if tcPr is not None else None
            if tcW is None or not tcW.get(qn("w:w")):
                continue
            cell_width = int(tcW.get(qn("w:w")))
            tcMar = tcPr.find(qn("w:tcMar"))
            tcMar_right_el = tcMar.find(qn("w:right")) if tcMar is not None else None
            tcMar_right = int(tcMar_right_el.get(qn("w:w"))) if tcMar_right_el is not None and tcMar_right_el.get(qn("w:w")) else 0
            for p in tc.findall(qn("w:p")):
                pPr = p.find(qn("w:pPr"))
                ind = pPr.find(qn("w:ind")) if pPr is not None else None
                if ind is None or not ind.get(qn("w:right")):
                    continue
                right_indent = int(ind.get(qn("w:right")))
                if right_indent <= 0:
                    continue
                usable = cell_width - tcMar_right - right_indent
                if usable < 200:
                    remaining += 1
        assert_(remaining == 0, "No cell anywhere in the document retains a negative/near-zero usable width")
    else:
        print("  (real document not present in this environment - skipped)")

    print("\n=== Test 2 (synthetic, exact reported numbers): tcW=1136, right-indent=1677, tcMar=400 -> reset to 0 ===")
    doc2 = Document()
    t2 = doc2.add_table(rows=1, cols=1)
    cell2 = t2.rows[0].cells[0]
    tcPr2 = cell2._tc.get_or_add_tcPr()
    tcPr2.find(qn("w:tcW")).set(qn("w:w"), "1136")
    tcMar2 = OxmlElement("w:tcMar")
    right_mar = OxmlElement("w:right")
    right_mar.set(qn("w:w"), "400")
    right_mar.set(qn("w:type"), "dxa")
    tcMar2.append(right_mar)
    tcPr2.append(tcMar2)
    cell2.text = "Brokerage Fee (Not included in total contract amount):"
    pPr2 = cell2.paragraphs[0]._p.get_or_add_pPr()
    ind2 = OxmlElement("w:ind")
    ind2.set(qn("w:right"), "1677")
    pPr2.append(ind2)

    fixed2 = asp._fix_negative_usable_width_from_paragraph_indent(doc2)
    assert_(fixed2 == 1, "Exactly 1 fix applied")
    ind2_after = cell2.paragraphs[0]._p.find(qn("w:pPr")).find(qn("w:ind"))
    assert_(ind2_after.get(qn("w:right")) == "0", "The paragraph's poisonous right-indent is now 0")

    print("\n=== Test 3: a normal, reasonable right-indent that leaves real usable width is left completely alone ===")
    doc3 = Document()
    t3 = doc3.add_table(rows=1, cols=1)
    cell3 = t3.rows[0].cells[0]
    tcPr3 = cell3._tc.get_or_add_tcPr()
    tcPr3.find(qn("w:tcW")).set(qn("w:w"), "5000")
    cell3.text = "Some normal cell text."
    pPr3 = cell3.paragraphs[0]._p.get_or_add_pPr()
    ind3 = OxmlElement("w:ind")
    ind3.set(qn("w:right"), "300")  # 5000 - 300 = 4700, plenty of usable width
    pPr3.append(ind3)

    fixed3 = asp._fix_negative_usable_width_from_paragraph_indent(doc3)
    assert_(fixed3 == 0, "Zero fixes - this cell had plenty of real usable width, nothing wrong to fix")
    ind3_after = cell3.paragraphs[0]._p.find(qn("w:pPr")).find(qn("w:ind"))
    assert_(ind3_after.get(qn("w:right")) == "300", "The reasonable right-indent is completely untouched")

    print("\n=== Test 4: a cell with NO right-indent at all is never touched (no false positive from an absent property) ===")
    doc4 = Document()
    t4 = doc4.add_table(rows=1, cols=1)
    cell4 = t4.rows[0].cells[0]
    cell4._tc.get_or_add_tcPr().find(qn("w:tcW")).set(qn("w:w"), "100")  # even a tiny cell
    cell4.text = "Text with no indent at all."
    fixed4 = asp._fix_negative_usable_width_from_paragraph_indent(doc4)
    assert_(fixed4 == 0, "Zero fixes - there is no right-indent to reset in the first place")

    print("\nALL TESTS PASSED")


if __name__ == "__main__":
    run()
