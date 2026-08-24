"""
Regression test for the REDESIGNED _fix_paragraph_direction margin
handling: an RTL->LTR paragraph's own original left/right margin is now
MIRRORED (new_left = old_right, new_right = old_left), replacing the
earlier "flatten every paragraph to one hardcoded constant, discard the
original values" approach - per explicit direction, after that earlier
approach's underlying assumption (that per-paragraph indent variance
was meaningless Aspose noise) was never actually verified against the
original RTL document.

Uses the user's own worked example directly: left=2in, right=3in
(width unaffected) -> left=3in, right=2in.

Run: python3 margin_mirror_test.py
"""
import os
import sys

sys.path.insert(0, os.path.dirname(__file__))
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import aspose_test_pipeline as asp


def assert_(cond, label):
    if not cond:
        print("FAIL:", label)
        sys.exit(1)
    print("PASS:", label)


def _make_rtl_paragraph(doc, text, left_twips, right_twips):
    p = doc.add_paragraph(text)
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:bidi"))  # marks it as an RTL source paragraph
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left_twips))
    ind.set(qn("w:right"), str(right_twips))
    pPr.append(ind)
    return p


def run():
    print("=== Test 1: the user's own worked example - left=2in/right=3in -> left=3in/right=2in ===")
    doc = Document()
    p1 = _make_rtl_paragraph(doc, "Some translated body text.", Inches(2).twips if hasattr(Inches(2), "twips") else int(Inches(2) / 635), int(Inches(3) / 635))
    asp._fix_paragraph_direction(doc, "English")
    ind = p1._p.find(qn("w:pPr")).find(qn("w:ind"))
    new_left = int(ind.get(qn("w:left")))
    new_right = int(ind.get(qn("w:right")))
    assert_(new_left == int(Inches(3) / 635), "New left indent equals the ORIGINAL right indent (3in) - " + str(new_left))
    assert_(new_right == int(Inches(2) / 635), "New right indent equals the ORIGINAL left indent (2in) - " + str(new_right))

    print("\n=== Test 2: no hardcoded constant is used anymore ===")
    src = open(os.path.join(os.path.dirname(__file__), "aspose_test_pipeline.py")).read()
    fn_start = src.find("def _fix_paragraph_direction(")
    fn_body = src[fn_start:fn_start + 8000]
    assert_('"446"' not in fn_body, "The old hardcoded 446-twip constant is gone from the function")
    assert_("ind.set(qn(\"w:left\"), \"0\")" not in fn_body, "The old hardcoded heading left=0 constant is gone")
    assert_("_adjacent_table_indent" not in src, "The now-unneeded adjacent-table-indent lookup is fully removed, not just unused")

    print("\n=== Test 3: original right/firstLine/hanging are no longer deleted outright ===")
    assert_("del ind.attrib[qn(attr)]" not in fn_body, "The old delete-right/firstLine/hanging loop is gone - values are preserved (mirrored), not discarded")

    print("\n=== Test 4: a paragraph with NO original margin is left untouched (nothing to mirror) ===")
    doc2 = Document()
    p2 = doc2.add_paragraph("Plain paragraph, no explicit w:ind at all.")
    pPr2 = p2._p.get_or_add_pPr()
    pPr2.append(OxmlElement("w:bidi"))
    asp._fix_paragraph_direction(doc2, "English")
    ind2 = pPr2.find(qn("w:ind"))
    assert_(ind2 is None, "No w:ind element was invented out of nothing for a paragraph that never had one")

    print("\n=== Test 5: equal left/right (already symmetric) is a correct no-op ===")
    doc3 = Document()
    p3 = _make_rtl_paragraph(doc3, "Symmetric margins already.", 500, 500)
    asp._fix_paragraph_direction(doc3, "English")
    ind3 = p3._p.find(qn("w:pPr")).find(qn("w:ind"))
    assert_(int(ind3.get(qn("w:left"))) == 500 and int(ind3.get(qn("w:right"))) == 500, "Symmetric left==right margins stay exactly as they were")

    print("\n=== Test 6: table-cell paragraphs are still excluded (handled by table fixes instead) ===")
    doc4 = Document()
    t4 = doc4.add_table(rows=1, cols=1)
    cell_p = t4.rows[0].cells[0].paragraphs[0]
    cell_p.text = "Cell content"
    cPPr = cell_p._p.get_or_add_pPr()
    cPPr.append(OxmlElement("w:bidi"))
    cind = OxmlElement("w:ind")
    cind.set(qn("w:left"), "300")
    cind.set(qn("w:right"), "900")
    cPPr.append(cind)
    asp._fix_paragraph_direction(doc4, "English")
    cind_after = cPPr.find(qn("w:ind"))
    assert_(int(cind_after.get(qn("w:left"))) == 300 and int(cind_after.get(qn("w:right"))) == 900,
            "A table-cell paragraph's margin is untouched by this function (left=300/right=900 unchanged)")

    print("\nALL TESTS PASSED")


if __name__ == "__main__":
    run()
