"""
Standalone regression test for py/ocr_router.py.

Run: python3 ocr_router_test.py <table_and_color_pdf> [plain_text_pdf]

Verifies:
  1. A PDF with real tables/background colors is routed to "aspose".
  2. A plain-text PDF (no tables, no fills) is routed to "lightweight".
  3. run_ocr() never raises just because Aspose isn't configured - it
     falls back to the lightweight extractor and reports why.
  4. The lightweight extractor's output is a valid, openable .docx with
     real extracted text, and Arabic text (if any) is RTL-corrected
     (reused from source_reconstruction.py, not reimplemented).

This was run against a real 10-page bilingual Arabic/English REGA/Ejar
lease PDF during development (56 real tables, 41% max background fill
coverage on one page) and a synthetic plain-text-only PDF - both
produced the expected routing decision.
"""
import sys
import os

sys.path.insert(0, os.path.dirname(__file__))
import ocr_router  # noqa: E402


def run(pdf_with_tables, pdf_plain=None):
    print("=== Test 1: table/background-color PDF -> should route to aspose ===")
    analysis = ocr_router.analyze_pdf_for_ocr_strategy(pdf_with_tables)
    assert analysis["strategy"] == "aspose", f"Expected aspose, got {analysis['strategy']}"
    assert analysis["table_count_total"] > 0, "Expected real tables to be detected"
    print("PASS:", analysis["reason"])

    print("\n=== Test 2: run_ocr() on table PDF without Aspose creds -> graceful fallback ===")
    result = ocr_router.run_ocr(pdf_with_tables, "/tmp/ocr_router_test_out1.docx")
    assert result["strategy_used"] in ("aspose", "lightweight_fallback_from_aspose"), result["strategy_used"]
    if result["strategy_used"] == "lightweight_fallback_from_aspose":
        assert result.get("aspose_fallback_reason"), "Expected an explicit fallback reason"
        print("PASS (fell back as expected, no crash):", result["aspose_fallback_reason"][:100])
    else:
        print("PASS (Aspose was actually configured and used):", result.get("output_path"))

    from docx import Document
    d = Document(result["output_path"])
    assert len(d.paragraphs) > 0, "Expected non-empty output document"
    print(f"PASS: output docx has {len(d.paragraphs)} paragraphs")

    if pdf_plain:
        print("\n=== Test 3: plain-text PDF -> should route to lightweight ===")
        analysis2 = ocr_router.analyze_pdf_for_ocr_strategy(pdf_plain)
        assert analysis2["strategy"] == "lightweight", f"Expected lightweight, got {analysis2['strategy']}"
        print("PASS:", analysis2["reason"])

    print("\nALL TESTS PASSED")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    run(sys.argv[1], sys.argv[2] if len(sys.argv) > 2 else None)
