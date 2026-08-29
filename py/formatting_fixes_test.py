"""
Regression tests for 3 real formatting bugs reported against a real
Aspose-translated document:
  Issue 1/4: body paragraphs carried a stray, essentially-random
             w:right indent (confirmed: 78 of ~120 real paragraphs had
             exactly -200 twips, letting text overrun the true right
             margin before wrapping) - fixed by resetting to 0.
  Issue 2:   two adjacent numbered sub-clauses (13-1-5, 14-1-5) were
             merged into one paragraph with no break - confirmed the
             merge was WITHIN A SINGLE RUN's text (not a run boundary,
             unlike the existing heading-merge fix) - fixed by
             splitting the run's text and the paragraph at that point.
  Issue 3:   "Article Four: Rent"'s background color was on the RUN's
             own rPr, not the paragraph's pPr, so only the exact
             characters were highlighted instead of the full row -
             fixed by promoting uniform run-level shading to the
             paragraph level.
  Follow-up (real reported issues 1/2/3, second round): 9 genuinely
             long body clause paragraphs carried w:jc="center" -
             confirmed via a real LibreOffice render this produces a
             "staircase" look (each wrapped line centers on its own
             width independently), which explains all three symptoms
             at once - fixed by resetting center to "both", matching
             this document's own dominant real norm.
  Follow-up round 2 (real reported margin/continuity issues, after
             direct comparison against a user-provided reference
             image): the original right-indent-only fix wasn't enough
             - w:left was still random per paragraph, w:firstLine was
             never reset, and what reads as one flowing clause was
             often stored as MULTIPLE separate paragraphs with no real
             sentence break between them. Fixed via the user's own
             explicit methodology (derive left/right margins from the
             real "Contract Data" table's own width/position),
             resetting firstLine, merging genuine continuation
             paragraphs back into one real paragraph, and normalizing
             the "before" spacing on clause-start paragraphs so
             adjacent clauses get a consistent visible gap.

Run: python3 formatting_fixes_test.py
"""
import os
import sys

sys.path.insert(0, os.path.dirname(__file__))
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Twips
import aspose_test_pipeline as asp


def assert_(cond, label):
    if not cond:
        print("FAIL:", label)
        sys.exit(1)
    print("PASS:", label)


def run():
    print("=== Test 1: real document - margin normalization (issues 1 and 4) ===")
    real_docx = "/home/claude/work/aspose_debug4/input7.docx"
    if os.path.exists(real_docx):
        doc = Document(real_docx)
        fixed = asp._fix_body_paragraph_right_indent(doc)
        assert_(fixed > 100, "A large, real number of body paragraphs had their margin normalized (got " + str(fixed) + ")")
        # spot-check: every non-table paragraph's right-indent should now
        # be a SINGLE consistent value (either the table-derived one, or
        # 0 if no reference table was found) - not scattered per-paragraph
        right_values = set()
        for p in doc.paragraphs:
            if not p.text.strip():
                continue
            if p._p.getparent().tag == qn("w:tc"):
                continue
            pPr = p._p.find(qn("w:pPr"))
            ind = pPr.find(qn("w:ind")) if pPr is not None else None
            if ind is not None and ind.get(qn("w:right")) is not None:
                right_values.add(ind.get(qn("w:right")))
        assert_(len(right_values) == 1, "Every non-table paragraph now shares the exact same right-indent value (got " + str(right_values) + ")")
    else:
        print("  (real document not present in this environment - using synthetic test below instead)")

    print("\n=== Test 2 (synthetic, exact mechanics): no reference table found -> falls back to right=0, table cells untouched ===")
    doc2 = Document()
    p2 = doc2.add_paragraph("Some real body text that should wrap at the true margin.")
    pPr2 = p2._p.get_or_add_pPr()
    ind2 = OxmlElement("w:ind")
    ind2.set(qn("w:left"), "771")
    ind2.set(qn("w:right"), "-200")
    pPr2.append(ind2)
    t2 = doc2.add_table(rows=1, cols=1)
    # A first-cell text this long deliberately does NOT look like a real
    # section-header label (e.g. "Contract Data") - keeps
    # _find_reference_margin_table from matching this unrelated table,
    # so this test genuinely exercises the "no reference table found"
    # fallback path, not the table-derived one.
    long_cell_text = "This is a long paragraph of body text sitting inside a table cell, not a real section header"
    cell_p = t2.rows[0].cells[0].paragraphs[0]
    cell_p.text = long_cell_text
    cPPr = cell_p._p.get_or_add_pPr()
    cind = OxmlElement("w:ind")
    cind.set(qn("w:right"), "-999")
    cPPr.append(cind)
    fixed2 = asp._fix_body_paragraph_right_indent(doc2)
    assert_(fixed2 == 1, "Exactly 1 body paragraph fixed (the table cell is untouched)")
    ind_after = p2._p.find(qn("w:pPr")).find(qn("w:ind"))
    assert_(ind_after.get(qn("w:right")) == "0", "No reference table found - falls back to resetting right-indent to 0")
    assert_(ind_after.get(qn("w:left")) == "771", "No reference table found - left-indent is left untouched (no basis to change it)")
    cell_ind_after = cell_p._p.find(qn("w:pPr")).find(qn("w:ind"))
    assert_(cell_ind_after.get(qn("w:right")) == "-999", "The table-cell paragraph's right-indent is NOT touched")

    print("\n=== Test 2b (THE REAL REPORTED CASE): a document with an UNRELATED short-labeled table (no real Contract Data table at all) correctly returns None, not the wrong table ===")
    doc2b = Document()
    t2b_wrong = doc2b.add_table(rows=1, cols=1)
    t2b_wrong.rows[0].cells[0].text = "Issuer:"  # the exact real false-match this bug produced
    t2b_right = doc2b.add_table(rows=1, cols=1)
    t2b_right.rows[0].cells[0].text = "Property Data"  # also a real, but unrelated, short-labeled table
    found2b = asp._find_reference_margin_table(doc2b)
    assert_(found2b is None, "No table matches 'Contract Data' - correctly returns None instead of the first unrelated short-labeled table")
    left2b, right2b = asp._compute_table_derived_margins(doc2b)
    assert_(left2b is None and right2b is None, "Margin computation correctly reports 'no reference available' rather than silently using the wrong table's numbers")

    print("\n=== Test 2c: a real 'Contract Data' table IS found even with translation/OCR variance (leading number, double spaces) ===")
    doc2c = Document()
    t2c_wrong = doc2c.add_table(rows=1, cols=1)
    t2c_wrong.rows[0].cells[0].text = "Issuer:"
    t2c_right = doc2c.add_table(rows=1, cols=2)
    t2c_right.rows[0].cells[0].text = "1  Contract  Data"  # real OCR-style variance: leading number + double spaces
    grid2c = t2c_right._tbl.find(qn("w:tblGrid"))
    for c, w in zip(grid2c.findall(qn("w:gridCol")), ["867", "8571"]):
        c.set(qn("w:w"), w)
    tblPr2c = t2c_right._tbl.find(qn("w:tblPr"))
    tblInd2c = OxmlElement("w:tblInd")
    tblInd2c.set(qn("w:w"), "867")
    tblInd2c.set(qn("w:type"), "dxa")
    tblPr2c.append(tblInd2c)
    found2c = asp._find_reference_margin_table(doc2c)
    assert_(found2c is not None, "The real Contract Data table (with real-world OCR variance) is correctly found")
    assert_(found2c.rows[0].cells[0].text.strip() == "1  Contract  Data", "It's genuinely the Contract Data table, not the earlier unrelated 'Issuer:' one")

    print("\n=== Test 3: merged sub-clause split (issue 2) - the exact reported case ===")
    doc3 = Document()
    p3 = doc3.add_paragraph(
        "13-1-5 The Tenant and its employees shall not smoke in the corridors and lobbies of the property and "
        "shall comply with smoking in designated areas. 14-1-5 The Tenant shall dispose of waste in the "
        "designated place and shall not throw any solid materials or oils into the sewage drains."
    )
    before_count = len(doc3.paragraphs)
    fixed3 = asp._fix_merged_numbered_subclause(doc3)
    assert_(fixed3 == 1, "Exactly 1 paragraph split")
    after_count = len(doc3.paragraphs)
    assert_(after_count == before_count + 1, "Paragraph count increased by exactly 1")
    texts = [p.text for p in doc3.paragraphs if p.text.strip()]
    assert_(texts[0].startswith("13-1-5") and "14-1-5" not in texts[0], "First paragraph is ONLY clause 13-1-5")
    assert_(texts[1].startswith("14-1-5"), "Second paragraph is clause 14-1-5, on its own")

    print("\n=== Test 4: merged sub-clause split - a normal, ALREADY-separate clause is left alone (no false positive) ===")
    doc4 = Document()
    doc4.add_paragraph("12-1-5 The Tenant shall maintain all mechanical and electrical equipment.")
    doc4.add_paragraph("13-1-5 The Tenant and its employees shall not smoke in the corridors.")
    fixed4 = asp._fix_merged_numbered_subclause(doc4)
    assert_(fixed4 == 0, "Zero splits - these were already two separate, correctly-formed paragraphs")

    print("\n=== Test 3b (THE REAL FOLLOW-UP GAP, found via a systematic re-check): a 2-segment clause code (N-N, not N-N-N) is also caught ===")
    doc3b = Document()
    p3b = doc3b.add_paragraph(
        "10-1 This contract shall be effective from the date of its signing and shall be binding on both "
        "parties, and shall be renewed according to what is stipulated in Article 3 of the parties' "
        "obligations under this contract. 10-2 Non-compliance by one of the parties with any clause or "
        "article of this contract does not affect the validity of its remaining provisions."
    )
    before_count3b = len(doc3b.paragraphs)
    fixed3b = asp._fix_merged_numbered_subclause(doc3b)
    assert_(fixed3b == 1, "Exactly 1 split (the original regex only matched N-N-N, missing this real N-N case)")
    after_count3b = len(doc3b.paragraphs)
    assert_(after_count3b == before_count3b + 1, "Paragraph count increased by exactly 1")
    texts3b = [p.text for p in doc3b.paragraphs if p.text.strip()]
    assert_(texts3b[0].startswith("10-1") and "10-2" not in texts3b[0], "First paragraph is ONLY clause 10-1")
    assert_(texts3b[1].startswith("10-2"), "Second paragraph is clause 10-2, on its own")

    print("\n=== Test 5: full-width shading promotion (issue 3) - the exact reported case ===")
    doc5 = Document()
    p5 = doc5.add_paragraph()
    run5 = p5.add_run("Article Four: Rent")
    rPr5 = run5._r.get_or_add_rPr()
    shd5 = OxmlElement("w:shd")
    shd5.set(qn("w:val"), "clear")
    shd5.set(qn("w:color"), "auto")
    shd5.set(qn("w:fill"), "DDDDDD")
    rPr5.append(shd5)
    fixed5 = asp._promote_uniform_run_shading_to_paragraph(doc5)
    assert_(fixed5 == 1, "Exactly 1 paragraph promoted")
    pPr5 = p5._p.find(qn("w:pPr"))
    para_shd = pPr5.find(qn("w:shd")) if pPr5 is not None else None
    assert_(para_shd is not None and para_shd.get(qn("w:fill")) == "DDDDDD", "Paragraph-level shading now carries the DDDDDD fill (spans full width)")
    run_rpr_after = run5._r.find(qn("w:rPr"))
    run_shd_after = run_rpr_after.find(qn("w:shd")) if run_rpr_after is not None else None
    assert_(run_shd_after is None, "The now-redundant run-level shading was removed")

    print("\n=== Test 6: shading promotion - MIXED colors in one paragraph are left alone (no false positive) ===")
    doc6 = Document()
    p6 = doc6.add_paragraph()
    run6a = p6.add_run("Red part")
    rPr6a = run6a._r.get_or_add_rPr()
    shd6a = OxmlElement("w:shd")
    shd6a.set(qn("w:fill"), "FF0000")
    rPr6a.append(shd6a)
    run6b = p6.add_run(" normal part, no shading")
    fixed6 = asp._promote_uniform_run_shading_to_paragraph(doc6)
    assert_(fixed6 == 0, "Zero promotions - the paragraph has mixed shading (one run shaded, one not), correctly left alone")

    print("\n=== Test 7: center-alignment fix (follow-up issues 1/2/3) - the exact reported case ===")
    real_docx2 = "/home/claude/work/margin_debug/final_output.docx"
    if os.path.exists(real_docx2):
        doc7 = Document(real_docx2)
        fixed7 = asp._fix_body_paragraph_center_alignment(doc7)
        assert_(fixed7 == 9, "Exactly 9 paragraphs had center alignment fixed (matches the real confirmed count)")
        remaining_center = 0
        for p in doc7.paragraphs:
            if not p.text.strip():
                continue
            if p._p.getparent().tag == qn("w:tc"):
                continue
            pPr = p._p.find(qn("w:pPr"))
            jc = pPr.find(qn("w:jc")) if pPr is not None else None
            if jc is not None and jc.get(qn("w:val")) == "center":
                remaining_center += 1
        assert_(remaining_center == 0, "No non-table paragraph retains center alignment after the fix")
    else:
        print("  (real document not present in this environment - using synthetic test below instead)")

    print("\n=== Test 8 (synthetic, exact mechanics): center -> both, left/both untouched, table cells untouched ===")
    doc8 = Document()
    p8a = doc8.add_paragraph("A long clause paragraph that was wrongly center-aligned by Aspose's conversion process here.")
    pPr8a = p8a._p.get_or_add_pPr()
    jc8a = OxmlElement("w:jc")
    jc8a.set(qn("w:val"), "center")
    pPr8a.append(jc8a)

    p8b = doc8.add_paragraph("A paragraph that was already correctly justified.")
    pPr8b = p8b._p.get_or_add_pPr()
    jc8b = OxmlElement("w:jc")
    jc8b.set(qn("w:val"), "both")
    pPr8b.append(jc8b)

    p8c = doc8.add_paragraph("A left-aligned paragraph.")
    pPr8c = p8c._p.get_or_add_pPr()
    jc8c = OxmlElement("w:jc")
    jc8c.set(qn("w:val"), "left")
    pPr8c.append(jc8c)

    t8 = doc8.add_table(rows=1, cols=1)
    cell_p8 = t8.rows[0].cells[0].paragraphs[0]
    cell_p8.text = "Cell text"
    cPPr8 = cell_p8._p.get_or_add_pPr()
    cjc8 = OxmlElement("w:jc")
    cjc8.set(qn("w:val"), "center")
    cPPr8.append(cjc8)

    fixed8 = asp._fix_body_paragraph_center_alignment(doc8)
    assert_(fixed8 == 1, "Exactly 1 paragraph fixed (only the wrongly-centered one)")
    assert_(pPr8a.find(qn("w:jc")).get(qn("w:val")) == "both", "The wrongly-centered paragraph is now 'both' (justify)")
    assert_(pPr8b.find(qn("w:jc")).get(qn("w:val")) == "both", "An already-'both' paragraph is untouched")
    assert_(pPr8c.find(qn("w:jc")).get(qn("w:val")) == "left", "A 'left' paragraph is untouched - not forced to 'both'")
    assert_(cjc8.get(qn("w:val")) == "center", "A table-cell paragraph's center alignment is NOT touched")

    print("\n=== Test 9: table-derived margin computation - the exact reported case ===")
    real_docx3 = "/home/claude/work/margin_debug2/final_output.docx"
    if os.path.exists(real_docx3):
        doc9 = Document(real_docx3)
        left, right = asp._compute_table_derived_margins(doc9)
        assert_(left == 867, "Computed target left margin is 867 twips (matches the real Contract Data table's own indent)")
        assert_(right == 555, "Computed target right margin is 555 twips (matches where the real table ends)")
    else:
        print("  (real document not present in this environment - skipped)")

    print("\n=== Test 10: right-indent fix now also fixes left + firstLine, using table-derived values ===")
    doc10 = Document()
    sec10 = doc10.sections[0]
    sec10.page_width = Twips(11900)
    sec10.left_margin = Twips(540)
    sec10.right_margin = Twips(500)
    t10 = doc10.add_table(rows=1, cols=1)
    t10.rows[0].cells[0].text = "Contract Data"
    tblPr10 = t10._tbl.find(qn("w:tblPr"))
    tblInd10 = OxmlElement("w:tblInd")
    tblInd10.set(qn("w:w"), "867")
    tblInd10.set(qn("w:type"), "dxa")
    tblPr10.append(tblInd10)
    grid10 = t10._tbl.find(qn("w:tblGrid"))
    gc10 = grid10.find(qn("w:gridCol"))
    gc10.set(qn("w:w"), "9438")

    p10 = doc10.add_paragraph("Some real body text with a wrong random indent.")
    pPr10 = p10._p.get_or_add_pPr()
    ind10 = OxmlElement("w:ind")
    ind10.set(qn("w:left"), "1742")
    ind10.set(qn("w:right"), "-200")
    ind10.set(qn("w:firstLine"), "229")
    pPr10.append(ind10)

    fixed10 = asp._fix_body_paragraph_right_indent(doc10)
    assert_(fixed10 == 1, "Exactly 1 body paragraph fixed")
    ind10_after = p10._p.find(qn("w:pPr")).find(qn("w:ind"))
    assert_(ind10_after.get(qn("w:left")) == "867", "Left indent is now the table-derived value (867)")
    assert_(ind10_after.get(qn("w:right")) == "555", "Right indent is now the table-derived value (555)")
    assert_(ind10_after.get(qn("w:firstLine")) == "0", "firstLine is reset to 0")

    print("\n=== Test 11: continuation-paragraph merge - the exact reported case (14-1 split across 2 paragraphs) ===")
    doc11 = Document()
    doc11.add_paragraph(
        "14-1 The lessor shall return the security deposit amount specified in clause number (12) of this "
        "contract after vacating the rental unit, provided that there are no damages or defects caused by the "
        "tenant to the rental unit and that there are no outstanding bills that have not been paid or any "
        "amounts due from the tenant relating to general services or rent, and in these cases"
    )
    doc11.add_paragraph(
        "the value of bills or rent or general services or the cost of repairing damages shall be deducted "
        "from the paid security deposit amount and the remainder, if any, shall be returned to the tenant."
    )
    doc11.add_paragraph("14-2 The lessor bears the burden of proving the amounts due or damages claimed.")
    before_count11 = len(doc11.paragraphs)
    merged11 = asp._merge_continuation_paragraphs(doc11)
    assert_(merged11 == 1, "Exactly 1 continuation paragraph merged")
    after_count11 = len(doc11.paragraphs)
    assert_(after_count11 == before_count11 - 1, "Paragraph count decreased by exactly 1")
    texts11 = [p.text for p in doc11.paragraphs if p.text.strip()]
    assert_("and in these cases the value of bills" in texts11[0], "The merged paragraph reads as one continuous sentence")
    assert_(texts11[1].startswith("14-2"), "14-2 remains its own separate paragraph, untouched")

    print("\n=== Test 12: continuation-merge - real 'Label: Value' fields are NEVER merged (no false positive) ===")
    doc12 = Document()
    doc12.add_paragraph("This contract is considered an authenticated contract dated 3/4/1435 AH")
    doc12.add_paragraph("Name: Al-Anoud Sulaiman Ali Al-Masoud")
    merged12 = asp._merge_continuation_paragraphs(doc12)
    assert_(merged12 == 0, "Zero merges - a real 'Name:' field is never merged into unrelated preceding text")

    print("\n=== Test 13: continuation-merge - adjacent short 'Label:' fields are never merged into each other ===")
    doc13 = Document()
    doc13.add_paragraph("Unit Type: Office")
    doc13.add_paragraph("Special sign")
    merged13 = asp._merge_continuation_paragraphs(doc13)
    assert_(merged13 == 0, "Zero merges - 'Unit Type: Office' itself looks like a field, so nothing merges into it")

    print("\n=== Test 14: clause-start spacing normalization ===")
    doc14 = Document()
    p14a = doc14.add_paragraph("14-1 First clause.")
    pPr14a = p14a._p.get_or_add_pPr()
    sp14a = OxmlElement("w:spacing")
    sp14a.set(qn("w:before"), "22")
    pPr14a.append(sp14a)
    p14b = doc14.add_paragraph("14-2 Second clause.")
    pPr14b = p14b._p.get_or_add_pPr()
    sp14b = OxmlElement("w:spacing")
    sp14b.set(qn("w:before"), "4")
    pPr14b.append(sp14b)
    fixed14 = asp._fix_clause_start_spacing(doc14)
    assert_(fixed14 == 2, "Both clause-start paragraphs got their spacing normalized")
    assert_(sp14a.get(qn("w:before")) == "134", "First clause now has the consistent gap value")
    assert_(sp14b.get(qn("w:before")) == "134", "Second clause now has the consistent gap value")

    print("\nALL TESTS PASSED")


if __name__ == "__main__":
    run()
