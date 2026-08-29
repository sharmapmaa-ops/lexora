"""
Regression test for _fix_within_cell_self_duplicated_label, per a real
user-provided screenshot: several table header cells carried their own
label text duplicated, concatenated with no separator - 'Rent
ValueRent value', 'VATVAT', 'Total ValueTotal value', 'Issue Date
(AD)Issued Date(AD)', etc. A systematic whole-document scan (not
scoped to just the examples shown) found 9 real instances in the
reported document.

Run: python3 within_cell_duplicate_label_test.py
"""
import os
import sys

sys.path.insert(0, os.path.dirname(__file__))
from docx import Document
from docx.oxml.ns import qn
import aspose_test_pipeline as asp


def assert_(cond, label):
    if not cond:
        print("FAIL:", label)
        sys.exit(1)
    print("PASS:", label)


def _set_cell_text(cell, text):
    cell.text = text


def run():
    print("=== Test 1 (THE REAL REPORTED DOCUMENT): comprehensive whole-document scan ===")
    real_docx = "/home/claude/work/two_docs_review/doc1.docx"
    if os.path.exists(real_docx):
        doc = Document(real_docx)
        fixed = asp._fix_within_cell_self_duplicated_label(doc)
        assert_(fixed == 9, "Exactly 9 real cells fixed (matches the real confirmed whole-document count)")

        expected_clean = {
            "Expiry Date", "Rent value", "VAT", "Total value",
            "Issued Date(AD)", "Due Date(AD)", "Issued Date(AH)", "Due Date(AH)",
        }
        found = set()
        for tc in doc.element.body.iter(qn("w:tc")):
            ts = tc.findall(".//" + qn("w:t"))
            text = "".join(t.text or "" for t in ts).strip()
            if text in expected_clean:
                found.add(text)
        assert_(expected_clean <= found, "Every real reported example cleanly resolved to its de-duplicated form")

        # false-positive protection: genuinely different-content cells untouched
        untouched_found = []
        for tc in doc.element.body.iter(qn("w:tc")):
            ts = tc.findall(".//" + qn("w:t"))
            text = "".join(t.text or "" for t in ts).strip()
            if text in ("Serial Number.No", "Fixed Amounts ValueServices"):
                untouched_found.append(text)
        assert_(len(untouched_found) == 2, "Both genuinely-different-content cells (not self-duplicates) remain completely untouched")
    else:
        print("  (real document not present in this environment - skipped)")

    print("\n=== Test 2 (synthetic, exact reported case): the tricky bracket-boundary split ===")
    doc2 = Document()
    t2 = doc2.add_table(rows=1, cols=1)
    _set_cell_text(t2.rows[0].cells[0], "Issue Date (AD)Issued Date(AD)")
    fixed2 = asp._fix_within_cell_self_duplicated_label(doc2)
    assert_(fixed2 == 1, "One cell fixed")
    assert_(t2.rows[0].cells[0].text == "Issued Date(AD)", "Resolves to the clean, bracket-balanced form with no stray leading punctuation")

    print("\n=== Test 3 (synthetic): exact duplication (VATVAT) ===")
    doc3 = Document()
    t3 = doc3.add_table(rows=1, cols=1)
    _set_cell_text(t3.rows[0].cells[0], "VATVAT")
    fixed3 = asp._fix_within_cell_self_duplicated_label(doc3)
    assert_(fixed3 == 1, "One cell fixed")
    assert_(t3.rows[0].cells[0].text == "VAT", "Resolves to a single clean 'VAT'")

    print("\n=== Test 4 (synthetic): case-varying duplication (Rent ValueRent value) ===")
    doc4 = Document()
    t4 = doc4.add_table(rows=1, cols=1)
    _set_cell_text(t4.rows[0].cells[0], "Rent ValueRent value")
    fixed4 = asp._fix_within_cell_self_duplicated_label(doc4)
    assert_(fixed4 == 1, "One cell fixed")
    assert_(t4.rows[0].cells[0].text in ("Rent Value", "Rent value"), "Resolves to a single clean copy of the label")

    print("\n=== Test 5: a normal, non-duplicated cell is never touched ===")
    doc5 = Document()
    t5 = doc5.add_table(rows=1, cols=1)
    _set_cell_text(t5.rows[0].cells[0], "This is a genuinely long sentence with no self-repetition at all.")
    fixed5 = asp._fix_within_cell_self_duplicated_label(doc5)
    assert_(fixed5 == 0, "Zero fixes - ordinary content is left completely alone")

    print("\n=== Test 6: two genuinely DIFFERENT concatenated labels (not a self-duplicate) are left alone ===")
    doc6 = Document()
    t6 = doc6.add_table(rows=1, cols=1)
    _set_cell_text(t6.rows[0].cells[0], "Fixed Amounts ValueServices")
    fixed6 = asp._fix_within_cell_self_duplicated_label(doc6)
    assert_(fixed6 == 0, "Zero fixes - these are two different label fragments, not a self-repeat")

    print("\nALL TESTS PASSED")


if __name__ == "__main__":
    run()
