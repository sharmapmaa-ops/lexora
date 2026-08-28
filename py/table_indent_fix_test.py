"""
Regression test for _fix_table_overflow_indent, covering its full real
history:
  V1/V2 bug (older, still covered by Tests 1-5 via the fallback path
  when no reference table exists): the correction only triggered when
  a table's indent+width literally overflowed the page, and even then
  a formula bug made it a no-op exactly at the boundary.

  THE REAL, CURRENT BUG (Test 6, the case that matters most): a real
  user reported a table sitting at 50%+ of the page width in a real
  MS Word screenshot that NEITHER earlier version touched at all -
  because that table's own width was narrow enough that indent+width
  never reached the overflow threshold. The bug was the CONDITION
  itself ("only fix it if it overflows"), not the formula. Current
  fix: when a real reference table (e.g. "Contract Data") exists,
  EVERY table's position is forced to match it unconditionally - not
  gated on whether it currently overflows.

This test uses synthetic docs built to reproduce the exact real
numbers found in reported documents, rather than requiring the actual
uploaded files to be present, so it can run standalone and keep
catching this regression.

Run: python3 table_indent_fix_test.py
"""
import os
import sys

sys.path.insert(0, os.path.dirname(__file__))
from docx import Document
from docx.shared import Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import aspose_test_pipeline as asp


def assert_(cond, label):
    if not cond:
        print("FAIL:", label)
        sys.exit(1)
    print("PASS:", label)


def _make_table(doc, indent_twips, col_widths_twips):
    tbl = doc.add_table(rows=1, cols=len(col_widths_twips))
    tbl_el = tbl._tbl
    tblPr = tbl_el.find(qn("w:tblPr"))
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:type"), "dxa")
    tblInd.set(qn("w:w"), str(indent_twips))
    tblPr.append(tblInd)
    grid = tbl_el.find(qn("w:tblGrid"))
    for gc, w in zip(grid.findall(qn("w:gridCol")), col_widths_twips):
        gc.set(qn("w:w"), str(w))
    return tbl_el


def run():
    doc = Document()
    sec = doc.sections[0]
    # Match the real document's geometry closely enough that
    # content_width_twips comes out to the same real value (10860).
    sec.page_width = Twips(11900)
    sec.page_height = Twips(16840)
    sec.left_margin = Twips(540)
    sec.right_margin = Twips(500)

    # 41 normal tables (real widths/indents from the actual document,
    # a representative sample of the observed 691/771/867 values).
    normal_indents = [867, 771, 867, 867, 771, 691, 691, 691, 867, 771,
                       691, 691, 771, 771, 691, 691, 867, 867, 771, 771,
                       771, 771, 771, 771, 771, 771, 867, 867, 771, 771,
                       861, 867, 771, 771, 771, 691, 771, 691, 691, 830, 830]
    for ind in normal_indents:
        _make_table(doc, ind, [940, 1986, 6497])  # sums to 9423, same shape as the real anomalous tables

    # 4 real anomalous tables (the exact reported case).
    anomalous_els = []
    for ind in (1437, 1437, 1438, 1438):
        anomalous_els.append(_make_table(doc, ind, [940, 1986, 6497]))

    print("=== Test 1: content width matches the real document's real value ===")
    content_width_twips = round((sec.page_width - sec.left_margin - sec.right_margin) / 635)
    assert_(content_width_twips == 10860, "content_width_twips is 10860, matching the real reported document")

    print("\n=== Test 2: the fix reports exactly 4 tables fixed ===")
    fixed = asp._fix_table_overflow_indent(doc)
    assert_(fixed == 4, "Exactly 4 tables reported as fixed (got " + str(fixed) + ")")

    print("\n=== Test 3: the fix produces a REAL, visible change (not the no-op from the original bug) ===")
    for i, tbl_el in enumerate(anomalous_els):
        tblPr = tbl_el.find(qn("w:tblPr"))
        tblInd = tblPr.find(qn("w:tblInd"))
        new_val = int(tblInd.get(qn("w:w")))
        assert_(new_val != 1437 and new_val != 1438, "Anomalous table " + str(i) + "'s indent genuinely changed (is now " + str(new_val) + ", was 1437/1438)")

    print("\n=== Test 4: anomalous tables now match their siblings' real median indent ===")
    for i, tbl_el in enumerate(anomalous_els):
        tblPr = tbl_el.find(qn("w:tblPr"))
        tblInd = tblPr.find(qn("w:tblInd"))
        new_val = int(tblInd.get(qn("w:w")))
        assert_(new_val == 771, "Anomalous table " + str(i) + "'s indent is now 771 (the real median of the 41 normal tables) - got " + str(new_val))

    print("\n=== Test 5: normal tables are completely untouched ===")
    all_tbls = list(doc.element.body.iter(qn("w:tbl")))
    normal_tbl_els = all_tbls[:41]
    all_unchanged = True
    for i, (tbl_el, expected) in enumerate(zip(normal_tbl_els, normal_indents)):
        tblPr = tbl_el.find(qn("w:tblPr"))
        tblInd = tblPr.find(qn("w:tblInd"))
        actual = int(tblInd.get(qn("w:w")))
        if actual != expected:
            all_unchanged = False
            print("  unexpected change at table", i, ":", expected, "->", actual)
    assert_(all_unchanged, "All 41 normal tables retain their original indent exactly")

    print("\n=== Test 6 (THE REAL REPORTED CASE): with a real 'Contract Data' reference table present, EVERY table's position is forced to match it - unconditionally, not only when it overflows ===")
    doc2 = Document()
    sec2 = doc2.sections[0]
    sec2.page_width = Twips(11900)
    sec2.page_height = Twips(16840)
    sec2.left_margin = Twips(540)
    sec2.right_margin = Twips(500)

    # The real reference table, exactly as found in the real document.
    ref_table = doc2.add_table(rows=1, cols=2)
    ref_table.rows[0].cells[0].text = "Contract Data"
    ref_tblPr = ref_table._tbl.find(qn("w:tblPr"))
    ref_tblInd = OxmlElement("w:tblInd")
    ref_tblInd.set(qn("w:w"), "867")
    ref_tblInd.set(qn("w:type"), "dxa")
    ref_tblPr.append(ref_tblInd)
    ref_grid = ref_table._tbl.find(qn("w:tblGrid"))
    ref_cols = ref_grid.findall(qn("w:gridCol"))
    ref_cols[0].set(qn("w:w"), "4695")
    ref_cols[1].set(qn("w:w"), "4743")

    # A NARROW table sitting anomalously far right - the exact real bug:
    # indent=6000 (way more than the reference's 867), but its own width
    # (2000) is narrow enough that indent+width=8000 never reaches the
    # content-width overflow threshold (10860) - the old fix would
    # NEVER have touched this table at all.
    narrow_el = _make_table(doc2, 6000, [1000, 1000])

    # A table that ALREADY happens to look "normal" under the OLD
    # anomaly definition (small indent, real buffer) but still doesn't
    # match the reference exactly - must ALSO be corrected now, since
    # the new rule is "match the reference", not "don't overflow".
    almost_right_el = _make_table(doc2, 700, [3467, 3226, 2730])

    fixed6 = asp._fix_table_overflow_indent(doc2)
    assert_(fixed6 == 2, "Both non-matching tables were fixed (got " + str(fixed6) + ") - the reference table itself needs no change and isn't counted")

    narrow_ind_after = narrow_el.find(qn("w:tblPr")).find(qn("w:tblInd"))
    assert_(narrow_ind_after.get(qn("w:w")) == "867", "The narrow, far-right table (real reported bug: 50%+ indent, never overflowed) is now at the reference position 867")

    almost_ind_after = almost_right_el.find(qn("w:tblPr")).find(qn("w:tblInd"))
    assert_(almost_ind_after.get(qn("w:w")) == "867", "A table that was already 'not overflowing' but still didn't match the reference is ALSO corrected to 867")

    ref_ind_after = ref_table._tbl.find(qn("w:tblPr")).find(qn("w:tblInd"))
    assert_(ref_ind_after.get(qn("w:w")) == "867", "The reference table itself is untouched (already correct)")

    print("\n=== Test 7: re-running the fix again is a safe no-op (every table already matches) ===")
    fixed7 = asp._fix_table_overflow_indent(doc2)
    assert_(fixed7 == 0, "Zero tables reported as changed on a second run - nothing left to fix")

    print("\nALL TESTS PASSED")


if __name__ == "__main__":
    run()
