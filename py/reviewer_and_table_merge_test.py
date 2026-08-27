"""
Regression tests for 2 real bugs reported against real documents:
  Issue 1/2: the document reviewer (Step 3) was reverting a LEGITIMATE
             center->both alignment fix from Step 2 back to "center",
             because it blindly compared against the ORIGINAL
             (untranslated) document's alignment, which still had the
             SAME uncorrected pre-existing "center" bug. Confirmed
             directly: calling _fix_body_paragraph_center_alignment
             again on a real "Final Output" (post-review) document
             found 6 more paragraphs that had been reverted back to
             center by the reviewer.
  Issue 4:   a large Appendix table was actually split into 3 SEPARATE
             <w:tbl> elements (16/14/15 rows), each with its own
             "Clause Number | Field | Explanation" header row, with
             nothing but empty filler between them - confirmed
             directly in a real document's XML. Word/LibreOffice
             paginates each fragment as its own table, so a later
             fragment's header can start mid-page instead of only ever
             appearing at the true top of the logical table. Fixed by
             merging adjacent same-header-style tables into one real
             table and marking the surviving header row to repeat via
             the standard OOXML <w:tblHeader/> mechanism.

Run: python3 reviewer_and_table_merge_test.py
"""
import os
import sys

sys.path.insert(0, os.path.dirname(__file__))
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import aspose_test_pipeline as asp


def assert_(cond, label):
    if not cond:
        print("FAIL:", label)
        sys.exit(1)
    print("PASS:", label)


def _add_shaded_header_row(table, texts, fill="666666"):
    row = table.rows[0] if len(table.rows) == 1 and not table.rows[0].cells[0].text else table.add_row()
    for cell, text in zip(row.cells, texts):
        cell.text = text
        pPr = cell.paragraphs[0]._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        pPr.append(shd)
    return row


def run():
    print("=== Test 1 (the user's own reported case): reviewer no longer reverts a legitimate center->both fix ===")
    orig = Document()
    p_orig = orig.add_paragraph("A genuinely long clause paragraph that was wrongly center-aligned by Aspose here today.")
    pPr_o = p_orig._p.get_or_add_pPr()
    jc_o = OxmlElement("w:jc")
    jc_o.set(qn("w:val"), "center")
    pPr_o.append(jc_o)
    orig.save("/tmp/rt_orig_align.docx")

    trans = Document()
    p_trans = trans.add_paragraph("A genuinely long clause paragraph that was wrongly center-aligned by Aspose here today.")
    pPr_t = p_trans._p.get_or_add_pPr()
    jc_t = OxmlElement("w:jc")
    jc_t.set(qn("w:val"), "both")  # step 2 already correctly fixed this
    pPr_t.append(jc_t)
    trans.save("/tmp/rt_trans_align.docx")

    issues, out = asp.review_and_fix_translation("/tmp/rt_orig_align.docx", "/tmp/rt_trans_align.docx", "English", "/tmp/rt_out_align.docx")
    alignment_issues = [i for i in issues if "Alignment" in i["issue"]]
    assert_(len(alignment_issues) == 0, "No alignment issue is raised - the reviewer recognizes center->both as a real fix, not a regression")

    fixed_doc = Document(out)
    jc_after = fixed_doc.paragraphs[0]._p.find(qn("w:pPr")).find(qn("w:jc"))
    assert_(jc_after.get(qn("w:val")) == "both", "The paragraph stays 'both' (justify) - the reviewer did NOT revert it back to center")

    print("\n=== Test 2: a genuine, real alignment regression (not center-related) is still caught and fixed ===")
    orig2 = Document()
    p_orig2 = orig2.add_paragraph("Some paragraph.")
    pPr_o2 = p_orig2._p.get_or_add_pPr()
    jc_o2 = OxmlElement("w:jc")
    jc_o2.set(qn("w:val"), "left")
    pPr_o2.append(jc_o2)
    orig2.save("/tmp/rt_orig_align2.docx")

    trans2 = Document()
    p_trans2 = trans2.add_paragraph("Some translated paragraph.")
    pPr_t2 = p_trans2._p.get_or_add_pPr()
    jc_t2 = OxmlElement("w:jc")
    jc_t2.set(qn("w:val"), "right")  # a genuine, unexpected regression
    pPr_t2.append(jc_t2)
    trans2.save("/tmp/rt_trans_align2.docx")

    issues2, out2 = asp.review_and_fix_translation("/tmp/rt_orig_align2.docx", "/tmp/rt_trans_align2.docx", "English", "/tmp/rt_out_align2.docx")
    alignment_issues2 = [i for i in issues2 if "Alignment" in i["issue"]]
    assert_(len(alignment_issues2) == 1, "A genuine left->right regression (unrelated to the center-alignment fix) is still caught")
    fixed_doc2 = Document(out2)
    jc_after2 = fixed_doc2.paragraphs[0]._p.find(qn("w:pPr")).find(qn("w:jc"))
    assert_(jc_after2.get(qn("w:val")) == "left", "The genuine regression is still corrected back to the original ('left')")

    print("\n=== Test 3 (the user's own reported case): 3 fragmented appendix tables merge into 1, header repeats ===")
    doc3 = Document()
    t3a = doc3.add_table(rows=1, cols=3)
    _add_shaded_header_row(t3a, ["Clause Number", "Field", "Explanation"])
    row3a2 = t3a.add_row()
    for cell, text in zip(row3a2.cells, ["1", "Contract Type", "The type of the contract"]):
        cell.text = text

    doc3.add_paragraph("")  # empty filler between fragments, matching the real document

    t3b = doc3.add_table(rows=1, cols=3)
    _add_shaded_header_row(t3b, ["Clause Number", "Field", "Explanation"])
    row3b2 = t3b.add_row()
    for cell, text in zip(row3b2.cells, ["2", "Lessor", "The lessor's details"]):
        cell.text = text

    doc3.add_paragraph("")

    t3c = doc3.add_table(rows=1, cols=3)
    _add_shaded_header_row(t3c, ["Clause Number", "Field", "Explanation"])
    row3c2 = t3c.add_row()
    for cell, text in zip(row3c2.cells, ["3", "Tenant", "The tenant's details"]):
        cell.text = text

    tables_before = len(doc3.tables)
    merged3 = asp._merge_adjacent_header_tables(doc3)
    assert_(merged3 == 2, "Exactly 2 merges happened (3 fragments -> 1 real table)")
    tables_after = len(doc3.tables)
    assert_(tables_after == 1, "Only 1 table remains (was " + str(tables_before) + ")")

    final_table = doc3.tables[0]
    assert_(len(final_table.rows) == 4, "The merged table has 4 rows: 1 header + 3 real data rows (got " + str(len(final_table.rows)) + ")")
    header_texts = [c.text for c in final_table.rows[0].cells]
    assert_(header_texts == ["Clause Number", "Field", "Explanation"], "The surviving header row is the real header text")
    data_row_texts = [final_table.rows[i].cells[0].text for i in range(1, 4)]
    assert_(data_row_texts == ["1", "2", "3"], "All 3 real data rows survived, in order, across the merge")

    trPr = final_table.rows[0]._tr.find(qn("w:trPr"))
    tblHeader = trPr.find(qn("w:tblHeader")) if trPr is not None else None
    assert_(tblHeader is not None, "The surviving header row is marked with <w:tblHeader/> to repeat on every page")

    print("\n=== Test 4: tables with DIFFERENT header styles are never merged together ===")
    doc4 = Document()
    t4a = doc4.add_table(rows=1, cols=2)
    _add_shaded_header_row(t4a, ["Name", "Value"], fill="666666")
    doc4.add_paragraph("")
    t4b = doc4.add_table(rows=1, cols=2)
    _add_shaded_header_row(t4b, ["Field", "Amount"], fill="00A886")  # different shading color - a genuinely different table family
    merged4 = asp._merge_adjacent_header_tables(doc4)
    assert_(merged4 == 0, "Zero merges - different header shading colors mean these are genuinely different table families")

    print("\nALL TESTS PASSED")


if __name__ == "__main__":
    run()
