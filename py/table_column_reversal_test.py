"""
Regression test for _reverse_table_column_order, per an explicit
user-provided reference image showing the exact expected before/after:
an RTL-source table's ENTIRE column sequence (not just top-level
groups) is reversed end-to-end when translating into an LTR target
language - "Translation(T2,T1) | Data(D2,D1) | SR.No" (Arabic order)
becomes "SR.No | Data(D1,D2) | Translation(T1,T2)" (English order),
applied uniformly to header AND data rows, and merged/spanned cells
move as one whole block, keeping their own gridSpan value unchanged.

Run: python3 table_column_reversal_test.py
"""
import os
import sys

sys.path.insert(0, os.path.dirname(__file__))
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import aspose_test_pipeline as asp


def assert_(cond, label):
    if not cond:
        print("FAIL:", label)
        sys.exit(1)
    print("PASS:", label)


def _cell_texts(tr_el):
    out = []
    for tc in tr_el.findall(qn("w:tc")):
        ts = tc.findall(".//" + qn("w:t"))
        out.append("".join(t.text or "" for t in ts))
    return out


def run():
    print("=== Test 1 (THE EXACT REFERENCE IMAGE CASE): full column-sequence reversal, header + data row ===")
    doc = Document()
    t = doc.add_table(rows=2, cols=5)
    grid = t._tbl.find(qn("w:tblGrid"))
    cols = grid.findall(qn("w:gridCol"))
    for c, w in zip(cols, ["1000", "1000", "1000", "1000", "800"]):
        c.set(qn("w:w"), w)
    for cell, text in zip(t.rows[0].cells, ["T2", "T1", "D2", "D1", "SR"]):
        cell.text = text
    for cell, text in zip(t.rows[1].cells, ["T2_1", "T1_1", "D2_1", "D1_1", "1"]):
        cell.text = text

    fixed = asp._reverse_table_column_order(doc, "English")
    assert_(fixed == 2, "Both rows (header + data) were reversed")
    assert_([c.text for c in t.rows[0].cells] == ["SR", "D1", "D2", "T1", "T2"], "Header row matches the reference image's exact expected order")
    assert_([c.text for c in t.rows[1].cells] == ["1", "D1_1", "D2_1", "T1_1", "T2_1"], "Data row matches the reference image's exact expected order")
    grid_after = [c.get(qn("w:w")) for c in grid.findall(qn("w:gridCol"))]
    assert_(grid_after == ["800", "1000", "1000", "1000", "1000"], "Column widths (gridCol) follow the same reversal as the cells")

    print("\n=== Test 2: target language is ALSO RTL - no reversal happens (columns are already correct for an RTL reader) ===")
    doc2 = Document()
    t2 = doc2.add_table(rows=1, cols=3)
    for cell, text in zip(t2.rows[0].cells, ["A", "B", "C"]):
        cell.text = text
    fixed2 = asp._reverse_table_column_order(doc2, "Arabic")
    assert_(fixed2 == 0, "Zero rows touched - target is RTL, existing order is already correct")
    assert_([c.text for c in t2.rows[0].cells] == ["A", "B", "C"], "Row is completely unchanged")

    print("\n=== Test 3: a merged/spanned cell (real reference image case - 'Translation' spanning 2 sub-columns) moves as ONE block, keeping its own gridSpan ===")
    doc3 = Document()
    t3 = doc3.add_table(rows=1, cols=3)
    grid3 = t3._tbl.find(qn("w:tblGrid"))
    for c, w in zip(grid3.findall(qn("w:gridCol")), ["1000", "1000", "800"]):
        c.set(qn("w:w"), w)
    tr3 = t3.rows[0]._tr
    tcs3 = tr3.findall(qn("w:tc"))

    def _set(tc_el, text, span=None):
        for p in tc_el.findall(qn("w:p")):
            for r in p.findall(qn("w:r")):
                p.remove(r)
            run = OxmlElement("w:r")
            t_el = OxmlElement("w:t")
            t_el.text = text
            run.append(t_el)
            p.append(run)
        if span:
            tcPr = tc_el.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                tc_el.insert(0, tcPr)
            gs = OxmlElement("w:gridSpan")
            gs.set(qn("w:val"), str(span))
            tcPr.append(gs)

    tr3.remove(tcs3[1])
    _set(tcs3[0], "Translation-header", span=2)
    _set(tcs3[2], "SR-header")

    assert_(_cell_texts(tr3) == ["Translation-header", "SR-header"], "Setup: merged Translation header is first, SR-header is second")
    asp._reverse_table_column_order(doc3, "English")
    assert_(_cell_texts(tr3) == ["SR-header", "Translation-header"], "After reversal: SR-header is now first, the merged Translation header moved as one block to last")
    remaining = tr3.findall(qn("w:tc"))
    span_after = remaining[-1].find(".//" + qn("w:gridSpan"))
    assert_(span_after is not None and span_after.get(qn("w:val")) == "2", "The merged cell's own gridSpan (2) survives unchanged after moving")

    print("\n=== Test 4: applies uniformly to EVERY table in the document, not just one ===")
    doc4 = Document()
    t4a = doc4.add_table(rows=1, cols=2)
    for cell, text in zip(t4a.rows[0].cells, ["X2", "X1"]):
        cell.text = text
    t4b = doc4.add_table(rows=1, cols=2)
    for cell, text in zip(t4b.rows[0].cells, ["Y2", "Y1"]):
        cell.text = text
    fixed4 = asp._reverse_table_column_order(doc4, "English")
    assert_(fixed4 == 2, "One row fixed per table, across both tables")
    assert_([c.text for c in t4a.rows[0].cells] == ["X1", "X2"], "First table reversed")
    assert_([c.text for c in t4b.rows[0].cells] == ["Y1", "Y2"], "Second table ALSO reversed - no table is special-cased or skipped")

    print("\nALL TESTS PASSED")


if __name__ == "__main__":
    run()
