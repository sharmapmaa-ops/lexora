"""
Lexora - OCR SERVICE ROUTER
================================================================
Decides, per PDF, whether the OCR service needs Aspose.Words Cloud's
paid conversion or whether a free, local pdfplumber-based extraction
is enough - per the explicit rule this was built against:

    "table aur background color hone par aspose ka use karna he"
    (use Aspose only when the source has tables or background colors)

This module does NOT replace the existing vision-LLM OCR pipeline
(js/translation-offline.js's buildHybridDocxBlob, which is what the
OCR service currently always uses for scanned/photographed pages that
have no text layer at all). This router is for a DIFFERENT case: a
PDF that already HAS a text layer (so no vision model is needed to
read it), where the only question is whether it also needs Aspose's
table/formatting-aware reconstruction, or whether a much cheaper plain
text extraction is enough.

DECISION RULE (evidence-based, tested against a real 10-page lease
PDF in project history - see ocr_router_test.py):
    - Any page with >=1 pdfplumber-detected table -> Aspose
    - Any page with filled-rectangle (background color) coverage over
      BACKGROUND_COLOR_AREA_THRESHOLD of the page area -> Aspose
    - Otherwise -> lightweight (pdfplumber text extraction only)

Nothing about this is document-specific: no hardcoded page counts,
coordinates, or text. The thresholds below are the only tunable
constants, and they operate on whatever geometry pdfplumber reports
for ANY input PDF.
"""

import os

BACKGROUND_COLOR_AREA_THRESHOLD = 0.05  # 5% of a page's area covered by fills


class OcrRouterError(Exception):
    pass


# ---------------------------------------------------------------------------
# STEP 1: decide which engine a given PDF needs
# ---------------------------------------------------------------------------

def analyze_pdf_for_ocr_strategy(pdf_path):
    """Opens the PDF once and inspects EVERY page (not just the first,
    since a table or colored header can appear on any page) for:
      - tables (pdfplumber's own table-line/text-alignment detection)
      - background/fill color coverage (sum of filled rect areas as a
        fraction of the page area)

    Returns a dict:
        {
          "strategy": "aspose" | "lightweight",
          "reason": "<human-readable, page-specific evidence>",
          "page_count": int,
          "pages_with_tables": [1-based page numbers],
          "pages_with_background_color": [1-based page numbers],
          "table_count_total": int,
          "max_background_coverage_pct": float,
        }

    Never invents geometry - every field here comes directly from
    pdfplumber's own analysis of the actual file."""
    import pdfplumber

    pages_with_tables = []
    pages_with_bg = []
    table_count_total = 0
    max_bg_pct = 0.0

    with pdfplumber.open(pdf_path) as pdf:
        page_count = len(pdf.pages)
        for i, page in enumerate(pdf.pages):
            page_num = i + 1
            try:
                tables = page.find_tables()
            except Exception:  # noqa: BLE001
                tables = []  # a single page's table-detection failure shouldn't abort analysis
            if tables:
                pages_with_tables.append(page_num)
                table_count_total += len(tables)

            filled_rects = [r for r in (page.rects or []) if r.get("fill")]
            page_area = max(1.0, page.width * page.height)
            filled_area = sum(
                max(0.0, r["x1"] - r["x0"]) * max(0.0, r["bottom"] - r["top"])
                for r in filled_rects
            )
            bg_pct = filled_area / page_area
            max_bg_pct = max(max_bg_pct, bg_pct)
            if bg_pct >= BACKGROUND_COLOR_AREA_THRESHOLD:
                pages_with_bg.append(page_num)

    needs_aspose = bool(pages_with_tables or pages_with_bg)
    if needs_aspose:
        reasons = []
        if pages_with_tables:
            reasons.append(f"{len(pages_with_tables)} page(s) with tables ({table_count_total} total)")
        if pages_with_bg:
            reasons.append(
                f"{len(pages_with_bg)} page(s) with background color coverage "
                f">= {BACKGROUND_COLOR_AREA_THRESHOLD*100:.0f}% (max {max_bg_pct*100:.1f}%)"
            )
        reason = "Aspose needed: " + "; ".join(reasons)
    else:
        reason = (
            f"No tables and no significant background color found across all "
            f"{page_count} page(s) (max background coverage {max_bg_pct*100:.1f}%, "
            f"threshold {BACKGROUND_COLOR_AREA_THRESHOLD*100:.0f}%) - lightweight "
            f"pdfplumber extraction is sufficient."
        )

    return {
        "strategy": "aspose" if needs_aspose else "lightweight",
        "reason": reason,
        "page_count": page_count,
        "pages_with_tables": pages_with_tables,
        "pages_with_background_color": pages_with_bg,
        "table_count_total": table_count_total,
        "max_background_coverage_pct": round(max_bg_pct * 100, 2),
    }


# ---------------------------------------------------------------------------
# STEP 2a: lightweight path - pdfplumber text extraction, no tables/colors
# ---------------------------------------------------------------------------

def run_lightweight_ocr(pdf_path, output_path):
    """Plain text-layer extraction for PDFs with no tables and no
    meaningful background color - i.e. nothing Aspose's paid
    conversion would meaningfully add over just reading the text.
    Builds a simple DOCX: one paragraph per source line, a page break
    between source pages, page dimensions matched to the source (same
    approach already used and tested in source_reconstruction.py,
    reused here rather than reinvented).

    RTL: reuses source_reconstruction._fix_visual_order_to_logical
    (already tested against real Arabic PDF text in this project's own
    history - see test_rtl_fix.py) rather than re-solving the same
    visual-order-vs-logical-order problem a second time here."""
    import pdfplumber
    from docx import Document
    from docx.shared import Emu, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    try:
        from source_reconstruction import _fix_visual_order_to_logical, _detect_text_direction
    except Exception:  # noqa: BLE001
        _fix_visual_order_to_logical = lambda t: t  # noqa: E731 - degrade gracefully if that module ever moves
        _detect_text_direction = lambda t: None  # noqa: E731

    doc = Document()
    doc.styles["Normal"].font.size = Pt(10)

    pages_extracted = 0
    total_lines = 0
    with pdfplumber.open(pdf_path) as pdf:
        if pdf.pages:
            first = pdf.pages[0]
            section = doc.sections[0]
            section.page_width = Emu(int(first.width * 12700))
            section.page_height = Emu(int(first.height * 12700))

        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            lines = [ln for ln in text.split("\n") if ln.strip()]
            for line in lines:
                fixed_line = _fix_visual_order_to_logical(line)
                p = doc.add_paragraph()
                run = p.add_run(fixed_line)
                if _detect_text_direction(fixed_line) == "rtl":
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    pPr = p._p.get_or_add_pPr()
                    from docx.oxml import OxmlElement
                    pPr.append(OxmlElement("w:bidi"))
            total_lines += len(lines)
            pages_extracted += 1
            if i < len(pdf.pages) - 1:
                doc.add_page_break()

    doc.save(output_path)
    return {
        "engine": "lightweight_pdfplumber",
        "output_path": output_path,
        "pages_extracted": pages_extracted,
        "lines_extracted": total_lines,
        "aspose_calls": 0,
    }


# ---------------------------------------------------------------------------
# STEP 2b: Aspose path - reuses the EXISTING, already-working Aspose
# structure-only conversion from aspose_test_pipeline.py rather than
# reimplementing PDF->DOCX conversion a second time.
# ---------------------------------------------------------------------------

def run_aspose_ocr(pdf_path, output_path):
    """Table/background-color-preserving conversion via Aspose.Words
    Cloud. Deliberately reuses aspose_test_pipeline.run_structure_only_test
    (already implemented, already tested in this project's own history -
    see that function's docstring) instead of writing a second Aspose
    client here - one Aspose call site to maintain, not two."""
    import aspose_test_pipeline as asp

    result = asp.run_structure_only_test(pdf_path, output_path)
    result["engine"] = "aspose_words_cloud"
    return result


# ---------------------------------------------------------------------------
# STEP 3: top-level entry point
# ---------------------------------------------------------------------------

def run_ocr(pdf_path, output_path, force_strategy=None):
    """Full router: analyze -> pick engine -> run it -> return
    diagnostics (never raises for "Aspose not configured" - falls back
    to lightweight and says so explicitly, since an OCR request should
    still produce SOMETHING usable rather than fail outright just
    because paid credentials aren't set up yet).

    force_strategy: "aspose" | "lightweight" | None - lets a caller
    (e.g. an admin test route) override the automatic decision for
    comparison purposes; None (default) uses analyze_pdf_for_ocr_strategy.
    """
    analysis = analyze_pdf_for_ocr_strategy(pdf_path)
    strategy = force_strategy or analysis["strategy"]

    diagnostics = {"analysis": analysis, "requested_strategy": strategy}

    if strategy == "aspose":
        import aspose_test_pipeline as asp
        if not asp.is_configured():
            diagnostics["aspose_fallback_reason"] = (
                "Aspose was the chosen strategy (tables/background color present) "
                "but ASPOSE_CLIENT_ID/ASPOSE_CLIENT_SECRET are not configured in "
                "this environment - falling back to the lightweight extractor so "
                "the request still produces a usable document. Tables/background "
                "shading will NOT be preserved in this fallback output."
            )
            result = run_lightweight_ocr(pdf_path, output_path)
            result["strategy_used"] = "lightweight_fallback_from_aspose"
        else:
            try:
                result = run_aspose_ocr(pdf_path, output_path)
                result["strategy_used"] = "aspose"
            except Exception as err:  # noqa: BLE001
                diagnostics["aspose_error"] = str(err)
                diagnostics["aspose_fallback_reason"] = f"Aspose call failed: {err} - falling back to lightweight."
                result = run_lightweight_ocr(pdf_path, output_path)
                result["strategy_used"] = "lightweight_fallback_from_aspose_error"
    else:
        result = run_lightweight_ocr(pdf_path, output_path)
        result["strategy_used"] = "lightweight"

    diagnostics.update(result)
    return diagnostics
