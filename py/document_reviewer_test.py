"""
Regression test for the NEW document reviewer (Step 3), per explicit
direction: identifies every line/object in the original file, stores
their structural fingerprints in memory, compares against the
translated output to find real issues (formatting, style, background,
ordering, LTR/RTL), builds a full issue+solution list, then applies
every solution.

Uses the user's own two worked examples directly:
  Example 1: a row/paragraph missing its background color -> solution
             restores the original color.
  Example 2: source is one direction, target requires the other, but
             the translated line still shows the wrong direction ->
             solution corrects it to match the target language.

Run: python3 document_reviewer_test.py
"""
import os
import sys

sys.path.insert(0, os.path.dirname(__file__))
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import aspose_test_pipeline as asp


def assert_(cond, label):
    if not cond:
        print("FAIL:", label)
        sys.exit(1)
    print("PASS:", label)


def _add_shaded_paragraph(doc, text, fill):
    p = doc.add_paragraph(text)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)
    return p


def _add_rtl_paragraph(doc, text):
    p = doc.add_paragraph(text)
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:bidi"))
    return p


def run():
    print("=== Test 1 (the user's own Example 1): missing background color is detected and fixed ===")
    orig = Document()
    _add_shaded_paragraph(orig, "Row 3 content.", "0000FF")
    orig.save("/tmp/rt_orig1.docx")
    trans = Document()
    trans.add_paragraph("Contenu de la ligne 3.")  # shading lost
    trans.save("/tmp/rt_trans1.docx")

    issues, out = asp.review_and_fix_translation("/tmp/rt_orig1.docx", "/tmp/rt_trans1.docx", "French", "/tmp/rt_fixed1.docx")
    assert_(len(issues) == 1, "Exactly 1 issue found (got " + str(len(issues)) + ")")
    assert_("background" in issues[0]["issue"].lower(), "Issue correctly describes a background problem")
    assert_("0000FF" in issues[0]["solution"], "Solution names the correct original color to restore")

    fixed = Document(out)
    pPr = fixed.paragraphs[0]._p.find(qn("w:pPr"))
    shd = pPr.find(qn("w:shd")) if pPr is not None else None
    assert_(shd is not None and shd.get(qn("w:fill")) == "0000FF", "The fix was genuinely applied - background color restored to 0000FF")

    print("\n=== Test 2 (the user's own Example 2): direction mismatch is detected and fixed ===")
    orig2 = Document()
    _add_rtl_paragraph(orig2, "Arabic source line.")
    orig2.save("/tmp/rt_orig2.docx")
    trans2 = Document()
    _add_rtl_paragraph(trans2, "English text but still marked RTL (the bug).")
    trans2.save("/tmp/rt_trans2.docx")

    issues2, out2 = asp.review_and_fix_translation("/tmp/rt_orig2.docx", "/tmp/rt_trans2.docx", "English", "/tmp/rt_fixed2.docx")
    assert_(len(issues2) == 1, "Exactly 1 issue found for the direction mismatch")
    assert_("direction" in issues2[0]["issue"].lower(), "Issue correctly describes a direction problem")
    assert_("LTR" in issues2[0]["solution"], "Solution correctly says to set direction to LTR (target is English)")

    fixed2 = Document(out2)
    pPr2 = fixed2.paragraphs[0]._p.find(qn("w:pPr"))
    has_bidi = pPr2.find(qn("w:bidi")) is not None if pPr2 is not None else False
    assert_(has_bidi is False, "The fix was genuinely applied - bidi flag removed, paragraph is now LTR")

    print("\n=== Test 3: a document with NO issues reports zero issues (no false positives) ===")
    orig3 = Document()
    orig3.add_paragraph("Plain paragraph, nothing special.")
    orig3.save("/tmp/rt_orig3.docx")
    trans3 = Document()
    trans3.add_paragraph("Plain translated paragraph.")
    trans3.save("/tmp/rt_trans3.docx")
    issues3, _ = asp.review_and_fix_translation("/tmp/rt_orig3.docx", "/tmp/rt_trans3.docx", "French", "/tmp/rt_fixed3.docx")
    assert_(len(issues3) == 0, "No false-positive issues on a genuinely clean document (got " + str(len(issues3)) + ")")

    print("\n=== Test 4: paragraph count mismatch is flagged, not silently ignored ===")
    orig4 = Document()
    orig4.add_paragraph("One.")
    orig4.add_paragraph("Two.")
    orig4.save("/tmp/rt_orig4.docx")
    trans4 = Document()
    trans4.add_paragraph("Un.")  # missing the second paragraph entirely
    trans4.save("/tmp/rt_trans4.docx")
    issues4, _ = asp.review_and_fix_translation("/tmp/rt_orig4.docx", "/tmp/rt_trans4.docx", "French", "/tmp/rt_fixed4.docx")
    assert_(any("paragraph" in i["issue"].lower() and ("2" in i["issue"] or "1" in i["issue"]) for i in issues4),
            "A paragraph-count mismatch is flagged as an issue, not silently ignored")

    print("\n=== Test 5: server.py wiring for the new reviewer endpoint ===")
    server_src = open(os.path.join(os.path.dirname(__file__), "server.py")).read()
    assert_('"/api/translation/review"' in server_src, "A new /api/translation/review route is registered")

    print("\nALL TESTS PASSED")


if __name__ == "__main__":
    run()
