"""
Regression tests for a batch of 11 real reported issues, all confirmed
against a real uploaded document before implementing:
  1. Table width overflow - already covered by existing fixes
     (_fix_ambiguous_table_width / _fix_row_level_table_indent_override)
  2. Appendix table header labels mismatched vs their own column
     content, AND a genuine duplicate row - the duplicate is covered
     by the existing _remove_duplicate_table_rows; the label-mismatch
     is new here (_fix_mismatched_column_header_labels)
  3. Table row height should auto-fit, not use a fixed/exact height
  4. A numbered list's first item ("1. ...") merged onto the end of
     the preceding sentence, instead of starting its own line like
     every later item already does
  5. Table cell text: short -> centered, long -> left-aligned
  6. "Article <WordNumber>" headings should read "Article <digit>"
  7. The same field label duplicated twice within one row
  8. An Article heading's own spacing was a real outlier vs its
     siblings (visible as excess "white" space in the shaded bar), and
     a data cell was carrying its table's own header shading color
  9. Table cell vertical alignment should always be centered
  10. Table header/title rows should always be horizontally centered
  11. A cell containing "Label:   Value" (multi-space gap) should split
      into two separately-tab-stopped pieces within that one cell

Run: python3 eleven_issues_test.py
"""
import os
import sys

sys.path.insert(0, os.path.dirname(__file__))
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Twips
import aspose_test_pipeline as asp


def assert_(cond, label):
    if not cond:
        print("FAIL:", label)
        sys.exit(1)
    print("PASS:", label)


REAL_DOCX = "/home/claude/work/eleven_issues/final_output.docx"


def run():
    print("=== Issue 3: row-height auto-fit ===")
    if os.path.exists(REAL_DOCX):
        doc = Document(REAL_DOCX)
        fixed = asp._fix_table_row_height_autofit(doc)
        assert_(fixed > 0, "A real number of exact-height rows were reset to auto (got " + str(fixed) + ")")
        remaining = sum(1 for h in doc.element.body.iter(qn("w:trHeight")) if h.get(qn("w:hRule")) != "auto")
        assert_(remaining == 0, "No row retains hRule=exact after the fix")
    else:
        print("  (real document not present - skipped)")

    print("\n=== Issue 9: table cell vertical alignment always centered ===")
    doc9 = Document()
    t9 = doc9.add_table(rows=1, cols=2)
    t9.rows[0].cells[0].text = "A"
    tcPr9 = t9.rows[0].cells[1]._tc.get_or_add_tcPr()
    va9 = OxmlElement("w:vAlign")
    va9.set(qn("w:val"), "bottom")
    tcPr9.append(va9)
    fixed9 = asp._fix_table_vertical_alignment(doc9)
    assert_(fixed9 == 2, "Both cells (one with no vAlign, one with 'bottom') got fixed to center")
    for tc in t9._tbl.iter(qn("w:tc")):
        va = tc.find(qn("w:tcPr")).find(qn("w:vAlign"))
        assert_(va.get(qn("w:val")) == "center", "Cell vAlign is now 'center'")

    print("\n=== Issue 10: table header row always centered ===")
    doc10 = Document()
    t10 = doc10.add_table(rows=2, cols=2)
    hdr_p = t10.rows[0].cells[0].paragraphs[0]
    hdr_p.text = "Header"
    shd10 = OxmlElement("w:shd")
    shd10.set(qn("w:fill"), "666666")
    t10.rows[0].cells[0]._tc.get_or_add_tcPr().append(shd10)
    data_p = t10.rows[1].cells[0].paragraphs[0]
    data_p.text = "Data"
    fixed10 = asp._fix_table_header_alignment(doc10)
    assert_(fixed10 >= 1, "The header row's own paragraph got centered")
    hdr_jc = hdr_p._p.find(qn("w:pPr")).find(qn("w:jc"))
    assert_(hdr_jc.get(qn("w:val")) == "center", "Header paragraph is now centered")
    data_pPr = data_p._p.find(qn("w:pPr"))
    assert_(data_pPr is None or data_pPr.find(qn("w:jc")) is None, "The DATA row (not a header) is untouched by this fix")

    print("\n=== Issue 5: short table-cell text centered, long text left-aligned ===")
    doc5 = Document()
    t5 = doc5.add_table(rows=1, cols=2)
    t5.rows[0].cells[0].text = "Short"
    t5.rows[0].cells[1].text = "This is a genuinely long sentence of explanatory table-cell text that clearly exceeds the length threshold for a short label."
    fixed5 = asp._fix_table_cell_alignment_by_length(doc5)
    assert_(fixed5 == 2, "Both cells got an alignment set")
    short_jc = t5.rows[0].cells[0].paragraphs[0]._p.find(qn("w:pPr")).find(qn("w:jc"))
    long_jc = t5.rows[0].cells[1].paragraphs[0]._p.find(qn("w:pPr")).find(qn("w:jc"))
    assert_(short_jc.get(qn("w:val")) == "center", "Short cell text is centered")
    assert_(long_jc.get(qn("w:val")) == "left", "Long cell text is left-aligned")

    print("\n=== Issue 6 (THE REAL REPORTED CASE): 'Article Seven' -> 'Article 7' ===")
    doc6 = Document()
    p6 = doc6.add_paragraph("Article Seven: Contract Expiration")
    fixed6 = asp._convert_article_word_numbers_to_digits(doc6)
    assert_(fixed6 == 1, "One paragraph converted")
    assert_(p6.text == "Article 7: Contract Expiration", "Text now reads 'Article 7', matching the real reported requirement")

    print("\n=== Issue 6: real document - all 17 real Article headings converted ===")
    if os.path.exists(REAL_DOCX):
        doc6b = Document(REAL_DOCX)
        fixed6b = asp._convert_article_word_numbers_to_digits(doc6b)
        assert_(fixed6b == 22, "22 paragraphs touched (17 headings + some cross-references), matching the real confirmed count")
        remaining_word_articles = sum(1 for p in doc6b.paragraphs if re.search(r"Article (One|Two|Three|Four|Five|Six|Seven|Eight|Nine|Ten|Eleven|Twelve|Thirteen|Fourteen|Fifteen|Sixteen|Seventeen)\b", p.text))
        assert_(remaining_word_articles == 0, "No word-form Article numbering remains anywhere")
    else:
        print("  (real document not present - skipped)")

    print("\n=== Issue 7 (THE REAL REPORTED CASE): duplicate field label in the same row ===")
    doc7 = Document()
    t7 = doc7.add_table(rows=1, cols=3)
    t7.rows[0].cells[0].text = "Contract Sealing Date:"
    t7.rows[0].cells[1].text = "2025-06-17"
    t7.rows[0].cells[2].text = "Contract Sealing Date"
    fixed7 = asp._remove_duplicate_field_labels_in_row(doc7)
    assert_(fixed7 == 1, "One duplicate removed")
    assert_(t7.rows[0].cells[0].text == "Contract Sealing Date:", "The first (real) label is untouched")
    assert_(t7.rows[0].cells[2].text == "", "The duplicate (no-colon) label is emptied")
    assert_(t7.rows[0].cells[1].text == "2025-06-17", "The value cell in between is completely unaffected")

    print("\n=== Issue 7: a coincidentally-matching SHORT value is never treated as a duplicate label ===")
    doc7b = Document()
    t7b = doc7b.add_table(rows=1, cols=2)
    t7b.rows[0].cells[0].text = "12345"
    t7b.rows[0].cells[1].text = "12345"
    fixed7b = asp._remove_duplicate_field_labels_in_row(doc7b)
    assert_(fixed7b == 0, "Zero removals - short numeric matches are never treated as duplicate labels")

    print("\n=== Issue 8a: Article heading spacing normalized to the document's own real dominant value ===")
    if os.path.exists(REAL_DOCX):
        doc8 = Document(REAL_DOCX)
        asp._convert_article_word_numbers_to_digits(doc8)
        fixed8 = asp._fix_article_heading_spacing(doc8)
        assert_(fixed8 > 0, "A real number of outlier Article headings got their spacing normalized")
        before_values = set()
        for p in doc8.paragraphs:
            if re.match(r"^Article\s+\d+\s*:", p.text.strip()):
                pPr = p._p.find(qn("w:pPr"))
                spacing = pPr.find(qn("w:spacing")) if pPr is not None else None
                if spacing is not None:
                    before_values.add(spacing.get(qn("w:before")))
        assert_(len(before_values) == 1, "Every Article heading now shares the exact same before-spacing value")
    else:
        print("  (real document not present - skipped)")

    print("\n=== Issue 8b (THE REAL REPORTED CASE): data cell carrying its table's own header shading is reset ===")
    doc8b = Document()
    t8b = doc8b.add_table(rows=2, cols=1)
    hdr_shd = OxmlElement("w:shd")
    hdr_shd.set(qn("w:fill"), "666666")
    t8b.rows[0].cells[0]._tc.get_or_add_tcPr().append(hdr_shd)
    t8b.rows[0].cells[0].text = "Header"
    bad_shd = OxmlElement("w:shd")
    bad_shd.set(qn("w:fill"), "666666")  # same as header - the real bug
    t8b.rows[1].cells[0]._tc.get_or_add_tcPr().append(bad_shd)
    t8b.rows[1].cells[0].text = "1928550.00"
    fixed8b = asp._fix_anomalous_header_shading_on_data_cell(doc8b)
    assert_(fixed8b == 1, "The one anomalous data cell was fixed")
    assert_(bad_shd.get(qn("w:fill")) == "auto", "Its shading is now 'auto' (no shading), not the header's dark color")
    assert_(hdr_shd.get(qn("w:fill")) == "666666", "The real header row's own shading is completely untouched")

    print("\n=== Issue 8c (THE REAL FOLLOW-UP GAP, found via a systematic XML re-check, not just the visually-obvious case): PARAGRAPH-level shading (not cell-level) is also caught ===")
    doc8c = Document()
    t8c = doc8c.add_table(rows=2, cols=1)
    hdr_shd_c = OxmlElement("w:shd")
    hdr_shd_c.set(qn("w:fill"), "666666")
    t8c.rows[0].cells[0]._tc.get_or_add_tcPr().append(hdr_shd_c)
    t8c.rows[0].cells[0].text = "Header"
    data_p = t8c.rows[1].cells[0].paragraphs[0]
    data_p.text = "1928550.00"
    # the real confirmed bug: the CELL's own tcPr/shd is correctly
    # "auto" (untouched), but the PARAGRAPH inside carries the
    # header's dark fill - the original version of this fix never
    # checked here at all.
    para_shd = OxmlElement("w:shd")
    para_shd.set(qn("w:fill"), "666666")
    data_p._p.get_or_add_pPr().append(para_shd)
    fixed8c = asp._fix_anomalous_header_shading_on_data_cell(doc8c)
    assert_(fixed8c == 1, "The paragraph-level anomalous shading was found and fixed")
    assert_(para_shd.get(qn("w:fill")) == "auto", "The paragraph's shading is now 'auto'")
    assert_(hdr_shd_c.get(qn("w:fill")) == "666666", "The real header's own shading is still completely untouched")

    print("\n=== Issue 4 (THE REAL REPORTED CASE): first numbered-list item split onto its own line ===")
    doc4 = Document()
    t4 = doc4.add_table(rows=1, cols=1)
    t4.rows[0].cells[0].text = ("This field indicates the annual rent value agreed upon between the Lessor "
                                  "and the Tenant. 1. Cleaning.")
    fixed4 = asp._split_first_numbered_list_item_onto_new_line(doc4)
    assert_(fixed4 == 1, "One split happened")
    paras4 = [p for p in t4.rows[0].cells[0].paragraphs if p.text.strip()]
    assert_(len(paras4) == 2, "The cell now has 2 paragraphs")
    assert_(paras4[0].text.strip().endswith("Tenant."), "First paragraph ends at the real sentence boundary")
    assert_(paras4[1].text.strip().startswith("1. Cleaning"), "Second paragraph starts fresh with the list item")

    print("\n=== Issue 4: a bare mid-sentence digit that ISN'T list item 1 is never split (avoids false positives) ===")
    doc4b = Document()
    t4b = doc4b.add_table(rows=1, cols=1)
    t4b.rows[0].cells[0].text = "The meeting happened in 2024. 5 people attended."
    fixed4b = asp._split_first_numbered_list_item_onto_new_line(doc4b)
    assert_(fixed4b == 0, "Zero splits - '5 people' is not list-item '1.', so it's correctly left alone")

    print("\n=== Issue 2 (THE REAL REPORTED CASE): mismatched header labels swapped to match real column content ===")
    doc2 = Document()
    t2 = doc2.add_table(rows=4, cols=3)
    for cell, text in zip(t2.rows[0].cells, ["Explanation", "Field", "Item Number"]):
        cell.text = text
    rows_data = [
        ("1", "Contract Type", "The type of lease contract shall be one of the following options available to both parties involved."),
        ("2", "Contract Date", "This field indicates the date of documenting the lease contract by both parties to this agreement."),
        ("3", "Lease Start", "This is the actual start date of the lease and utilization of the rental unit by the tenant."),
    ]
    for row, (a, b, c) in zip(t2.rows[1:], rows_data):
        row.cells[0].text = a
        row.cells[1].text = b
        row.cells[2].text = c
    fixed2 = asp._fix_mismatched_column_header_labels(doc2)
    assert_(fixed2 == 1, "One table's header got corrected")
    header_texts = [c.text for c in t2.rows[0].cells]
    assert_(header_texts == ["Item Number", "Field", "Explanation"], "Header now reads Item Number/Field/Explanation, matching real column content")

    print("\n=== Issue 11 (THE REAL REPORTED EXAMPLE): 'Label:   Value' split into two tab-stopped pieces within one cell ===")
    doc11 = Document()
    t11 = doc11.add_table(rows=1, cols=1)
    cell11 = t11.rows[0].cells[0]
    cell11._tc.get_or_add_tcPr().find(qn("w:tcW")).set(qn("w:w"), "3000")
    cell11.text = "Nationality:     Indian"
    fixed11 = asp._split_label_value_pairs_within_cell(doc11)
    assert_(fixed11 == 1, "One pair split")
    parts = []
    for el in cell11._tc.findall(qn("w:p"))[0]:
        if el.tag == qn("w:r"):
            t_el = el.find(qn("w:t"))
            tab_el = el.find(qn("w:tab"))
            if tab_el is not None:
                parts.append("[TAB]")
            elif t_el is not None:
                parts.append(t_el.text)
    assert_(parts == ["Nationality:", "[TAB]", "Indian"], "Label, tab, value - three distinct pieces on one line, matching the real example exactly")
    tabs_el = cell11._tc.findall(qn("w:p"))[0].find(qn("w:pPr")).find(qn("w:tabs"))
    assert_(tabs_el.find(qn("w:tab")).get(qn("w:pos")) == "3000", "The tab stop lands at the cell's own real right edge")

    print("\n=== Issue 11: normal single-space 'Label: Value' text is never touched ===")
    doc11b = Document()
    t11b = doc11b.add_table(rows=1, cols=1)
    cell11b = t11b.rows[0].cells[0]
    cell11b._tc.get_or_add_tcPr().find(qn("w:tcW")).set(qn("w:w"), "3000")
    cell11b.text = "Nationality: Indian"
    fixed11b = asp._split_label_value_pairs_within_cell(doc11b)
    assert_(fixed11b == 0, "Zero splits - ordinary single-space text is not this bug")
    assert_(cell11b.text == "Nationality: Indian", "Text is completely unchanged")

    print("\n=== Issue 11: multiple pairs within one cell each become their own line ===")
    doc11c = Document()
    t11c = doc11c.add_table(rows=1, cols=1)
    cell11c = t11c.rows[0].cells[0]
    cell11c._tc.get_or_add_tcPr().find(qn("w:tcW")).set(qn("w:w"), "3000")
    cell11c.text = "Nationality:     Indian   Gender:     Male"
    fixed11c = asp._split_label_value_pairs_within_cell(doc11c)
    assert_(fixed11c == 2, "Two pairs split")
    real_paras = [p for p in cell11c._tc.findall(qn("w:p"))]
    assert_(len(real_paras) == 2, "Two separate paragraphs now exist - one line per pair")
    doc11c.save("/tmp/eleven_issues_multi_pair_check.docx")
    reopened = Document("/tmp/eleven_issues_multi_pair_check.docx")
    assert_(reopened.tables[0].rows[0].cells[0].text == "Nationality:\tIndian\nGender:\tMale", "Document saves and reopens with valid, correct structure")

    print("\nALL TESTS PASSED")


if __name__ == "__main__":
    import re
    run()
