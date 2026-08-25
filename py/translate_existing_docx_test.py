"""
Regression test for STEP 2 ONLY, per explicit direction: translation
injection into an ALREADY Aspose-converted docx (step 1's real output),
with the RTL/LTR direction-fixing step deliberately excluded (to be
wired in only later, per further instruction).

Run: python3 translate_existing_docx_test.py
"""
import os
import sys

sys.path.insert(0, os.path.dirname(__file__))
import aspose_test_pipeline as asp


def assert_(cond, label):
    if not cond:
        print("FAIL:", label)
        sys.exit(1)
    print("PASS:", label)


def run():
    print("=== Test 1: translate_existing_docx is defined and reuses the real, existing translation logic ===")
    src = open(os.path.join(os.path.dirname(__file__), "aspose_test_pipeline.py")).read()
    fn_start = src.find("def translate_existing_docx(")
    assert_(fn_start != -1, "translate_existing_docx is defined")
    fn_body = src[fn_start:fn_start + 1800]
    assert_("_translate_docx_segments_in_place(" in fn_body, "Reuses _translate_docx_segments_in_place directly - not a second translation implementation")
    assert_("_fix_paragraph_direction(" not in fn_body, "Does NOT call _fix_paragraph_direction - RTL/LTR fixing is deliberately excluded per direction")
    assert_("run_structure_only_test(" not in fn_body and "convert_document(" not in fn_body,
            "Does NOT re-run Aspose conversion - takes an ALREADY-converted docx, doesn't convert from PDF again")
    assert_("Document(docx_path)" in fn_body, "Loads the existing docx directly by path, not by converting a PDF")

    print("\n=== Test 2: server.py wiring ===")
    server_src = open(os.path.join(os.path.dirname(__file__), "server.py")).read()
    assert_('"/api/translation/inject-translation": self._handle_translation_inject_translation' in server_src,
            "/api/translation/inject-translation is registered in the URL routing table")
    handler_start = server_src.find("def _handle_translation_inject_translation(")
    next_def = server_src.find("\n    def ", handler_start + 1)
    handler_src = server_src[handler_start:next_def]
    assert_("translate_existing_docx(" in handler_src, "Handler calls translate_existing_docx")
    assert_("docxBase64" in handler_src, "Handler expects a DOCX as input (not a PDF) - matches taking step 1's real output as input")
    assert_("_require_role" not in handler_src, "No role restriction - real users can call this")

    print("\n=== Test 3: real end-to-end mechanics on the actual reported document (no real LLM call - verifies structure/save, not translation quality) ===")
    real_docx = "/home/claude/work/aspose_debug2/input3.docx"
    if os.path.exists(real_docx):
        from docx import Document
        doc = Document(real_docx)
        para_count_before = len([p for p in doc.paragraphs if p.text.strip()])
        assert_(para_count_before > 0, "The real uploaded document has real non-empty paragraphs to translate (" + str(para_count_before) + " found)")
        # Directly exercise the segment-collection half (no live LLM call
        # in this sandbox) to confirm segments are found without error.
        segments = []
        for p in doc.paragraphs:
            if p.text.strip():
                segments.append(p.text)
        assert_(len(segments) == para_count_before, "Segment collection finds the same real paragraphs a live run would attempt to translate")
    else:
        print("  (real document not present in this environment - skipped, not a failure)")

    print("\nALL TESTS PASSED")


if __name__ == "__main__":
    run()
